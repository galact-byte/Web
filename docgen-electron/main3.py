#!/usr/bin/env python
# -*- coding:utf-8 -*-
# author:qiao xiong
# datetime:2025/3/13 20:31
# name: main

from docx import Document
import json
from copy import deepcopy
import os
from dotenv import load_dotenv  # 新增：导入 dotenv
import sys


def load_env_variables():
    """加载 .env 文件中的环境变量到字典"""
    load_dotenv()  # 加载 .env 文件
    
    # 从环境变量中读取所有配置
    env_vars = {
        'CLIENT_COMPANY': os.getenv('CLIENT_COMPANY', ''),
        'EVAL_CONTACT': os.getenv('EVAL_CONTACT', ''),
        'EVAL_PHONE': os.getenv('EVAL_PHONE', ''),
    }
    
    print(f"已加载环境变量: {env_vars}")
    return env_vars


def copy_cell_00_format(table, dest_row, dest_col):
    """
    将表格中 (1,0) 位置的单元格全部格式复制到指定目标单元格（dest_row, dest_col）。
    注意：这将清除目标单元格原有内容。

    参数:
        table: 表格对象（Table）
        dest_row: 目标单元格所在的行索引（从 0 开始）
        dest_col: 目标单元格所在的列索引（从 0 开始）

    返回:
        复制格式后的目标单元格对象
    """
    src_cell = table.rows[1].cells[0]
    dest_cell = table.rows[dest_row].cells[dest_col]

    # 清除目标单元格原有内容
    dest_cell._tc.clear_content()

    # 复制源单元格的所有 XML 元素（格式、样式等）到目标单元格
    for element in src_cell._tc:
        dest_cell._tc.append(deepcopy(element))

    return dest_cell


def merge_system_information(project_info):
    """合并系统信息"""
    if '系统' in project_info and project_info['系统']:
        merge_info = project_info['系统']
        systems = [f"{system_info['系统名称']}（{system_info['系统级别']}）" for system_info in merge_info]
        project_info['合并系统信息'] = '、'.join(systems)
    else:
        project_info['合并系统信息'] = ''
    return project_info


def auto_add_project_info(project_info, env_vars):
    """自动添加项目信息的默认值，并合并环境变量"""
    # 合并环境变量到项目信息中
    project_info.update(env_vars)
    
    project_info = merge_system_information(project_info)

    # 如果用户没有填写"业务影响"，则使用默认值
    if '业务影响' not in project_info or not project_info['业务影响']:
        project_info['业务影响'] = '中断、停机等'

    # 如果用户没有填写"影响后果"，则使用默认值
    if '影响后果' not in project_info or not project_info['影响后果']:
        project_info['影响后果'] = '业务中断、服务器宕机等情况'

    # 用于清空某些字段
    project_info['取空'] = ""

    return project_info


def replace_doc_content(rules, project_info, output_subdir=None):
    """
    替换文档中的固定内容并保存
    :param rules: 替换规则(包含模板文档名称和替换规则)
    :param project_info: 项目信息
    :param output_subdir: 输出子目录名称（用于多项目时区分不同项目）
    :return: None
    """
    template_path = os.path.join("templates", rules['文档名称'])
    
    if not os.path.exists(template_path):
        print(f"警告: 模板文件不存在: {template_path}")
        return False
    
    try:
        doc = Document(template_path)
    except Exception as e:
        print(f"错误: 无法打开模板文件 {template_path}: {e}")
        return False

    # 新模式：docx 占位符模式
    if '文档替换规则' not in rules or not rules['文档替换规则']:
        replace_by_placeholder(doc, project_info)
    # 旧模式：坐标规则模式（完全保留）
    else:
        replace_by_rules(doc, rules, project_info)

    # 保存新文档
    output_dir = '授权文档输出目录'
    if not os.path.exists(output_dir):
        try:
            os.makedirs(output_dir)
            print(f"创建输出目录: {output_dir}")
        except Exception as e:
            print(f"创建输出目录失败: {e}")
            return False

    # 如果指定了子目录，则在输出目录下创建子目录
    if output_subdir:
        output_dir = os.path.join(output_dir, output_subdir)

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    output_filename = os.path.splitext(rules['文档名称'])[0] + ".docx"
    output_path = os.path.join(output_dir, output_filename)
    doc.save(output_path)
    print(f"已生成文档: {output_path}")
    return True


def process_single_project(project_info, rules_dir, env_vars):
    """处理单个项目的文档生成"""
    project_info = auto_add_project_info(project_info, env_vars)

    # 遍历rules文件夹下的所有json规则文件
    for rule_file in os.listdir(rules_dir):
        if rule_file.endswith('.json'):
            # 读取规则文件
            with open(os.path.join(rules_dir, rule_file), 'r', encoding='utf-8') as f:
                config = json.load(f)
            # 替换文档内容并保存
            replace_doc_content(config, project_info)


def process_multiple_projects(projects_data, rules_dir, env_vars):
    """处理多个项目的批量文档生成"""
    for idx, project_info in enumerate(projects_data['projects'], start=1):
        print(f"\n正在处理第 {idx} 个项目: {project_info.get('项目名称', '未命名项目')}")
        project_info = auto_add_project_info(project_info, env_vars)

        # 使用项目编号作为子目录名称
        project_number = project_info.get('项目编号', f'project_{idx}')
        project_name = project_info.get('项目名称', '未命名项目')
        output_subdir = f"{project_number}_{project_name}"

        # 遍历rules文件夹下的所有json规则文件
        for rule_file in os.listdir(rules_dir):
            if rule_file.endswith('.json'):
                # 读取规则文件
                with open(os.path.join(rules_dir, rule_file), 'r', encoding='utf-8') as f:
                    config = json.load(f)
                # 替换文档内容并保存到对应项目的子目录
                replace_doc_content(config, project_info, output_subdir)


def replace_by_placeholder(doc, mapping):
    """替换占位符，保留原有格式"""
    def replace_in_runs(paragraph, mapping):
        runs = paragraph.runs
        if not runs:
            return

        # 1. 拼全文 + run 索引映射
        full_text = ""
        index_map = []  # [(run, start, end)]
        cursor = 0

        for run in runs:
            start = cursor
            text = run.text or ""
            cursor += len(text)
            end = cursor
            full_text += text
            index_map.append((run, start, end))

        # 2. 逐个占位符处理（可以多个）
        for key, value in mapping.items():
            placeholder = f"{{{{{key}}}}}"
            
            if value is None:
                value = ""
            value = str(value)

            # 循环替换所有出现的占位符
            while placeholder in full_text:
                start_idx = full_text.index(placeholder)
                end_idx = start_idx + len(placeholder)

                # 3. 找占位符覆盖的 runs
                affected_runs = []
                for run, s, e in index_map:
                    if e <= start_idx or s >= end_idx:
                        continue
                    affected_runs.append((run, s, e))

                if not affected_runs:
                    break

                # 4. 参考样式 = 占位符第一个 run 的样式
                ref_run = affected_runs[0][0]

                # 5. 清空被影响的 runs，并计算新的文本内容
                for i, (run, s, e) in enumerate(affected_runs):
                    run_text = run.text or ""
                    
                    # 计算占位符在当前 run 中的相对位置
                    placeholder_start_in_run = max(0, start_idx - s)
                    placeholder_end_in_run = min(len(run_text), end_idx - s)
                    
                    if i == 0:
                        # 第一个 run: 保留占位符之前的部分
                        before_text = run_text[:placeholder_start_in_run]
                        if len(affected_runs) == 1:
                            # 只有一个 run 被影响：before + value + after
                            after_text = run_text[placeholder_end_in_run:]
                            run.text = before_text + value + after_text
                        else:
                            # 多个 runs：第一个 run 放 before + value
                            run.text = before_text + value
                    elif i == len(affected_runs) - 1:
                        # 最后一个 run: 保留占位符之后的部分
                        after_text = run_text[placeholder_end_in_run:]
                        run.text = after_text
                    else:
                        # 中间的 run: 完全清空
                        run.text = ""

                # 6. 保留原有样式
                target_run = affected_runs[0][0]
                target_run.font.name = ref_run.font.name
                target_run.font.size = ref_run.font.size
                target_run.font.bold = ref_run.font.bold
                target_run.font.italic = ref_run.font.italic
                target_run.font.underline = ref_run.font.underline

                # 7. 重建 full_text 和 index_map 以处理下一个匹配
                full_text = ""
                index_map = []
                cursor = 0
                for run in runs:
                    start = cursor
                    text = run.text or ""
                    cursor += len(text)
                    end = cursor
                    full_text += text
                    index_map.append((run, start, end))


    
    # 替换段落
    for paragraph in doc.paragraphs:
        replace_in_runs(paragraph, mapping)
    
    # 替换表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_runs(paragraph, mapping)


def replace_by_rules(doc, rules, project_info):
    """使用坐标规则替换（旧模式）"""
    for rule in rules['文档替换规则']:
        try:
            if len(rule) != 3:
                # 表格单元格替换
                table = doc.tables[rule[1]]
                dest_cell = copy_cell_00_format(table, rule[2], rule[3])
                for p in dest_cell.paragraphs:
                    p.text = project_info.get(rule[0], "")
            else:
                # 段落替换
                value = project_info.get(rule[0], "")
                if not isinstance(value, str):  # 确保值是字符串
                    value = str(value)
                doc.paragraphs[rule[1]].runs[rule[2]].text = value
        except (IndexError, AttributeError) as e:
            print(f"跳过无效规则{rule}: {e}")


import argparse

if __name__ == '__main__':
    # 首先加载环境变量
    env_vars = load_env_variables()
    
    parser = argparse.ArgumentParser(description='DocGen Pro Automation')
    parser.add_argument('--mode', type=str, choices=['1', '2'], help='Mode: 1 for Single Project, 2 for Multi Project')
    parser.add_argument('--config', type=str, help='Path to configuration JSON file (information.json or projects.json)')
    parser.add_argument('--rules', type=str, default='rules', help='Path to rules directory')
    
    # If no arguments provided, fall back to interactive mode (compatibility)
    if len(sys.argv) == 1:
        print("=" * 60)
        print("文档自动生成工具")
        print("=" * 60)
        print("\n请选择生成模式：")
        print("1. 单项目模式 (使用 information.json)")
        print("2. 多项目批量模式 (使用 projects.json)")
        print("=" * 60)

        mode = input("\n请输入模式编号 (1 或 2): ").strip()
        rules_dir = 'rules'
        
        if mode == '1':
            config_file = 'information.json'
        elif mode == '2':
            config_file = 'projects.json'
        else:
            print("无效的选择！请输入 1 或 2")
            exit(1)
    else:
        args = parser.parse_args()
        mode = args.mode
        config_file = args.config
        rules_dir = args.rules

    if mode == '1':
        # 单项目模式
        print(f"\n正在运行单项目模式... 配置文件: {config_file}")
        if not os.path.exists(config_file):
            print(f"错误: 找不到配置文件 {config_file}")
            exit(1)

        with open(config_file, 'r', encoding='utf-8') as f:
            project_info = json.load(f)

        process_single_project(project_info, rules_dir, env_vars)
        print("\n单项目文档生成完成！")

    elif mode == '2':
        # 多项目批量模式
        print(f"\n正在运行多项目批量模式... 配置文件: {config_file}")
        if not os.path.exists(config_file):
            print(f"错误: 找不到配置文件 {config_file}")
            exit(1)

        with open(config_file, 'r', encoding='utf-8') as f:
            projects_data = json.load(f)

        if 'projects' not in projects_data or len(projects_data['projects']) == 0:
            print("错误: projects.json 中没有项目数据！")
            exit(1)

        print(f"共找到 {len(projects_data['projects'])} 个项目")
        process_multiple_projects(projects_data, rules_dir, env_vars)
        print("\n多项目批量文档生成完成！")
    
    print("\n所有文档已生成到 '授权文档输出目录' 文件夹中")