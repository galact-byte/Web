# from docx import Document
#
# docPath = 'templates/20测评环境恢复确认表.docx'
# doc = Document(docPath)
# table = doc.tables[0]
#
# print("=== 表格单元格详细信息 ===")
# for i, row in enumerate(table.rows):
#     print(f"\n第 {i} 行:")
#     for j, cell in enumerate(row.cells):
#         # 显示单元格内容和是否与其他单元格合并
#         text = cell.text.strip().replace('\n', ' ')[:30]
#         print(f"  列{j}: '{text}'")


from docx import Document

path = 'templates/13自愿放弃验证测试声明（V4.5）.docx'
doc = Document(path)

print("第4段的所有run:")
for j, run in enumerate(doc.paragraphs[4].runs):
    print(f"run {j}: '{run.text}'")