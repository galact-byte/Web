#!/usr/bin/env python
# -*- coding:utf-8 -*-
# author:qiao xiong
# datetime:2025/3/13 17:20
# name: main

from docx import Document
import json
import os


def copy_cell_00_format(table, dest_row, dest_col):
    """
    将表格中 (1,0) 位置的单元格全部格式复制到指定目标单元格（dest_row, dest_col）。
    注意：这将清除目标单元格原有内容。

    参数:
        table: 表格对象（Table）
        dest_row: 目标单元格所在的行索引（从 0 开始）
        dest_col: 目标单元格所在的列索引（从 0 开始）

    返回:
        复制格式后的目标单元格对象
    """
    try:
        # 检查源单元格是否存在 (1,0)
        if len(table.rows) < 2 or len(table.rows[1].cells) < 1:
            print(f"警告：源单元格 (1,0) 不存在，使用目标单元格原格式")
            return table.rows[dest_row].cells[dest_col]

        src_cell = table.rows[1].cells[0]
        dest_cell = table.rows[dest_row].cells[dest_col]

        # 清除目标单元格原有内容
        dest_cell._tc.clear_content()

        # 复制源单元格的所有 XML 元素（格式、样式等）到目标单元格
        from copy import deepcopy
        for element in src_cell._tc:
            dest_cell._tc.append(deepcopy(element))

        return dest_cell
    except Exception as e:
        print(f"警告：复制单元格格式时发生异常: {e}")
        return table.rows[dest_row].cells[dest_col]


def merge_system_information(project_info):
    merge_info = project_info['系统']
    systems = [f"{system_info['系统名称']}（{system_info['系统级别']}）" for system_info in merge_info]
    project_info['合并系统信息'] = '、'.join(systems)
    return project_info


def auto_add_project_info(project_info):
    project_info = merge_system_information(project_info)
    project_info['业务影响'] = '中断、停机等'
    project_info['影响后果'] = '业务中断、服务器宕机等情况'
    project_info['取空'] = ""
    return project_info


def replace_doc_content(rules, project_info):
    """
    替换文档中的固定内容并保存
    :param rules: 替换规则(包含模板文档名称和替换规则)
    :param project_info: 项目信息
    :return: None
    """
    try:
        doc = Document("templates\\" + rules['文档名称'])
        print(f"成功打开文档: {rules['文档名称']}")
    except Exception as e:
        print(f"错误：无法打开文档 {rules['文档名称']}: {e}")
        return

    for i, rule in enumerate(rules['文档替换规则']):
        try:
            print(f"处理规则 {i}: {rule}")

            # 检查规则格式
            if len(rule) < 3:
                print(f"警告：规则 {i} 格式不正确，跳过: {rule}")
                continue

            # 检查项目信息中是否存在对应的键
            if rule[0] not in project_info:
                print(f"警告：项目信息中未找到键 '{rule[0]}'，跳过规则 {i}")
                continue

            if len(rule) == 4:  # 表格替换
                print(f"  表格替换: 表格{rule[1]}, 行{rule[2]}, 列{rule[3]}")

                # 检查表格索引是否存在
                if rule[1] >= len(doc.tables):
                    print(f"警告：表格索引 {rule[1]} 超出范围（总表格数：{len(doc.tables)}），跳过规则 {i}")
                    continue

                table = doc.tables[rule[1]]
                print(f"  表格有 {len(table.rows)} 行")

                # 检查行索引是否存在
                if rule[2] >= len(table.rows):
                    print(f"警告：表格 {rule[1]} 的行索引 {rule[2]} 超出范围（总行数：{len(table.rows)}），跳过规则 {i}")
                    continue

                # 检查列索引是否存在
                if rule[3] >= len(table.rows[rule[2]].cells):
                    print(
                        f"警告：表格 {rule[1]} 行 {rule[2]} 的列索引 {rule[3]} 超出范围（总列数：{len(table.rows[rule[2]].cells)}），跳过规则 {i}")
                    continue

                # 使用格式复制方法处理所有表格单元格
                dest_cell = copy_cell_00_format(table, dest_row=rule[2], dest_col=rule[3])

                if dest_cell is None:
                    print(f"错误：无法获取目标单元格，跳过规则 {i}")
                    continue

                # 在复制格式后的单元格中添加内容
                try:
                    if dest_cell.paragraphs:
                        # 在第一个段落中添加run
                        if dest_cell.paragraphs[0].runs:
                            # 如果有现有的runs，修改第一个run的文本
                            dest_cell.paragraphs[0].runs[0].text = project_info[rule[0]]
                            # 清空其他runs
                            for run in dest_cell.paragraphs[0].runs[1:]:
                                run.text = ""
                        else:
                            # 如果段落没有runs，添加一个新的run
                            dest_cell.paragraphs[0].add_run(project_info[rule[0]])
                        print(f"  成功设置单元格文本（保持格式）: '{project_info[rule[0]][:50]}...'")
                    else:
                        # 如果没有段落，创建一个新段落
                        p = dest_cell.add_paragraph()
                        p.add_run(project_info[rule[0]])
                        print(f"  创建新段落成功")

                except Exception as e:
                    print(f"错误：设置单元格内容时发生异常: {e}")
                    continue

            elif len(rule) == 3:  # 段落替换
                print(f"  段落替换: 段落{rule[1]}, run{rule[2]}")

                # 检查段落索引是否存在
                if rule[1] >= len(doc.paragraphs):
                    print(f"警告：段落索引 {rule[1]} 超出范围（总段落数：{len(doc.paragraphs)}），跳过规则 {i}")
                    continue

                paragraph = doc.paragraphs[rule[1]]

                # 检查 runs 索引是否存在
                if rule[2] >= len(paragraph.runs):
                    print(
                        f"警告：段落 {rule[1]} 的 run 索引 {rule[2]} 超出范围（总 runs 数：{len(paragraph.runs)}），跳过规则 {i}")
                    # 如果没有足够的 runs，可以尝试添加一个新的 run
                    if len(paragraph.runs) == 0:
                        # 如果段落没有 runs，添加一个
                        paragraph.add_run(project_info[rule[0]])
                        print(f"  添加新run成功")
                    else:
                        # 使用最后一个 run
                        paragraph.runs[-1].text = project_info[rule[0]]
                        print(f"  使用最后一个run成功")
                    continue

                # 正常替换
                paragraph.runs[rule[2]].text = project_info[rule[0]]
                print(f"  成功设置段落文本")
            else:
                print(f"警告：规则 {i} 长度不正确（应为 3 或 4），跳过: {rule}")

        except Exception as e:
            print(f"错误：处理规则 {i} 时发生异常: {e}")
            print(f"规则内容: {rule}")
            continue

    # 保存新文档到output文件夹
    output_dir = '授权文档输出目录'
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"创建输出目录: {output_dir}")

    output_filename = os.path.splitext(rules['文档名称'])[0] + ".docx"

    try:
        output_path = os.path.join(output_dir, output_filename)
        doc.save(output_path)
        print(f"成功保存文档: {output_path}")
    except Exception as e:
        print(f"错误：保存文档 {output_filename} 时发生异常: {e}")


if __name__ == '__main__':
    print("这个工具是为了快速的生成授权文档，13号文档目前有问题")

    # 检查必要的文件和文件夹是否存在
    if not os.path.exists('../information.json'):
        print("错误：找不到 information.json 文件")
        exit(1)

    if not os.path.exists('../rules'):
        print("错误：找不到 rules 文件夹")
        exit(1)

    if not os.path.exists('../templates'):
        print("错误：找不到 templates 文件夹")
        exit(1)

    try:
        # 从json文件读取项目信息
        with open('../information.json', 'r', encoding='utf-8') as f:
            project_info = json.load(f)
        project_info = auto_add_project_info(project_info)
        print("项目信息加载成功")
        print(f"处理的项目：{project_info.get('项目名称', '未知项目')}")
        print(f"合并系统信息：{project_info.get('合并系统信息', '无')}")
    except Exception as e:
        print(f"错误：读取 information.json 时发生异常: {e}")
        exit(1)

    # 遍历rules文件夹下的所有json规则文件
    rules_dir = '../rules'
    rule_files = [f for f in os.listdir(rules_dir) if f.endswith('.json')]

    if not rule_files:
        print("警告：rules 文件夹中没有找到 .json 规则文件")
        exit(1)

    print(f"找到 {len(rule_files)} 个规则文件")

    for rule_file in rule_files:
        print(f"\n{'=' * 50}")
        print(f"处理规则文件: {rule_file}")
        try:
            # 读取规则文件
            with open(os.path.join(rules_dir, rule_file), 'r', encoding='utf-8') as f:
                config = json.load(f)

            # 检查规则文件格式
            if '文档名称' not in config or '文档替换规则' not in config:
                print(f"警告：规则文件 {rule_file} 格式不正确，跳过")
                continue

            print(f"目标文档: {config['文档名称']}")
            print(f"替换规则数量: {len(config['文档替换规则'])}")

            # 替换文档内容并保存
            replace_doc_content(config, project_info)

        except Exception as e:
            print(f"错误：处理规则文件 {rule_file} 时发生异常: {e}")
            continue

    print(f"\n{'=' * 50}")
    print("处理完成！")