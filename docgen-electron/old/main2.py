#!/usr/bin/env python
# -*- coding:utf-8 -*-
# author:qiao xiong
# datetime:2025/3/13 17:20
# name: main

# !/usr/bin/env python
# -*- coding:utf-8 -*-
# author:qiao xiong
# datetime:2025/3/13 20:31
# name: main


from docx import Document
import json
from copy import deepcopy


def copy_cell_00_format(table, dest_row, dest_col):
    """
    将表格中 (0,0) 位置的单元格全部格式复制到指定目标单元格（dest_row, dest_col）。
    注意：这将清除目标单元格原有内容。

    参数:
        table: 表格对象（Table）
        dest_row: 目标单元格所在的行索引（从 0 开始）
        dest_col: 目标单元格所在的列索引（从 0 开始）

    返回:
        复制格式后的目标单元格对象
    """
    src_cell = table.rows[1].cells[0]
    dest_cell = table.rows[dest_row].cells[dest_col]

    # 清除目标单元格原有内容
    dest_cell._tc.clear_content()

    # 复制源单元格的所有 XML 元素（格式、样式等）到目标单元格
    for element in src_cell._tc:
        dest_cell._tc.append(deepcopy(element))

    return dest_cell


def merge_system_information(project_info):
    merge_info = project_info['系统']
    systems = [f"{system_info['系统名称']}（{system_info['系统级别']}）" for system_info in merge_info]
    project_info['合并系统信息'] = '、'.join(systems)
    return project_info


def auto_add_project_info(project_info):
    project_info = merge_system_information(project_info)

    if '业务影响' not in project_info or not project_info['业务影响']:
        project_info['业务影响'] = '中断、停机等'
    if '影响后果' not in project_info or not project_info['影响后果']:
        project_info['影响后果'] = '业务中断、服务器宕机等情况'
    project_info['取空'] = ""
    return project_info


def replace_doc_content(rules, project_info):
    """
    替换文档中的固定内容并保存
    :param rules: 替换规则(包含模板文档名称和替换规则)
    :param project_info: 项目信息
    :return: None
    """
    doc = Document("templates\\" + rules['文档名称'])

    for rule in rules['文档替换规则']:
        # doc.paragraphs[rule[1]].runs[rule[2]].text = project_info[rule[0]]

        if len(rule) != 3:

            table = doc.tables[rule[1]]
            dest_cell = copy_cell_00_format(table, dest_row=rule[2], dest_col=rule[3])
            # 修改目标单元格内容（注意：直接赋值 text 会重建段落，故这里修改已有 run 的文本）
            # dest_cell.paragraphs[0].runs[0].text = project_info[rule[0]]
            if dest_cell.paragraphs:
                for paragraph in dest_cell.paragraphs:
                    all_text = "".join(run.text for run in paragraph.runs)

                    # 替换文本
                    if all_text:
                        new_runs = paragraph.runs
                        new_runs[0].text = project_info[rule[0]]  # 修改第一个 run 的文本
                        for run in new_runs[1:]:  # 清空后续 run 的文本
                            run.text = ""
        else:
            doc.paragraphs[rule[1]].runs[rule[2]].text = project_info[rule[0]]

    # 保存新文档到output文件夹
    import os
    output_dir = '授权文档输出目录'
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    output_filename = os.path.splitext(rules['文档名称'])[0] + ".docx"
    doc.save(os.path.join(output_dir, output_filename))


if __name__ == '__main__':
    print("这个工具是为了快速的生成授权文档，13号文档目前有问题")

    # 从json文件读取项目信息
    with open('../information.json', 'r', encoding='utf-8') as f:
        project_info = json.load(f)
    project_info = auto_add_project_info(project_info)

    # 遍历rules文件夹下的所有json规则文件
    import os

    rules_dir = '../rules'
    for rule_file in os.listdir(rules_dir):
        if rule_file.endswith('.json'):
            # 读取规则文件
            with open(os.path.join(rules_dir, rule_file), 'r', encoding='utf-8') as f:
                config = json.load(f)
            # 替换文档内容并保存
            replace_doc_content(config, project_info)