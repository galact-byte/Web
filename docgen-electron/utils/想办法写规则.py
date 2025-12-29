#!/usr/bin/env python
# -*- coding:utf-8 -*-
# author:qiao xiong
# datetime:2025/3/13 17:52
# name: 01
#!/usr/bin/env python
# -*- coding:utf-8 -*-
# author:qiao xiong
# datetime:2025/3/13 17:21
# name: demo-csdn


from docx import Document

path = '17渗透测试授权书.docx'

doc = Document(path)
print(doc.paragraphs)
print(len(doc.paragraphs))


for i, paragraph in enumerate(doc.paragraphs, start=0):
    print(f"第 {i} 个段落：{paragraph.text}")

    for j, run in enumerate(paragraph.runs, start=0):
        print(f"    └─ 第 {j} 个 run：{run.text}")

doc.save("demo-csdn.docx")

# {
#     "文档名称": "20测评环境恢复确认表.docx",
#     "文档替换规则": [
#       ["项目编号",0,0,1],
# 		["项目名称",0,1,1],
# 		["合并系统信息",0,2,1]
#     ]
# }

# 第一个字段表示，要填充的信息
# 第一个数字是指第几段
# 第二个数字表示。第几个run
