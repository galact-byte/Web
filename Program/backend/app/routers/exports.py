"""
导出路由 - Excel和Word完结单导出
"""
import io
from typing import List
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session, joinedload
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, Border, Side, PatternFill
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.database import get_db
from app.models import User, Project, ProjectAssignment
from app.schemas import ExcelExportRequest
from app.services.auth import get_current_user

router = APIRouter(prefix="/api/exports", tags=["导出"])


# 业务类别简写映射
CATEGORY_SHORT = {
    "等保测评": "等保", "密码评估": "密评", "风险评估": "风评",
    "安全评估": "安评", "数据评估": "数评", "软件测试": "软测",
    "安全服务": "安服", "其他": "其他"
}

# 系统级别转数字
LEVEL_NUM = {
    "第一级": "1", "第二级": "2", "第三级": "3",
    "第四级": "4", "第五级": "5", "/": "/"
}

# 不需要系统级别和类型的业务类别
NO_LEVEL_CATEGORIES = ["风险评估", "安全评估", "数据评估", "软件测试", "安全服务"]


@router.post("/excel")
async def export_excel(
    request: ExcelExportRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """导出Excel格式的季度完结单"""
    projects = db.query(Project).options(
        joinedload(Project.systems),
        joinedload(Project.assignments).joinedload(ProjectAssignment.assignee)
    ).filter(Project.id.in_(request.project_ids)).all()
    
    if not projects:
        raise HTTPException(status_code=404, detail="未找到选中的项目")
    
    # 创建工作簿
    wb = Workbook()
    ws = wb.active
    ws.title = f"{request.year}年第{request.quarter}季度完结单"
    
    # 表头样式
    header_font = Font(bold=True, size=11)
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font_white = Font(bold=True, size=11, color="FFFFFF")
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin")
    )
    center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    
    # 写入表头
    headers = [
        "序号", "部门", "项目编号", "项目名称", "客户单位名称",
        "项目地点", "项目类型", "系统名称", "系统级别", "系统类型", "人员贡献率"
    ]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = header_font_white
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = center_align
    
    # 设置列宽
    widths = [6, 10, 18, 30, 25, 12, 8, 25, 8, 12, 35]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[chr(64 + i)].width = w
    
    # 写入数据
    row = 2
    seq = 1
    for project in projects:
        # 获取项目对应的分配信息
        assignment = project.assignments[0] if project.assignments else None
        department = assignment.department if assignment else request.department
        contribution = assignment.contribution if assignment else "/"
        
        category_short = CATEGORY_SHORT.get(project.business_category, project.business_category)
        
        if project.systems:
            for i, sys in enumerate(project.systems):
                # 特殊类别不需要系统级别和类型
                if project.business_category in NO_LEVEL_CATEGORIES:
                    level = "/"
                    sys_type = "/"
                else:
                    level = LEVEL_NUM.get(sys.system_level, sys.system_level)
                    sys_type = sys.system_type or "/"
                
                data = [
                    seq if i == 0 else "",
                    department if i == 0 else "",
                    project.project_code if i == 0 else "",
                    project.project_name if i == 0 else "",
                    project.client_name or "/" if i == 0 else "",
                    project.project_location or "/" if i == 0 else "",
                    category_short if i == 0 else "",
                    sys.system_name,
                    level,
                    sys_type,
                    contribution if i == 0 else ""
                ]
                
                for col, value in enumerate(data, 1):
                    cell = ws.cell(row=row, column=col, value=value)
                    cell.border = thin_border
                    cell.alignment = center_align
                
                row += 1
        else:
            # 无系统
            data = [
                seq, department, project.project_code, project.project_name,
                project.client_name or "/", project.project_location or "/",
                category_short, "/", "/", "/", contribution
            ]
            for col, value in enumerate(data, 1):
                cell = ws.cell(row=row, column=col, value=value)
                cell.border = thin_border
                cell.alignment = center_align
            row += 1
        
        seq += 1
    
    # 保存到内存
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    
    filename = f"{request.year}年第{request.quarter}季度项目完结单.xlsx"
    
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}"}
    )


@router.post("/word/{project_id}")
async def export_word(
    project_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """导出单个项目的Word完结单"""
    project = db.query(Project).options(
        joinedload(Project.systems),
        joinedload(Project.assignments).joinedload(ProjectAssignment.assignee),
        joinedload(Project.business_manager),
        joinedload(Project.implementation_manager)
    ).filter(Project.id == project_id).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="项目不存在")
    
    # 创建Word文档
    doc = Document()
    
    # 标题
    title = doc.add_heading(f"项目完结单", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息表格
    doc.add_heading("一、项目基本信息", level=1)
    
    basic_info = [
        ("项目编号", project.project_code),
        ("项目名称", project.project_name),
        ("被测单位名称", project.client_name or "/"),
        ("业务类别", project.business_category),
        ("项目地点", project.project_location or "/"),
        ("合同状态", project.contract_status),
        ("定级备案状态", project.filing_status),
        ("审批完成时间", project.approval_date or "/"),
        ("业务负责人", project.business_manager.display_name if project.business_manager else "/"),
        ("实施负责人", project.implementation_manager.display_name if project.implementation_manager else "/"),
    ]
    
    table1 = doc.add_table(rows=len(basic_info), cols=2)
    table1.style = 'Table Grid'
    for i, (label, value) in enumerate(basic_info):
        table1.rows[i].cells[0].text = label
        table1.rows[i].cells[1].text = str(value)
    
    # 系统信息
    doc.add_heading("二、系统信息", level=1)
    
    if project.systems:
        table2 = doc.add_table(rows=len(project.systems) + 1, cols=4)
        table2.style = 'Table Grid'
        
        headers = ["系统编号", "系统名称", "系统级别", "系统类型"]
        for i, h in enumerate(headers):
            table2.rows[0].cells[i].text = h
        
        for i, sys in enumerate(project.systems, 1):
            table2.rows[i].cells[0].text = sys.system_code or "/"
            table2.rows[i].cells[1].text = sys.system_name
            table2.rows[i].cells[2].text = sys.system_level
            table2.rows[i].cells[3].text = sys.system_type or "/"
    else:
        doc.add_paragraph("暂无系统信息")
    
    # 人员贡献
    doc.add_heading("三、人员贡献", level=1)
    
    if project.assignments:
        for a in project.assignments:
            p = doc.add_paragraph()
            p.add_run(f"姓名: ").bold = True
            p.add_run(a.assignee.display_name if a.assignee else "未知")
            p.add_run(f"  部门: ").bold = True
            p.add_run(a.department or "/")
            doc.add_paragraph(f"贡献率: {a.contribution or '/'}")
    else:
        doc.add_paragraph("暂无人员分配")
    
    # 保存到内存
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    filename = f"{project.project_code}完结单.docx"
    
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}"}
    )
