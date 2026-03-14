"""
导出路由 - Excel和Word完结单导出
"""
import io
import re
import zipfile
from copy import deepcopy
from typing import List
from pathlib import Path
from urllib.parse import quote
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session, selectinload
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, Border, Side, PatternFill
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.database import get_db
from app.models import User, UserRole, Project, ProjectAssignment
from app.schemas import ExcelExportRequest, WordExportRequest
from app.services.auth import get_current_user

router = APIRouter(prefix="/api/exports", tags=["导出"])


# 业务类别简写映射
CATEGORY_SHORT = {
    "等保测评": "等保", "密码评估": "密评", "风险评估": "风评",
    "安全评估": "安评", "数据评估": "数评", "软件测试": "软测",
    "安全服务": "安服", "其他": "其他"
}

# 系统级别转数字
LEVEL_NUM = {
    "第一级": "1", "第二级": "2", "第三级": "3",
    "第四级": "4", "第五级": "5", "/": "/"
}

# 不需要系统级别和类型的业务类别
NO_LEVEL_CATEGORIES = ["风险评估", "安全评估", "数据评估", "软件测试", "安全服务"]

# 模板目录
TEMPLATE_DIR = Path(__file__).resolve().parent.parent / "templates"

# 业务类别 → 模板文件映射
CATEGORY_TEMPLATE = {
    "等保测评": "等保完结单.docx",
    "密码评估": "等保完结单.docx",
    "风险评估": "风评完结单（包含风评和安评）正.docx",
    "安全评估": "风评完结单（包含风评和安评）正.docx",
    "数据评估": "等保完结单.docx",
    "软件测试": "软测完结单.docx",
    "安全服务": "安服完结单.docx",
    "其他": "等保完结单.docx",
}


def _sanitize_filename(name: str) -> str:
    """净化文件名，移除可能引发头注入的字符"""
    return re.sub(r'[^\w\u4e00-\u9fa5\-.]', '_', name)


@router.post("/excel")
def export_excel(
    request: ExcelExportRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """导出Excel格式的季度完结单"""
    # 员工只能导出自己被分配到的项目
    project_ids = request.project_ids
    if current_user.role == UserRole.employee:
        assigned_ids = {
            a.project_id for a in db.query(ProjectAssignment.project_id).filter(
                ProjectAssignment.assignee_id == current_user.id
            ).all()
        }
        project_ids = [pid for pid in project_ids if pid in assigned_ids]
        if not project_ids:
            raise HTTPException(status_code=403, detail="没有权限导出这些项目")

    projects = db.query(Project).options(
        selectinload(Project.systems),
        selectinload(Project.assignments).selectinload(ProjectAssignment.assignee)
    ).filter(Project.id.in_(project_ids)).all()

    if not projects:
        raise HTTPException(status_code=404, detail="未找到选中的项目")

    # 创建工作簿
    wb = Workbook()
    ws = wb.active
    ws.title = f"{request.year}年第{request.quarter}季度完结单"

    # 表头样式
    header_font = Font(bold=True, size=11)
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font_white = Font(bold=True, size=11, color="FFFFFF")
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin")
    )
    center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)

    # 写入表头
    headers = [
        "序号", "部门", "项目编号", "项目名称", "客户单位名称",
        "项目地点", "项目类型", "系统名称", "系统级别", "系统类型", "人员贡献率"
    ]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = header_font_white
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = center_align

    # 设置列宽
    widths = [6, 10, 18, 30, 25, 12, 8, 25, 8, 12, 35]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    # 写入数据
    row = 2
    seq = 1
    for project in projects:
        # 汇总所有分配记录的贡献率（修复：不再只取第一条）
        contributions = []
        for a in project.assignments:
            if a.contribution:
                contributions.append(a.contribution)
        contribution = "\n".join(contributions) if contributions else "/"
        # 部门使用经理在导出配置中指定的值
        department = request.department

        category_short = CATEGORY_SHORT.get(project.business_category, project.business_category)

        if project.systems:
            for i, sys in enumerate(project.systems):
                # 特殊类别不需要系统级别和类型
                if project.business_category in NO_LEVEL_CATEGORIES:
                    level = "/"
                    sys_type = "/"
                else:
                    level = LEVEL_NUM.get(sys.system_level, sys.system_level)
                    sys_type = sys.system_type or "/"

                data = [
                    seq if i == 0 else "",
                    department if i == 0 else "",
                    project.project_code if i == 0 else "",
                    project.project_name if i == 0 else "",
                    project.client_name or "/" if i == 0 else "",
                    project.project_location or "/" if i == 0 else "",
                    category_short if i == 0 else "",
                    sys.system_name,
                    level,
                    sys_type,
                    contribution if i == 0 else ""
                ]

                for col, value in enumerate(data, 1):
                    cell = ws.cell(row=row, column=col, value=value)
                    cell.border = thin_border
                    cell.alignment = center_align

                row += 1
        else:
            # 无系统
            data = [
                seq, department, project.project_code, project.project_name,
                project.client_name or "/", project.project_location or "/",
                category_short, "/", "/", "/", contribution
            ]
            for col, value in enumerate(data, 1):
                cell = ws.cell(row=row, column=col, value=value)
                cell.border = thin_border
                cell.alignment = center_align
            row += 1

        seq += 1

    # 保存到内存
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)

    filename = _sanitize_filename(f"{request.year}年第{request.quarter}季度项目完结单.xlsx")

    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}"}
    )


@router.post("/word/{project_id}")
def export_word(
    project_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """导出单个项目的Word完结单"""
    # 员工只能导出自己被分配到的项目
    if current_user.role == UserRole.employee:
        is_assigned = db.query(ProjectAssignment).filter(
            ProjectAssignment.project_id == project_id,
            ProjectAssignment.assignee_id == current_user.id
        ).first()
        if not is_assigned:
            raise HTTPException(status_code=403, detail="没有权限导出此项目")

    project = db.query(Project).options(
        selectinload(Project.systems),
        selectinload(Project.assignments).selectinload(ProjectAssignment.assignee),
        selectinload(Project.implementation_manager)
    ).filter(Project.id == project_id).first()

    if not project:
        raise HTTPException(status_code=404, detail="项目不存在")

    # 创建Word文档
    doc = Document()

    # 标题
    title = doc.add_heading(f"项目完结单", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 基本信息表格
    doc.add_heading("一、项目基本信息", level=1)

    basic_info = [
        ("项目编号", project.project_code),
        ("项目名称", project.project_name),
        ("被测单位名称", project.client_name or "/"),
        ("业务类别", project.business_category),
        ("项目地点", project.project_location or "/"),
        ("合同状态", project.contract_status),
        ("定级备案状态", project.filing_status),
        ("审批完成时间", project.approval_date or "/"),
        ("业务负责人", project.business_manager_name or "/"),
        ("实施负责人", project.implementation_manager.display_name if project.implementation_manager else "/"),
    ]

    table1 = doc.add_table(rows=len(basic_info), cols=2)
    table1.style = 'Table Grid'
    for i, (label, value) in enumerate(basic_info):
        table1.rows[i].cells[0].text = label
        table1.rows[i].cells[1].text = str(value)

    # 系统信息
    doc.add_heading("二、系统信息", level=1)

    if project.systems:
        table2 = doc.add_table(rows=len(project.systems) + 1, cols=4)
        table2.style = 'Table Grid'

        headers = ["系统编号", "系统名称", "系统级别", "系统类型"]
        for i, h in enumerate(headers):
            table2.rows[0].cells[i].text = h

        for i, sys in enumerate(project.systems, 1):
            table2.rows[i].cells[0].text = sys.system_code or "/"
            table2.rows[i].cells[1].text = sys.system_name
            table2.rows[i].cells[2].text = sys.system_level
            table2.rows[i].cells[3].text = sys.system_type or "/"
    else:
        doc.add_paragraph("暂无系统信息")

    # 人员贡献
    doc.add_heading("三、人员贡献", level=1)

    if project.assignments:
        for a in project.assignments:
            p = doc.add_paragraph()
            run_label1 = p.add_run("姓名: ")
            run_label1.bold = True
            p.add_run(a.assignee.display_name if a.assignee else "未知")
            run_label2 = p.add_run("  部门: ")
            run_label2.bold = True
            p.add_run(a.department or (a.assignee.department if a.assignee else None) or "/")
            doc.add_paragraph(f"贡献率: {a.contribution or '/'}")
    else:
        doc.add_paragraph("暂无人员分配")

    # 保存到内存
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    safe_code = _sanitize_filename(project.project_code)
    filename = f"{safe_code}完结单.docx"

    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}"}
    )


# ============ 基于模板的批量 Word 导出 ============

def _get_contribution_text(project) -> str:
    """获取项目的贡献率汇总文本"""
    contributions = []
    for a in project.assignments:
        if a.contribution:
            contributions.append(a.contribution)
    return "\n".join(contributions) if contributions else ""


def _set_cell_text(cell, text: str):
    """设置单元格文本，尽量保留原有段落/字体格式"""
    text = str(text)
    if cell.paragraphs and cell.paragraphs[0].runs:
        # 清空所有 run
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.text = ''
        cell.paragraphs[0].runs[0].text = text
    else:
        cell.text = text


def _fill_word_template(project, contribution_text: str):
    """基于模板填充项目数据，生成 Word 文档"""
    template_name = CATEGORY_TEMPLATE.get(project.business_category, "等保完结单.docx")
    template_path = TEMPLATE_DIR / template_name
    if not template_path.exists():
        template_path = TEMPLATE_DIR / "等保完结单.docx"

    doc = Document(str(template_path))
    table = doc.tables[0]
    total_cols = len(table.columns)

    # 项目名称列位置：9列模板→col 7，10列模板→col 8
    name_col = total_cols - 2

    # 填充项目基本信息
    _set_cell_text(table.cell(0, 1), project.project_code)
    _set_cell_text(table.cell(0, name_col), project.project_name)
    _set_cell_text(table.cell(1, 1), project.client_name or '')
    _set_cell_text(table.cell(3, name_col), project.project_location or '')

    # 更新编号段落
    if len(doc.paragraphs) > 1:
        doc.paragraphs[1].text = f"编号：{project.project_code}"

    # 系统数据行（row 7）
    data_row_idx = 7
    contribution_col = total_cols - 3

    if project.systems:
        # 多个系统时克隆数据行
        if len(project.systems) > 1:
            base_tr = table.rows[data_row_idx]._tr
            last_tr = base_tr
            for _ in range(len(project.systems) - 1):
                new_tr = deepcopy(base_tr)
                last_tr.addnext(new_tr)
                last_tr = new_tr

        # 填充每个系统行
        for i, sys in enumerate(project.systems):
            row = table.rows[data_row_idx + i]
            _set_cell_text(row.cells[0], sys.system_name)
            level = LEVEL_NUM.get(sys.system_level, sys.system_level)
            _set_cell_text(row.cells[2], str(level))
            _set_cell_text(row.cells[3], project.filing_status or '')
            _set_cell_text(row.cells[4], sys.system_type or '')
            _set_cell_text(row.cells[5], '是')
            if i == 0 and contribution_text:
                _set_cell_text(row.cells[contribution_col], contribution_text)
            elif i > 0:
                _set_cell_text(row.cells[contribution_col], '')
    else:
        # 无系统时清空数据行
        row = table.rows[data_row_idx]
        _set_cell_text(row.cells[0], '/')
        _set_cell_text(row.cells[2], '/')
        _set_cell_text(row.cells[3], project.filing_status or '/')
        _set_cell_text(row.cells[4], '/')
        _set_cell_text(row.cells[5], '/')
        if contribution_text:
            _set_cell_text(row.cells[contribution_col], contribution_text)

    return doc


@router.post("/word-batch")
def export_word_batch(
    request: WordExportRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """批量导出Word格式的完结单（多个项目返回ZIP压缩包）"""
    project_ids = request.project_ids
    if current_user.role == UserRole.employee:
        assigned_ids = {
            a.project_id for a in db.query(ProjectAssignment.project_id).filter(
                ProjectAssignment.assignee_id == current_user.id
            ).all()
        }
        project_ids = [pid for pid in project_ids if pid in assigned_ids]
        if not project_ids:
            raise HTTPException(status_code=403, detail="没有权限导出这些项目")

    projects = db.query(Project).options(
        selectinload(Project.systems),
        selectinload(Project.assignments).selectinload(ProjectAssignment.assignee),
        selectinload(Project.implementation_manager)
    ).filter(Project.id.in_(project_ids)).all()

    if not projects:
        raise HTTPException(status_code=404, detail="未找到选中的项目")

    # 单个项目直接返回 docx
    if len(projects) == 1:
        project = projects[0]
        contribution_text = _get_contribution_text(project)
        doc = _fill_word_template(project, contribution_text)
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        safe_code = _sanitize_filename(project.project_code)
        filename = f"{safe_code}完结单.docx"
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}"}
        )

    # 多个项目打包为 ZIP
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        for project in projects:
            contribution_text = _get_contribution_text(project)
            doc = _fill_word_template(project, contribution_text)
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            safe_code = _sanitize_filename(project.project_code)
            zf.writestr(f"{safe_code}完结单.docx", doc_buffer.getvalue())

    zip_buffer.seek(0)
    filename = _sanitize_filename("项目完结单.zip")
    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}"}
    )
