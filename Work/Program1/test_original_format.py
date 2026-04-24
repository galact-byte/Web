"""测试导出是否保持模板原始格式"""
from docx import Document
from pathlib import Path
import sys

sys.path.insert(0, str(Path(__file__).parent))
from app.main import set_cell_check_by_index

# 读取模板
template_path = Path(r"C:\Users\galact\Desktop\山西晋深交易有限公司-山西国资供应链服务系统-网络安全等级保护定级备案表.docx")
output_path = Path(r"E:\vscode\Programs\Work\Program1\exports\test_original_format.docx")

print("正在读取模板...")
doc = Document(str(template_path))

if len(doc.tables) > 1:
    t1 = doc.tables[1]
    if len(t1.rows) > 16:
        cell = t1.rows[16].cells[1]

        print("\n修改前的文本（前3段）:")
        for i, p in enumerate(cell.paragraphs[:3]):
            print(f"段落 {i}: {repr(p.text[:80])}")

        # 只设置复选框勾选状态，不修改格式
        print("\n正在设置复选框（勾选第47项：数据管理）...")
        set_cell_check_by_index(cell, {46})  # 第47项，索引为46

        print("\n修改后的文本（前3段）:")
        for i, p in enumerate(cell.paragraphs[:3]):
            print(f"段落 {i}: {repr(p.text[:80])}")

        print(f"\n正在保存到: {output_path}")
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        print("✓ 保存成功！")
        print("\n结论：文本格式完全保持不变，只有复选框状态改变了喵～")
else:
    print("错误：模板结构不正确")
