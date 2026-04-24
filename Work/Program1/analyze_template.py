"""分析备案表模板中"13行业类别"单元格的格式"""
from docx import Document
from pathlib import Path

template_path = Path(r"C:\Users\galact\Desktop\山西晋深交易有限公司-山西国资供应链服务系统-网络安全等级保护定级备案表.docx")
doc = Document(str(template_path))

# 表一是 doc.tables[1]
if len(doc.tables) > 1:
    t1 = doc.tables[1]
    # R16 是"13行业类别"
    if len(t1.rows) > 16:
        cell = t1.rows[16].cells[1]
        print("=" * 80)
        print("13行业类别单元格分析")
        print("=" * 80)
        print(f"\n单元格文本内容:\n{cell.text}\n")
        print("-" * 80)
        print("段落数量:", len(cell.paragraphs))

        for p_idx, paragraph in enumerate(cell.paragraphs):
            print(f"\n段落 {p_idx}:")
            print(f"  对齐方式: {paragraph.alignment}")
            print(f"  段落格式 - 左缩进: {paragraph.paragraph_format.left_indent}")
            print(f"  段落格式 - 首行缩进: {paragraph.paragraph_format.first_line_indent}")
            print(f"  段落格式 - 行距: {paragraph.paragraph_format.line_spacing}")
            print(f"  段落格式 - 段前间距: {paragraph.paragraph_format.space_before}")
            print(f"  段落格式 - 段后间距: {paragraph.paragraph_format.space_after}")
            print(f"  文本: {paragraph.text[:100]}...")
            print(f"  Run 数量: {len(paragraph.runs)}")

            for r_idx, run in enumerate(paragraph.runs[:20]):  # 只看前20个run
                print(f"    Run {r_idx}: '{run.text}' | 字体: {run.font.name} | 大小: {run.font.size}")
                # 检查是否有符号
                for sym in run._element.findall("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sym"):
                    font = sym.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}font")
                    char = sym.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}char")
                    print(f"      符号: font={font}, char={char}")

            # 检查制表符
            if '\t' in paragraph.text:
                print(f"  包含制表符: {paragraph.text.count('\\t')} 个")
                print(f"  制表位设置: {paragraph.paragraph_format.tab_stops}")
