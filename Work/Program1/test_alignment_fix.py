"""测试行业类别对齐修复功能"""
from docx import Document
from pathlib import Path
import sys
import re

# 添加 app 目录到路径
sys.path.insert(0, str(Path(__file__).parent))

from app.main import normalize_industry_cell_alignment

# 读取模板
template_path = Path(r"C:\Users\galact\Desktop\山西晋深交易有限公司-山西国资供应链服务系统-网络安全等级保护定级备案表.docx")
from datetime import datetime
output_path = Path(rf"E:\vscode\Programs\Work\Program1\exports\test_alignment_fixed_{datetime.now():%Y%m%d_%H%M%S}.docx")

print("正在读取模板...")
doc = Document(str(template_path))

if len(doc.tables) > 1:
    t1 = doc.tables[1]
    if len(t1.rows) > 16:
        cell = t1.rows[16].cells[1]

        print("\n修复前的文本（前3段）:")
        for i, p in enumerate(cell.paragraphs[:3]):
            print(f"段落 {i}: {repr(p.text[:80])}")

        print("\n正在应用对齐修复...")
        normalize_industry_cell_alignment(cell)

        print("\n修复后的文本（前3段）:")
        for i, p in enumerate(cell.paragraphs[:3]):
            print(f"段落 {i}: {repr(p.text[:80])}")
            # 检查是否包含制表符
            if '\t' in p.text:
                print(f"  ✓ 包含制表符: {p.text.count(chr(9))} 个")

        print(f"\n正在保存到: {output_path}")
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        print("✓ 保存成功！")
        print(f"\n请打开文件查看效果: {output_path}")
else:
    print("错误：模板结构不正确")
