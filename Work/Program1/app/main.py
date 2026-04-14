import io
import hmac
import hashlib
import json
import logging
import mimetypes
import os
import re
import secrets
import shutil
import smtplib
import sqlite3
import threading
import time
import uuid
import zipfile
import xml.etree.ElementTree as ET
from contextlib import asynccontextmanager
from datetime import date, datetime, timedelta
from email.mime.text import MIMEText
from pathlib import Path
from typing import Any
from urllib.parse import quote

from dotenv import load_dotenv
load_dotenv()

from docx import Document
from fastapi import Depends, FastAPI, File, Form, HTTPException, Query, Request, Response, UploadFile
from fastapi.openapi.docs import get_redoc_html, get_swagger_ui_html
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse, RedirectResponse
from starlette.background import BackgroundTask
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from openpyxl import Workbook, load_workbook
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfgen import canvas

    _reportlab_canvas_cls = getattr(canvas, "Canvas", None)
    _reportlab_canvas_methods = ("setFont", "drawString", "showPage", "save")
    if A4 is None or UnicodeCIDFont is None:
        raise RuntimeError("reportlab API incomplete")
    if not callable(getattr(pdfmetrics, "registerFont", None)):
        raise RuntimeError("reportlab pdfmetrics incomplete")
    if _reportlab_canvas_cls is None or any(not hasattr(_reportlab_canvas_cls, name) for name in _reportlab_canvas_methods):
        raise RuntimeError("reportlab canvas incomplete")

    HAS_REPORTLAB = True
except Exception:
    A4 = None
    pdfmetrics = None
    UnicodeCIDFont = None
    canvas = None
    HAS_REPORTLAB = False
from sqlalchemy import extract, func, inspect, or_, text
from sqlalchemy.orm import Session

from .db import SessionLocal, engine, get_db, init_db
from .models import (
    Attachment,
    DeleteRequest,
    AuthSession,
    KnowledgeDocument,
    KnowledgeDocumentVersion,
    KnowledgeDownloadLog,
    KnowledgePin,
    OrganizationCollectionLink,
    Organization,
    OrganizationHistory,
    OrganizationSubmission,
    Report,
    ReportTemplate,
    ReportTemplateVersion,
    ReviewRecord,
    SystemHistory,
    SystemInfo,
    UserAccount,
    WorkflowAction,
    WorkflowConfig,
    WorkflowInstance,
    WorkflowReminder,
    WorkflowStepRule,
)
from .schemas import (
    OrganizationCreate,
    OrganizationUpdate,
    ReportEdit,
    SystemCreate,
    SystemUpdate,
    WorkflowConfigUpdate,
)
from .services.reporting import (
    export_report_docx,
    export_report_docx_with_template,
    export_report_pdf,
    generate_report_payload,
    sanitize_template_docx_content,
)
from .validators import (
    is_placeholder_value,
    validate_credit_code_format_only,
    validate_email,
    validate_mobile_phone,
    validate_office_phone,
)

BASE_DIR = Path(__file__).resolve().parent.parent
UPLOAD_DIR = BASE_DIR / "uploads"
EXPORT_DIR = BASE_DIR / "exports"
BACKUP_DIR = EXPORT_DIR / "backups"
LOCAL_OFFICIAL_TEMPLATE_DIR = BASE_DIR / "template_docs"
DESKTOP_FILING_TEMPLATE_PATH = Path(
    r"C:\Users\galact\Desktop\山西晋深交易有限公司-山西国资供应链服务系统-网络安全等级保护定级备案表.docx"
)
MAX_ORG_ATTACHMENT = 100 * 1024 * 1024
MAX_SYS_ATTACHMENT = 200 * 1024 * 1024
MAX_KNOWLEDGE_FILE = 300 * 1024 * 1024
MAX_BACKUP_FILE = 1024 * 1024 * 1024
MAX_BACKUP_ZIP_ENTRIES = max(1, int(os.getenv("MAX_BACKUP_ZIP_ENTRIES", "5000")))
MAX_BACKUP_UNCOMPRESSED = max(
    MAX_BACKUP_FILE,
    int(os.getenv("MAX_BACKUP_UNCOMPRESSED", str(MAX_BACKUP_FILE * 2))),
)
DEFAULT_WORKFLOW_STEPS = ["信息收集", "信息审核", "报告生成", "报告审核", "报告终稿", "归档"]
VALID_REPORT_TYPES = {"filing_form", "grading_report", "expert_review_form"}
LOCAL_OFFICIAL_TEMPLATE_RULES = [
    {
        "pattern": "01-*.docx",
        "report_type": "filing_form",
        "template_name": "官方备案表模板",
        "category": "官方模板",
        "default_title": "网络安全等级保护备案表",
    },
    {
        "pattern": "02-*.docx",
        "report_type": "grading_report",
        "template_name": "官方定级报告模板",
        "category": "官方模板",
        "default_title": "网络安全等级保护定级报告",
    },
    {
        "pattern": "03-*.docx",
        "report_type": "expert_review_form",
        "template_name": "官方专家评审意见表模板",
        "category": "官方模板",
        "default_title": "专家评审意见表",
    },
]
PUBLIC_BASE_URL = os.getenv("PUBLIC_BASE_URL", "http://127.0.0.1:8011").rstrip("/")
STRICT_AUTH = str(os.getenv("STRICT_AUTH", "0")).strip().lower() in {"1", "true", "yes", "on", "y"}
AUTH_SESSION_HOURS = int(os.getenv("AUTH_SESSION_HOURS", "12"))
SMTP_HOST = (os.getenv("SMTP_HOST") or "").strip()
SMTP_PORT = int(os.getenv("SMTP_PORT", "25"))
SMTP_USER = (os.getenv("SMTP_USER") or "").strip()
SMTP_PASS = os.getenv("SMTP_PASS", "")
SMTP_FROM = (os.getenv("SMTP_FROM") or "").strip()
logger = logging.getLogger(__name__)


def env_to_bool(value: str | None, default: bool) -> bool:
    if value is None:
        return default
    return value.strip().lower() in {"1", "true", "yes", "on", "y"}


def parse_optional_bool(value: Any, field_name: str) -> bool | None:
    if value is None:
        return None
    if isinstance(value, bool):
        return value
    if isinstance(value, int) and value in (0, 1):
        return bool(value)
    if isinstance(value, str):
        text = value.strip().lower()
        if text in {"1", "true", "yes", "on", "y"}:
            return True
        if text in {"0", "false", "no", "off", "n"}:
            return False
    raise HTTPException(status_code=400, detail=f"{field_name} 仅支持布尔值 true/false。")


SMTP_TLS = env_to_bool(os.getenv("SMTP_TLS"), False)
SMTP_SSL = env_to_bool(os.getenv("SMTP_SSL"), False)


APP_ENV = os.getenv("APP_ENV", "development").strip().lower()
ENABLE_API_DOCS = env_to_bool(os.getenv("ENABLE_API_DOCS"), False)
APP_LITE_MODE = env_to_bool(os.getenv("APP_LITE_MODE"), False)
API_AUTH_REQUIRED = env_to_bool(os.getenv("API_AUTH_REQUIRED"), not APP_LITE_MODE)
FORCE_HTTPS = env_to_bool(os.getenv("FORCE_HTTPS"), False)
LOGIN_MAX_FAILS = max(1, int(os.getenv("LOGIN_MAX_FAILS", "5")))
LOGIN_FAIL_WINDOW_MINUTES = max(1, int(os.getenv("LOGIN_FAIL_WINDOW_MINUTES", "10")))
LOGIN_LOCK_MINUTES = max(1, int(os.getenv("LOGIN_LOCK_MINUTES", "15")))
PASSWORD_PBKDF2_ITERATIONS = max(200000, int(os.getenv("PASSWORD_PBKDF2_ITERATIONS", "390000")))
WORKFLOW_EMAIL_DEFAULT_TO = (os.getenv("WORKFLOW_EMAIL_DEFAULT_TO") or "").strip()
SCRIPT_TAG_RE = re.compile(r"<\s*/?\s*[a-zA-Z!]|on\w+\s*=|javascript\s*:", re.IGNORECASE)
LOGIN_ATTEMPTS: dict[str, dict[str, Any]] = {}
LOGIN_ATTEMPTS_LOCK = threading.Lock()
DEFAULT_BOOTSTRAP_ACCOUNTS = (
    ("admin", "DEFAULT_ADMIN_PASSWORD", "admin", "admin123"),
    ("tester", "DEFAULT_TESTER_PASSWORD", "evaluator", "tester123"),
    ("leader", "DEFAULT_LEADER_PASSWORD", "reviewer", "leader123"),
)


def escape_like(value: str) -> str:
    """转义 SQL LIKE 通配符，防止用户输入中的 % 和 _ 被当作通配符。"""
    return value.replace("\\", "\\\\").replace("%", "\\%").replace("_", "\\_")


def _cleanup_export_file(path: Path) -> None:
    """后台任务：响应完成后删除导出的临时文件。"""
    # Windows 下文件句柄释放可能存在短暂延迟，重试可避免临时导出文件残留。
    for _ in range(20):
        try:
            if not path.exists():
                return
            path.unlink()
            return
        except OSError:
            time.sleep(0.1)


@asynccontextmanager
async def app_lifespan(_: FastAPI):
    run_startup_tasks()
    yield


app = FastAPI(
    title="定级备案管理系统",
    version="1.2.0",
    docs_url=None,
    redoc_url=None,
    openapi_url=None,
    lifespan=app_lifespan,
)
app.mount("/static", StaticFiles(directory=str(BASE_DIR / "app" / "static")), name="static")
templates = Jinja2Templates(directory=str(BASE_DIR / "app" / "templates"))
templates.env.globals["show_api_docs"] = ENABLE_API_DOCS
templates.env.globals["lite_mode"] = APP_LITE_MODE


if APP_LITE_MODE:

    @app.middleware("http")
    async def lite_mode_guard(request: Request, call_next):  # type: ignore[override]
        path = request.url.path
        blocked_prefixes = [
            "/workflow",
            "/templates",
            "/knowledge",
            "/backup",
            "/login",
            "/register",
            "/api/workflow",
            "/api/templates",
            "/api/knowledge",
            "/api/backup",
            "/api/auth",
        ]
        if any(path.startswith(prefix) for prefix in blocked_prefixes):
            if path.startswith("/api/"):
                return JSONResponse(status_code=404, content={"detail": "精简交付版未开放该功能。"})
            return HTMLResponse(status_code=404, content="精简交付版未开放该页面。")
        return await call_next(request)


def ensure_dirs() -> None:
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    (UPLOAD_DIR / "attachments").mkdir(parents=True, exist_ok=True)
    (UPLOAD_DIR / "knowledge").mkdir(parents=True, exist_ok=True)
    (UPLOAD_DIR / "templates").mkdir(parents=True, exist_ok=True)
    LOCAL_OFFICIAL_TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)


def resolve_sqlite_db_file() -> Path:
    if engine.url.get_backend_name() != "sqlite":
        raise HTTPException(status_code=400, detail="当前仅支持 SQLite 数据库备份与恢复。")
    db_name = (engine.url.database or "").strip()
    if not db_name or db_name == ":memory:":
        raise HTTPException(status_code=400, detail="内存数据库不支持备份与恢复。")
    db_file = Path(db_name)
    if not db_file.is_absolute():
        db_file = (BASE_DIR / db_file).resolve()
    return db_file


def assert_safe_backup_file_name(file_name: str) -> str:
    clean = Path(file_name or "").name
    if clean != file_name:
        raise HTTPException(status_code=400, detail="备份文件名非法。")
    if not clean.endswith(".zip"):
        raise HTTPException(status_code=400, detail="仅支持 .zip 备份文件。")
    return clean


def safe_extract_zip(zf: zipfile.ZipFile, target_dir: Path) -> None:
    file_count = 0
    total_uncompressed = 0
    for member in zf.infolist():
        member_path = Path(member.filename)
        if member_path.is_absolute() or ".." in member_path.parts:
            raise HTTPException(status_code=400, detail="备份压缩包包含非法路径。")
        if member.is_dir():
            continue
        file_count += 1
        if file_count > MAX_BACKUP_ZIP_ENTRIES:
            raise HTTPException(status_code=400, detail="备份压缩包文件数量超限，拒绝解压。")
        file_size = int(member.file_size or 0)
        if file_size < 0:
            raise HTTPException(status_code=400, detail="备份压缩包文件大小非法。")
        total_uncompressed += file_size
        if total_uncompressed > MAX_BACKUP_UNCOMPRESSED:
            raise HTTPException(status_code=400, detail="备份压缩包解压后总大小超限，拒绝解压。")
    zf.extractall(target_dir)


def create_backup_archive(actor: str, trigger: str) -> Path:
    ensure_dirs()
    db_file = resolve_sqlite_db_file()
    if not db_file.exists():
        raise HTTPException(status_code=404, detail="数据库文件不存在，无法备份。")

    now = datetime.now()
    backup_name = f"backup_{now.strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}.zip"
    backup_path = BACKUP_DIR / backup_name
    temp_db = BACKUP_DIR / f".tmp_{uuid.uuid4().hex}.db"

    src_conn = sqlite3.connect(str(db_file))
    dst_conn = sqlite3.connect(str(temp_db))
    try:
        with dst_conn:
            src_conn.backup(dst_conn)
    finally:
        src_conn.close()
        dst_conn.close()

    manifest = {
        "created_at": now.isoformat(),
        "actor": actor,
        "trigger": trigger,
        "database_file": db_file.name,
        "app_version": app.version,
    }

    try:
        with zipfile.ZipFile(backup_path, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
            zf.write(temp_db, arcname="data/app.db")
            if UPLOAD_DIR.exists():
                for file_path in UPLOAD_DIR.rglob("*"):
                    if not file_path.is_file():
                        continue
                    arcname = (Path("uploads") / file_path.relative_to(UPLOAD_DIR)).as_posix()
                    zf.write(file_path, arcname=arcname)
            zf.writestr("manifest.json", json.dumps(manifest, ensure_ascii=False, indent=2))
    finally:
        if temp_db.exists():
            temp_db.unlink()
    return backup_path


def hash_password(password: str) -> str:
    salt = secrets.token_hex(16)
    digest = hashlib.pbkdf2_hmac(
        "sha256",
        password.encode("utf-8"),
        salt.encode("utf-8"),
        PASSWORD_PBKDF2_ITERATIONS,
    ).hex()
    return f"pbkdf2_sha256${PASSWORD_PBKDF2_ITERATIONS}${salt}${digest}"


def verify_password(password: str, stored_hash: str) -> tuple[bool, bool]:
    stored = (stored_hash or "").strip()
    if stored.startswith("pbkdf2_sha256$"):
        parts = stored.split("$")
        if len(parts) != 4:
            return False, False
        _, iter_str, salt, digest_hex = parts
        try:
            iterations = int(iter_str)
        except ValueError:
            return False, False
        calc = hashlib.pbkdf2_hmac(
            "sha256",
            password.encode("utf-8"),
            salt.encode("utf-8"),
            iterations,
        ).hex()
        return hmac.compare_digest(calc, digest_hex), False
    legacy = hashlib.sha256(password.encode("utf-8")).hexdigest()
    return hmac.compare_digest(legacy, stored), True


def assert_safe_text(value: Any, field_name: str) -> None:
    if not isinstance(value, str):
        return
    text = value.strip()
    if not text:
        return
    if SCRIPT_TAG_RE.search(text):
        raise HTTPException(status_code=400, detail=f"{field_name} 包含潜在危险脚本内容，请移除后重试。")


def assert_safe_payload(payload: Any, field_name: str) -> None:
    if isinstance(payload, dict):
        for key, val in payload.items():
            child_name = f"{field_name}.{key}"
            assert_safe_payload(val, child_name)
        return
    if isinstance(payload, list):
        for idx, val in enumerate(payload):
            child_name = f"{field_name}[{idx}]"
            assert_safe_payload(val, child_name)
        return
    assert_safe_text(payload, field_name)


def _login_attempt_key(username: str, request: Request) -> str:
    client_host = request.client.host if request.client and request.client.host else "unknown"
    return f"{client_host}:{(username or '').strip().lower()}"


def check_login_locked(key: str) -> None:
    with LOGIN_ATTEMPTS_LOCK:
        item = LOGIN_ATTEMPTS.get(key)
        if not item:
            return
        lock_until = item.get("lock_until")
        if isinstance(lock_until, datetime) and datetime.now() < lock_until:
            wait_seconds = int((lock_until - datetime.now()).total_seconds())
            wait_seconds = max(1, wait_seconds)
            raise HTTPException(status_code=429, detail=f"登录失败次数过多，请在 {wait_seconds} 秒后重试。")
        window_start = item.get("window_start")
        if isinstance(window_start, datetime) and datetime.now() - window_start > timedelta(minutes=LOGIN_FAIL_WINDOW_MINUTES):
            LOGIN_ATTEMPTS.pop(key, None)


def record_login_failure(key: str) -> None:
    now = datetime.now()
    with LOGIN_ATTEMPTS_LOCK:
        item = LOGIN_ATTEMPTS.get(key)
        if not item or now - item.get("window_start", now) > timedelta(minutes=LOGIN_FAIL_WINDOW_MINUTES):
            item = {"window_start": now, "fails": 0, "lock_until": None}
            LOGIN_ATTEMPTS[key] = item
        item["fails"] = int(item.get("fails", 0)) + 1
        if item["fails"] >= LOGIN_MAX_FAILS:
            item["lock_until"] = now + timedelta(minutes=LOGIN_LOCK_MINUTES)


def clear_login_failure(key: str) -> None:
    with LOGIN_ATTEMPTS_LOCK:
        LOGIN_ATTEMPTS.pop(key, None)


def ensure_default_accounts() -> None:
    db = SessionLocal()
    try:
        realigned_users: list[str] = []
        for username, env_key, role, fallback_password in DEFAULT_BOOTSTRAP_ACCOUNTS:
            desired_password = (os.getenv(env_key, "") or "").strip() or fallback_password
            exists = db.query(UserAccount).filter(UserAccount.username == username).first()
            if exists:
                password_updated_at = getattr(exists, "password_updated_at", None)
                last_login_at = getattr(exists, "last_login_at", None)
                should_realign = bool(getattr(exists, "must_change_password", False)) and (
                    password_updated_at is not None or last_login_at is None
                )
                if should_realign:
                    ok, _ = verify_password(desired_password, exists.password_hash)
                    if not ok:
                        exists.password_hash = hash_password(desired_password)
                        exists.password_updated_at = datetime.now()
                        realigned_users.append(username)
                continue
            db.add(
                UserAccount(
                    username=username,
                    password_hash=hash_password(desired_password),
                    role=role,
                    enabled=True,
                    must_change_password=True,
                    password_updated_at=datetime.now(),
                )
            )
        db.commit()
        if realigned_users:
            logger.warning("默认账号密码已按当前配置自动对齐：%s", ", ".join(realigned_users))
    finally:
        db.close()


def ensure_user_account_schema() -> None:
    insp = inspect(engine)
    tables = set(insp.get_table_names())
    if "user_accounts" not in tables:
        return
    cols = {c["name"] for c in insp.get_columns("user_accounts")}
    with engine.begin() as conn:
        if "must_change_password" not in cols:
            conn.execute(text("ALTER TABLE user_accounts ADD COLUMN must_change_password BOOLEAN NOT NULL DEFAULT 0"))
        if "password_updated_at" not in cols:
            conn.execute(text("ALTER TABLE user_accounts ADD COLUMN password_updated_at DATETIME NULL"))


def ensure_filing_workspace_schema() -> None:
    insp = inspect(engine)
    tables = set(insp.get_table_names())
    with engine.begin() as conn:
        if "organizations" in tables:
            org_cols = {c["name"] for c in insp.get_columns("organizations")}
            if "filing_detail" not in org_cols:
                conn.execute(text("ALTER TABLE organizations ADD COLUMN filing_detail JSON NULL"))
        if "systems" in tables:
            sys_cols = {c["name"] for c in insp.get_columns("systems")}
            if "filing_detail" not in sys_cols:
                conn.execute(text("ALTER TABLE systems ADD COLUMN filing_detail JSON NULL"))


def run_startup_tasks() -> None:
    ensure_dirs()
    init_db()
    ensure_user_account_schema()
    ensure_filing_workspace_schema()
    ensure_default_accounts()


def obj_to_dict(obj: Any, fields: list[str]) -> dict[str, Any]:
    def _normalize(value: Any) -> Any:
        if isinstance(value, (datetime, date)):
            return value.isoformat()
        if isinstance(value, dict):
            return {k: _normalize(v) for k, v in value.items()}
        if isinstance(value, list):
            return [_normalize(i) for i in value]
        return value

    return {f: _normalize(getattr(obj, f)) for f in fields}


def template_snapshot(tpl: ReportTemplate) -> dict[str, Any]:
    return {
        "template_name": tpl.template_name,
        "report_type": tpl.report_type,
        "category": tpl.category,
        "city": tpl.city,
        "protection_level": tpl.protection_level,
        "description": tpl.description,
        "status": tpl.status,
        "is_default": tpl.is_default,
        "version_no": tpl.version_no,
        "file_name": tpl.file_name,
        "file_path": tpl.file_path,
        "file_size": tpl.file_size,
        "config_json": tpl.config_json,
        "created_by": tpl.created_by,
    }


def log_template_version(db: Session, tpl: ReportTemplate, action: str, changed_by: str) -> None:
    db.add(
        ReportTemplateVersion(
            template_id=tpl.id,
            action=action,
            snapshot=template_snapshot(tpl),
            changed_by=changed_by,
        )
    )


def get_auth_token_from_request(request: Request) -> str:
    auth = (request.headers.get("Authorization") or "").strip()
    if auth.lower().startswith("bearer "):
        return auth.split(" ", 1)[1].strip()
    header_token = (request.headers.get("X-Auth-Token") or "").strip()
    if header_token:
        return header_token
    cookie_token = (request.cookies.get("auth_token") or "").strip()
    return cookie_token


def get_user_by_token(db: Session, token: str) -> UserAccount | None:
    if not token:
        return None
    now = datetime.now()
    session = (
        db.query(AuthSession, UserAccount)
        .join(UserAccount, UserAccount.id == AuthSession.user_id)
        .filter(AuthSession.token == token, AuthSession.expires_at > now, UserAccount.enabled.is_(True))
        .first()
    )
    if not session:
        return None
    _, user = session
    return user


def _is_api_auth_exempt(path: str) -> bool:
    if not path.startswith("/api/"):
        return True
    if path == "/api/auth/login":
        return True
    if path == "/api/auth/register":
        return True
    if path.startswith("/api/public/organizations/collect/"):
        return True
    return False


def _is_web_auth_exempt(path: str) -> bool:
    if path.startswith("/api/"):
        return True
    if path.startswith("/static/"):
        return True
    if path.startswith("/organizations/collect/"):
        return True
    if path in {"/login", "/register", "/health", "/favicon.ico"}:
        return True
    return False


def _is_password_change_exempt(path: str) -> bool:
    return path in {"/api/auth/change-password", "/api/auth/logout", "/api/auth/me"}


@app.middleware("http")
async def security_middleware(request: Request, call_next):  # type: ignore[override]
    req_scheme = (request.headers.get("x-forwarded-proto") or request.url.scheme or "").strip().lower()
    if FORCE_HTTPS and req_scheme and req_scheme != "https":
        target = request.url.replace(scheme="https")
        return JSONResponse(status_code=307, content={"detail": "请使用 HTTPS 访问。", "redirect": str(target)})
    path = request.url.path

    if API_AUTH_REQUIRED and not APP_LITE_MODE and request.method.upper() != "OPTIONS":
        if not _is_api_auth_exempt(path):
            token = get_auth_token_from_request(request)
            if not token:
                return JSONResponse(status_code=401, content={"detail": "未登录或令牌缺失。"})
            db = SessionLocal()
            try:
                user = get_user_by_token(db, token)
                if not user:
                    return JSONResponse(status_code=401, content={"detail": "登录令牌无效或已过期。"})
                if bool(getattr(user, "must_change_password", False)) and not _is_password_change_exempt(path):
                    return JSONResponse(status_code=403, content={"detail": "当前账号需先修改密码后再继续操作。"})
                request.state.current_user = user.username
                request.state.current_user_role = getattr(user, "role", "")
                request.state.current_user_must_change_password = bool(getattr(user, "must_change_password", False))
            finally:
                db.close()
        if not _is_web_auth_exempt(path):
            token = get_auth_token_from_request(request)
            if not token:
                return RedirectResponse(url=f"/login?next={quote(path)}", status_code=307)
            db = SessionLocal()
            try:
                user = get_user_by_token(db, token)
                if not user:
                    return RedirectResponse(url=f"/login?next={quote(path)}", status_code=307)
                request.state.current_user = user.username
                request.state.current_user_role = getattr(user, "role", "")
                request.state.current_user_must_change_password = bool(getattr(user, "must_change_password", False))
            finally:
                db.close()

    response = await call_next(request)
    response.headers["X-Frame-Options"] = "DENY"
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["Referrer-Policy"] = "no-referrer"
    response.headers["Permissions-Policy"] = "camera=(), microphone=(), geolocation=()"
    response.headers["Content-Security-Policy"] = (
        "default-src 'self'; "
        "img-src 'self' data: https:; "
        "style-src 'self' 'unsafe-inline' https:; "
        "script-src 'self' 'unsafe-inline' https:; "
        "font-src 'self' https: data:; "
        "connect-src 'self'; "
        "object-src 'none'; "
        "base-uri 'self'; "
        "frame-ancestors 'none'"
    )
    if req_scheme == "https":
        response.headers["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains"
    return response


def get_current_user_optional(request: Request, db: Session) -> UserAccount | None:
    token = get_auth_token_from_request(request)
    return get_user_by_token(db, token)


def require_docs_admin(request: Request, db: Session, token: str | None = None) -> str:
    auth_token = (token or "").strip() or get_auth_token_from_request(request)
    if not auth_token:
        raise HTTPException(status_code=401, detail="访问API文档需要管理员登录令牌。")
    user = get_user_by_token(db, auth_token)
    if not user:
        raise HTTPException(status_code=401, detail="登录令牌无效或已过期。")
    if user.role != "admin":
        raise HTTPException(status_code=403, detail="仅管理员可访问API文档。")
    return auth_token


if ENABLE_API_DOCS:

    @app.get("/openapi.json", include_in_schema=False)
    def secured_openapi(
        request: Request,
        token: str | None = Query(None),
        db: Session = Depends(get_db),
    ) -> JSONResponse:
        require_docs_admin(request, db, token=token)
        return JSONResponse(app.openapi())

    @app.get("/docs", include_in_schema=False)
    def secured_docs(
        request: Request,
        token: str | None = Query(None),
        db: Session = Depends(get_db),
    ) -> HTMLResponse:
        auth_token = require_docs_admin(request, db, token=token)
        openapi_url = f"/openapi.json?token={quote(auth_token)}"
        return get_swagger_ui_html(openapi_url=openapi_url, title=f"{app.title} - Swagger UI")

    @app.get("/redoc", include_in_schema=False)
    def secured_redoc(
        request: Request,
        token: str | None = Query(None),
        db: Session = Depends(get_db),
    ) -> HTMLResponse:
        auth_token = require_docs_admin(request, db, token=token)
        openapi_url = f"/openapi.json?token={quote(auth_token)}"
        return get_redoc_html(openapi_url=openapi_url, title=f"{app.title} - ReDoc")


def require_roles(
    request: Request,
    db: Session,
    allowed_roles: set[str],
    legacy_admin: bool = False,
) -> tuple[str, str]:
    user = get_current_user_optional(request, db)
    if user:
        if bool(getattr(user, "must_change_password", False)) and not _is_password_change_exempt(request.url.path):
            raise HTTPException(status_code=403, detail="当前账号需先修改密码后再继续操作。")
        if user.role not in allowed_roles:
            raise HTTPException(status_code=403, detail="当前账号无权限执行此操作。")
        return user.username, user.role
    if legacy_admin and "admin" in allowed_roles and not STRICT_AUTH and (APP_LITE_MODE or not API_AUTH_REQUIRED):
        return "legacy_admin", "admin"
    if STRICT_AUTH or allowed_roles:
        raise HTTPException(status_code=401, detail="请先登录。")
    return "legacy_user", "evaluator"


def resolve_effective_actor_name(request: Request, db: Session, fallback: str | None = None) -> str:
    user = get_current_user_optional(request, db)
    if user and getattr(user, "username", ""):
        return str(user.username).strip()
    text = str(fallback or "").strip()
    return text or "system"


def finalize_pending_delete_requests(
    db: Session,
    *,
    entity_type: str,
    entity_id: int,
    reviewer: str,
    reviewed_at: datetime | None = None,
    comment: str | None = None,
) -> None:
    rows = (
        db.query(DeleteRequest)
        .filter(
            DeleteRequest.entity_type == entity_type,
            DeleteRequest.entity_id == entity_id,
            DeleteRequest.status == "pending",
        )
        .all()
    )
    if not rows:
        return
    resolved_at = reviewed_at or datetime.now()
    default_comment = comment or "实体已由管理员直接删除，申请自动结案。"
    for row in rows:
        row.status = "approved"
        row.reviewed_by = reviewer
        row.reviewed_at = resolved_at
        row.review_comment = row.review_comment or default_comment


def reconcile_pending_delete_requests(db: Session) -> None:
    rows = db.query(DeleteRequest).filter(DeleteRequest.status == "pending").all()
    changed = False
    for row in rows:
        deleted_by = ""
        deleted_at: datetime | None = None
        if row.entity_type == "organization":
            org = db.query(Organization).filter(Organization.id == row.entity_id).first()
            if not org or not org.deleted_at:
                continue
            deleted_by = str(org.deleted_by or "").strip()
            deleted_at = org.deleted_at
        elif row.entity_type == "system":
            system = db.query(SystemInfo).filter(SystemInfo.id == row.entity_id).first()
            if not system or not system.deleted_at:
                continue
            deleted_by = str(system.deleted_by or "").strip()
            deleted_at = system.deleted_at
        else:
            continue
        row.status = "approved"
        row.reviewed_by = deleted_by or "system"
        row.reviewed_at = deleted_at or datetime.now()
        row.review_comment = row.review_comment or "实体已由管理员直接删除，申请自动结案。"
        changed = True
    if changed:
        db.commit()


def backfill_delete_request_review_comments(db: Session) -> None:
    rows = (
        db.query(DeleteRequest)
        .filter(
            DeleteRequest.status == "approved",
            DeleteRequest.reviewed_by.is_not(None),
            or_(DeleteRequest.review_comment.is_(None), DeleteRequest.review_comment == ""),
        )
        .all()
    )
    changed = False
    for row in rows:
        entity_deleted_by = ""
        entity_deleted_at: datetime | None = None
        if row.entity_type == "organization":
            org = db.query(Organization).filter(Organization.id == row.entity_id).first()
            if not org or not org.deleted_at:
                continue
            entity_deleted_by = str(org.deleted_by or "").strip()
            entity_deleted_at = org.deleted_at
        elif row.entity_type == "system":
            system = db.query(SystemInfo).filter(SystemInfo.id == row.entity_id).first()
            if not system or not system.deleted_at:
                continue
            entity_deleted_by = str(system.deleted_by or "").strip()
            entity_deleted_at = system.deleted_at
        else:
            continue
        if not entity_deleted_by or str(row.reviewed_by or "").strip() != entity_deleted_by:
            continue
        reviewed_at = row.reviewed_at or entity_deleted_at
        if not reviewed_at or not entity_deleted_at:
            continue
        if abs((reviewed_at - entity_deleted_at).total_seconds()) > 5:
            continue
        row.review_comment = "实体已由管理员直接删除，申请自动结案。"
        if row.reviewed_at is None:
            row.reviewed_at = entity_deleted_at
        changed = True
    if changed:
        db.commit()


ORG_FIELDS = [
    "id",
    "name",
    "credit_code",
    "legal_representative",
    "address",
    "office_phone",
    "mobile_phone",
    "email",
    "industry",
    "organization_type",
    "cybersecurity_dept",
    "cybersecurity_owner_name",
    "cybersecurity_owner_title",
    "cybersecurity_owner_phone",
    "cybersecurity_owner_email",
    "data_security_dept",
    "data_security_owner_name",
    "data_security_owner_title",
    "data_security_owner_phone",
    "data_security_owner_email",
    "supervising_department",
    "filing_region",
    "filing_detail",
    "involves_state_secret",
    "is_cii",
    "remark",
    "created_by",
    "archived",
    "locked",
    "deleted_at",
    "deleted_by",
    "created_at",
    "updated_at",
]

ORG_CREATE_FIELDS = [
    "name",
    "credit_code",
    "legal_representative",
    "address",
    "office_phone",
    "mobile_phone",
    "email",
    "industry",
    "organization_type",
    "cybersecurity_dept",
    "cybersecurity_owner_name",
    "cybersecurity_owner_title",
    "cybersecurity_owner_phone",
    "cybersecurity_owner_email",
    "data_security_dept",
    "data_security_owner_name",
    "data_security_owner_title",
    "data_security_owner_phone",
    "data_security_owner_email",
    "supervising_department",
    "filing_region",
    "involves_state_secret",
    "is_cii",
    "remark",
    "created_by",
]

ORG_UPDATE_FIELDS = [f for f in ORG_CREATE_FIELDS if f != "created_by"]

SYSTEM_FIELDS = [
    "id",
    "organization_id",
    "system_name",
    "system_code",
    "proposed_level",
    "filing_detail",
    "business_description",
    "system_type",
    "deployment_mode",
    "go_live_date",
    "boundary",
    "subsystems",
    "service_object",
    "service_scope",
    "network_topology",
    "data_name",
    "data_level",
    "data_category",
    "data_security_dept",
    "data_security_owner",
    "contains_personal_info",
    "data_total",
    "monthly_growth",
    "data_source",
    "data_flow",
    "data_interaction",
    "storage_location",
    "level_basis",
    "impact_scope",
    "level_reason",
    "needs_expert_review",
    "expert_review_info",
    "ops_unit",
    "ops_personnel",
    "created_by",
    "archived",
    "locked",
    "deleted_at",
    "deleted_by",
    "created_at",
    "updated_at",
]

SYSTEM_TEXT_FIELDS = [
    "system_name",
    "business_description",
    "system_type",
    "deployment_mode",
    "boundary",
    "subsystems",
    "service_object",
    "service_scope",
    "network_topology",
    "data_name",
    "data_level",
    "data_category",
    "data_security_dept",
    "data_security_owner",
    "data_source",
    "data_flow",
    "data_interaction",
    "storage_location",
    "level_basis",
    "impact_scope",
    "level_reason",
    "expert_review_info",
    "ops_unit",
    "ops_personnel",
    "created_by",
]


def validate_org_payload(data: dict[str, Any]) -> None:
    required_fields = [
        ("name", "单位名称"),
        ("credit_code", "统一社会信用代码"),
        ("legal_representative", "单位负责人"),
        ("address", "单位地址"),
        ("mobile_phone", "移动电话"),
        ("email", "邮箱"),
        ("industry", "所属行业"),
        ("organization_type", "单位类型"),
        ("filing_region", "备案地区"),
    ]
    for key, label in required_fields:
        if not str(data.get(key) or "").strip():
            raise HTTPException(status_code=400, detail=f"{label}为必填项，不能为空。")
    for key in ORG_CREATE_FIELDS:
        assert_safe_text(data.get(key), key)
    name = str(data.get("name") or "").strip()
    if not name or is_placeholder_value(name):
        raise HTTPException(status_code=400, detail="单位名称不能为空，且不能仅填写“/”。")
    if not validate_credit_code_format_only(data["credit_code"]):
        raise HTTPException(status_code=400, detail="统一社会信用代码格式错误，应为18位大写字母或数字。")
    if not validate_mobile_phone(data["mobile_phone"]):
        raise HTTPException(status_code=400, detail="手机号格式错误，应为11位中国大陆手机号。")
    if not validate_office_phone(data.get("office_phone") or ""):
        raise HTTPException(status_code=400, detail="办公电话格式错误。")
    if not validate_email(data["email"]):
        raise HTTPException(status_code=400, detail="邮箱格式错误。")


def validate_org_partial(data: dict[str, Any]) -> None:
    required_labels = {
        "name": "单位名称",
        "credit_code": "统一社会信用代码",
        "legal_representative": "单位负责人",
        "address": "单位地址",
        "mobile_phone": "移动电话",
        "email": "邮箱",
        "industry": "所属行业",
        "organization_type": "单位类型",
        "filing_region": "备案地区",
    }
    for key, label in required_labels.items():
        if key in data and not str(data.get(key) or "").strip():
            raise HTTPException(status_code=400, detail=f"{label}为必填项，不能为空。")
    for key in ORG_UPDATE_FIELDS:
        if key in data:
            assert_safe_text(data.get(key), key)
    if "name" in data:
        name = str(data.get("name") or "").strip()
        if not name or is_placeholder_value(name):
            raise HTTPException(status_code=400, detail="单位名称不能为空，且不能仅填写“/”。")
    if "credit_code" in data and not validate_credit_code_format_only(str(data.get("credit_code") or "")):
        raise HTTPException(status_code=400, detail="统一社会信用代码格式错误，应为18位大写字母或数字。")
    if "mobile_phone" in data and not validate_mobile_phone(data["mobile_phone"]):
        raise HTTPException(status_code=400, detail="手机号格式错误，应为11位中国大陆手机号。")
    if "office_phone" in data and not validate_office_phone(data.get("office_phone") or ""):
        raise HTTPException(status_code=400, detail="办公电话格式错误。")
    if "email" in data and not validate_email(data["email"]):
        raise HTTPException(status_code=400, detail="邮箱格式错误。")


def validate_system_payload(data: dict[str, Any], partial: bool = False) -> None:
    if not partial:
        name = str(data.get("system_name") or "").strip()
        if not name:
            raise HTTPException(status_code=400, detail="系统名称为必填项，不能为空。")
    for key in SYSTEM_TEXT_FIELDS:
        if key in data:
            assert_safe_text(data.get(key), key)


def normalize_org_payload(data: dict[str, Any]) -> None:
    if "name" in data and isinstance(data["name"], str):
        data["name"] = data["name"].strip()
    if "credit_code" in data and isinstance(data["credit_code"], str):
        data["credit_code"] = data["credit_code"].strip().upper()
    for key in ("mobile_phone", "office_phone", "email"):
        if key in data and isinstance(data[key], str):
            data[key] = data[key].strip()


def deep_copy_json(value: Any) -> Any:
    return json.loads(json.dumps(value, ensure_ascii=False))


def default_org_filing_detail() -> dict[str, Any]:
    return {
        "table1": {
            "address_parts": {"province": "", "city": "", "district": "", "detail": ""},
            "postal_code": "",
            "district_code": "",
            "unit_internet_addresses": "无",
            "responsible_person": {"name": "", "title": "", "office_phone": "", "email": ""},
            "affiliation": {"code": "", "other": ""},
            "organization_type_detail": {"code": "", "other": ""},
            "industry_category": {"code": "", "label": "", "other": ""},
            "current_filing_counts": {"2": 0, "3": 0, "4": 0, "5": 0},
            "total_filing_counts": {"1": 0, "2": 0, "3": 0, "4": 0, "5": 0},
        }
    }


def default_system_filing_detail() -> dict[str, Any]:
    return {
        "table2": {
            "object_types": [],
            "technology_types": [],
            "technology_other": "",
            "business_types": [],
            "business_other": "",
            "business_description": "",
            "network_service": {
                "scope_code": "",
                "cross_province_count": "",
                "cross_city_count": "",
                "other": "",
                "service_targets": [],
                "service_target_other": "",
            },
            "network_platform": {
                "coverage": {"code": "", "other": ""},
                "network_nature": {
                    "code": "",
                    "other": "",
                    "source_ip_range": "",
                    "domain": "",
                    "protocol_ports": "",
                },
                "interconnection": {"items": [], "other": ""},
            },
            "go_live_date": "",
            "is_sub_system": False,
            "parent_system_name": "",
            "parent_organization_name": "",
        },
        "table3": {
            "business_security_damage_items": [],
            "business_security_level": 0,
            "service_security_damage_items": [],
            "service_security_level": 0,
            "final_level": 0,
            "grading_date": "",
            "grading_report": {"has_file": False, "file_name": "", "attachment_ids": []},
            "expert_review": {"status": "unreviewed", "file_name": "", "attachment_ids": []},
            "has_supervisor": False,
            "supervisor_name": "",
            "supervisor_review": {"status": "unreviewed", "file_name": "", "attachment_ids": []},
            "filler_name": "",
            "filled_date": "",
        },
        "table4": {
            "cloud": {
                "enabled": False,
                "responsibility_types": [],
                "service_modes": [],
                "service_mode_other": "",
                "deployment_modes": [],
                "deployment_mode_other": "",
                "customer_count": "",
                "infra_location": "",
                "ops_location": "",
                "provider_name": "",
                "provider_level": "",
                "provider_platform_name": "",
                "provider_record_no": "",
                "customer_ops_location": "",
                "record_file_name": "",
                "attachment_ids": [],
            },
            "mobile": {
                "enabled": False,
                "app_name": "",
                "wireless_channels": [],
                "terminal_types": [],
            },
            "iot": {
                "enabled": False,
                "perception_layers": [],
                "perception_other": "",
                "transport_layers": [],
                "transport_other": "",
            },
            "industrial_control": {
                "enabled": False,
                "function_layers": [],
                "components": [],
                "component_other": "",
            },
            "big_data": {
                "enabled": False,
                "system_components": [],
                "cross_border_status": "",
                "application_count": "",
                "infra_location": "",
                "ops_location": "",
                "provider_name": "",
                "provider_level": "",
                "provider_platform_name": "",
                "provider_record_no": "",
                "record_file_name": "",
                "attachment_ids": [],
            },
        },
        "table5": {
            "network_topology": {"status": "none", "file_name": "", "attachment_ids": []},
            "security_org_and_rules": {"status": "none", "file_name": "", "attachment_ids": []},
            "security_design_plan": {"status": "none", "file_name": "", "attachment_ids": []},
            "security_products": {"status": "none", "file_name": "", "attachment_ids": []},
            "security_services": {"status": "none", "file_name": "", "attachment_ids": []},
            "supervisor_guidance": {"status": "none", "file_name": "", "attachment_ids": []},
        },
        "table6": {"items": []},
    }


def merged_org_filing_detail(org: Organization | None) -> dict[str, Any]:
    data = default_org_filing_detail()
    if org and isinstance(org.filing_detail, dict):
        data = deep_merge_dicts(data, org.filing_detail)
    return data


def merged_system_filing_detail(system: SystemInfo | None) -> dict[str, Any]:
    data = default_system_filing_detail()
    if system and isinstance(system.filing_detail, dict):
        data = deep_merge_dicts(data, system.filing_detail)
    return data


def deep_merge_dicts(base: dict[str, Any], override: dict[str, Any]) -> dict[str, Any]:
    merged = deep_copy_json(base)
    for key, value in (override or {}).items():
        if isinstance(value, dict) and isinstance(merged.get(key), dict):
            merged[key] = deep_merge_dicts(merged[key], value)
            continue
        merged[key] = deep_copy_json(value)
    return merged


def to_int_or_zero(value: Any) -> int:
    try:
        return max(0, int(value))
    except (TypeError, ValueError):
        return 0


def coerce_bool(value: Any) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return bool(value)
    if isinstance(value, str):
        return value.strip().lower() in {"1", "true", "yes", "on", "y", "是", "有"}
    return False


def normalize_attachment_refs(refs: Any) -> list[int]:
    if not isinstance(refs, list):
        return []
    result: list[int] = []
    for item in refs:
        try:
            value = int(item)
        except (TypeError, ValueError):
            continue
        if value > 0:
            result.append(value)
    return result


def normalize_slot_attachment_map(mapping: Any) -> dict[str, dict[str, Any]]:
    result: dict[str, dict[str, Any]] = {}
    if not isinstance(mapping, dict):
        return result
    for slot_key, slot_value in mapping.items():
        if not isinstance(slot_value, dict):
            continue
        result[str(slot_key)] = {
            "file_name": str(slot_value.get("file_name") or "").strip(),
            "attachment_ids": normalize_attachment_refs(slot_value.get("attachment_ids")),
        }
    return result


def get_attachment_briefs(db: Session, attachment_ids: list[int]) -> list[dict[str, Any]]:
    if not attachment_ids:
        return []
    rows = db.query(Attachment).filter(Attachment.id.in_(attachment_ids)).order_by(Attachment.uploaded_at.desc()).all()
    index = {
        row.id: {
            "id": row.id,
            "file_name": row.file_name,
            "file_size": row.file_size,
            "uploaded_at": row.uploaded_at,
            "download_url": f"/api/attachment-files/{row.id}/download",
            "preview_url": f"/api/attachment-files/{row.id}/preview",
        }
        for row in rows
    }
    return [index[attachment_id] for attachment_id in attachment_ids if attachment_id in index]


def assert_credit_code_available(db: Session, credit_code: str, current_org_id: int | None = None) -> None:
    exists = db.query(Organization).filter(Organization.credit_code == credit_code).first()
    if not exists:
        return
    if current_org_id is not None and exists.id == current_org_id:
        return
    if exists.deleted_at:
        raise HTTPException(
            status_code=409,
            detail=f"统一社会信用代码已存在于回收站(ID={exists.id})，请先恢复该单位。",
        )
    raise HTTPException(status_code=409, detail="统一社会信用代码已存在。")


def recycle_meta(deleted_at: datetime | None) -> dict[str, Any]:
    if not deleted_at:
        return {"expires_at": None, "days_left": 0, "expired": False}
    expire_at = deleted_at + timedelta(days=30)
    seconds_left = (expire_at - datetime.now()).total_seconds()
    days_left = int(max(0, (seconds_left + 86399) // 86400))
    return {"expires_at": expire_at, "days_left": days_left, "expired": seconds_left <= 0}


def send_workflow_email(to_email: str, subject: str, body: str) -> tuple[bool, str]:
    if not SMTP_HOST:
        return False, "SMTP_HOST 未配置"
    sender = SMTP_FROM or SMTP_USER
    if not sender:
        return False, "SMTP_FROM/SMTP_USER 未配置"
    msg = MIMEText(body, _subtype="plain", _charset="utf-8")
    msg["Subject"] = subject
    msg["From"] = sender
    msg["To"] = to_email
    try:
        if SMTP_SSL:
            server = smtplib.SMTP_SSL(SMTP_HOST, SMTP_PORT, timeout=10)
        else:
            server = smtplib.SMTP(SMTP_HOST, SMTP_PORT, timeout=10)
        with server:
            if SMTP_TLS and not SMTP_SSL:
                server.starttls()
            if SMTP_USER:
                server.login(SMTP_USER, SMTP_PASS)
            server.sendmail(sender, [to_email], msg.as_string())
        return True, "ok"
    except Exception as exc:
        return False, str(exc)


def save_attachment_row(
    db: Session,
    *,
    entity_type: str,
    entity_id: int,
    actor: str,
    file_name: str,
    content: bytes,
) -> Attachment:
    ext = Path(file_name).suffix.lower().lstrip(".")
    safe_name = Path(file_name).name
    save_name = f"{uuid.uuid4().hex}_{safe_name}"
    target_dir = UPLOAD_DIR / "attachments" / entity_type
    target_dir.mkdir(parents=True, exist_ok=True)
    target_path = target_dir / save_name
    target_path.write_bytes(content)
    row = Attachment(
        entity_type=entity_type,
        entity_id=entity_id,
        file_name=safe_name,
        file_path=str(target_path),
        file_ext=ext,
        file_size=len(content),
        uploaded_by=actor,
    )
    db.add(row)
    db.flush()
    return row


def assert_org_deletable(db: Session, org: Organization, is_admin: bool) -> None:
    if org.archived and not is_admin:
        raise HTTPException(status_code=400, detail="单位已归档，仅管理员可删除。")
    related_systems = (
        db.query(SystemInfo)
        .filter(SystemInfo.organization_id == org.id, SystemInfo.deleted_at.is_(None))
        .order_by(SystemInfo.created_at.desc())
        .limit(20)
        .all()
    )
    if related_systems:
        details = ", ".join([f"{s.id}:{s.system_name}" for s in related_systems])
        raise HTTPException(status_code=400, detail=f"存在关联系统，暂不可删除。请先删除系统：{details}")


def assert_system_deletable(db: Session, system: SystemInfo, is_admin: bool) -> None:
    if system.archived and not is_admin:
        raise HTTPException(status_code=400, detail="系统已归档，仅管理员可删除。")
    related_reports = db.query(Report).filter(Report.system_id == system.id).count()
    if related_reports > 0 and not is_admin:
        raise HTTPException(status_code=400, detail=f"系统已有关联报告({related_reports})，暂不可删除。")


def purge_entity_attachments(db: Session, entity_type: str, entity_id: int) -> int:
    rows = (
        db.query(Attachment)
        .filter(Attachment.entity_type == entity_type.lower(), Attachment.entity_id == entity_id)
        .all()
    )
    for row in rows:
        safe_delete_uploaded_file(row.file_path)
        db.delete(row)
    return len(rows)


def purge_system_cascade(db: Session, system: SystemInfo) -> dict[str, int]:
    system_id = system.id
    report_ids = [row[0] for row in db.query(Report.id).filter(Report.system_id == system_id).all()]
    workflow_instance_ids = [row[0] for row in db.query(WorkflowInstance.id).filter(WorkflowInstance.system_id == system_id).all()]

    attachments_deleted = purge_entity_attachments(db, "system", system_id)
    review_records_deleted = 0
    reports_deleted = 0
    workflow_actions_deleted = 0
    workflow_reminders_deleted = 0
    workflow_instances_deleted = 0

    if report_ids:
        review_records_deleted = (
            db.query(ReviewRecord)
            .filter(ReviewRecord.report_id.in_(report_ids))
            .delete(synchronize_session=False)
        )
        reports_deleted = (
            db.query(Report)
            .filter(Report.id.in_(report_ids))
            .delete(synchronize_session=False)
        )

    if workflow_instance_ids:
        workflow_actions_deleted = (
            db.query(WorkflowAction)
            .filter(WorkflowAction.instance_id.in_(workflow_instance_ids))
            .delete(synchronize_session=False)
        )
        workflow_reminders_deleted = (
            db.query(WorkflowReminder)
            .filter(WorkflowReminder.instance_id.in_(workflow_instance_ids))
            .delete(synchronize_session=False)
        )
        workflow_instances_deleted = (
            db.query(WorkflowInstance)
            .filter(WorkflowInstance.id.in_(workflow_instance_ids))
            .delete(synchronize_session=False)
        )

    histories_deleted = (
        db.query(SystemHistory)
        .filter(SystemHistory.system_id == system_id)
        .delete(synchronize_session=False)
    )
    delete_requests_deleted = (
        db.query(DeleteRequest)
        .filter(DeleteRequest.entity_type == "system", DeleteRequest.entity_id == system_id)
        .delete(synchronize_session=False)
    )
    db.delete(system)

    return {
        "systems_deleted": 1,
        "reports_deleted": reports_deleted,
        "review_records_deleted": review_records_deleted,
        "attachments_deleted": attachments_deleted,
        "workflow_actions_deleted": workflow_actions_deleted,
        "workflow_reminders_deleted": workflow_reminders_deleted,
        "workflow_instances_deleted": workflow_instances_deleted,
        "histories_deleted": histories_deleted,
        "delete_requests_deleted": delete_requests_deleted,
    }


def purge_organization_cascade(db: Session, org: Organization) -> dict[str, int]:
    org_id = org.id
    totals = {
        "organizations_deleted": 1,
        "systems_deleted": 0,
        "reports_deleted": 0,
        "review_records_deleted": 0,
        "attachments_deleted": 0,
        "workflow_actions_deleted": 0,
        "workflow_reminders_deleted": 0,
        "workflow_instances_deleted": 0,
        "histories_deleted": 0,
        "delete_requests_deleted": 0,
        "collection_links_deleted": 0,
        "submissions_deleted": 0,
    }

    systems = db.query(SystemInfo).filter(SystemInfo.organization_id == org_id).order_by(SystemInfo.id.asc()).all()
    for system in systems:
        result = purge_system_cascade(db, system)
        for key, value in result.items():
            totals[key] = totals.get(key, 0) + value

    link_ids = [row[0] for row in db.query(OrganizationCollectionLink.id).filter(OrganizationCollectionLink.organization_id == org_id).all()]
    submission_filters = [OrganizationSubmission.organization_id == org_id]
    if link_ids:
        submission_filters.append(OrganizationSubmission.link_id.in_(link_ids))
    submission_ids = [
        row[0]
        for row in db.query(OrganizationSubmission.id).filter(or_(*submission_filters)).all()
    ]
    if submission_ids:
        totals["submissions_deleted"] += (
            db.query(OrganizationSubmission)
            .filter(OrganizationSubmission.id.in_(submission_ids))
            .delete(synchronize_session=False)
        )
    if link_ids:
        totals["collection_links_deleted"] += (
            db.query(OrganizationCollectionLink)
            .filter(OrganizationCollectionLink.id.in_(link_ids))
            .delete(synchronize_session=False)
        )

    totals["attachments_deleted"] += purge_entity_attachments(db, "organization", org_id)
    totals["histories_deleted"] += (
        db.query(OrganizationHistory)
        .filter(OrganizationHistory.organization_id == org_id)
        .delete(synchronize_session=False)
    )
    totals["delete_requests_deleted"] += (
        db.query(DeleteRequest)
        .filter(DeleteRequest.entity_type == "organization", DeleteRequest.entity_id == org_id)
        .delete(synchronize_session=False)
    )
    db.delete(org)
    return totals


def record_org_history(
    db: Session,
    organization_id: int,
    changed_by: str,
    change_type: str,
    before_data: dict[str, Any] | None,
    after_data: dict[str, Any] | None,
) -> None:
    db.add(
        OrganizationHistory(
            organization_id=organization_id,
            changed_by=changed_by,
            change_type=change_type,
            before_data=before_data,
            after_data=after_data,
        )
    )


def record_system_history(
    db: Session,
    system_id: int,
    changed_by: str,
    change_type: str,
    before_data: dict[str, Any] | None,
    after_data: dict[str, Any] | None,
) -> None:
    db.add(
        SystemHistory(
            system_id=system_id,
            changed_by=changed_by,
            change_type=change_type,
            before_data=before_data,
            after_data=after_data,
        )
    )


def parse_workspace_date(raw_value: Any) -> date | None:
    value = str(raw_value or "").strip()
    if not value:
        return None
    value = value.replace("年", "-").replace("月", "-").replace("日", "").replace("/", "-").replace(".", "-")
    try:
        return date.fromisoformat(value)
    except ValueError:
        return None


def format_workspace_date(raw_value: Any) -> str:
    if isinstance(raw_value, date):
        return raw_value.isoformat()
    parsed = parse_workspace_date(raw_value)
    return parsed.isoformat() if parsed else ""


def summarize_table6_item(item: dict[str, Any]) -> dict[str, Any]:
    return {
        "data_name": str(item.get("data_name") or "").strip(),
        "data_level": str(item.get("data_level") or "").strip(),
        "data_category": str(item.get("data_category") or "").strip(),
    }


def hydrate_workspace_attachments(db: Session, detail: dict[str, Any]) -> dict[str, Any]:
    data = deep_copy_json(detail)
    table3 = data.get("table3") if isinstance(data.get("table3"), dict) else {}
    for slot in ("grading_report", "expert_review", "supervisor_review"):
        slot_data = table3.get(slot)
        if isinstance(slot_data, dict):
            slot_data["attachments"] = get_attachment_briefs(db, normalize_attachment_refs(slot_data.get("attachment_ids")))
    table4 = data.get("table4") if isinstance(data.get("table4"), dict) else {}
    for slot in ("cloud", "big_data"):
        slot_data = table4.get(slot)
        if isinstance(slot_data, dict):
            slot_data["attachments"] = get_attachment_briefs(db, normalize_attachment_refs(slot_data.get("attachment_ids")))
    table5 = data.get("table5") if isinstance(data.get("table5"), dict) else {}
    for slot_data in table5.values():
        if isinstance(slot_data, dict):
            slot_data["attachments"] = get_attachment_briefs(db, normalize_attachment_refs(slot_data.get("attachment_ids")))
    return data


def sync_org_from_workspace(db: Session, org: Organization, payload: dict[str, Any]) -> tuple[dict[str, Any], dict[str, Any]]:
    before = obj_to_dict(org, ORG_FIELDS)
    table1 = merged_org_filing_detail(org).get("table1", {})
    incoming_detail = payload.get("filing_detail") if isinstance(payload.get("filing_detail"), dict) else {}
    merged_detail = deep_merge_dicts(merged_org_filing_detail(org), incoming_detail)
    table1 = merged_detail.get("table1", {})
    address_parts = table1.get("address_parts") if isinstance(table1.get("address_parts"), dict) else {}
    address = str(payload.get("address") or "").strip()
    if not address:
        address = "".join(
            [
                str(address_parts.get("province") or "").strip(),
                str(address_parts.get("city") or "").strip(),
                str(address_parts.get("district") or "").strip(),
                str(address_parts.get("detail") or "").strip(),
            ]
        )
    industry_meta = table1.get("industry_category") if isinstance(table1.get("industry_category"), dict) else {}
    org_type_meta = table1.get("organization_type_detail") if isinstance(table1.get("organization_type_detail"), dict) else {}
    responsible_person = table1.get("responsible_person") if isinstance(table1.get("responsible_person"), dict) else {}
    org_data = {
        "name": str(payload.get("name") or org.name).strip(),
        "credit_code": str(payload.get("credit_code") or org.credit_code).strip().upper(),
        "legal_representative": str(responsible_person.get("name") or payload.get("legal_representative") or org.legal_representative).strip(),
        "address": address or org.address,
        "office_phone": str(responsible_person.get("office_phone") or payload.get("office_phone") or org.office_phone or "").strip() or None,
        "mobile_phone": str(payload.get("mobile_phone") or org.mobile_phone).strip(),
        "email": str(responsible_person.get("email") or payload.get("email") or org.email).strip(),
        "industry": str(industry_meta.get("label") or payload.get("industry") or org.industry).strip(),
        "organization_type": str(org_type_meta.get("label") or payload.get("organization_type") or org.organization_type).strip(),
        "cybersecurity_dept": str(payload.get("cybersecurity_dept") or org.cybersecurity_dept or "").strip() or None,
        "cybersecurity_owner_name": str(payload.get("cybersecurity_owner_name") or org.cybersecurity_owner_name or "").strip() or None,
        "cybersecurity_owner_title": str(payload.get("cybersecurity_owner_title") or org.cybersecurity_owner_title or "").strip() or None,
        "cybersecurity_owner_phone": str(payload.get("cybersecurity_owner_phone") or org.cybersecurity_owner_phone or "").strip() or None,
        "cybersecurity_owner_email": str(payload.get("cybersecurity_owner_email") or org.cybersecurity_owner_email or "").strip() or None,
        "data_security_dept": str(payload.get("data_security_dept") or org.data_security_dept or "").strip() or None,
        "data_security_owner_name": str(payload.get("data_security_owner_name") or org.data_security_owner_name or "").strip() or None,
        "data_security_owner_title": str(payload.get("data_security_owner_title") or org.data_security_owner_title or "").strip() or None,
        "data_security_owner_phone": str(payload.get("data_security_owner_phone") or org.data_security_owner_phone or "").strip() or None,
        "data_security_owner_email": str(payload.get("data_security_owner_email") or org.data_security_owner_email or "").strip() or None,
        "supervising_department": str(payload.get("supervising_department") or org.supervising_department or "").strip() or None,
        "filing_region": str(payload.get("filing_region") or org.filing_region).strip(),
        "involves_state_secret": coerce_bool(payload.get("involves_state_secret", org.involves_state_secret)),
        "is_cii": coerce_bool(payload.get("is_cii", org.is_cii)),
        "remark": str(payload.get("remark") or org.remark or "").strip() or None,
    }
    normalize_org_payload(org_data)
    validate_org_partial(org_data)
    assert_credit_code_available(db, org_data["credit_code"], current_org_id=org.id)
    for key, value in org_data.items():
        setattr(org, key, value)
    org.filing_detail = merged_detail
    return before, obj_to_dict(org, ORG_FIELDS)


def build_org_address_parts(org: Organization, detail: dict[str, Any]) -> dict[str, Any]:
    table1 = detail.get("table1") if isinstance(detail.get("table1"), dict) else {}
    address_parts = table1.get("address_parts") if isinstance(table1.get("address_parts"), dict) else {}
    if any(str(address_parts.get(key) or "").strip() for key in ("province", "city", "district", "detail")):
        return address_parts
    return {"province": "", "city": "", "district": "", "detail": org.address or ""}


def sync_system_from_workspace(system: SystemInfo, payload: dict[str, Any]) -> tuple[dict[str, Any], dict[str, Any]]:
    before = obj_to_dict(system, SYSTEM_FIELDS)
    incoming_detail = payload.get("filing_detail") if isinstance(payload.get("filing_detail"), dict) else {}
    merged_detail = deep_merge_dicts(merged_system_filing_detail(system), incoming_detail)
    table2 = merged_detail.get("table2", {})
    table3 = merged_detail.get("table3", {})
    table6 = merged_detail.get("table6", {})
    final_level = to_int_or_zero(table3.get("final_level")) or max(
        to_int_or_zero(table3.get("business_security_level")),
        to_int_or_zero(table3.get("service_security_level")),
        system.proposed_level or 0,
    )
    go_live = format_workspace_date(table2.get("go_live_date"))
    first_data_item = {}
    if isinstance(table6, dict) and isinstance(table6.get("items"), list) and table6["items"]:
        first = table6["items"][0]
        if isinstance(first, dict):
            first_data_item = summarize_table6_item(first)
    system_data = {
        "system_name": str(payload.get("system_name") or system.system_name).strip(),
        "proposed_level": min(max(final_level or 1, 1), 5),
        "business_description": str(table2.get("business_description") or payload.get("business_description") or "").strip() or None,
        "system_type": str(payload.get("system_type") or system.system_type or "").strip() or None,
        "deployment_mode": str(payload.get("deployment_mode") or system.deployment_mode or "").strip() or None,
        "go_live_date": parse_workspace_date(go_live) or system.go_live_date,
        "service_scope": str(payload.get("service_scope") or system.service_scope or "").strip() or None,
        "service_object": str(payload.get("service_object") or system.service_object or "").strip() or None,
        "boundary": str(payload.get("boundary") or system.boundary or "").strip() or None,
        "network_topology": str(payload.get("network_topology") or system.network_topology or "").strip() or None,
        "data_name": str(first_data_item.get("data_name") or payload.get("data_name") or system.data_name or "").strip() or None,
        "data_level": str(first_data_item.get("data_level") or payload.get("data_level") or system.data_level or "").strip() or None,
        "data_category": str(first_data_item.get("data_category") or payload.get("data_category") or system.data_category or "").strip() or None,
        "level_basis": str(payload.get("level_basis") or system.level_basis or "").strip() or None,
        "level_reason": str(payload.get("level_reason") or system.level_reason or "").strip() or None,
        "impact_scope": str(payload.get("impact_scope") or system.impact_scope or "").strip() or None,
        "needs_expert_review": table3.get("expert_review", {}).get("status") == "reviewed" if isinstance(table3.get("expert_review"), dict) else system.needs_expert_review,
    }
    validate_system_payload(system_data, partial=True)
    for key, value in system_data.items():
        setattr(system, key, value)
    system.filing_detail = merged_detail
    return before, obj_to_dict(system, SYSTEM_FIELDS)


def build_filing_overview_item(org: Organization, systems: list[SystemInfo]) -> dict[str, Any]:
    org_detail = merged_org_filing_detail(org)
    return {
        "organization": {
            "id": org.id,
            "name": org.name,
            "credit_code": org.credit_code,
            "industry": org.industry,
            "filing_region": org.filing_region,
            "archived": org.archived,
            "locked": org.locked,
            "filing_detail": org_detail,
        },
        "system_count": len(systems),
        "systems": [
            {
                "id": system.id,
                "system_name": system.system_name,
                "system_code": system.system_code,
                "proposed_level": system.proposed_level,
                "archived": system.archived,
                "locked": system.locked,
                "go_live_date": system.go_live_date.isoformat() if system.go_live_date else "",
                "updated_at": system.updated_at,
            }
            for system in systems
        ],
    }


def build_workspace_detail_response(db: Session, org: Organization, system: SystemInfo) -> dict[str, Any]:
    org_detail = merged_org_filing_detail(org)
    table1 = org_detail.get("table1", {})
    if isinstance(table1, dict):
        if not str(table1.get("unit_internet_addresses") or "").strip():
            table1["unit_internet_addresses"] = "无"
        table1["address_parts"] = build_org_address_parts(org, org_detail)
        responsible_person = table1.get("responsible_person") if isinstance(table1.get("responsible_person"), dict) else {}
        responsible_person.setdefault("name", org.legal_representative or "")
        responsible_person.setdefault("office_phone", org.office_phone or "")
        responsible_person.setdefault("email", org.email or "")
        table1["responsible_person"] = responsible_person
    system_detail = hydrate_workspace_attachments(db, merged_system_filing_detail(system))
    table2 = system_detail.get("table2", {})
    if isinstance(table2, dict):
        table2["go_live_date"] = format_workspace_date(table2.get("go_live_date") or system.go_live_date)
    table3 = system_detail.get("table3", {})
    if isinstance(table3, dict):
        table3["final_level"] = to_int_or_zero(table3.get("final_level")) or system.proposed_level
        table3["filled_date"] = str(table3.get("filled_date") or "").strip()
    return {
        "organization": {
            "id": org.id,
            "name": org.name,
            "credit_code": org.credit_code,
            "address": org.address,
            "office_phone": org.office_phone,
            "mobile_phone": org.mobile_phone,
            "email": org.email,
            "industry": org.industry,
            "organization_type": org.organization_type,
            "cybersecurity_dept": org.cybersecurity_dept,
            "cybersecurity_owner_name": org.cybersecurity_owner_name,
            "cybersecurity_owner_title": org.cybersecurity_owner_title,
            "cybersecurity_owner_phone": org.cybersecurity_owner_phone,
            "cybersecurity_owner_email": org.cybersecurity_owner_email,
            "data_security_dept": org.data_security_dept,
            "data_security_owner_name": org.data_security_owner_name,
            "data_security_owner_title": org.data_security_owner_title,
            "data_security_owner_phone": org.data_security_owner_phone,
            "data_security_owner_email": org.data_security_owner_email,
            "supervising_department": org.supervising_department,
            "filing_region": org.filing_region,
            "involves_state_secret": org.involves_state_secret,
            "is_cii": org.is_cii,
            "remark": org.remark,
            "archived": org.archived,
            "locked": org.locked,
            "filing_detail": org_detail,
        },
        "system": {
            "id": system.id,
            "organization_id": system.organization_id,
            "system_name": system.system_name,
            "system_code": system.system_code,
            "proposed_level": system.proposed_level,
            "business_description": system.business_description,
            "system_type": system.system_type,
            "deployment_mode": system.deployment_mode,
            "go_live_date": system.go_live_date.isoformat() if system.go_live_date else "",
            "service_scope": system.service_scope,
            "service_object": system.service_object,
            "boundary": system.boundary,
            "network_topology": system.network_topology,
            "data_name": system.data_name,
            "data_level": system.data_level,
            "data_category": system.data_category,
            "level_basis": system.level_basis,
            "level_reason": system.level_reason,
            "impact_scope": system.impact_scope,
            "needs_expert_review": system.needs_expert_review,
            "archived": system.archived,
            "locked": system.locked,
            "filing_detail": system_detail,
        },
    }

def get_workflow_config(db: Session) -> WorkflowConfig:
    config = db.query(WorkflowConfig).filter(WorkflowConfig.name == "default").first()
    if config:
        return config
    config = WorkflowConfig(name="default", steps_json=DEFAULT_WORKFLOW_STEPS, updated_by="system")
    db.add(config)
    db.commit()
    db.refresh(config)
    return config


def get_system_or_404(db: Session, system_id: int) -> SystemInfo:
    system = db.query(SystemInfo).filter(SystemInfo.id == system_id, SystemInfo.deleted_at.is_(None)).first()
    if not system:
        raise HTTPException(status_code=404, detail="系统不存在。")
    return system


def get_org_or_404(db: Session, org_id: int) -> Organization:
    org = db.query(Organization).filter(Organization.id == org_id, Organization.deleted_at.is_(None)).first()
    if not org:
        raise HTTPException(status_code=404, detail="单位不存在。")
    return org


def generate_system_code(db: Session) -> str:
    for _ in range(8):
        candidate = f"SYS-{datetime.now():%Y%m%d%H%M%S%f}-{secrets.token_hex(2).upper()}"
        exists = db.query(SystemInfo.id).filter(SystemInfo.system_code == candidate).first()
        if not exists:
            return candidate
    raise HTTPException(status_code=500, detail="系统编号生成失败，请重试。")


def build_report_title(report_type: str, system_name: str, version: int) -> str:
    title_map = {
        "filing_form": "定级备案表",
        "grading_report": "定级报告",
        "expert_review_form": "专家评审意见表",
    }
    return f"{title_map.get(report_type, '报告')}-{system_name}-V{version}"


def choose_report_template(
    db: Session,
    report_type: str,
    city: str,
    level: int | None,
    template_id: int | None = None,
) -> ReportTemplate | None:
    if template_id:
        row = db.query(ReportTemplate).filter(ReportTemplate.id == template_id, ReportTemplate.status == "enabled").first()
        return row
    query = db.query(ReportTemplate).filter(ReportTemplate.report_type == report_type, ReportTemplate.status == "enabled")
    all_rows = query.order_by(ReportTemplate.is_default.desc(), ReportTemplate.updated_at.desc()).all()
    if not all_rows:
        return None
    city_rows: list[ReportTemplate] = []
    level_rows: list[ReportTemplate] = []
    for row in all_rows:
        city_ok = not row.city or (city and row.city in city)
        level_ok = not row.protection_level or (level is not None and str(level) == str(row.protection_level).replace("级", ""))
        if city_ok and level_ok:
            return row
        if city_ok:
            city_rows.append(row)
        if level_ok:
            level_rows.append(row)
    if city_rows:
        return city_rows[0]
    if level_rows:
        return level_rows[0]
    return all_rows[0]


def apply_template_config_to_payload(payload: dict[str, Any], template: ReportTemplate | None) -> dict[str, Any]:
    if not template:
        return payload
    result = dict(payload)
    cfg = dict(template.config_json or {})
    custom_title = cfg.get("title")
    if custom_title:
        result["标题"] = str(custom_title)
    section_cfg = cfg.get("sections")
    if isinstance(section_cfg, list) and section_cfg:
        result["章节"] = section_cfg
    sign_cfg = cfg.get("signatures")
    if isinstance(sign_cfg, dict):
        result["签章"] = sign_cfg
    result["模板信息"] = {
        "template_id": template.id,
        "template_name": template.template_name,
        "city": template.city,
        "protection_level": template.protection_level,
        "version_no": template.version_no,
    }
    return result


def local_official_template_config(report_type: str, default_title: str) -> dict[str, Any]:
    if report_type == "filing_form":
        return {
            "title": default_title,
            "sections": [
                {
                    "名称": "单位信息",
                    "内容": {
                        "单位名称": "{{单位名称}}",
                        "统一社会信用代码": "{{统一社会信用代码}}",
                        "单位地址": "{{单位地址}}",
                        "单位负责人": "{{单位负责人}}",
                        "联系电话": "{{联系电话}}",
                        "邮箱": "{{邮箱}}",
                    },
                },
                {
                    "名称": "定级对象信息",
                    "内容": {
                        "系统名称": "{{系统名称}}",
                        "拟定等级": "{{拟定等级}}",
                        "部署方式": "{{部署方式}}",
                        "备案地区": "{{备案地区}}",
                    },
                },
            ],
            "source": "local_official_docx",
        }
    if report_type == "expert_review_form":
        return {
            "title": default_title,
            "sections": [
                {
                    "名称": "基础信息",
                    "内容": {
                        "单位名称": "{{单位名称}}",
                        "系统名称": "{{系统名称}}",
                        "系统自定安全级别": "{{拟定等级}}",
                        "项目负责人": "{{联系人}}",
                        "联系电话": "{{联系电话}}",
                    },
                },
                {
                    "名称": "专家意见",
                    "内容": {"专家结论": "", "评审专家组组长（签字）": "", "评审专家组成员（签字）": ""},
                },
            ],
            "source": "local_official_docx",
        }
    return {"title": default_title, "source": "local_official_docx"}


def find_latest_docx_by_pattern(pattern: str) -> Path | None:
    candidates = sorted(LOCAL_OFFICIAL_TEMPLATE_DIR.glob(pattern), key=lambda p: p.stat().st_mtime, reverse=True)
    return candidates[0] if candidates else None


def parse_docx_key_values(content: bytes) -> dict[str, str]:
    result: dict[str, str] = {}

    def parse_line(text_line: str) -> None:
        text = (text_line or "").strip()
        if not text:
            return
        delimiter = "：" if "：" in text else (":" if ":" in text else None)
        if not delimiter:
            return
        key, value = text.split(delimiter, 1)
        k = key.strip()
        v = value.strip()
        if k and v:
            result[k] = v

    def parse_table_pair(key_text: str, value_text: str) -> None:
        k = (key_text or "").strip()
        v = (value_text or "").strip()
        if not k or not v:
            return
        if "：" in k or ":" in k:
            return
        if len(v) <= 10 and normalize_word_field_key(v) in {"姓名", "办公电话", "移动电话", "电子邮件", "电子邮箱", "职务/职称", "职务职称"}:
            return
        result[k] = v

    def dedup_texts(values: list[str]) -> list[str]:
        compact: list[str] = []
        for raw in values:
            text = (raw or "").replace("\u3000", " ").strip()
            if not text:
                continue
            if not compact or compact[-1] != text:
                compact.append(text)
        return compact

    def is_cell_label(text: str) -> bool:
        normalized = normalize_word_field_key(text)
        if not normalized:
            return False
        return normalized in {
            "姓名",
            "办公电话",
            "移动电话",
            "电子邮件",
            "电子邮箱",
            "职务/职称",
            "职务职称",
            "行政区划代码",
        }

    def parse_with_python_docx() -> bool:
        doc = Document(io.BytesIO(content))
        for para in doc.paragraphs:
            parse_line(para.text)
        for table in doc.tables:
            for row in table.rows:
                cells = row.cells
                row_values = dedup_texts([cell.text for cell in cells])
                if len(row_values) >= 2:
                    parse_table_pair(row_values[0], row_values[1])
                    lead = row_values[0]
                    idx = 1
                    while idx + 1 < len(row_values):
                        if is_cell_label(row_values[idx]):
                            parse_table_pair(f"{lead}{row_values[idx]}", row_values[idx + 1])
                            idx += 2
                            continue
                        idx += 1
                for cell in cells:
                    lines = [line.strip() for line in (cell.text or "").splitlines() if line.strip()]
                    for line in lines:
                        parse_line(line)
        return True

    def parse_with_raw_xml_fallback() -> bool:
        try:
            zf = zipfile.ZipFile(io.BytesIO(content), mode="r")
        except zipfile.BadZipFile:
            return False
        with zf:
            if "word/document.xml" not in zf.namelist():
                return False
            try:
                root = ET.fromstring(zf.read("word/document.xml"))
            except Exception:
                return False
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        tag_p = ns + "p"
        tag_tbl = ns + "tbl"
        tag_tr = ns + "tr"
        tag_tc = ns + "tc"
        tag_t = ns + "t"
        tag_body = ns + "body"
        body = root.find(tag_body)
        if body is None:
            return False

        def node_text(node: ET.Element) -> str:
            return "".join((t.text or "") for t in node.iter(tag_t)).strip()

        for child in list(body):
            if child.tag == tag_p:
                parse_line(node_text(child))
            elif child.tag == tag_tbl:
                for tr in child.findall(".//" + tag_tr):
                    cells = tr.findall(tag_tc)
                    row_values = dedup_texts([node_text(tc) for tc in cells])
                    if len(row_values) >= 2:
                        parse_table_pair(row_values[0], row_values[1])
                        lead = row_values[0]
                        idx = 1
                        while idx + 1 < len(row_values):
                            if is_cell_label(row_values[idx]):
                                parse_table_pair(f"{lead}{row_values[idx]}", row_values[idx + 1])
                                idx += 2
                                continue
                            idx += 1
                    for tc in cells:
                        for p in tc.findall(".//" + tag_p):
                            parse_line(node_text(p))
        return True

    try:
        parse_with_python_docx()
    except Exception:
        if not parse_with_raw_xml_fallback():
            raise HTTPException(status_code=400, detail="Word文件无法解析，请确认使用 .docx 且文件未损坏。")
    return result


def normalize_word_field_key(raw_key: str) -> str:
    key = (raw_key or "").strip()
    key = re.sub(r"[（(].*?[）)]", "", key)
    key = key.replace(" ", "").replace("\u3000", "")
    return key.strip()


def infer_filing_region_from_address(address_text: str | None) -> str | None:
    text = (address_text or "").strip()
    if not text:
        return None
    text = text.replace("\n", " ")
    m = re.search(r"([\u4e00-\u9fa5]{2,20})\s*地\(区、市、州、盟\)", text)
    if m:
        return m.group(1)
    for ptn in [r"([\u4e00-\u9fa5]{2,20})市", r"([\u4e00-\u9fa5]{2,20})州", r"([\u4e00-\u9fa5]{2,20})盟", r"([\u4e00-\u9fa5]{2,20})地区"]:
        m = re.search(ptn, text)
        if m:
            return m.group(1)
    return None


def to_bool_text(value: str | None) -> bool:
    if not value:
        return False
    return value.strip().lower() in {"1", "true", "yes", "y", "是", "有"}


def parse_proposed_level(raw_value: Any) -> int:
    level_text = ("" if raw_value is None else str(raw_value)).strip().replace("级", "")
    if not level_text:
        raise ValueError("拟定等级不能为空。")
    try:
        level = int(level_text)
    except Exception as exc:
        raise ValueError(f"拟定等级格式错误: {exc}") from exc
    if level < 1 or level > 5:
        raise ValueError("拟定等级必须在1-5级范围内。")
    return level


def parse_optional_go_live_date(raw_value: Any) -> date | None:
    if raw_value is None:
        return None
    if isinstance(raw_value, datetime):
        return raw_value.date()
    if isinstance(raw_value, date):
        return raw_value
    date_text = str(raw_value).strip()
    if not date_text:
        return None
    date_text = date_text.replace("/", "-")
    if " " in date_text:
        date_text = date_text.split(" ", 1)[0]
    try:
        return date.fromisoformat(date_text)
    except Exception as exc:
        raise ValueError(f"上线时间格式错误(应为YYYY-MM-DD): {exc}") from exc


def ensure_file_exists(path_value: str | None, not_found_detail: str) -> Path:
    path = Path(str(path_value or "")).resolve()
    allowed_bases = [UPLOAD_DIR.resolve(), EXPORT_DIR.resolve()]
    if not any(path == base or path.is_relative_to(base) for base in allowed_bases):
        raise HTTPException(status_code=403, detail="文件路径非法。")
    if not path.exists() or not path.is_file():
        raise HTTPException(status_code=404, detail=not_found_detail)
    return path


def safe_delete_uploaded_file(path_value: str | None) -> None:
    """删除已上传的文件，仅允许删除 UPLOAD_DIR / EXPORT_DIR 下的文件。"""
    if not path_value:
        return
    path = Path(path_value).resolve()
    allowed_bases = [UPLOAD_DIR.resolve(), EXPORT_DIR.resolve()]
    if not any(path == base or path.is_relative_to(base) for base in allowed_bases):
        return
    if path.exists() and path.is_file():
        path.unlink()


def get_or_create_workflow_step_rules(db: Session, config: WorkflowConfig) -> dict[str, WorkflowStepRule]:
    rows = (
        db.query(WorkflowStepRule)
        .filter(WorkflowStepRule.config_name == config.name)
        .order_by(WorkflowStepRule.id.desc())
        .all()
    )
    rule_map: dict[str, WorkflowStepRule] = {}
    duplicate_rows: list[WorkflowStepRule] = []
    for row in rows:
        # 同一节点可能因历史缺陷出现多条配置，保留最新一条并清理旧数据。
        if row.step_name in rule_map:
            duplicate_rows.append(row)
            continue
        rule_map[row.step_name] = row
    changed = False
    for dup in duplicate_rows:
        db.delete(dup)
        changed = True
    for step in config.steps_json:
        if step not in rule_map:
            row = WorkflowStepRule(
                config_name=config.name,
                step_name=step,
                owner="system",
                time_limit_hours=24,
                enabled=True,
                updated_by="system",
            )
            db.add(row)
            db.flush()
            rule_map[step] = row
            changed = True
    if changed:
        db.flush()
    return rule_map


def recalc_workflow_due_at(db: Session, instance: WorkflowInstance, config: WorkflowConfig) -> None:
    steps = config.steps_json or []
    if instance.status != "in_progress":
        instance.due_at = None
        return
    if instance.current_step_index >= len(steps):
        instance.due_at = None
        return
    current_step = steps[instance.current_step_index]
    rule_map = get_or_create_workflow_step_rules(db, config)
    hours = 24
    if current_step in rule_map:
        hours = max(1, int(rule_map[current_step].time_limit_hours or 24))
    instance.due_at = datetime.now() + timedelta(hours=hours)


def get_workflow_step_owner(
    db: Session,
    config: WorkflowConfig,
    step_index: int,
    rule_map: dict[str, WorkflowStepRule] | None = None,
) -> str:
    steps = config.steps_json or []
    if step_index < 0 or step_index >= len(steps):
        return ""
    step_name = steps[step_index]
    current_rule_map = rule_map if rule_map is not None else get_or_create_workflow_step_rules(db, config)
    if step_name in current_rule_map:
        return current_rule_map[step_name].owner
    return "system"


def knowledge_snapshot(doc: KnowledgeDocument) -> dict[str, Any]:
    return {
        "title": doc.title,
        "keywords": doc.keywords,
        "city": doc.city,
        "district": doc.district,
        "doc_type": doc.doc_type,
        "protection_level": doc.protection_level,
        "version_no": doc.version_no,
        "status": doc.status,
        "file_name": doc.file_name,
        "file_path": doc.file_path,
        "file_size": doc.file_size,
        "uploaded_by": doc.uploaded_by,
    }


def log_knowledge_version(db: Session, doc: KnowledgeDocument, action: str, changed_by: str) -> None:
    db.add(
        KnowledgeDocumentVersion(
            document_id=doc.id,
            action=action,
            snapshot=knowledge_snapshot(doc),
            changed_by=changed_by,
        )
    )


ORG_WORD_MAP = {
    "单位名称": "name",
    "单位全称": "name",
    "备案单位": "name",
    "统一社会信用代码": "credit_code",
    "社会信用代码": "credit_code",
    "统一社会代码": "credit_code",
    "单位统一社会信用代码": "credit_code",
    "单位负责人": "legal_representative",
    "单位负责人姓名": "legal_representative",
    "单位主要负责人": "legal_representative",
    "主要负责人": "legal_representative",
    "负责人": "legal_representative",
    "单位地址": "address",
    "通讯地址": "address",
    "单位通讯地址": "address",
    "办公电话": "office_phone",
    "联系电话": "office_phone",
    "固定电话": "office_phone",
    "单位负责人办公电话": "office_phone",
    "网络安全责任部门联系人办公电话": "office_phone",
    "移动电话": "mobile_phone",
    "手机号": "mobile_phone",
    "手机号码": "mobile_phone",
    "联系电话手机": "mobile_phone",
    "单位负责人移动电话": "mobile_phone",
    "网络安全责任部门联系人移动电话": "mobile_phone",
    "邮箱": "email",
    "电子邮箱": "email",
    "电子邮件": "email",
    "单位负责人电子邮件": "email",
    "网络安全责任部门联系人电子邮件": "email",
    "所属行业": "industry",
    "行业类别": "industry",
    "单位类型": "organization_type",
    "单位性质": "organization_type",
    "备案地区": "filing_region",
    "备案地市": "filing_region",
    "属地": "filing_region",
    "主管部门": "supervising_department",
    "是否涉及国家秘密": "involves_state_secret",
    "是否关键信息基础设施": "is_cii",
    "备注": "remark",
}

SYSTEM_WORD_MAP = {
    "系统名称": "system_name",
    "系统编号": "system_code",
    "单位ID": "organization_id",
    "拟定等级": "proposed_level",
    "业务描述": "business_description",
    "部署方式": "deployment_mode",
    "系统类型": "system_type",
    "上线时间": "go_live_date",
    "定级依据": "level_basis",
    "定级理由": "level_reason",
    "系统边界": "boundary",
    "子系统": "subsystems",
    "服务对象": "service_object",
    "服务范围": "service_scope",
}


@app.get("/health")
def health() -> dict[str, str]:
    return {"status": "ok"}


@app.get("/", response_class=HTMLResponse)
def dashboard_page(request: Request) -> HTMLResponse:
    return templates.TemplateResponse(request, "index.html", {"request": request})


@app.get("/organizations", response_class=HTMLResponse)
def organizations_page(request: Request) -> HTMLResponse:
    return templates.TemplateResponse(request, "organizations.html", {"request": request})


@app.get("/organizations/systems/{system_id}", response_class=HTMLResponse)
def organization_system_detail_page(system_id: int, request: Request) -> HTMLResponse:
    return templates.TemplateResponse(
        request,
        "organization_system_detail.html",
        {
            "request": request,
            "system_id": system_id,
            "page_layout_variant": "management-tight",
        },
    )


@app.get("/login", response_class=HTMLResponse)
def login_page(request: Request) -> HTMLResponse:
    next_path = (request.query_params.get("next") or "/").strip() or "/"
    return templates.TemplateResponse(request, "auth.html", {"request": request, "mode": "login", "next_path": next_path})


@app.get("/register", response_class=HTMLResponse)
def register_page() -> RedirectResponse:
    return RedirectResponse(url="/login", status_code=307)


@app.get("/organizations/collect/{token}", response_class=HTMLResponse)
def organizations_collect_page(token: str, request: Request, db: Session = Depends(get_db)) -> HTMLResponse:
    link = db.query(OrganizationCollectionLink).filter(OrganizationCollectionLink.token == token).first()
    if not link:
        raise HTTPException(status_code=404, detail="链接不存在。")
    if link.disabled:
        raise HTTPException(status_code=403, detail="链接已停用。")
    if datetime.now() > link.expires_at:
        raise HTTPException(status_code=403, detail="链接已过期。")
    initial = {}
    if link.organization_id:
        org = db.query(Organization).filter(Organization.id == link.organization_id, Organization.deleted_at.is_(None)).first()
        if org:
            initial = obj_to_dict(org, ORG_FIELDS)
    return templates.TemplateResponse(request, "organization_collect.html", {"request": request, "token": token, "initial": initial})


@app.get("/systems", response_class=HTMLResponse)
def systems_page(request: Request) -> HTMLResponse:
    return RedirectResponse(url="/organizations?focus=systems", status_code=307)


@app.get("/users", response_class=HTMLResponse)
def users_page(request: Request, db: Session = Depends(get_db)) -> HTMLResponse:
    require_roles(request, db, {"admin"})
    return templates.TemplateResponse(request, "users.html", {"request": request})


@app.get("/reports", response_class=HTMLResponse)
def reports_page(request: Request) -> HTMLResponse:
    return templates.TemplateResponse(request, "reports.html", {"request": request})


@app.get("/workflow", response_class=HTMLResponse)
def workflow_page(request: Request) -> HTMLResponse:
    return templates.TemplateResponse(request, "workflow.html", {"request": request})


@app.get("/knowledge", response_class=HTMLResponse)
def knowledge_page(request: Request) -> HTMLResponse:
    return templates.TemplateResponse(request, "knowledge.html", {"request": request})


@app.get("/templates", response_class=HTMLResponse)
def templates_page(request: Request) -> HTMLResponse:
    return templates.TemplateResponse(request, "templates.html", {"request": request})


@app.get("/backup", response_class=HTMLResponse)
def backup_page(request: Request, db: Session = Depends(get_db)) -> HTMLResponse:
    require_roles(request, db, {"admin"})
    return templates.TemplateResponse(request, "backup.html", {"request": request})


@app.get("/api/backup/list")
def list_backups(request: Request, db: Session = Depends(get_db)) -> dict[str, Any]:
    require_roles(request, db, {"admin"})
    ensure_dirs()
    items: list[dict[str, Any]] = []
    for path in BACKUP_DIR.glob("*.zip"):
        if not path.is_file():
            continue
        stat = path.stat()
        items.append(
            {
                "file_name": path.name,
                "file_size": stat.st_size,
                "updated_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
            }
        )
    items.sort(key=lambda x: x["updated_at"], reverse=True)
    return {"items": items, "total": len(items)}


@app.post("/api/backup/create")
def create_backup(request: Request, db: Session = Depends(get_db)) -> dict[str, Any]:
    actor, _ = require_roles(request, db, {"admin"})
    backup_path = create_backup_archive(actor=actor, trigger="manual")
    stat = backup_path.stat()
    return {
        "message": "备份创建成功",
        "file_name": backup_path.name,
        "file_size": stat.st_size,
        "download_url": f"/api/backup/download/{quote(backup_path.name)}",
    }


@app.get("/api/backup/download/{file_name}")
def download_backup(file_name: str, request: Request, db: Session = Depends(get_db)) -> FileResponse:
    require_roles(request, db, {"admin"})
    clean_name = assert_safe_backup_file_name(file_name)
    backup_path = BACKUP_DIR / clean_name
    if not backup_path.exists() or not backup_path.is_file():
        raise HTTPException(status_code=404, detail="备份文件不存在。")
    return FileResponse(path=backup_path, filename=backup_path.name, media_type="application/zip")


@app.post("/api/backup/restore")
def restore_backup(
    request: Request,
    file: UploadFile = File(...),
    confirm: bool = Form(False),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor, _ = require_roles(request, db, {"admin"})
    if not confirm:
        raise HTTPException(status_code=400, detail="恢复操作需要 confirm=true。")
    if not file.filename:
        raise HTTPException(status_code=400, detail="请选择备份压缩包。")
    if not str(file.filename).lower().endswith(".zip"):
        raise HTTPException(status_code=400, detail="仅支持 .zip 备份文件。")

    ensure_dirs()
    temp_root = BACKUP_DIR / f"_restore_{uuid.uuid4().hex}"
    payload_dir = temp_root / "payload"
    uploaded_zip = temp_root / "uploaded.zip"
    temp_root.mkdir(parents=True, exist_ok=True)

    upload_size = 0
    try:
        with uploaded_zip.open("wb") as out:
            while True:
                chunk = file.file.read(1024 * 1024)
                if not chunk:
                    break
                upload_size += len(chunk)
                if upload_size > MAX_BACKUP_FILE:
                    raise HTTPException(status_code=400, detail="备份文件过大，已超过系统限制。")
                out.write(chunk)
        try:
            with zipfile.ZipFile(uploaded_zip, mode="r") as zf:
                safe_extract_zip(zf, payload_dir)
        except zipfile.BadZipFile as ex:
            raise HTTPException(status_code=400, detail="备份压缩包损坏或格式非法。") from ex

        incoming_db = payload_dir / "data" / "app.db"
        if not incoming_db.exists() or not incoming_db.is_file():
            raise HTTPException(status_code=400, detail="备份包缺少 data/app.db，无法恢复。")
        incoming_uploads = payload_dir / "uploads"

        pre_restore = create_backup_archive(actor=actor, trigger="pre_restore")

        db_file = resolve_sqlite_db_file()
        db.close()
        engine.dispose()

        db_file.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(incoming_db, db_file)

        moved_old_uploads = False
        old_uploads_dir = BACKUP_DIR / f"_uploads_before_restore_{uuid.uuid4().hex[:8]}"
        if incoming_uploads.exists() and incoming_uploads.is_dir():
            if UPLOAD_DIR.exists():
                UPLOAD_DIR.rename(old_uploads_dir)
                moved_old_uploads = True
            try:
                shutil.copytree(incoming_uploads, UPLOAD_DIR)
                if moved_old_uploads and old_uploads_dir.exists():
                    shutil.rmtree(old_uploads_dir, ignore_errors=True)
            except Exception:
                if UPLOAD_DIR.exists():
                    shutil.rmtree(UPLOAD_DIR, ignore_errors=True)
                if moved_old_uploads and old_uploads_dir.exists():
                    old_uploads_dir.rename(UPLOAD_DIR)
                raise

        ensure_dirs()
        init_db()
    finally:
        file.file.close()
        if temp_root.exists():
            shutil.rmtree(temp_root, ignore_errors=True)

    return {
        "message": "恢复成功",
        "pre_restore_backup": pre_restore.name,
    }


@app.post("/api/auth/login")
def auth_login(payload: dict[str, Any], request: Request, response: Response, db: Session = Depends(get_db)) -> dict[str, Any]:
    username = str(payload.get("username") or "").strip()
    password = str(payload.get("password") or "")
    assert_safe_text(username, "用户名")
    if not username or not password:
        raise HTTPException(status_code=400, detail="用户名和密码不能为空。")
    attempt_key = _login_attempt_key(username, request)
    check_login_locked(attempt_key)
    user = db.query(UserAccount).filter(UserAccount.username == username, UserAccount.enabled.is_(True)).first()
    ok, need_upgrade = verify_password(password, user.password_hash if user else "")
    if not user or not ok:
        record_login_failure(attempt_key)
        raise HTTPException(status_code=401, detail="用户名或密码错误。")
    clear_login_failure(attempt_key)
    if need_upgrade:
        user.password_hash = hash_password(password)
    token = secrets.token_hex(32)
    expires_at = datetime.now() + timedelta(hours=max(1, AUTH_SESSION_HOURS))
    db.add(AuthSession(user_id=user.id, token=token, expires_at=expires_at))
    user.last_login_at = datetime.now()
    db.commit()
    response.set_cookie(
        key="auth_token",
        value=token,
        httponly=True,
        samesite="lax",
        max_age=max(1, AUTH_SESSION_HOURS) * 3600,
        path="/",
    )
    return {
        "message": "登录成功",
        "token": token,
        "expires_at": expires_at,
        "user": {
            "username": user.username,
            "role": user.role,
            "must_change_password": bool(getattr(user, "must_change_password", False)),
        },
        "must_change_password": bool(getattr(user, "must_change_password", False)),
    }


@app.post("/api/auth/register")
def auth_register() -> dict[str, Any]:
    raise HTTPException(status_code=403, detail="系统未开放自助注册，请联系管理员分发账号。")


@app.post("/api/auth/logout")
def auth_logout(request: Request, db: Session = Depends(get_db)) -> JSONResponse:
    token = get_auth_token_from_request(request)
    response = JSONResponse(content={"message": "已退出"})
    response.delete_cookie("auth_token", path="/")
    if not token:
        return response
    row = db.query(AuthSession).filter(AuthSession.token == token).first()
    if row:
        db.delete(row)
        db.commit()
    return response


@app.get("/api/auth/me")
def auth_me(request: Request, db: Session = Depends(get_db)) -> dict[str, Any]:
    user = get_current_user_optional(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="未登录。")
    return {
        "username": user.username,
        "role": user.role,
        "enabled": user.enabled,
        "last_login_at": user.last_login_at,
        "must_change_password": bool(getattr(user, "must_change_password", False)),
        "password_updated_at": user.password_updated_at,
    }


@app.get("/api/auth/users")
def list_users(request: Request, db: Session = Depends(get_db)) -> dict[str, Any]:
    require_roles(request, db, {"admin"})
    rows = db.query(UserAccount).order_by(UserAccount.created_at.desc()).all()
    return {
        "total": len(rows),
        "items": [
            {
                "id": r.id,
                "username": r.username,
                "role": r.role,
                "enabled": r.enabled,
                "must_change_password": bool(getattr(r, "must_change_password", False)),
                "created_at": r.created_at,
                "last_login_at": r.last_login_at,
                "password_updated_at": r.password_updated_at,
            }
            for r in rows
        ],
    }


@app.post("/api/auth/users")
def create_user(request: Request, payload: dict[str, Any], db: Session = Depends(get_db)) -> dict[str, Any]:
    actor, _ = require_roles(request, db, {"admin"})
    username = str(payload.get("username") or "").strip()
    password = str(payload.get("password") or "")
    role = str(payload.get("role") or "evaluator").strip().lower()
    require_password_change = parse_optional_bool(payload.get("require_password_change"), "require_password_change")
    assert_safe_text(username, "用户名")
    if role not in {"admin", "reviewer", "evaluator"}:
        raise HTTPException(status_code=400, detail="role 仅支持 admin/reviewer/evaluator。")
    if not username or len(password) < 8:
        raise HTTPException(status_code=400, detail="用户名不能为空，密码至少8位。")
    if require_password_change is None:
        require_password_change = role != "admin"
    exists = db.query(UserAccount).filter(UserAccount.username == username).first()
    if exists:
        raise HTTPException(status_code=409, detail="用户名已存在。")
    row = UserAccount(
        username=username,
        password_hash=hash_password(password),
        role=role,
        enabled=True,
        must_change_password=require_password_change,
        password_updated_at=None if require_password_change else datetime.now(),
    )
    db.add(row)
    db.commit()
    return {
        "message": "账号创建成功",
        "data": {
            "id": row.id,
            "username": row.username,
            "role": row.role,
            "must_change_password": bool(getattr(row, "must_change_password", False)),
            "created_by": actor,
        },
    }


@app.post("/api/auth/change-password")
def change_password(request: Request, payload: dict[str, Any], db: Session = Depends(get_db)) -> dict[str, Any]:
    user = get_current_user_optional(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="请先登录。")
    current_password = str(payload.get("current_password") or "")
    new_password = str(payload.get("new_password") or "")
    if not current_password or not new_password:
        raise HTTPException(status_code=400, detail="当前密码和新密码不能为空。")
    if len(new_password) < 8:
        raise HTTPException(status_code=400, detail="新密码至少8位。")
    ok, _ = verify_password(current_password, user.password_hash)
    if not ok:
        raise HTTPException(status_code=400, detail="当前密码不正确。")
    if current_password == new_password:
        raise HTTPException(status_code=400, detail="新密码不能与当前密码相同。")
    user.password_hash = hash_password(new_password)
    user.must_change_password = False
    user.password_updated_at = datetime.now()
    db.query(AuthSession).filter(AuthSession.user_id == user.id).delete(synchronize_session=False)
    db.commit()
    return {"message": "密码修改成功，请使用新密码登录。"}


@app.post("/api/auth/users/{user_id}/reset-password")
def reset_user_password(request: Request, user_id: int, payload: dict[str, Any], db: Session = Depends(get_db)) -> dict[str, Any]:
    actor, _ = require_roles(request, db, {"admin"})
    user = db.query(UserAccount).filter(UserAccount.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="用户不存在。")
    new_password = str(payload.get("new_password") or "").strip()
    if len(new_password) < 8:
        raise HTTPException(status_code=400, detail="重置密码至少8位。")
    user.password_hash = hash_password(new_password)
    user.must_change_password = True
    user.password_updated_at = None
    db.query(AuthSession).filter(AuthSession.user_id == user.id).delete(synchronize_session=False)
    db.commit()
    return {
        "message": "密码重置成功，用户下次登录需修改密码。",
        "data": {
            "user_id": user.id,
            "username": user.username,
            "reset_by": actor,
            "must_change_password": True,
        },
    }


@app.post("/api/auth/users/{user_id}/toggle")
def toggle_user(request: Request, user_id: int, enabled: bool = Query(True), db: Session = Depends(get_db)) -> dict[str, Any]:
    require_roles(request, db, {"admin"})
    user = db.query(UserAccount).filter(UserAccount.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="用户不存在。")
    user.enabled = enabled
    db.commit()
    return {"message": "状态更新成功", "enabled": enabled}


@app.post("/api/templates/upload")
async def upload_report_template(
    request: Request,
    template_name: str = Form(...),
    report_type: str = Form(...),
    category: str = Form(""),
    city: str = Form(""),
    protection_level: str = Form(""),
    description: str = Form(""),
    is_default: bool = Form(False),
    config_json: str = Form(""),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor, _ = require_roles(request, db, {"admin"})
    report_type = report_type.strip()
    assert_safe_text(template_name, "template_name")
    assert_safe_text(category, "category")
    assert_safe_text(city, "city")
    assert_safe_text(protection_level, "protection_level")
    assert_safe_text(description, "description")
    if report_type not in VALID_REPORT_TYPES:
        raise HTTPException(status_code=400, detail="report_type 无效。")
    content = await file.read()
    if len(content) > MAX_KNOWLEDGE_FILE:
        raise HTTPException(status_code=400, detail="模板文件过大。")
    safe_name = Path(file.filename).name
    save_name = f"{uuid.uuid4().hex}_{safe_name}"
    target_dir = UPLOAD_DIR / "templates"
    target_dir.mkdir(parents=True, exist_ok=True)
    path = target_dir / save_name
    cfg = None
    if config_json.strip():
        try:
            import json as _json

            cfg = _json.loads(config_json)
            assert_safe_payload(cfg, "config_json")
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"config_json 解析失败: {exc}") from exc
    path.write_bytes(content)
    if is_default:
        exists = db.query(ReportTemplate).filter(ReportTemplate.report_type == report_type, ReportTemplate.is_default.is_(True)).all()
        for row in exists:
            row.is_default = False
    row = ReportTemplate(
        template_name=template_name.strip(),
        report_type=report_type,
        category=category.strip() or None,
        city=city.strip() or None,
        protection_level=protection_level.strip() or None,
        description=description.strip() or None,
        status="enabled",
        is_default=is_default,
        version_no=1,
        file_name=safe_name,
        file_path=str(path),
        file_size=len(content),
        config_json=cfg,
        created_by=actor,
    )
    db.add(row)
    db.flush()
    log_template_version(db, row, "upload", actor)
    db.commit()
    return {"message": "模板上传成功", "data": {"id": row.id, "template_name": row.template_name}}


@app.post("/api/templates/import-local-official")
def import_local_official_templates(
    request: Request,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor_name, _ = require_roles(request, db, {"admin"}, legacy_admin=True)
    imported: list[dict[str, Any]] = []
    missing: list[str] = []
    target_dir = UPLOAD_DIR / "templates"
    target_dir.mkdir(parents=True, exist_ok=True)

    for rule in LOCAL_OFFICIAL_TEMPLATE_RULES:
        src = find_latest_docx_by_pattern(rule["pattern"])
        if not src or not src.exists():
            missing.append(rule["pattern"])
            continue

        content = src.read_bytes()
        cleaned_content = sanitize_template_docx_content(content)
        safe_name = src.name
        save_name = f"{uuid.uuid4().hex}_{safe_name}"
        path = target_dir / save_name
        path.write_bytes(cleaned_content)
        config = local_official_template_config(rule["report_type"], rule["default_title"])
        config["source_file"] = safe_name

        defaults = (
            db.query(ReportTemplate)
            .filter(ReportTemplate.report_type == rule["report_type"], ReportTemplate.is_default.is_(True))
            .all()
        )
        for old in defaults:
            old.is_default = False

        row = (
            db.query(ReportTemplate)
            .filter(
                ReportTemplate.report_type == rule["report_type"],
                ReportTemplate.template_name == rule["template_name"],
                ReportTemplate.category == rule["category"],
            )
            .first()
        )

        if row:
            log_template_version(db, row, "import_local_official_before", actor_name)
            row.status = "enabled"
            row.is_default = True
            row.version_no = int(row.version_no or 0) + 1
            row.file_name = safe_name
            row.file_path = str(path)
            row.file_size = len(cleaned_content)
            row.config_json = config
            row.description = f"来源本地模板文件：{safe_name}（已去除测试数据）"
            row.created_by = actor_name
        else:
            row = ReportTemplate(
                template_name=rule["template_name"],
                report_type=rule["report_type"],
                category=rule["category"],
                city=None,
                protection_level=None,
                description=f"来源本地模板文件：{safe_name}（已去除测试数据）",
                status="enabled",
                is_default=True,
                version_no=1,
                file_name=safe_name,
                file_path=str(path),
                file_size=len(cleaned_content),
                config_json=config,
                created_by=actor_name,
            )
            db.add(row)
            db.flush()
            log_template_version(db, row, "import_local_official", actor_name)

        imported.append(
            {
                "report_type": rule["report_type"],
                "template_name": rule["template_name"],
                "source_file": safe_name,
                "is_default": True,
            }
        )

    db.commit()
    return {"message": "本地官方模板导入完成", "imported": imported, "missing": missing}


@app.get("/api/templates")
def list_report_templates(
    report_type: str | None = None,
    category: str | None = None,
    city: str | None = None,
    protection_level: str | None = None,
    enabled_only: bool = True,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    query = db.query(ReportTemplate)
    if enabled_only:
        query = query.filter(ReportTemplate.status == "enabled")
    if report_type:
        query = query.filter(ReportTemplate.report_type == report_type)
    if category:
        query = query.filter(ReportTemplate.category.like(f"%{escape_like(category)}%"))
    if city:
        query = query.filter(ReportTemplate.city.like(f"%{escape_like(city)}%"))
    if protection_level:
        query = query.filter(ReportTemplate.protection_level.like(f"%{escape_like(protection_level)}%"))
    rows = query.order_by(ReportTemplate.is_default.desc(), ReportTemplate.updated_at.desc()).all()
    return {
        "total": len(rows),
        "items": [
            {
                "id": r.id,
                "template_name": r.template_name,
                "report_type": r.report_type,
                "category": r.category,
                "city": r.city,
                "protection_level": r.protection_level,
                "description": r.description,
                "status": r.status,
                "is_default": r.is_default,
                "version_no": r.version_no,
                "file_name": r.file_name,
                "created_by": r.created_by,
                "created_at": r.created_at,
                "updated_at": r.updated_at,
            }
            for r in rows
        ],
    }


@app.put("/api/templates/{template_id}")
def update_report_template(
    request: Request,
    template_id: int,
    payload: dict[str, Any],
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor, _ = require_roles(request, db, {"admin"})
    assert_safe_payload(payload, "payload")
    row = db.query(ReportTemplate).filter(ReportTemplate.id == template_id).first()
    if not row:
        raise HTTPException(status_code=404, detail="模板不存在。")
    log_template_version(db, row, "update_before", actor)
    for key in ["template_name", "category", "city", "protection_level", "description", "status", "config_json"]:
        if key in payload:
            value = payload[key]
            if value is None and key in {"template_name", "status"}:
                raise HTTPException(status_code=400, detail=f"{key} 不能为空。")
            if isinstance(value, str):
                value = value.strip()
            if value == "":
                if key in {"template_name", "status"}:
                    raise HTTPException(status_code=400, detail=f"{key} 不能为空。")
                value = None
            if key == "template_name" and value is not None:
                assert_safe_text(str(value), "template_name")
            if key == "status" and value is not None and str(value) not in {"enabled", "disabled"}:
                raise HTTPException(status_code=400, detail="status 仅支持 enabled/disabled。")
            setattr(row, key, value)
    if "is_default" in payload:
        parsed_default = parse_optional_bool(payload.get("is_default"), "is_default")
        if parsed_default is True:
            exists = (
                db.query(ReportTemplate)
                .filter(ReportTemplate.report_type == row.report_type, ReportTemplate.is_default.is_(True))
                .all()
            )
            for old in exists:
                old.is_default = False
            row.is_default = True
        elif parsed_default is False:
            row.is_default = False
    db.commit()
    return {"message": "模板更新成功"}


@app.post("/api/templates/{template_id}/new-version")
async def new_report_template_version(
    request: Request,
    template_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor, _ = require_roles(request, db, {"admin"})
    row = db.query(ReportTemplate).filter(ReportTemplate.id == template_id).first()
    if not row:
        raise HTTPException(status_code=404, detail="模板不存在。")
    content = await file.read()
    if len(content) > MAX_KNOWLEDGE_FILE:
        raise HTTPException(status_code=400, detail="模板文件过大。")
    log_template_version(db, row, "new_version_before", actor)
    safe_name = Path(file.filename).name
    save_name = f"{uuid.uuid4().hex}_{safe_name}"
    target_dir = UPLOAD_DIR / "templates"
    target_dir.mkdir(parents=True, exist_ok=True)
    path = target_dir / save_name
    path.write_bytes(content)
    row.file_name = safe_name
    row.file_path = str(path)
    row.file_size = len(content)
    row.version_no = int(row.version_no or 0) + 1
    db.commit()
    return {"message": "模板新版本上传成功", "version_no": row.version_no}


@app.get("/api/templates/{template_id}/versions")
def template_versions(template_id: int, db: Session = Depends(get_db)) -> dict[str, Any]:
    tpl = db.query(ReportTemplate).filter(ReportTemplate.id == template_id).first()
    if not tpl:
        raise HTTPException(status_code=404, detail="模板不存在。")
    rows = (
        db.query(ReportTemplateVersion)
        .filter(ReportTemplateVersion.template_id == template_id)
        .order_by(ReportTemplateVersion.changed_at.desc())
        .all()
    )
    return {
        "total": len(rows),
        "items": [
            {
                "id": r.id,
                "action": r.action,
                "changed_by": r.changed_by,
                "changed_at": r.changed_at,
                "snapshot": r.snapshot,
            }
            for r in rows
        ],
    }


@app.post("/api/templates/{template_id}/restore/{version_id}")
def restore_template_version(
    request: Request,
    template_id: int,
    version_id: int,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor, _ = require_roles(request, db, {"admin"})
    tpl = db.query(ReportTemplate).filter(ReportTemplate.id == template_id).first()
    if not tpl:
        raise HTTPException(status_code=404, detail="模板不存在。")
    row = (
        db.query(ReportTemplateVersion)
        .filter(ReportTemplateVersion.id == version_id, ReportTemplateVersion.template_id == template_id)
        .first()
    )
    if not row:
        raise HTTPException(status_code=404, detail="模板版本不存在。")
    snap = dict(row.snapshot or {})
    if not snap:
        raise HTTPException(status_code=400, detail="版本快照为空。")
    target_file = ensure_file_exists(snap.get("file_path"), "版本文件不存在，无法回滚。")
    log_template_version(db, tpl, "restore_before", actor)
    for key in [
        "template_name",
        "category",
        "city",
        "protection_level",
        "description",
        "status",
        "is_default",
        "file_name",
        "file_path",
        "file_size",
        "config_json",
    ]:
        if key in snap:
            setattr(tpl, key, snap[key])
    if bool(tpl.is_default):
        others = (
            db.query(ReportTemplate)
            .filter(
                ReportTemplate.report_type == tpl.report_type,
                ReportTemplate.id != tpl.id,
                ReportTemplate.is_default.is_(True),
            )
            .all()
        )
        for old in others:
            old.is_default = False
    tpl.version_no = int(tpl.version_no or 0) + 1
    db.commit()
    return {"message": "模板版本恢复成功", "version_no": tpl.version_no}


@app.post("/api/templates/{template_id}/set-default")
def set_default_template(request: Request, template_id: int, db: Session = Depends(get_db)) -> dict[str, Any]:
    require_roles(request, db, {"admin"})
    tpl = db.query(ReportTemplate).filter(ReportTemplate.id == template_id).first()
    if not tpl:
        raise HTTPException(status_code=404, detail="模板不存在。")
    old = db.query(ReportTemplate).filter(ReportTemplate.report_type == tpl.report_type, ReportTemplate.is_default.is_(True)).all()
    for row in old:
        row.is_default = False
    tpl.is_default = True
    db.commit()
    return {"message": "默认模板已更新", "template_id": tpl.id}


@app.get("/api/templates/{template_id}/download")
def download_template(template_id: int, db: Session = Depends(get_db)) -> FileResponse:
    tpl = db.query(ReportTemplate).filter(ReportTemplate.id == template_id).first()
    if not tpl:
        raise HTTPException(status_code=404, detail="模板不存在。")
    path = ensure_file_exists(tpl.file_path, "模板文件不存在或已被移除。")
    return FileResponse(path=str(path), filename=tpl.file_name)


@app.post("/api/templates/{template_id}/test-fill")
def test_fill_template(
    request: Request,
    template_id: int,
    system_id: int | None = Query(None),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    require_roles(request, db, {"admin", "evaluator"}, legacy_admin=True)
    tpl = db.query(ReportTemplate).filter(ReportTemplate.id == template_id).first()
    if not tpl:
        raise HTTPException(status_code=404, detail="模板不存在。")
    sample = {
        "单位名称": "示例单位",
        "系统名称": "示例系统",
        "拟定等级": "三级",
        "备案地区": tpl.city or "示例地区",
    }
    if system_id:
        system = get_system_or_404(db, system_id)
        org = get_org_or_404(db, system.organization_id)
        sample.update(
            {
                "单位名称": org.name,
                "系统名称": system.system_name,
                "拟定等级": f"{system.proposed_level}级",
                "备案地区": org.filing_region,
            }
        )
    cfg = tpl.config_json or {}
    preview = {"template_name": tpl.template_name, "sample": sample, "config": cfg}
    return {"message": "模板测试填充完成", "preview": preview}


@app.delete("/api/templates/{template_id}")
def delete_template(request: Request, template_id: int, db: Session = Depends(get_db)) -> dict[str, Any]:
    actor, _ = require_roles(request, db, {"admin"})
    tpl = db.query(ReportTemplate).filter(ReportTemplate.id == template_id).first()
    if not tpl:
        raise HTTPException(status_code=404, detail="模板不存在。")
    log_template_version(db, tpl, "delete_before", actor)
    safe_delete_uploaded_file(tpl.file_path)
    db.delete(tpl)
    db.commit()
    return {"message": "模板已删除"}


@app.post("/api/organizations")
def create_organization(request: Request, payload: OrganizationCreate, db: Session = Depends(get_db)) -> dict[str, Any]:
    actor_name, _ = require_roles(request, db, {"admin", "evaluator"}, legacy_admin=True)
    data = payload.model_dump()
    data["created_by"] = resolve_effective_actor_name(request, db, payload.created_by or actor_name)
    normalize_org_payload(data)
    validate_org_payload(data)
    assert_credit_code_available(db, data["credit_code"])
    org = Organization(**data)
    db.add(org)
    db.flush()
    db.refresh(org)
    record_org_history(db, org.id, data["created_by"], "create", None, obj_to_dict(org, ORG_FIELDS))
    db.commit()
    return {"message": "创建成功", "data": obj_to_dict(org, ORG_FIELDS)}


@app.get("/api/organizations")
def list_organizations(
    name: str | None = None,
    credit_code: str | None = None,
    industry: str | None = None,
    filing_region: str | None = None,
    created_by: str | None = None,
    include_deleted: bool = False,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    query = db.query(Organization)
    if not include_deleted:
        query = query.filter(Organization.deleted_at.is_(None))
    if name:
        query = query.filter(Organization.name.like(f"%{escape_like(name)}%"))
    if credit_code:
        query = query.filter(Organization.credit_code.like(f"%{escape_like(credit_code)}%"))
    if industry:
        query = query.filter(Organization.industry.like(f"%{escape_like(industry)}%"))
    if filing_region:
        query = query.filter(Organization.filing_region.like(f"%{escape_like(filing_region)}%"))
    if created_by:
        query = query.filter(Organization.created_by.like(f"%{escape_like(created_by)}%"))
    items = query.order_by(Organization.created_at.desc()).all()
    return {"total": len(items), "items": [obj_to_dict(i, ORG_FIELDS) for i in items]}


@app.post("/api/organizations/collection-links")
def create_org_collection_link(
    request: Request,
    organization_id: int | None = Query(None),
    expires_days: int = Query(7, ge=1, le=30),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor_name, _ = require_roles(request, db, {"admin", "reviewer", "evaluator"}, legacy_admin=True)
    if organization_id:
        get_org_or_404(db, organization_id)
    token = uuid.uuid4().hex
    link = OrganizationCollectionLink(
        token=token,
        organization_id=organization_id,
        created_by=actor_name,
        expires_at=datetime.now() + timedelta(days=expires_days),
        disabled=False,
    )
    db.add(link)
    db.commit()
    db.refresh(link)
    collect_url = f"{PUBLIC_BASE_URL}/organizations/collect/{link.token}"
    return {
        "message": "采集链接创建成功",
        "data": {
            "id": link.id,
            "token": link.token,
            "organization_id": link.organization_id,
            "expires_at": link.expires_at,
            "collect_url": collect_url,
            "qrcode_url": f"https://quickchart.io/qr?size=220&text={quote(collect_url)}",
        },
    }


@app.get("/api/organizations/collection-links")
def list_org_collection_links(
    request: Request,
    include_expired: bool = Query(False),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    require_roles(request, db, {"admin", "reviewer", "evaluator"})
    query = db.query(OrganizationCollectionLink)
    if not include_expired:
        query = query.filter(OrganizationCollectionLink.expires_at > datetime.now(), OrganizationCollectionLink.disabled.is_(False))
    rows = query.order_by(OrganizationCollectionLink.created_at.desc()).all()
    return {
        "total": len(rows),
        "items": [
            {
                "id": row.id,
                "token": row.token,
                "organization_id": row.organization_id,
                "created_by": row.created_by,
                "expires_at": row.expires_at,
                "disabled": row.disabled,
                "created_at": row.created_at,
                "collect_url": f"{PUBLIC_BASE_URL}/organizations/collect/{row.token}",
            }
            for row in rows
        ],
    }


@app.post("/api/organizations/collection-links/{link_id}/toggle")
def toggle_org_collection_link(
    request: Request,
    link_id: int,
    enabled: bool = Query(True),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    require_roles(request, db, {"admin", "reviewer", "evaluator"})
    row = db.query(OrganizationCollectionLink).filter(OrganizationCollectionLink.id == link_id).first()
    if not row:
        raise HTTPException(status_code=404, detail="采集链接不存在。")
    row.disabled = not enabled
    db.commit()
    return {"message": "状态更新成功", "enabled": enabled}


@app.post("/api/public/organizations/collect/{token}")
def submit_org_collection(token: str, payload: dict[str, Any], db: Session = Depends(get_db)) -> dict[str, Any]:
    link = db.query(OrganizationCollectionLink).filter(OrganizationCollectionLink.token == token).first()
    if not link:
        raise HTTPException(status_code=404, detail="链接不存在。")
    if link.disabled:
        raise HTTPException(status_code=403, detail="链接已停用。")
    if datetime.now() > link.expires_at:
        raise HTTPException(status_code=403, detail="链接已过期。")
    assert_safe_payload(payload, "payload")
    normalize_org_payload(payload)
    if link.organization_id:
        validate_org_partial(payload)
    else:
        required = ["name", "credit_code", "legal_representative", "address", "mobile_phone", "email", "industry", "organization_type", "filing_region"]
        missing = [k for k in required if not payload.get(k)]
        if missing:
            raise HTTPException(status_code=400, detail=f"缺少必填项: {', '.join(missing)}")
        normalize_org_payload(payload)
        validate_org_payload(payload)
    row = OrganizationSubmission(
        link_id=link.id,
        organization_id=link.organization_id,
        payload=payload,
        submitter=str(payload.get("submitter") or "customer"),
        status="pending",
    )
    db.add(row)
    db.commit()
    db.refresh(row)
    return {"message": "提交成功，待测评师审核。", "data": {"submission_id": row.id, "status": row.status}}


@app.get("/api/organizations/submissions")
def list_org_submissions(
    request: Request,
    status: str | None = None,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    require_roles(request, db, {"admin", "reviewer", "evaluator"})
    query = db.query(OrganizationSubmission)
    if status:
        query = query.filter(OrganizationSubmission.status == status)
    rows = query.order_by(OrganizationSubmission.submitted_at.desc()).all()
    return {
        "total": len(rows),
        "items": [
            {
                "id": r.id,
                "link_id": r.link_id,
                "organization_id": r.organization_id,
                "status": r.status,
                "payload": r.payload,
                "submitter": r.submitter,
                "review_comment": r.review_comment,
                "reviewed_by": r.reviewed_by,
                "reviewed_at": r.reviewed_at,
                "submitted_at": r.submitted_at,
            }
            for r in rows
        ],
    }


@app.post("/api/organizations/submissions/{submission_id}/review")
def review_org_submission(
    request: Request,
    submission_id: int,
    action: str = Query(...),
    comment: str = Query(""),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor_name, _ = require_roles(request, db, {"admin", "reviewer", "evaluator"}, legacy_admin=True)
    assert_safe_text(comment, "comment")
    row = db.query(OrganizationSubmission).filter(OrganizationSubmission.id == submission_id).first()
    if not row:
        raise HTTPException(status_code=404, detail="提交记录不存在。")
    if row.status != "pending":
        raise HTTPException(status_code=400, detail="仅待审核记录可处理。")
    action = action.lower()
    if action not in {"approve", "reject"}:
        raise HTTPException(status_code=400, detail="action 仅支持 approve/reject。")
    if action == "reject":
        row.status = "rejected"
        row.review_comment = comment
        row.reviewed_by = actor_name
        row.reviewed_at = datetime.now()
        db.commit()
        return {"message": "已驳回", "status": row.status}

    payload = dict(row.payload or {})
    payload.pop("submitter", None)
    assert_safe_payload(payload, "payload")
    if row.organization_id:
        org = get_org_or_404(db, row.organization_id)
        before = obj_to_dict(org, ORG_FIELDS)
        payload = {k: v for k, v in payload.items() if k in ORG_UPDATE_FIELDS}
        normalize_org_payload(payload)
        validate_org_partial(payload)
        if "credit_code" in payload:
            assert_credit_code_available(db, payload["credit_code"], current_org_id=org.id)
        for key, value in payload.items():
            setattr(org, key, value)
        db.flush()
        record_org_history(db, org.id, actor_name, "customer_approve_update", before, obj_to_dict(org, ORG_FIELDS))
        target_org_id = org.id
    else:
        required = ["name", "credit_code", "legal_representative", "address", "mobile_phone", "email", "industry", "organization_type", "filing_region"]
        missing = [k for k in required if not payload.get(k)]
        if missing:
            raise HTTPException(status_code=400, detail=f"缺少必填项: {', '.join(missing)}")
        payload = {k: v for k, v in payload.items() if k in ORG_CREATE_FIELDS}
        payload["created_by"] = actor_name
        normalize_org_payload(payload)
        validate_org_payload(payload)
        assert_credit_code_available(db, payload["credit_code"])
        org = Organization(**payload)
        db.add(org)
        db.flush()
        record_org_history(db, org.id, actor_name, "customer_approve_create", None, obj_to_dict(org, ORG_FIELDS))
        target_org_id = org.id
        row.organization_id = org.id

    row.status = "approved"
    row.review_comment = comment
    row.reviewed_by = actor_name
    row.reviewed_at = datetime.now()
    db.commit()
    return {"message": "审核通过并已入库", "status": row.status, "organization_id": target_org_id}


@app.get("/api/organizations/recycle-bin/list")
def list_organization_recycle_bin(
    include_expired: bool = Query(False),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    rows = (
        db.query(Organization)
        .filter(Organization.deleted_at.is_not(None))
        .order_by(Organization.deleted_at.desc())
        .all()
    )
    items: list[dict[str, Any]] = []
    for row in rows:
        meta = recycle_meta(row.deleted_at)
        if meta["expired"] and not include_expired:
            continue
        data = obj_to_dict(row, ORG_FIELDS)
        data.update(meta)
        items.append(data)
    return {"total": len(items), "items": items}


@app.post("/api/organizations/recycle-bin/cleanup")
def cleanup_organization_recycle_bin(
    request: Request,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor_name, _ = require_roles(request, db, {"admin"}, legacy_admin=True)
    threshold = datetime.now() - timedelta(days=30)
    rows = db.query(Organization).filter(Organization.deleted_at.is_not(None), Organization.deleted_at <= threshold).all()
    purged = 0
    skipped: list[str] = []
    for row in rows:
        linked = db.query(SystemInfo).filter(SystemInfo.organization_id == row.id).count()
        if linked > 0:
            skipped.append(f"{row.id}:{row.name} 仍有关联系统({linked})")
            continue
        db.delete(row)
        purged += 1
    db.commit()
    return {"message": "单位回收站清理完成", "purged": purged, "skipped": skipped, "actor": actor_name}


@app.get("/api/organizations/export/excel")
def export_organizations_excel(request: Request, db: Session = Depends(get_db)) -> FileResponse:
    require_roles(request, db, {"admin", "evaluator"}, legacy_admin=True)
    rows = db.query(Organization).filter(Organization.deleted_at.is_(None)).all()
    wb = Workbook()
    ws = wb.active
    ws.title = "单位信息"
    headers = [
        "单位名称",
        "统一社会信用代码",
        "单位负责人",
        "单位地址",
        "办公电话",
        "移动电话",
        "邮箱",
        "所属行业",
        "单位类型",
        "备案地区",
    ]
    ws.append(headers)
    for i in rows:
        ws.append(
            [
                i.name,
                i.credit_code,
                i.legal_representative,
                i.address,
                i.office_phone,
                i.mobile_phone,
                i.email,
                i.industry,
                i.organization_type,
                i.filing_region,
            ]
        )
    path = EXPORT_DIR / f"organizations_{datetime.now():%Y%m%d%H%M%S}.xlsx"
    wb.save(path)
    return FileResponse(path=str(path), filename=path.name, background=BackgroundTask(_cleanup_export_file, path))


@app.post("/api/organizations/import/excel")
async def import_organizations_excel(
    request: Request, file: UploadFile = File(...), db: Session = Depends(get_db)
) -> dict[str, Any]:
    actor, _ = require_roles(request, db, {"admin"}, legacy_admin=True)
    if not file.filename or not file.filename.lower().endswith(".xlsx"):
        raise HTTPException(status_code=400, detail="仅支持xlsx格式。")
    content = await file.read()
    # 限制导入文件大小，防止恶意上传超大文件
    if len(content) > 50 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="导入文件过大，请限制在50MB以内。")
    wb = load_workbook(io.BytesIO(content))
    ws = wb.active
    imported = 0
    skipped: list[str] = []
    for idx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
        if not row or not row[0]:
            continue
        if len(row) < 10:
            skipped.append(f"第{idx}行：列数不足（需要至少10列）")
            continue
        data = {
            "name": str(row[0] or "").strip(),
            "credit_code": str(row[1] or "").strip().upper(),
            "legal_representative": str(row[2] or "").strip(),
            "address": str(row[3] or "").strip(),
            "office_phone": str(row[4]).strip() if row[4] else None,
            "mobile_phone": str(row[5] or "").strip(),
            "email": str(row[6] or "").strip(),
            "industry": str(row[7] or "").strip(),
            "organization_type": str(row[8] or "").strip(),
            "filing_region": str(row[9] or "").strip(),
            "created_by": actor,
        }
        try:
            with db.begin_nested():
                normalize_org_payload(data)
                validate_org_payload(data)
                assert_credit_code_available(db, data["credit_code"])
                org = Organization(**data)
                db.add(org)
                db.flush()
                record_org_history(db, org.id, actor, "import", None, obj_to_dict(org, ORG_FIELDS))
            imported += 1
        except Exception as exc:
            skipped.append(f"第{idx}行：{exc}")
    db.commit()
    return {"message": "导入完成", "imported": imported, "skipped": skipped}


@app.post("/api/organizations/import/word")
async def import_organization_word(
    request: Request,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor, _ = require_roles(request, db, {"admin"}, legacy_admin=True)
    if not file.filename or not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="仅支持docx格式。")
    content = await file.read()
    kv = parse_docx_key_values(content)
    raw: dict[str, Any] = {}

    def value_quality(text: str) -> int:
        value = (text or "").strip()
        if not value:
            return 0
        if value in {"/", "／"}:
            return 1
        norm = normalize_word_field_key(value)
        if norm in {"姓名", "办公电话", "移动电话", "电子邮件", "电子邮箱", "职务/职称", "职务职称"}:
            return 1
        return 2

    for raw_key, raw_value in kv.items():
        norm_key = normalize_word_field_key(raw_key)
        model_key = ORG_WORD_MAP.get(norm_key)
        if not model_key:
            continue
        if raw_value is None:
            continue
        value_text = str(raw_value).strip()
        if not value_text:
            continue
        current = str(raw.get(model_key) or "").strip()
        if value_quality(value_text) < value_quality(current):
            continue
        raw[model_key] = value_text

    if not str(raw.get("filing_region") or "").strip():
        inferred_region = infer_filing_region_from_address(str(raw.get("address") or ""))
        if inferred_region:
            raw["filing_region"] = inferred_region
    if "involves_state_secret" in raw:
        raw["involves_state_secret"] = to_bool_text(str(raw["involves_state_secret"]))
    if "is_cii" in raw:
        raw["is_cii"] = to_bool_text(str(raw["is_cii"]))
    if str(raw.get("office_phone") or "").strip() in {"/", "／", "-", "无", "暂无"}:
        raw["office_phone"] = None
    raw["created_by"] = actor
    required = ["name", "credit_code", "legal_representative", "address", "mobile_phone", "email", "industry", "organization_type", "filing_region"]
    missing = [k for k in required if not raw.get(k)]
    if missing:
        label_map = {
            "name": "单位名称",
            "credit_code": "统一社会信用代码",
            "legal_representative": "单位负责人",
            "address": "单位地址",
            "mobile_phone": "移动电话",
            "email": "邮箱",
            "industry": "所属行业",
            "organization_type": "单位类型",
            "filing_region": "备案地区",
        }
        missing_labels = [label_map.get(k, k) for k in missing]
        raise HTTPException(status_code=400, detail=f"Word缺少必填字段: {', '.join(missing_labels)}")
    raw = {k: v for k, v in raw.items() if k in ORG_CREATE_FIELDS}
    normalize_org_payload(raw)
    validate_org_payload(raw)
    assert_credit_code_available(db, raw["credit_code"])
    org = Organization(**raw)
    db.add(org)
    db.flush()
    record_org_history(db, org.id, actor, "import_word", None, obj_to_dict(org, ORG_FIELDS))
    db.commit()
    db.refresh(org)
    return {"message": "Word导入成功", "data": obj_to_dict(org, ORG_FIELDS)}


@app.get("/api/organizations/{org_id}")
def get_organization(org_id: int, db: Session = Depends(get_db)) -> dict[str, Any]:
    org = get_org_or_404(db, org_id)
    return obj_to_dict(org, ORG_FIELDS)


@app.put("/api/organizations/{org_id}")
def update_organization(
    request: Request,
    org_id: int,
    payload: OrganizationUpdate,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor, role = require_roles(request, db, {"admin", "evaluator"}, legacy_admin=True)
    org = get_org_or_404(db, org_id)
    effective_is_admin = role == "admin"
    if org.archived and org.locked and not effective_is_admin:
        raise HTTPException(status_code=403, detail="已归档单位默认不可编辑，请管理员解锁。")
    data = payload.model_dump(exclude_unset=True)
    normalize_org_payload(data)
    validate_org_partial(data)
    if "credit_code" in data:
        assert_credit_code_available(db, data["credit_code"], current_org_id=org.id)
    before = obj_to_dict(org, ORG_FIELDS)
    for key, value in data.items():
        setattr(org, key, value)
    db.flush()
    db.refresh(org)
    record_org_history(db, org.id, actor, "update", before, obj_to_dict(org, ORG_FIELDS))
    db.commit()
    return {"message": "更新成功", "data": obj_to_dict(org, ORG_FIELDS)}


@app.get("/api/organizations/{org_id}/history")
def organization_history(org_id: int, db: Session = Depends(get_db)) -> dict[str, Any]:
    get_org_or_404(db, org_id)
    rows = (
        db.query(OrganizationHistory)
        .filter(OrganizationHistory.organization_id == org_id)
        .order_by(OrganizationHistory.changed_at.desc())
        .all()
    )
    return {
        "total": len(rows),
        "items": [
            {
                "changed_by": r.changed_by,
                "change_type": r.change_type,
                "before_data": r.before_data,
                "after_data": r.after_data,
                "changed_at": r.changed_at,
            }
            for r in rows
        ],
    }


@app.post("/api/organizations/{org_id}/archive")
def archive_organization(request: Request, org_id: int, db: Session = Depends(get_db)) -> dict[str, Any]:
    actor, _ = require_roles(request, db, {"admin"}, legacy_admin=True)
    org = get_org_or_404(db, org_id)
    before = obj_to_dict(org, ORG_FIELDS)
    org.archived = True
    org.locked = True
    db.flush()
    db.refresh(org)
    record_org_history(db, org.id, actor, "archive", before, obj_to_dict(org, ORG_FIELDS))
    db.commit()
    return {"message": "归档成功", "data": obj_to_dict(org, ORG_FIELDS)}


@app.post("/api/organizations/{org_id}/unlock")
def unlock_organization(
    request: Request,
    org_id: int,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor_name, _ = require_roles(request, db, {"admin"}, legacy_admin=True)
    org = get_org_or_404(db, org_id)
    org.locked = False
    db.flush()
    db.refresh(org)
    record_org_history(db, org.id, actor_name, "unlock", None, {"locked": False})
    db.commit()
    return {"message": "解锁成功", "data": obj_to_dict(org, ORG_FIELDS)}


@app.delete("/api/organizations/{org_id}")
def delete_organization(
    request: Request,
    org_id: int,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    org = get_org_or_404(db, org_id)
    actor_name, role = require_roles(request, db, {"admin", "evaluator"}, legacy_admin=True)
    if role != "admin":
        raise HTTPException(status_code=403, detail="请先提交删除申请，由管理员审核后删除。")
    assert_org_deletable(db, org, is_admin=True)
    before = obj_to_dict(org, ORG_FIELDS)
    org.deleted_at = datetime.now()
    org.deleted_by = actor_name
    finalize_pending_delete_requests(
        db,
        entity_type="organization",
        entity_id=org.id,
        reviewer=actor_name,
        reviewed_at=org.deleted_at,
    )
    db.flush()
    db.refresh(org)
    record_org_history(db, org.id, actor_name, "delete", before, obj_to_dict(org, ORG_FIELDS))
    db.commit()
    return {"message": "已移入回收站（保留30天）。"}


@app.post("/api/organizations/{org_id}/restore")
def restore_organization(request: Request, org_id: int, db: Session = Depends(get_db)) -> dict[str, Any]:
    actor, _ = require_roles(request, db, {"admin"}, legacy_admin=True)
    org = db.query(Organization).filter(Organization.id == org_id).first()
    if not org:
        raise HTTPException(status_code=404, detail="单位不存在。")
    if not org.deleted_at:
        raise HTTPException(status_code=400, detail="该单位未删除。")
    if datetime.now() - org.deleted_at > timedelta(days=30):
        raise HTTPException(status_code=400, detail="超出30天回收站恢复期。")
    org.deleted_at = None
    org.deleted_by = None
    db.flush()
    db.refresh(org)
    record_org_history(db, org.id, actor, "restore", None, obj_to_dict(org, ORG_FIELDS))
    db.commit()
    return {"message": "恢复成功", "data": obj_to_dict(org, ORG_FIELDS)}


@app.post("/api/organizations/{org_id}/purge")
def purge_organization(
    request: Request,
    org_id: int,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor_name, _ = require_roles(request, db, {"admin"}, legacy_admin=True)
    org = db.query(Organization).filter(Organization.id == org_id).first()
    if not org:
        raise HTTPException(status_code=404, detail="单位不存在。")
    if not org.deleted_at:
        raise HTTPException(status_code=400, detail="仅可永久删除回收站中的单位。")
    summary = purge_organization_cascade(db, org)
    db.commit()
    return {
        "message": "单位已永久删除，并已级联清理关联数据。",
        "actor": actor_name,
        "summary": summary,
    }


@app.post("/api/organizations/{org_id}/delete-request")
def create_org_delete_request(
    request: Request,
    org_id: int,
    reason: str = Query(""),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor, _ = require_roles(request, db, {"admin", "evaluator"}, legacy_admin=True)
    org = get_org_or_404(db, org_id)
    assert_org_deletable(db, org, is_admin=False)
    exists = (
        db.query(DeleteRequest)
        .filter(
            DeleteRequest.entity_type == "organization",
            DeleteRequest.entity_id == org_id,
            DeleteRequest.status == "pending",
        )
        .first()
    )
    if exists:
        raise HTTPException(status_code=409, detail="已存在待处理删除申请。")
    row = DeleteRequest(
        entity_type="organization",
        entity_id=org_id,
        reason=reason.strip() or None,
        status="pending",
        requested_by=actor,
    )
    db.add(row)
    db.commit()
    db.refresh(row)
    return {"message": "删除申请已提交", "data": {"request_id": row.id, "status": row.status}}


@app.get("/api/organizations/{org_id}/export/word")
def export_organization_word(request: Request, org_id: int, db: Session = Depends(get_db)) -> FileResponse:
    require_roles(request, db, {"admin", "evaluator"}, legacy_admin=True)
    org = get_org_or_404(db, org_id)
    path = EXPORT_DIR / f"organization_{org.id}_{datetime.now():%Y%m%d%H%M%S}.docx"
    doc = Document()
    doc.add_heading("单位基础信息", level=1)
    for key, value in [
        ("单位名称", org.name),
        ("统一社会信用代码", org.credit_code),
        ("单位负责人", org.legal_representative),
        ("单位地址", org.address),
        ("办公电话", org.office_phone or ""),
        ("移动电话", org.mobile_phone),
        ("邮箱", org.email),
        ("所属行业", org.industry),
        ("单位类型", org.organization_type),
        ("备案地区", org.filing_region),
    ]:
        doc.add_paragraph(f"{key}: {value}")
    doc.save(path)
    return FileResponse(path=str(path), filename=path.name, background=BackgroundTask(_cleanup_export_file, path))


@app.post("/api/systems")
def create_system(request: Request, payload: SystemCreate, db: Session = Depends(get_db)) -> dict[str, Any]:
    actor_name, _ = require_roles(request, db, {"admin", "evaluator"}, legacy_admin=True)
    org = get_org_or_404(db, payload.organization_id)
    if org.archived and org.locked:
        raise HTTPException(status_code=403, detail="单位已归档锁定，不可新增系统。")
    data = payload.model_dump()
    data["created_by"] = resolve_effective_actor_name(request, db, payload.created_by or actor_name)
    validate_system_payload(data, partial=False)
    data["system_code"] = generate_system_code(db)
    system = SystemInfo(**data)
    db.add(system)
    db.flush()
    record_system_history(db, system.id, data["created_by"], "create", None, obj_to_dict(system, SYSTEM_FIELDS))
    instance = WorkflowInstance(system_id=system.id, current_step_index=0, status="in_progress")
    db.add(instance)
    recalc_workflow_due_at(db, instance, get_workflow_config(db))
    db.commit()
    db.refresh(system)
    return {"message": "创建成功", "data": obj_to_dict(system, SYSTEM_FIELDS)}


@app.get("/api/systems")
def list_systems(
    system_name: str | None = None,
    system_code: str | None = None,
    organization_name: str | None = None,
    proposed_level: int | None = None,
    deployment_mode: str | None = None,
    created_by: str | None = None,
    include_deleted: bool = False,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    query = db.query(SystemInfo, Organization.name.label("organization_name")).join(
        Organization, Organization.id == SystemInfo.organization_id
    )
    if not include_deleted:
        query = query.filter(SystemInfo.deleted_at.is_(None))
    if system_name:
        query = query.filter(SystemInfo.system_name.like(f"%{escape_like(system_name)}%"))
    if system_code:
        query = query.filter(SystemInfo.system_code.like(f"%{escape_like(system_code)}%"))
    if organization_name:
        query = query.filter(Organization.name.like(f"%{escape_like(organization_name)}%"))
    if proposed_level:
        query = query.filter(SystemInfo.proposed_level == proposed_level)
    if deployment_mode:
        query = query.filter(SystemInfo.deployment_mode.like(f"%{escape_like(deployment_mode)}%"))
    if created_by:
        query = query.filter(SystemInfo.created_by.like(f"%{escape_like(created_by)}%"))
    rows = query.order_by(SystemInfo.created_at.desc()).all()
    items = []
    for system, org_name in rows:
        data = obj_to_dict(system, SYSTEM_FIELDS)
        data["organization_name"] = org_name
        items.append(data)
    return {"total": len(items), "items": items}


@app.get("/api/filing-workspace/overview")
def filing_workspace_overview(
    keyword: str | None = None,
    include_archived: bool = False,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    org_query = db.query(Organization).filter(Organization.deleted_at.is_(None))
    if keyword:
        token = f"%{escape_like(keyword.strip())}%"
        org_query = org_query.filter(
            or_(
                Organization.name.like(token),
                Organization.credit_code.like(token),
                Organization.industry.like(token),
            )
        )
    if not include_archived:
        org_query = org_query.filter(Organization.archived.is_(False))
    organizations = org_query.order_by(Organization.archived.asc(), Organization.updated_at.desc()).all()
    org_ids = [org.id for org in organizations]
    systems_by_org: dict[int, list[SystemInfo]] = {org_id: [] for org_id in org_ids}
    if org_ids:
        sys_query = db.query(SystemInfo).filter(SystemInfo.organization_id.in_(org_ids), SystemInfo.deleted_at.is_(None))
        if not include_archived:
            sys_query = sys_query.filter(SystemInfo.archived.is_(False))
        rows = sys_query.order_by(SystemInfo.archived.asc(), SystemInfo.updated_at.desc()).all()
        for row in rows:
            systems_by_org.setdefault(row.organization_id, []).append(row)
    items = [build_filing_overview_item(org, systems_by_org.get(org.id, [])) for org in organizations]
    return {"total": len(items), "items": items}


@app.get("/api/filing-workspace/systems/{system_id}")
def filing_workspace_detail(system_id: int, db: Session = Depends(get_db)) -> dict[str, Any]:
    system = get_system_or_404(db, system_id)
    org = get_org_or_404(db, system.organization_id)
    return build_workspace_detail_response(db, org, system)


@app.put("/api/filing-workspace/systems/{system_id}")
def save_filing_workspace(
    request: Request,
    system_id: int,
    payload: dict[str, Any],
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor, role = require_roles(request, db, {"admin", "evaluator"}, legacy_admin=True)
    system = get_system_or_404(db, system_id)
    org = get_org_or_404(db, system.organization_id)
    effective_is_admin = role == "admin"
    if (org.archived and org.locked) or (system.archived and system.locked and not effective_is_admin):
        raise HTTPException(status_code=403, detail="当前对象已归档锁定，不可编辑。")
    if not isinstance(payload, dict):
        raise HTTPException(status_code=400, detail="payload 必须为对象。")
    org_payload = payload.get("organization") if isinstance(payload.get("organization"), dict) else {}
    system_payload = payload.get("system") if isinstance(payload.get("system"), dict) else {}
    assert_safe_payload(org_payload, "organization")
    assert_safe_payload(system_payload, "system")
    before_org, _ = sync_org_from_workspace(db, org, org_payload)
    before_system, _ = sync_system_from_workspace(system, system_payload)
    db.flush()
    db.refresh(org)
    db.refresh(system)
    record_org_history(db, org.id, actor, "workspace_update", before_org, obj_to_dict(org, ORG_FIELDS))
    record_system_history(db, system.id, actor, "workspace_update", before_system, obj_to_dict(system, SYSTEM_FIELDS))
    db.commit()
    db.refresh(org)
    db.refresh(system)
    return {"message": "备案信息已保存", "data": build_workspace_detail_response(db, org, system)}


@app.post("/api/filing-workspace/systems/{system_id}/attachments/{slot_key}")
async def upload_filing_workspace_attachment(
    request: Request,
    system_id: int,
    slot_key: str,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor, _ = require_roles(request, db, {"admin", "evaluator"}, legacy_admin=True)
    system = get_system_or_404(db, system_id)
    content = await file.read()
    if len(content) > MAX_SYS_ATTACHMENT:
        raise HTTPException(status_code=400, detail=f"文件大小超过限制({MAX_SYS_ATTACHMENT // 1024 // 1024}MB)。")
    row = save_attachment_row(
        db,
        entity_type="system",
        entity_id=system_id,
        actor=actor,
        file_name=file.filename,
        content=content,
    )
    detail = merged_system_filing_detail(system)
    slot_key = slot_key.strip()
    slot_map = detail.setdefault("_attachment_slots", {})
    if not isinstance(slot_map, dict):
        slot_map = {}
        detail["_attachment_slots"] = slot_map
    slot_item = slot_map.setdefault(slot_key, {"file_name": "", "attachment_ids": []})
    if not isinstance(slot_item, dict):
        slot_item = {"file_name": "", "attachment_ids": []}
        slot_map[slot_key] = slot_item
    slot_item["file_name"] = Path(file.filename).name
    slot_item["attachment_ids"] = normalize_attachment_refs(slot_item.get("attachment_ids")) + [row.id]
    if slot_key.startswith("table3."):
        name = slot_key.split(".", 1)[1]
        target = detail.setdefault("table3", {}).setdefault(name, {})
        if isinstance(target, dict):
            target["file_name"] = slot_item["file_name"]
            target["attachment_ids"] = slot_item["attachment_ids"]
    elif slot_key.startswith("table4.cloud"):
        target = detail.setdefault("table4", {}).setdefault("cloud", {})
        if isinstance(target, dict):
            target["record_file_name"] = slot_item["file_name"]
            target["attachment_ids"] = slot_item["attachment_ids"]
    elif slot_key.startswith("table4.big_data"):
        target = detail.setdefault("table4", {}).setdefault("big_data", {})
        if isinstance(target, dict):
            target["record_file_name"] = slot_item["file_name"]
            target["attachment_ids"] = slot_item["attachment_ids"]
    elif slot_key.startswith("table5."):
        name = slot_key.split(".", 1)[1]
        target = detail.setdefault("table5", {}).setdefault(name, {})
        if isinstance(target, dict):
            target["file_name"] = slot_item["file_name"]
            target["attachment_ids"] = slot_item["attachment_ids"]
    system.filing_detail = detail
    db.commit()
    db.refresh(system)
    return {"message": "附件上传成功", "data": build_workspace_detail_response(db, system.organization, system)}


@app.get("/api/systems/recycle-bin/list")
def list_system_recycle_bin(
    include_expired: bool = Query(False),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    rows = (
        db.query(SystemInfo)
        .filter(SystemInfo.deleted_at.is_not(None))
        .order_by(SystemInfo.deleted_at.desc())
        .all()
    )
    items: list[dict[str, Any]] = []
    for row in rows:
        meta = recycle_meta(row.deleted_at)
        if meta["expired"] and not include_expired:
            continue
        data = obj_to_dict(row, SYSTEM_FIELDS)
        data["organization_name"] = row.organization.name if row.organization else ""
        data.update(meta)
        items.append(data)
    return {"total": len(items), "items": items}


@app.post("/api/systems/recycle-bin/cleanup")
def cleanup_system_recycle_bin(
    request: Request,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor_name, _ = require_roles(request, db, {"admin"}, legacy_admin=True)
    threshold = datetime.now() - timedelta(days=30)
    rows = db.query(SystemInfo).filter(SystemInfo.deleted_at.is_not(None), SystemInfo.deleted_at <= threshold).all()
    purged = 0
    skipped: list[str] = []
    for row in rows:
        report_count = db.query(Report).filter(Report.system_id == row.id).count()
        if report_count > 0:
            skipped.append(f"{row.id}:{row.system_name} 仍有关联报告({report_count})")
            continue
        db.delete(row)
        purged += 1
    db.commit()
    return {"message": "系统回收站清理完成", "purged": purged, "skipped": skipped, "actor": actor_name}


@app.get("/api/systems/export/excel")
def export_systems_excel(request: Request, db: Session = Depends(get_db)) -> FileResponse:
    require_roles(request, db, {"admin", "evaluator"}, legacy_admin=True)
    rows = db.query(SystemInfo).filter(SystemInfo.deleted_at.is_(None)).all()
    wb = Workbook()
    ws = wb.active
    ws.title = "系统信息"
    headers = ["系统名称", "系统编号", "单位ID", "拟定等级", "部署方式", "系统类型", "上线时间", "录入人"]
    ws.append(headers)
    for i in rows:
        ws.append(
            [
                i.system_name,
                i.system_code,
                i.organization_id,
                i.proposed_level,
                i.deployment_mode,
                i.system_type,
                i.go_live_date.isoformat() if i.go_live_date else "",
                i.created_by,
            ]
        )
    path = EXPORT_DIR / f"systems_{datetime.now():%Y%m%d%H%M%S}.xlsx"
    wb.save(path)
    return FileResponse(path=str(path), filename=path.name, background=BackgroundTask(_cleanup_export_file, path))


@app.post("/api/systems/import/excel")
async def import_systems_excel(
    request: Request, file: UploadFile = File(...), db: Session = Depends(get_db)
) -> dict[str, Any]:
    actor, _ = require_roles(request, db, {"admin"}, legacy_admin=True)
    if not file.filename or not file.filename.lower().endswith(".xlsx"):
        raise HTTPException(status_code=400, detail="仅支持xlsx格式。")
    content = await file.read()
    # 限制导入文件大小，防止恶意上传超大文件
    if len(content) > 50 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="导入文件过大，请限制在50MB以内。")
    wb = load_workbook(io.BytesIO(content))
    ws = wb.active
    imported = 0
    skipped: list[str] = []
    for idx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
        if not row or not row[0]:
            continue
        if len(row) < 7:
            skipped.append(f"第{idx}行：列数不足（需要至少7列）")
            continue
        try:
            with db.begin_nested():
                org_id = int(row[2])
                org = get_org_or_404(db, org_id)
                if org.archived and org.locked:
                    skipped.append(f"第{idx}行：单位已归档锁定")
                    continue
                sys = SystemInfo(
                    organization_id=org_id,
                    system_name=str(row[0] or "").strip(),
                    system_code=generate_system_code(db),
                    proposed_level=parse_proposed_level(row[3]),
                    deployment_mode=str(row[4]).strip() if row[4] else None,
                    system_type=str(row[5]).strip() if row[5] else None,
                    go_live_date=parse_optional_go_live_date(row[6]),
                    created_by=actor,
                )
                db.add(sys)
                db.flush()
                instance = WorkflowInstance(system_id=sys.id, current_step_index=0, status="in_progress")
                db.add(instance)
                recalc_workflow_due_at(db, instance, get_workflow_config(db))
                record_system_history(db, sys.id, actor, "import", None, obj_to_dict(sys, SYSTEM_FIELDS))
            imported += 1
        except Exception as exc:
            skipped.append(f"第{idx}行：{exc}")
    db.commit()
    return {"message": "导入完成", "imported": imported, "skipped": skipped}


@app.post("/api/systems/import/word")
async def import_system_word(
    request: Request,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor, _ = require_roles(request, db, {"admin"}, legacy_admin=True)
    if not file.filename or not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="仅支持docx格式。")
    content = await file.read()
    kv = parse_docx_key_values(content)
    raw: dict[str, Any] = {}
    for cn_key, model_key in SYSTEM_WORD_MAP.items():
        if cn_key in kv:
            raw[model_key] = kv[cn_key]
    if "organization_id" not in raw:
        raise HTTPException(status_code=400, detail="Word缺少单位ID。")
    try:
        org_id = int(str(raw["organization_id"]).strip())
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"单位ID格式错误: {exc}") from exc
    org = get_org_or_404(db, org_id)
    if org.archived and org.locked:
        raise HTTPException(status_code=403, detail="单位已归档锁定，不可新增系统。")
    level_text = str(raw.get("proposed_level", "")).strip()
    if not level_text:
        raise HTTPException(status_code=400, detail="Word缺少拟定等级。")
    try:
        proposed_level = parse_proposed_level(level_text)
    except Exception as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    go_live_date = None
    if raw.get("go_live_date"):
        try:
            go_live_date = parse_optional_go_live_date(raw["go_live_date"])
        except Exception as exc:
            raise HTTPException(status_code=400, detail=str(exc)) from exc
    name = str(raw.get("system_name", "")).strip()
    if not name:
        raise HTTPException(status_code=400, detail="Word缺少系统名称。")
    try:
        system = SystemInfo(
            organization_id=org_id,
            system_name=name,
            system_code=generate_system_code(db),
            proposed_level=proposed_level,
            business_description=str(raw.get("business_description") or "").strip() or None,
            deployment_mode=str(raw.get("deployment_mode") or "").strip() or None,
            system_type=str(raw.get("system_type") or "").strip() or None,
            go_live_date=go_live_date,
            level_basis=str(raw.get("level_basis") or "").strip() or None,
            level_reason=str(raw.get("level_reason") or "").strip() or None,
            boundary=str(raw.get("boundary") or "").strip() or None,
            subsystems=str(raw.get("subsystems") or "").strip() or None,
            service_object=str(raw.get("service_object") or "").strip() or None,
            service_scope=str(raw.get("service_scope") or "").strip() or None,
            created_by=actor,
        )
        db.add(system)
        db.flush()
        instance = WorkflowInstance(system_id=system.id, current_step_index=0, status="in_progress")
        db.add(instance)
        recalc_workflow_due_at(db, instance, get_workflow_config(db))
        record_system_history(db, system.id, actor, "import_word", None, obj_to_dict(system, SYSTEM_FIELDS))
        db.commit()
        db.refresh(system)
    except HTTPException:
        db.rollback()
        raise
    except Exception as exc:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"系统Word导入失败: {exc}") from exc
    return {"message": "Word导入成功", "data": obj_to_dict(system, SYSTEM_FIELDS)}


@app.get("/api/systems/{system_id}")
def get_system(system_id: int, db: Session = Depends(get_db)) -> dict[str, Any]:
    system = get_system_or_404(db, system_id)
    data = obj_to_dict(system, SYSTEM_FIELDS)
    data["organization_name"] = system.organization.name if system.organization else ""
    return data


@app.put("/api/systems/{system_id}")
def update_system(
    request: Request,
    system_id: int,
    payload: SystemUpdate,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor, role = require_roles(request, db, {"admin", "evaluator"}, legacy_admin=True)
    system = get_system_or_404(db, system_id)
    effective_is_admin = role == "admin"
    if system.archived and system.locked and not effective_is_admin:
        raise HTTPException(status_code=403, detail="已归档系统默认不可编辑。")
    data = payload.model_dump(exclude_unset=True)
    validate_system_payload(data, partial=True)
    before = obj_to_dict(system, SYSTEM_FIELDS)
    for key, value in data.items():
        setattr(system, key, value)
    db.flush()
    db.refresh(system)
    record_system_history(db, system.id, actor, "update", before, obj_to_dict(system, SYSTEM_FIELDS))
    db.commit()
    return {"message": "更新成功", "data": obj_to_dict(system, SYSTEM_FIELDS)}


@app.get("/api/systems/{system_id}/history")
def system_history(system_id: int, db: Session = Depends(get_db)) -> dict[str, Any]:
    get_system_or_404(db, system_id)
    rows = db.query(SystemHistory).filter(SystemHistory.system_id == system_id).order_by(SystemHistory.changed_at.desc()).all()
    return {
        "total": len(rows),
        "items": [
            {
                "changed_by": r.changed_by,
                "change_type": r.change_type,
                "before_data": r.before_data,
                "after_data": r.after_data,
                "changed_at": r.changed_at,
            }
            for r in rows
        ],
    }


@app.post("/api/systems/{system_id}/copy")
def copy_system(request: Request, system_id: int, db: Session = Depends(get_db)) -> dict[str, Any]:
    actor, _ = require_roles(request, db, {"admin", "evaluator"}, legacy_admin=True)
    source = get_system_or_404(db, system_id)
    data = obj_to_dict(source, SYSTEM_FIELDS)
    for key in ["id", "system_code", "created_at", "updated_at", "deleted_at"]:
        data.pop(key, None)
    data["system_name"] = f"{source.system_name}-副本"
    data["system_code"] = generate_system_code(db)
    data["archived"] = False
    data["locked"] = False
    data["deleted_by"] = None
    data["created_by"] = actor
    copied = SystemInfo(**data)
    db.add(copied)
    db.flush()
    record_system_history(db, copied.id, actor, "copy", None, obj_to_dict(copied, SYSTEM_FIELDS))
    instance = WorkflowInstance(system_id=copied.id, current_step_index=0, status="in_progress")
    db.add(instance)
    recalc_workflow_due_at(db, instance, get_workflow_config(db))
    db.commit()
    db.refresh(copied)
    return {"message": "复制成功", "data": obj_to_dict(copied, SYSTEM_FIELDS)}


@app.post("/api/systems/{system_id}/archive")
def archive_system(request: Request, system_id: int, db: Session = Depends(get_db)) -> dict[str, Any]:
    actor, _ = require_roles(request, db, {"admin"}, legacy_admin=True)
    system = get_system_or_404(db, system_id)
    before = obj_to_dict(system, SYSTEM_FIELDS)
    system.archived = True
    system.locked = True
    db.flush()
    db.refresh(system)
    record_system_history(db, system.id, actor, "archive", before, obj_to_dict(system, SYSTEM_FIELDS))
    db.commit()
    return {"message": "归档成功", "data": obj_to_dict(system, SYSTEM_FIELDS)}


@app.post("/api/systems/{system_id}/unlock")
def unlock_system(
    request: Request,
    system_id: int,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor_name, _ = require_roles(request, db, {"admin"}, legacy_admin=True)
    system = get_system_or_404(db, system_id)
    system.locked = False
    db.flush()
    db.refresh(system)
    record_system_history(db, system.id, actor_name, "unlock", None, {"locked": False})
    db.commit()
    return {"message": "解锁成功", "data": obj_to_dict(system, SYSTEM_FIELDS)}


@app.delete("/api/systems/{system_id}")
def delete_system(
    request: Request,
    system_id: int,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    system = get_system_or_404(db, system_id)
    actor_name, role = require_roles(request, db, {"admin", "evaluator"}, legacy_admin=True)
    if role != "admin":
        raise HTTPException(status_code=403, detail="请先提交删除申请，由管理员审核后删除。")
    assert_system_deletable(db, system, is_admin=True)
    before = obj_to_dict(system, SYSTEM_FIELDS)
    system.deleted_at = datetime.now()
    system.deleted_by = actor_name
    finalize_pending_delete_requests(
        db,
        entity_type="system",
        entity_id=system.id,
        reviewer=actor_name,
        reviewed_at=system.deleted_at,
    )
    db.flush()
    db.refresh(system)
    record_system_history(db, system.id, actor_name, "delete", before, obj_to_dict(system, SYSTEM_FIELDS))
    db.commit()
    return {"message": "已移入回收站（保留30天）。"}


@app.post("/api/systems/{system_id}/restore")
def restore_system(request: Request, system_id: int, db: Session = Depends(get_db)) -> dict[str, Any]:
    actor, _ = require_roles(request, db, {"admin"}, legacy_admin=True)
    system = db.query(SystemInfo).filter(SystemInfo.id == system_id).first()
    if not system:
        raise HTTPException(status_code=404, detail="系统不存在。")
    if not system.deleted_at:
        raise HTTPException(status_code=400, detail="该系统未删除。")
    if datetime.now() - system.deleted_at > timedelta(days=30):
        raise HTTPException(status_code=400, detail="超出30天回收站恢复期。")
    org = db.query(Organization).filter(Organization.id == system.organization_id).first()
    if not org:
        raise HTTPException(status_code=400, detail="所属单位不存在，无法恢复系统。")
    if org.deleted_at:
        raise HTTPException(status_code=400, detail="所属单位已删除，请先恢复单位后再恢复系统。")
    system.deleted_at = None
    system.deleted_by = None
    db.flush()
    db.refresh(system)
    record_system_history(db, system.id, actor, "restore", None, obj_to_dict(system, SYSTEM_FIELDS))
    db.commit()
    return {"message": "恢复成功", "data": obj_to_dict(system, SYSTEM_FIELDS)}


@app.post("/api/systems/{system_id}/purge")
def purge_system(
    request: Request,
    system_id: int,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor_name, _ = require_roles(request, db, {"admin"}, legacy_admin=True)
    system = db.query(SystemInfo).filter(SystemInfo.id == system_id).first()
    if not system:
        raise HTTPException(status_code=404, detail="系统不存在。")
    if not system.deleted_at:
        raise HTTPException(status_code=400, detail="仅可永久删除回收站中的系统。")
    summary = purge_system_cascade(db, system)
    db.commit()
    return {
        "message": "系统已永久删除，并已级联清理关联数据。",
        "actor": actor_name,
        "summary": summary,
    }


@app.post("/api/systems/{system_id}/delete-request")
def create_system_delete_request(
    request: Request,
    system_id: int,
    reason: str = Query(""),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor, _ = require_roles(request, db, {"admin", "evaluator"}, legacy_admin=True)
    system = get_system_or_404(db, system_id)
    assert_system_deletable(db, system, is_admin=False)
    exists = (
        db.query(DeleteRequest)
        .filter(
            DeleteRequest.entity_type == "system",
            DeleteRequest.entity_id == system_id,
            DeleteRequest.status == "pending",
        )
        .first()
    )
    if exists:
        raise HTTPException(status_code=409, detail="已存在待处理删除申请。")
    row = DeleteRequest(
        entity_type="system",
        entity_id=system_id,
        reason=reason.strip() or None,
        status="pending",
        requested_by=actor,
    )
    db.add(row)
    db.commit()
    db.refresh(row)
    return {"message": "删除申请已提交", "data": {"request_id": row.id, "status": row.status}}


@app.get("/api/delete-requests")
def list_delete_requests(
    request: Request,
    entity_type: str | None = None,
    status: str | None = None,
    requested_by: str | None = None,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    reconcile_pending_delete_requests(db)
    backfill_delete_request_review_comments(db)
    current_user = get_current_user_optional(request, db)
    q = db.query(DeleteRequest)
    if entity_type:
        q = q.filter(DeleteRequest.entity_type == entity_type.strip().lower())
    if status:
        q = q.filter(DeleteRequest.status == status.strip().lower())
    if requested_by:
        q = q.filter(DeleteRequest.requested_by.like(f"%{escape_like(requested_by)}%"))
    if current_user and current_user.role != "admin":
        q = q.filter(DeleteRequest.requested_by == current_user.username)
    rows = q.order_by(DeleteRequest.requested_at.desc()).all()
    return {
        "total": len(rows),
        "items": [
            {
                "id": r.id,
                "entity_type": r.entity_type,
                "entity_id": r.entity_id,
                "reason": r.reason,
                "status": r.status,
                "requested_by": r.requested_by,
                "requested_at": r.requested_at,
                "reviewed_by": r.reviewed_by,
                "reviewed_at": r.reviewed_at,
                "review_comment": r.review_comment,
            }
            for r in rows
        ],
    }


@app.post("/api/delete-requests/{request_id}/review")
def review_delete_request(
    request: Request,
    request_id: int,
    action: str = Query(...),
    comment: str = Query(""),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor_name, _ = require_roles(request, db, {"admin"}, legacy_admin=True)
    row = db.query(DeleteRequest).filter(DeleteRequest.id == request_id).first()
    if not row:
        raise HTTPException(status_code=404, detail="删除申请不存在。")
    if row.status != "pending":
        raise HTTPException(status_code=400, detail="仅待处理申请可审核。")
    action = action.strip().lower()
    if action not in {"approve", "reject"}:
        raise HTTPException(status_code=400, detail="action 仅支持 approve/reject。")

    if action == "approve":
        if row.entity_type == "organization":
            org = get_org_or_404(db, row.entity_id)
            assert_org_deletable(db, org, is_admin=True)
            before = obj_to_dict(org, ORG_FIELDS)
            org.deleted_at = datetime.now()
            org.deleted_by = actor_name
            record_org_history(db, org.id, actor_name, "delete_by_request", before, obj_to_dict(org, ORG_FIELDS))
        elif row.entity_type == "system":
            system = get_system_or_404(db, row.entity_id)
            assert_system_deletable(db, system, is_admin=True)
            before = obj_to_dict(system, SYSTEM_FIELDS)
            system.deleted_at = datetime.now()
            system.deleted_by = actor_name
            record_system_history(
                db,
                system.id,
                actor_name,
                "delete_by_request",
                before,
                obj_to_dict(system, SYSTEM_FIELDS),
            )
        else:
            raise HTTPException(status_code=400, detail=f"不支持的实体类型: {row.entity_type}")

    row.status = "approved" if action == "approve" else "rejected"
    row.reviewed_by = actor_name
    row.reviewed_at = datetime.now()
    row.review_comment = comment.strip() or None
    db.commit()
    return {"message": "删除申请处理完成", "status": row.status}


def replace_cell_text(cell: Any, text: str, *, bold: bool | None = None) -> None:
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    if paragraph.runs:
        paragraph.runs[0].text = text
        if bold is not None:
            paragraph.runs[0].bold = bold
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        run = paragraph.add_run(text)
        if bold is not None:
            run.bold = bold
    for extra in cell.paragraphs[1:]:
        for run in extra.runs:
            run.text = ""


def render_marked_choices(options: list[str], selected: list[str] | str | None, *, other_value: str = "") -> str:
    picked = {str(item).strip() for item in (selected if isinstance(selected, list) else [selected] if selected else []) if str(item).strip()}
    parts: list[str] = []
    for option in options:
        label = str(option).strip()
        mark = "√" if label in picked else "□"
        parts.append(f"{mark}{label}")
    if other_value:
        mark = "√" if ("其他" in picked or other_value) else "□"
        parts.append(f"{mark}其他 {other_value}".strip())
    return "  ".join(parts)


def joined_text(values: Any, *, empty_text: str = "无", sep: str = " / ", other_value: str = "") -> str:
    if isinstance(values, list):
        items = [str(item).strip() for item in values if str(item).strip()]
        if other_value:
            items.append(f"其他 {other_value}".strip())
        return sep.join(items) if items else empty_text
    text = str(values or "").strip()
    if other_value and text:
        return f"{text}{sep}其他 {other_value}".strip()
    if other_value and not text:
        return f"其他 {other_value}".strip()
    return text or empty_text


def extract_workspace_export_context(org: Organization, system: SystemInfo) -> dict[str, Any]:
    org_detail = merged_org_filing_detail(org)
    sys_detail = merged_system_filing_detail(system)
    table1 = org_detail.get("table1", {})
    table2 = sys_detail.get("table2", {})
    table3 = sys_detail.get("table3", {})
    table4 = sys_detail.get("table4", {})
    table5 = sys_detail.get("table5", {})
    table6 = sys_detail.get("table6", {})
    return {
        "table1": table1 if isinstance(table1, dict) else {},
        "table2": table2 if isinstance(table2, dict) else {},
        "table3": table3 if isinstance(table3, dict) else {},
        "table4": table4 if isinstance(table4, dict) else {},
        "table5": table5 if isinstance(table5, dict) else {},
        "table6": table6 if isinstance(table6, dict) else {"items": []},
    }


def fill_filing_template_document(doc: Document, org: Organization, system: SystemInfo) -> None:
    ctx = extract_workspace_export_context(org, system)
    table1 = ctx["table1"]
    table2 = ctx["table2"]
    table3 = ctx["table3"]
    table4 = ctx["table4"]
    table5 = ctx["table5"]
    table6 = ctx["table6"]
    address_parts = table1.get("address_parts") if isinstance(table1.get("address_parts"), dict) else {}
    responsible_person = table1.get("responsible_person") if isinstance(table1.get("responsible_person"), dict) else {}
    affiliation = table1.get("affiliation") if isinstance(table1.get("affiliation"), dict) else {}
    org_type = table1.get("organization_type_detail") if isinstance(table1.get("organization_type_detail"), dict) else {}
    industry = table1.get("industry_category") if isinstance(table1.get("industry_category"), dict) else {}
    current_counts = table1.get("current_filing_counts") if isinstance(table1.get("current_filing_counts"), dict) else {}
    total_counts = table1.get("total_filing_counts") if isinstance(table1.get("total_filing_counts"), dict) else {}
    date_text = datetime.now().strftime("%Y-%m-%d")

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("备 案 单 位："):
            paragraph.text = f"备 案 单 位：    {org.name}"
        elif text.startswith("备 案 日 期："):
            paragraph.text = f"备 案 日 期：    {date_text}"

    if len(doc.tables) < 7:
        raise HTTPException(status_code=500, detail="备案模板结构不完整。")

    t1 = doc.tables[1]
    replace_cell_text(t1.rows[0].cells[1], org.name)
    replace_cell_text(t1.rows[1].cells[1], org.credit_code)
    address_text = (
        f"{str(address_parts.get('province') or '').strip()} "
        f"省(自治区、直辖市)    {str(address_parts.get('city') or '').strip()} 地(区、市、州、盟)\n"
        f"{str(address_parts.get('district') or '').strip()} 县(区、市、旗)  详细地址 {str(address_parts.get('detail') or org.address).strip()}"
    ).strip()
    replace_cell_text(t1.rows[2].cells[1], address_text)
    replace_cell_text(t1.rows[3].cells[1], str(table1.get("postal_code") or ""))
    replace_cell_text(t1.rows[3].cells[6], str(table1.get("district_code") or ""))
    replace_cell_text(t1.rows[4].cells[2], str(responsible_person.get("name") or org.legal_representative or ""))
    replace_cell_text(t1.rows[4].cells[6], str(responsible_person.get("title") or ""))
    replace_cell_text(t1.rows[5].cells[2], str(responsible_person.get("office_phone") or org.office_phone or "/"))
    replace_cell_text(t1.rows[5].cells[6], str(responsible_person.get("email") or org.email or "/"))
    replace_cell_text(t1.rows[6].cells[1], str(org.cybersecurity_dept or ""))
    replace_cell_text(t1.rows[7].cells[2], str(org.cybersecurity_owner_name or ""))
    replace_cell_text(t1.rows[7].cells[6], str(org.cybersecurity_owner_title or ""))
    replace_cell_text(t1.rows[8].cells[2], str(org.cybersecurity_owner_phone or "/"))
    replace_cell_text(t1.rows[8].cells[6], str(org.cybersecurity_owner_email or "/"))
    replace_cell_text(t1.rows[9].cells[2], str(org.mobile_phone or "/"))
    replace_cell_text(t1.rows[9].cells[6], str(org.cybersecurity_owner_email or org.email or "/"))
    replace_cell_text(t1.rows[10].cells[1], str(org.data_security_dept or ""))
    replace_cell_text(t1.rows[11].cells[2], str(org.data_security_owner_name or ""))
    replace_cell_text(t1.rows[11].cells[6], str(org.data_security_owner_title or ""))
    replace_cell_text(t1.rows[12].cells[2], str(org.data_security_owner_phone or "/"))
    replace_cell_text(t1.rows[12].cells[6], str(org.data_security_owner_email or "/"))
    replace_cell_text(t1.rows[13].cells[2], str(org.mobile_phone or "/"))
    replace_cell_text(t1.rows[13].cells[6], str(org.data_security_owner_email or org.email or "/"))
    replace_cell_text(
        t1.rows[14].cells[1],
        render_marked_choices(
            ["中央", "省(自治区、直辖市)", "地(区、市、州、盟)", "县（区、市、旗）"],
            [str(affiliation.get("label") or "").strip()],
            other_value=str(affiliation.get("other") or "").strip(),
        ),
    )
    replace_cell_text(
        t1.rows[15].cells[1],
        render_marked_choices(
            ["党委机关", "政府机关", "事业单位", "企业"],
            [str(org_type.get("label") or org.organization_type or "").strip()],
            other_value=str(org_type.get("other") or "").strip(),
        ),
    )
    replace_cell_text(
        t1.rows[16].cells[1],
        str(industry.get("label") or org.industry or "").strip(),
    )
    replace_cell_text(t1.rows[17].cells[1], f"{to_int_or_zero(sum(current_counts.values()))} 个")
    replace_cell_text(t1.rows[17].cells[3], str(to_int_or_zero(current_counts.get("2"))))
    replace_cell_text(t1.rows[17].cells[7], str(to_int_or_zero(current_counts.get("3"))))
    replace_cell_text(t1.rows[18].cells[3], str(to_int_or_zero(current_counts.get("4"))))
    replace_cell_text(t1.rows[18].cells[7], str(to_int_or_zero(current_counts.get("5"))))
    replace_cell_text(t1.rows[19].cells[1], f"{to_int_or_zero(sum(total_counts.values()))} 个")
    replace_cell_text(t1.rows[19].cells[3], str(to_int_or_zero(total_counts.get("1"))))
    replace_cell_text(t1.rows[19].cells[7], str(to_int_or_zero(total_counts.get("2"))))
    replace_cell_text(t1.rows[20].cells[3], str(to_int_or_zero(total_counts.get("3"))))
    replace_cell_text(t1.rows[20].cells[7], str(to_int_or_zero(total_counts.get("4"))))
    replace_cell_text(t1.rows[21].cells[3], str(to_int_or_zero(total_counts.get("5"))))

    t2 = doc.tables[2]
    replace_cell_text(t2.rows[0].cells[2], system.system_name)
    replace_cell_text(t2.rows[1].cells[2], joined_text(table2.get("object_types"), empty_text="信息系统"))
    replace_cell_text(t2.rows[2].cells[2], joined_text(table2.get("business_types")))
    replace_cell_text(t2.rows[3].cells[2], str(table2.get("business_description") or system.business_description or ""))
    network_service = table2.get("network_service") if isinstance(table2.get("network_service"), dict) else {}
    network_platform = table2.get("network_platform") if isinstance(table2.get("network_platform"), dict) else {}
    service_scope_text = str(network_service.get("scope_code") or "").strip()
    if network_service.get("other"):
        service_scope_text = f"{service_scope_text} / 其他 {network_service.get('other')}".strip()
    replace_cell_text(t2.rows[4].cells[2], service_scope_text)
    replace_cell_text(t2.rows[5].cells[2], joined_text(network_service.get("service_targets")))
    coverage = network_platform.get("coverage") if isinstance(network_platform.get("coverage"), dict) else {}
    replace_cell_text(t2.rows[6].cells[2], str(coverage.get("label") or coverage.get("code") or ""))
    network_nature = network_platform.get("network_nature") if isinstance(network_platform.get("network_nature"), dict) else {}
    network_nature_text = str(network_nature.get("label") or network_nature.get("code") or "").strip()
    if str(network_nature.get("label") or network_nature.get("code") or "").strip() == "互联网":
        network_nature_text = (
            f"{network_nature_text}  源站IP地址范围 {str(network_nature.get('source_ip_range') or '')}  "
            f"域名 {str(network_nature.get('domain') or '')}  主要协议/端口 {str(network_nature.get('protocol_ports') or '')}"
        ).strip()
    elif network_nature.get("other"):
        network_nature_text = f"{network_nature_text} / 其他 {network_nature.get('other')}".strip()
    replace_cell_text(t2.rows[7].cells[2], network_nature_text)
    interconnection = network_platform.get("interconnection") if isinstance(network_platform.get("interconnection"), dict) else {}
    replace_cell_text(t2.rows[8].cells[2], joined_text(interconnection.get("items"), other_value=str(interconnection.get("other") or "")))
    replace_cell_text(t2.rows[9].cells[2], format_workspace_date(table2.get("go_live_date") or system.go_live_date))
    replace_cell_text(t2.rows[10].cells[2], "是" if coerce_bool(table2.get("is_sub_system")) else "否")
    replace_cell_text(t2.rows[11].cells[2], str(table2.get("parent_system_name") or "无"))
    replace_cell_text(t2.rows[12].cells[2], str(table2.get("parent_organization_name") or ""))

    t3 = doc.tables[3]
    business_items = {str(item).strip() for item in table3.get("business_security_damage_items", []) if str(item).strip()}
    service_items = {str(item).strip() for item in table3.get("service_security_damage_items", []) if str(item).strip()}
    for idx in range(1, 6):
        label = t3.rows[idx].cells[1].text.replace("\n", " / ").strip()
        replace_cell_text(t3.rows[idx].cells[2], f"{'√' if label in business_items else '□'} {label}")
    for idx in range(6, 11):
        label = t3.rows[idx].cells[1].text.replace("\n", " / ").strip()
        replace_cell_text(t3.rows[idx].cells[2], f"{'√' if label in service_items else '□'} {label}")
    final_level = to_int_or_zero(table3.get("final_level")) or system.proposed_level
    replace_cell_text(t3.rows[11].cells[2], render_marked_choices(["第二级", "第三级", "第四级", "第五级"], [f"第{final_level}级"]))
    replace_cell_text(t3.rows[12].cells[2], format_workspace_date(table3.get("grading_date")))
    grading_report = table3.get("grading_report") if isinstance(table3.get("grading_report"), dict) else {}
    replace_cell_text(t3.rows[13].cells[2], f"{'有' if coerce_bool(grading_report.get('has_file')) else '无'}    附件名称 {grading_report.get('file_name') or ''}", bold=True)
    expert_review = table3.get("expert_review") if isinstance(table3.get("expert_review"), dict) else {}
    replace_cell_text(t3.rows[14].cells[2], f"{'已评审' if expert_review.get('status') == 'reviewed' else '未评审'}    附件名称 {expert_review.get('file_name') or ''}", bold=True)
    has_supervisor = coerce_bool(table3.get("has_supervisor"))
    replace_cell_text(t3.rows[15].cells[2], "有" if has_supervisor else "无")
    replace_cell_text(t3.rows[16].cells[2], str(table3.get("supervisor_name") or "无"))
    supervisor_review = table3.get("supervisor_review") if isinstance(table3.get("supervisor_review"), dict) else {}
    replace_cell_text(t3.rows[17].cells[2], f"{'已审核' if supervisor_review.get('status') == 'reviewed' else '未审核'}    附件名称 {supervisor_review.get('file_name') or ''}", bold=True)
    replace_cell_text(t3.rows[18].cells[0], f"填表人： {str(table3.get('filler_name') or '').strip()}")
    replace_cell_text(t3.rows[18].cells[3], f"填表日期：{format_workspace_date(table3.get('filled_date'))}")

    t4 = doc.tables[4]
    cloud = table4.get("cloud") if isinstance(table4.get("cloud"), dict) else {}
    replace_cell_text(t4.rows[0].cells[2], "是" if coerce_bool(cloud.get("enabled")) else "否")
    replace_cell_text(t4.rows[1].cells[2], joined_text(cloud.get("responsibility_types")))
    replace_cell_text(t4.rows[2].cells[2], joined_text(cloud.get("service_modes"), other_value=str(cloud.get("service_mode_other") or "")))
    replace_cell_text(t4.rows[3].cells[2], joined_text(cloud.get("deployment_modes"), other_value=str(cloud.get("deployment_mode_other") or "")))
    replace_cell_text(t4.rows[5].cells[2], str(cloud.get("customer_count") or ""))
    replace_cell_text(t4.rows[6].cells[2], str(cloud.get("infra_location") or ""))
    replace_cell_text(t4.rows[7].cells[2], str(cloud.get("ops_location") or ""))
    replace_cell_text(
        t4.rows[9].cells[2],
        (
            f"云服务商为 {str(cloud.get('provider_name') or '')} / 平台安全等级 {str(cloud.get('provider_level') or '')} / "
            f"平台名称 {str(cloud.get('provider_platform_name') or '')} / 平台备案编号 {str(cloud.get('provider_record_no') or '')}"
        ).strip(),
    )
    replace_cell_text(t4.rows[10].cells[2], str(cloud.get("customer_ops_location") or ""))
    replace_cell_text(t4.rows[11].cells[2], str(cloud.get("record_file_name") or ""))
    mobile = table4.get("mobile") if isinstance(table4.get("mobile"), dict) else {}
    replace_cell_text(t4.rows[12].cells[2], "是" if coerce_bool(mobile.get("enabled")) else "否")
    replace_cell_text(t4.rows[13].cells[2], str(mobile.get("app_name") or ""))
    replace_cell_text(t4.rows[14].cells[2], joined_text(mobile.get("wireless_channels")))
    replace_cell_text(t4.rows[15].cells[2], joined_text(mobile.get("terminal_types")))
    iot = table4.get("iot") if isinstance(table4.get("iot"), dict) else {}
    replace_cell_text(t4.rows[16].cells[2], "是" if coerce_bool(iot.get("enabled")) else "否")
    replace_cell_text(t4.rows[17].cells[2], joined_text(iot.get("perception_layers"), other_value=str(iot.get("perception_other") or "")))
    replace_cell_text(t4.rows[18].cells[2], joined_text(iot.get("transport_layers"), other_value=str(iot.get("transport_other") or "")))
    industrial = table4.get("industrial_control") if isinstance(table4.get("industrial_control"), dict) else {}
    replace_cell_text(t4.rows[19].cells[2], "是" if coerce_bool(industrial.get("enabled")) else "否")
    replace_cell_text(t4.rows[20].cells[2], joined_text(industrial.get("function_layers")))
    replace_cell_text(t4.rows[21].cells[2], joined_text(industrial.get("components"), other_value=str(industrial.get("component_other") or "")))
    big_data = table4.get("big_data") if isinstance(table4.get("big_data"), dict) else {}
    replace_cell_text(t4.rows[22].cells[2], "是" if coerce_bool(big_data.get("enabled")) else "否")
    replace_cell_text(t4.rows[23].cells[2], joined_text(big_data.get("system_components")))
    replace_cell_text(t4.rows[24].cells[2], str(big_data.get("cross_border_status") or ""))
    replace_cell_text(t4.rows[26].cells[2], str(big_data.get("application_count") or ""))
    replace_cell_text(t4.rows[27].cells[2], str(big_data.get("infra_location") or ""))
    replace_cell_text(t4.rows[28].cells[2], str(big_data.get("ops_location") or ""))
    replace_cell_text(
        t4.rows[30].cells[2],
        (
            f"大数据平台服务商 {str(big_data.get('provider_name') or '')} 平台安全等级 {str(big_data.get('provider_level') or '')} / "
            f"平台名称 {str(big_data.get('provider_platform_name') or '')} 平台备案编号 {str(big_data.get('provider_record_no') or '')}"
        ).strip(),
    )
    replace_cell_text(t4.rows[31].cells[2], str(big_data.get("record_file_name") or ""))

    t5 = doc.tables[5]
    table5_rows = [
        ("network_topology", 0),
        ("security_org_and_rules", 1),
        ("security_design_plan", 2),
        ("security_products", 3),
        ("security_services", 4),
        ("supervisor_guidance", 5),
    ]
    for slot_name, row_index in table5_rows:
        slot = table5.get(slot_name) if isinstance(table5.get(slot_name), dict) else {}
        status = str(slot.get("status") or "none").strip()
        label = "有" if status == "has" else "无"
        replace_cell_text(t5.rows[row_index].cells[1], f"{label}    附件名称  {str(slot.get('file_name') or '')}")

    t6 = doc.tables[6]
    items = table6.get("items") if isinstance(table6.get("items"), list) else []
    if items and isinstance(items[0], dict):
        item = items[0]
        replace_cell_text(t6.rows[0].cells[1], str(item.get("data_name") or ""))
        replace_cell_text(t6.rows[0].cells[3], str(item.get("data_level") or ""))
        replace_cell_text(t6.rows[1].cells[1], str(item.get("data_category") or ""))
        replace_cell_text(t6.rows[2].cells[1], str(item.get("data_security_dept") or ""))
        replace_cell_text(t6.rows[2].cells[3], str(item.get("data_security_owner") or ""))
        replace_cell_text(t6.rows[3].cells[1], joined_text(item.get("personal_info_flags")))
        replace_cell_text(t6.rows[4].cells[1], str(item.get("data_total") or ""))
        replace_cell_text(t6.rows[5].cells[1], str(item.get("monthly_growth") or ""))
        replace_cell_text(t6.rows[6].cells[1], joined_text(item.get("data_sources"), other_value=str(item.get("data_source_other") or "")))
        replace_cell_text(t6.rows[7].cells[1], str(item.get("source_units") or ""))
        replace_cell_text(t6.rows[8].cells[1], str(item.get("target_units") or ""))
        replace_cell_text(t6.rows[9].cells[1], str(item.get("interaction") or ""))
        replace_cell_text(t6.rows[10].cells[1], str(item.get("storage_platform") or ""))
        replace_cell_text(t6.rows[11].cells[1], str(item.get("storage_machine_room") or ""))
        replace_cell_text(t6.rows[12].cells[1], str(item.get("storage_location") or ""))


@app.get("/api/systems/{system_id}/export/word")
def export_system_word(request: Request, system_id: int, db: Session = Depends(get_db)) -> FileResponse:
    require_roles(request, db, {"admin", "evaluator"}, legacy_admin=True)
    system = get_system_or_404(db, system_id)
    org = get_org_or_404(db, system.organization_id)
    template_path = DESKTOP_FILING_TEMPLATE_PATH if DESKTOP_FILING_TEMPLATE_PATH.exists() else find_latest_docx_by_pattern("01-*.docx")
    if not template_path or not Path(template_path).exists():
        raise HTTPException(status_code=404, detail="备案表模板不存在，无法导出。")
    path = EXPORT_DIR / f"filing_{system.id}_{datetime.now():%Y%m%d%H%M%S}.docx"
    doc = Document(str(template_path))
    fill_filing_template_document(doc, org, system)
    doc.save(path)
    return FileResponse(path=str(path), filename=path.name, background=BackgroundTask(_cleanup_export_file, path))


@app.post("/api/attachments/{entity_type}/{entity_id}")
async def upload_attachment(
    request: Request,
    entity_type: str,
    entity_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor, _ = require_roles(request, db, {"admin", "evaluator"}, legacy_admin=True)
    entity_type = entity_type.lower()
    if entity_type not in {"organization", "system"}:
        raise HTTPException(status_code=400, detail="entity_type 仅支持 organization/system。")
    if entity_type == "organization":
        get_org_or_404(db, entity_id)
        allowed = {"jpg", "jpeg", "png", "pdf"}
        limit = MAX_ORG_ATTACHMENT
    else:
        get_system_or_404(db, entity_id)
        allowed = {"jpg", "jpeg", "png", "pdf", "vsd", "vsdx"}
        limit = MAX_SYS_ATTACHMENT
    ext = Path(file.filename).suffix.lower().lstrip(".")
    if ext not in allowed:
        raise HTTPException(status_code=400, detail=f"不支持的附件格式: {ext}")
    content = await file.read()
    if len(content) > limit:
        raise HTTPException(status_code=400, detail=f"文件大小超过限制({limit // 1024 // 1024}MB)。")
    row = save_attachment_row(
        db,
        entity_type=entity_type,
        entity_id=entity_id,
        actor=actor,
        file_name=file.filename,
        content=content,
    )
    db.commit()
    db.refresh(row)
    return {"message": "上传成功", "data": {"id": row.id, "file_name": row.file_name, "file_size": row.file_size}}


@app.post("/api/attachments/{entity_type}/{entity_id}/batch")
async def upload_attachment_batch(
    request: Request,
    entity_type: str,
    entity_id: int,
    files: list[UploadFile] = File(...),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor, _ = require_roles(request, db, {"admin", "evaluator"}, legacy_admin=True)
    entity_type = entity_type.lower()
    if entity_type not in {"organization", "system"}:
        raise HTTPException(status_code=400, detail="entity_type 仅支持 organization/system。")
    if entity_type == "organization":
        get_org_or_404(db, entity_id)
        allowed = {"jpg", "jpeg", "png", "pdf"}
        limit = MAX_ORG_ATTACHMENT
    else:
        get_system_or_404(db, entity_id)
        allowed = {"jpg", "jpeg", "png", "pdf", "vsd", "vsdx"}
        limit = MAX_SYS_ATTACHMENT
    if not files:
        raise HTTPException(status_code=400, detail="files 不能为空。")
    uploaded = 0
    skipped: list[str] = []
    for f in files:
        ext = Path(f.filename).suffix.lower().lstrip(".")
        if ext not in allowed:
            skipped.append(f"{f.filename}: 不支持的附件格式 {ext}")
            continue
        content = await f.read()
        if len(content) > limit:
            skipped.append(f"{f.filename}: 文件大小超过限制({limit // 1024 // 1024}MB)")
            continue
        save_attachment_row(
            db,
            entity_type=entity_type,
            entity_id=entity_id,
            actor=actor,
            file_name=f.filename,
            content=content,
        )
        uploaded += 1
    db.commit()
    return {"message": "批量上传完成", "uploaded": uploaded, "skipped": skipped}


@app.get("/api/attachments/{entity_type}/{entity_id}")
def list_attachments(entity_type: str, entity_id: int, db: Session = Depends(get_db)) -> dict[str, Any]:
    rows = (
        db.query(Attachment)
        .filter(Attachment.entity_type == entity_type.lower(), Attachment.entity_id == entity_id)
        .order_by(Attachment.uploaded_at.desc())
        .all()
    )
    return {
        "total": len(rows),
        "items": [
            {
                "id": r.id,
                "file_name": r.file_name,
                "file_ext": r.file_ext,
                "file_size": r.file_size,
                "uploaded_by": r.uploaded_by,
                "uploaded_at": r.uploaded_at,
            }
            for r in rows
        ],
    }


@app.get("/api/attachment-files/{attachment_id}/download")
def download_attachment(attachment_id: int, db: Session = Depends(get_db)) -> FileResponse:
    row = db.query(Attachment).filter(Attachment.id == attachment_id).first()
    if not row:
        raise HTTPException(status_code=404, detail="附件不存在。")
    path = ensure_file_exists(row.file_path, "附件文件不存在或已被移除。")
    return FileResponse(path=str(path), filename=row.file_name)


@app.get("/api/attachment-files/{attachment_id}/preview")
def preview_attachment(attachment_id: int, db: Session = Depends(get_db)) -> FileResponse:
    row = db.query(Attachment).filter(Attachment.id == attachment_id).first()
    if not row:
        raise HTTPException(status_code=404, detail="附件不存在。")
    path = ensure_file_exists(row.file_path, "附件文件不存在或已被移除。")
    guessed_type = mimetypes.guess_type(row.file_name)[0] or "application/octet-stream"
    return FileResponse(
        path=str(path),
        media_type=guessed_type,
        filename=row.file_name,
        content_disposition_type="inline",
    )


@app.delete("/api/attachments/{attachment_id}")
def delete_attachment(request: Request, attachment_id: int, db: Session = Depends(get_db)) -> dict[str, Any]:
    require_roles(request, db, {"admin"}, legacy_admin=True)
    row = db.query(Attachment).filter(Attachment.id == attachment_id).first()
    if not row:
        raise HTTPException(status_code=404, detail="附件不存在。")
    safe_delete_uploaded_file(row.file_path)
    db.delete(row)
    db.commit()
    return {"message": "附件已删除。"}


@app.post("/api/reports/generate")
def generate_report(
    request: Request,
    system_id: int = Query(...),
    report_type: str = Query("grading_report"),
    template_id: int | None = Query(None),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor, _ = require_roles(request, db, {"admin", "evaluator"}, legacy_admin=True)
    report_type = report_type.lower()
    if report_type not in VALID_REPORT_TYPES:
        raise HTTPException(status_code=400, detail="report_type 无效。")
    system = get_system_or_404(db, system_id)
    org = get_org_or_404(db, system.organization_id)
    max_version = (
        db.query(func.max(Report.version_no))
        .filter(Report.system_id == system_id, Report.report_type == report_type)
        .scalar()
        or 0
    )
    version = max_version + 1
    content = generate_report_payload(report_type, org, system)
    template = choose_report_template(db, report_type, org.filing_region or "", system.proposed_level, template_id=template_id)
    content = apply_template_config_to_payload(content, template)
    if "签章" not in content:
        content["签章"] = {"公司签章": "", "测评师签章": "", "签章日期": ""}
    if template and report_type == "expert_review_form":
        content["模板信息"] = content.get("模板信息", {})
        content["模板信息"]["匹配说明"] = f"按地市[{org.filing_region}]匹配模板"
    report = Report(
        organization_id=org.id,
        system_id=system.id,
        report_type=report_type,
        version_no=version,
        title=build_report_title(report_type, system.system_name, version),
        status="draft",
        content=content,
        generated_by=actor,
    )
    db.add(report)
    db.commit()
    db.refresh(report)
    return {
        "message": "报告生成成功",
        "data": {
            "id": report.id,
            "title": report.title,
            "report_type": report.report_type,
            "version_no": report.version_no,
            "status": report.status,
            "content": report.content,
            "template_id": template.id if template else None,
            "template_name": template.template_name if template else None,
        },
    }


@app.get("/api/reports")
def list_reports(
    system_id: int | None = None,
    report_type: str | None = None,
    status: str | None = None,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    query = db.query(Report)
    if system_id:
        query = query.filter(Report.system_id == system_id)
    if report_type:
        query = query.filter(Report.report_type == report_type)
    if status:
        query = query.filter(Report.status == status)
    items = query.order_by(Report.generated_at.desc()).all()
    return {
        "total": len(items),
        "items": [
            {
                "id": r.id,
                "title": r.title,
                "report_type": r.report_type,
                "version_no": r.version_no,
                "status": r.status,
                "generated_by": r.generated_by,
                "generated_at": r.generated_at,
            }
            for r in items
        ],
    }


@app.get("/api/reports/{report_id}")
def get_report(report_id: int, db: Session = Depends(get_db)) -> dict[str, Any]:
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="报告不存在。")
    return {
        "id": report.id,
        "title": report.title,
        "report_type": report.report_type,
        "version_no": report.version_no,
        "status": report.status,
        "content": report.content,
        "generated_by": report.generated_by,
        "generated_at": report.generated_at,
    }


@app.put("/api/reports/{report_id}")
def edit_report(
    request: Request,
    report_id: int,
    payload: ReportEdit,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor, role = require_roles(request, db, {"admin", "evaluator"}, legacy_admin=True)
    assert_safe_payload(payload.content, "content")
    if payload.title:
        assert_safe_text(payload.title, "title")
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="报告不存在。")
    effective_is_admin = role == "admin"
    if report.status in {"submitted", "approved"} and not effective_is_admin:
        raise HTTPException(status_code=403, detail="当前状态不可编辑。")
    report.content = payload.content
    if payload.title:
        report.title = payload.title
    db.add(
        ReviewRecord(
            report_id=report.id,
            reviewer=actor,
            action="edit",
            comment="编辑报告内容",
        )
    )
    db.commit()
    db.refresh(report)
    return {"message": "保存成功", "data": {"id": report.id, "status": report.status}}


def _editable_report_or_403(db: Session, report_id: int, is_admin: bool) -> Report:
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="报告不存在。")
    if report.status in {"submitted", "approved"} and not is_admin:
        raise HTTPException(status_code=403, detail="当前状态不可编辑。")
    return report


@app.post("/api/reports/{report_id}/sections")
def add_report_section(
    request: Request,
    report_id: int,
    payload: dict[str, Any],
    index: int | None = Query(None),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor, role = require_roles(request, db, {"admin", "evaluator"}, legacy_admin=True)
    report = _editable_report_or_403(db, report_id, role == "admin")
    content = dict(report.content or {})
    sections = list(content.get("章节") or [])
    name = str(payload.get("name") or "").strip()
    if not name:
        raise HTTPException(status_code=400, detail="name 不能为空。")
    assert_safe_text(name, "name")
    section_content = payload.get("content")
    if section_content is None:
        section_content = {}
    assert_safe_payload(section_content, "section_content")
    new_section = {"名称": name, "内容": section_content}
    if index is None or index < 0 or index >= len(sections):
        sections.append(new_section)
    else:
        sections.insert(index, new_section)
    content["章节"] = sections
    report.content = content
    db.add(
        ReviewRecord(
            report_id=report.id,
            reviewer=actor,
            action="section_add",
            comment=f"新增章节: {name}",
        )
    )
    db.commit()
    return {"message": "章节新增成功", "section_count": len(sections)}


@app.delete("/api/reports/{report_id}/sections/{section_index}")
def delete_report_section(
    request: Request,
    report_id: int,
    section_index: int,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor, role = require_roles(request, db, {"admin", "evaluator"}, legacy_admin=True)
    report = _editable_report_or_403(db, report_id, role == "admin")
    content = dict(report.content or {})
    sections = list(content.get("章节") or [])
    if section_index < 0 or section_index >= len(sections):
        raise HTTPException(status_code=400, detail="section_index 超出范围。")
    removed = sections.pop(section_index)
    content["章节"] = sections
    report.content = content
    db.add(
        ReviewRecord(
            report_id=report.id,
            reviewer=actor,
            action="section_delete",
            comment=f"删除章节: {removed.get('名称', '')}",
        )
    )
    db.commit()
    return {"message": "章节删除成功", "section_count": len(sections)}


@app.post("/api/reports/{report_id}/sections/reorder")
def reorder_report_section(
    request: Request,
    report_id: int,
    from_index: int = Query(..., ge=0),
    to_index: int = Query(..., ge=0),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor_name, role = require_roles(request, db, {"admin", "reviewer", "evaluator"}, legacy_admin=True)
    report = _editable_report_or_403(db, report_id, role == "admin")
    content = dict(report.content or {})
    sections = list(content.get("章节") or [])
    if from_index >= len(sections) or to_index >= len(sections):
        raise HTTPException(status_code=400, detail="索引超出范围。")
    section = sections.pop(from_index)
    sections.insert(to_index, section)
    content["章节"] = sections
    report.content = content
    db.add(
        ReviewRecord(
            report_id=report.id,
            reviewer=actor_name,
            action="section_reorder",
            comment=f"章节顺序调整 {from_index}->{to_index}",
        )
    )
    db.commit()
    return {"message": "章节顺序调整成功"}


@app.post("/api/reports/{report_id}/signature")
def set_report_signature(
    request: Request,
    report_id: int,
    payload: dict[str, Any],
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor_name, role = require_roles(request, db, {"admin", "reviewer", "evaluator"}, legacy_admin=True)
    assert_safe_payload(payload, "payload")
    report = _editable_report_or_403(db, report_id, role == "admin")
    content = dict(report.content or {})
    sign = content.get("签章")
    if not isinstance(sign, dict):
        sign = {}
    for key in ["公司签章", "测评师签章", "签章日期", "签章位置", "备注"]:
        if key in payload:
            sign[key] = payload[key]
    content["签章"] = sign
    report.content = content
    db.add(
        ReviewRecord(
            report_id=report.id,
            reviewer=actor_name,
            action="signature",
            comment="更新签章信息",
        )
    )
    db.commit()
    return {"message": "签章信息已更新", "signature": sign}


@app.post("/api/reports/{report_id}/submit")
def submit_report(
    request: Request,
    report_id: int,
    reviewer: str = Query("reviewer"),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor_name, _ = require_roles(request, db, {"admin", "reviewer", "evaluator"}, legacy_admin=True)
    assert_safe_text(reviewer, "reviewer")
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="报告不存在。")
    if report.status not in {"draft", "rejected"}:
        raise HTTPException(status_code=400, detail="当前状态不可提交。")
    report.status = "submitted"
    db.add(
        ReviewRecord(
            report_id=report.id,
            reviewer=actor_name,
            action="submit",
            comment=f"提交审核给 {reviewer}",
        )
    )
    db.commit()
    return {"message": "提交成功", "status": report.status}


@app.post("/api/reports/{report_id}/review")
def review_report(
    request: Request,
    report_id: int,
    action: str = Query(...),
    comment: str = Query(""),
    position: str = Query(""),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor_name, _ = require_roles(request, db, {"admin", "reviewer"}, legacy_admin=True)
    assert_safe_text(comment, "comment")
    assert_safe_text(position, "position")
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="报告不存在。")
    if report.status != "submitted":
        raise HTTPException(status_code=400, detail="仅已提交报告可审核。")
    action = action.lower()
    if action not in {"approve", "reject"}:
        raise HTTPException(status_code=400, detail="action 仅支持 approve/reject。")
    report.status = "approved" if action == "approve" else "rejected"
    db.add(
        ReviewRecord(
            report_id=report.id,
            reviewer=actor_name,
            action=action,
            comment=comment,
            position=position,
        )
    )
    db.commit()
    return {"message": "审核完成", "status": report.status}


@app.get("/api/reports/{report_id}/reviews")
def report_reviews(report_id: int, db: Session = Depends(get_db)) -> dict[str, Any]:
    rows = db.query(ReviewRecord).filter(ReviewRecord.report_id == report_id).order_by(ReviewRecord.created_at.asc()).all()
    return {
        "total": len(rows),
        "items": [
            {
                "reviewer": r.reviewer,
                "action": r.action,
                "comment": r.comment,
                "position": r.position,
                "created_at": r.created_at,
            }
            for r in rows
        ],
    }


@app.get("/api/reports/{report_id}/versions")
def list_report_versions(report_id: int, db: Session = Depends(get_db)) -> dict[str, Any]:
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="报告不存在。")
    rows = (
        db.query(Report)
        .filter(Report.system_id == report.system_id, Report.report_type == report.report_type)
        .order_by(Report.version_no.desc())
        .all()
    )
    return {
        "total": len(rows),
        "items": [{"id": r.id, "version_no": r.version_no, "title": r.title, "status": r.status} for r in rows],
    }


def _normalize_for_diff(value: Any) -> str:
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False, sort_keys=True)
    return "" if value is None else str(value)


def _section_name_map(content: dict[str, Any]) -> dict[str, Any]:
    items = list(content.get("章节") or [])
    result: dict[str, Any] = {}
    for idx, row in enumerate(items):
        if not isinstance(row, dict):
            result[f"#{idx}"] = row
            continue
        name = str(row.get("名称") or f"#{idx}")
        result[name] = row.get("内容")
    return result


@app.get("/api/reports/{report_id}/compare/{target_id}")
def compare_report_versions(report_id: int, target_id: int, db: Session = Depends(get_db)) -> dict[str, Any]:
    left = db.query(Report).filter(Report.id == report_id).first()
    right = db.query(Report).filter(Report.id == target_id).first()
    if not left or not right:
        raise HTTPException(status_code=404, detail="报告不存在。")
    if left.system_id != right.system_id or left.report_type != right.report_type:
        raise HTTPException(status_code=400, detail="仅可对比同系统同类型报告。")
    left_content = dict(left.content or {})
    right_content = dict(right.content or {})
    fields = sorted(set(left_content.keys()) | set(right_content.keys()))
    field_changes = []
    for key in fields:
        lv = left_content.get(key)
        rv = right_content.get(key)
        if _normalize_for_diff(lv) != _normalize_for_diff(rv):
            field_changes.append({"field": key, "from": lv, "to": rv})

    left_sections = _section_name_map(left_content)
    right_sections = _section_name_map(right_content)
    left_names = set(left_sections.keys())
    right_names = set(right_sections.keys())
    section_added = sorted(right_names - left_names)
    section_removed = sorted(left_names - right_names)
    section_changed = sorted(
        [name for name in (left_names & right_names) if _normalize_for_diff(left_sections[name]) != _normalize_for_diff(right_sections[name])]
    )
    return {
        "left": {"id": left.id, "version_no": left.version_no, "title": left.title},
        "right": {"id": right.id, "version_no": right.version_no, "title": right.title},
        "summary": {
            "field_change_count": len(field_changes),
            "section_added_count": len(section_added),
            "section_removed_count": len(section_removed),
            "section_changed_count": len(section_changed),
        },
        "field_changes": field_changes,
        "section_added": section_added,
        "section_removed": section_removed,
        "section_changed": section_changed,
    }


@app.post("/api/reports/{report_id}/restore/{target_id}")
def restore_report_version(
    request: Request, report_id: int, target_id: int, db: Session = Depends(get_db)
) -> dict[str, Any]:
    actor, _ = require_roles(request, db, {"admin"}, legacy_admin=True)
    current = db.query(Report).filter(Report.id == report_id).first()
    target = db.query(Report).filter(Report.id == target_id).first()
    if not current or not target:
        raise HTTPException(status_code=404, detail="报告不存在。")
    if current.system_id != target.system_id or current.report_type != target.report_type:
        raise HTTPException(status_code=400, detail="仅可恢复同系统同类型报告。")
    max_version = (
        db.query(func.max(Report.version_no))
        .filter(Report.system_id == current.system_id, Report.report_type == current.report_type)
        .scalar()
        or 0
    )
    new_report = Report(
        organization_id=current.organization_id,
        system_id=current.system_id,
        report_type=current.report_type,
        version_no=max_version + 1,
        title=f"{current.title}-恢复V{max_version + 1}",
        status="draft",
        content=target.content,
        generated_by=actor,
    )
    db.add(new_report)
    db.commit()
    db.refresh(new_report)
    return {"message": "恢复成功", "new_report_id": new_report.id, "version_no": new_report.version_no}


def prepare_export_content(content: dict[str, Any]) -> dict[str, Any]:
    output = dict(content or {})
    sections = list(output.get("章节") or [])
    sign = output.get("签章")
    if isinstance(sign, dict):
        has_sign_section = any(str(s.get("名称") or "") == "签章信息" for s in sections if isinstance(s, dict))
        if not has_sign_section:
            sections.append({"名称": "签章信息", "内容": sign})
    tpl = output.get("模板信息")
    if isinstance(tpl, dict):
        has_tpl_section = any(str(s.get("名称") or "") == "模板信息" for s in sections if isinstance(s, dict))
        if not has_tpl_section:
            sections.append({"名称": "模板信息", "内容": tpl})
    toc_names = []
    for item in sections:
        if not isinstance(item, dict):
            continue
        name = str(item.get("名称") or "").strip()
        if not name or name == "目录":
            continue
        toc_names.append(name)
    toc_section = {"名称": "目录", "内容": {f"{idx + 1}": name for idx, name in enumerate(toc_names)}}
    sections = [s for s in sections if not (isinstance(s, dict) and str(s.get("名称") or "") == "目录")]
    sections.insert(0, toc_section)
    output["章节"] = sections
    return output


def build_report_template_field_map(report: Report, content: dict[str, Any], db: Session) -> dict[str, str]:
    field_map: dict[str, str] = {}
    org = db.query(Organization).filter(Organization.id == report.organization_id).first()
    system = db.query(SystemInfo).filter(SystemInfo.id == report.system_id).first()
    if org:
        field_map.update(
            {
                "单位名称": org.name or "",
                "统一社会信用代码": org.credit_code or "",
                "单位地址": org.address or "",
                "单位负责人": org.legal_representative or "",
                "邮箱": org.email or "",
                "备案地区": org.filing_region or "",
                "联系人": org.cybersecurity_owner_name or org.legal_representative or "",
                "联系电话": org.mobile_phone or org.office_phone or "",
            }
        )
    if system:
        level = f"第{system.proposed_level}级" if system.proposed_level else ""
        field_map.update(
            {
                "系统名称": system.system_name or "",
                "信息系统名称": system.system_name or "",
                "系统自定安全级别": level,
                "拟定等级": level,
                "部署方式": system.deployment_mode or "",
                "系统编号": system.system_code or "",
            }
        )
    field_map["报告标题"] = report.title or ""
    field_map["生成日期"] = (report.generated_at or datetime.now()).strftime("%Y-%m-%d")
    for section in content.get("章节", []):
        if not isinstance(section, dict):
            continue
        body = section.get("内容")
        if not isinstance(body, dict):
            continue
        for k, v in body.items():
            if v is None:
                continue
            key = str(k).strip()
            if not key:
                continue
            value = str(v).strip()
            if not value:
                continue
            field_map.setdefault(key, value)
    return field_map


@app.get("/api/reports/{report_id}/export/word")
def export_report_word(request: Request, report_id: int, db: Session = Depends(get_db)) -> FileResponse:
    require_roles(request, db, {"admin", "evaluator"}, legacy_admin=True)
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="报告不存在。")
    path = EXPORT_DIR / f"report_{report.id}_{datetime.now():%Y%m%d%H%M%S}.docx"
    export_content = prepare_export_content(report.content)
    tpl_info = export_content.get("模板信息")
    template_id = None
    if isinstance(tpl_info, dict):
        try:
            template_id = int(tpl_info.get("template_id")) if tpl_info.get("template_id") is not None else None
        except Exception:
            template_id = None
    if template_id:
        tpl = db.query(ReportTemplate).filter(ReportTemplate.id == template_id, ReportTemplate.status == "enabled").first()
        if tpl and tpl.file_path and Path(tpl.file_path).exists():
            field_map = build_report_template_field_map(report, export_content, db)
            try:
                export_report_docx_with_template(Path(tpl.file_path), field_map, path)
                return FileResponse(path=str(path), filename=path.name, background=BackgroundTask(_cleanup_export_file, path))
            except Exception as exc:
                logger.exception("模板导出失败，已回退普通导出。report_id=%s template_id=%s error=%s", report.id, template_id, exc)
    export_report_docx(report.title, export_content, path)
    return FileResponse(path=str(path), filename=path.name, background=BackgroundTask(_cleanup_export_file, path))


@app.get("/api/reports/{report_id}/export/pdf")
def export_report_pdf_file(
    request: Request,
    report_id: int,
    password: str | None = Query(None, min_length=4, max_length=64),
    db: Session = Depends(get_db),
) -> FileResponse:
    require_roles(request, db, {"admin", "evaluator"}, legacy_admin=True)
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="报告不存在。")
    path = EXPORT_DIR / f"report_{report.id}_{datetime.now():%Y%m%d%H%M%S}.pdf"
    export_report_pdf(report.title, prepare_export_content(report.content), path, password=password)
    return FileResponse(path=str(path), filename=path.name, background=BackgroundTask(_cleanup_export_file, path))


@app.get("/api/workflow/config")
def get_workflow(db: Session = Depends(get_db)) -> dict[str, Any]:
    config = get_workflow_config(db)
    return {"name": config.name, "steps": config.steps_json, "updated_by": config.updated_by, "updated_at": config.updated_at}


@app.put("/api/workflow/config")
def update_workflow(request: Request, payload: WorkflowConfigUpdate, db: Session = Depends(get_db)) -> dict[str, Any]:
    actor, _ = require_roles(request, db, {"admin"}, legacy_admin=(not STRICT_AUTH))
    if not payload.steps:
        raise HTTPException(status_code=400, detail="流程步骤不能为空。")
    for step in payload.steps:
        assert_safe_text(step, "workflow.step")
    config = get_workflow_config(db)
    config.steps_json = payload.steps
    config.updated_by = resolve_effective_actor_name(request, db, payload.updated_by or actor)
    get_or_create_workflow_step_rules(db, config)
    in_progress_rows = db.query(WorkflowInstance).filter(WorkflowInstance.status == "in_progress").all()
    for row in in_progress_rows:
        if row.current_step_index >= len(config.steps_json):
            row.current_step_index = max(0, len(config.steps_json) - 1)
        recalc_workflow_due_at(db, row, config)
    db.commit()
    db.refresh(config)
    return {"message": "流程配置更新成功", "steps": config.steps_json}


@app.get("/api/workflow/rules")
def list_workflow_rules(db: Session = Depends(get_db)) -> dict[str, Any]:
    config = get_workflow_config(db)
    rule_map = get_or_create_workflow_step_rules(db, config)
    rows = []
    for step in config.steps_json:
        r = rule_map[step]
        rows.append(
            {
                "step_name": r.step_name,
                "owner": r.owner,
                "time_limit_hours": r.time_limit_hours,
                "enabled": r.enabled,
                "updated_by": r.updated_by,
                "updated_at": r.updated_at,
            }
        )
    db.commit()
    return {"config_name": config.name, "rules": rows}


@app.put("/api/workflow/rules")
def update_workflow_rules(request: Request, payload: dict[str, Any], db: Session = Depends(get_db)) -> dict[str, Any]:
    actor, _ = require_roles(request, db, {"admin"}, legacy_admin=(not STRICT_AUTH))
    config = get_workflow_config(db)
    rules = payload.get("rules") or []
    updated_by = resolve_effective_actor_name(request, db, str(payload.get("updated_by") or actor or "admin"))
    assert_safe_text(updated_by, "updated_by")
    if not isinstance(rules, list) or not rules:
        raise HTTPException(status_code=400, detail="rules 不能为空。")
    step_set = set(config.steps_json or [])
    for item in rules:
        step_name = str(item.get("step_name") or "").strip()
        if not step_name:
            raise HTTPException(status_code=400, detail="step_name 不能为空。")
        assert_safe_text(step_name, "step_name")
        if step_name not in step_set:
            raise HTTPException(status_code=400, detail=f"非法流程节点: {step_name}")
        owner = str(item.get("owner") or "system").strip()
        assert_safe_text(owner, "owner")
        raw_limit = item.get("time_limit_hours", 24)
        try:
            limit = int(raw_limit)
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"{step_name} 的 time_limit_hours 必须为整数。") from exc
        enabled_raw = item.get("enabled", True)
        enabled_parsed = parse_optional_bool(enabled_raw, f"{step_name}.enabled")
        enabled = True if enabled_parsed is None else enabled_parsed
        row = (
            db.query(WorkflowStepRule)
            .filter(WorkflowStepRule.config_name == config.name, WorkflowStepRule.step_name == step_name)
            .order_by(WorkflowStepRule.id.desc())
            .first()
        )
        if not row:
            row = WorkflowStepRule(
                config_name=config.name,
                step_name=step_name,
                owner=owner,
                time_limit_hours=max(1, limit),
                enabled=enabled,
                updated_by=updated_by,
            )
            db.add(row)
        else:
            row.owner = owner
            row.time_limit_hours = max(1, limit)
            row.enabled = enabled
            row.updated_by = updated_by
    in_progress_rows = db.query(WorkflowInstance).filter(WorkflowInstance.status == "in_progress").all()
    for row in in_progress_rows:
        recalc_workflow_due_at(db, row, config)
    db.commit()
    return {"message": "流程规则更新成功"}


@app.post("/api/workflow/instances/{system_id}/extend")
def workflow_extend_due(
    request: Request,
    system_id: int,
    hours: int = Query(24, ge=1, le=720),
    reason: str = Query(""),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor_name, _ = require_roles(request, db, {"admin"}, legacy_admin=True)
    assert_safe_text(reason, "reason")
    get_system_or_404(db, system_id)
    instance = db.query(WorkflowInstance).filter(WorkflowInstance.system_id == system_id).first()
    if not instance:
        raise HTTPException(status_code=404, detail="流程实例不存在。")
    if instance.status != "in_progress":
        raise HTTPException(status_code=400, detail="仅进行中流程可延时。")
    base = instance.due_at or datetime.now()
    instance.due_at = base + timedelta(hours=hours)
    steps = get_workflow_config(db).steps_json
    current_step = steps[instance.current_step_index] if instance.current_step_index < len(steps) else "未知节点"
    db.add(
        WorkflowAction(
            instance_id=instance.id,
            step_name=current_step,
            actor=actor_name,
            action="extend_due",
            comment=f"延长{hours}小时 {reason}".strip(),
        )
    )
    db.commit()
    db.refresh(instance)
    return {"message": "已延长时限", "due_at": instance.due_at}


@app.get("/api/workflow/reminders")
def workflow_reminders(
    mode: str = Query("all"),
    within_hours: int = Query(4, ge=1, le=72),
    send: bool = Query(False),
    channel: str = Query("in_app"),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    mode = mode.lower()
    if mode not in {"all", "due", "overdue"}:
        raise HTTPException(status_code=400, detail="mode 仅支持 all/due/overdue。")
    channel = (channel or "in_app").strip().lower()
    if channel not in {"in_app", "email", "both"}:
        raise HTTPException(status_code=400, detail="channel 仅支持 in_app/email/both。")
    config = get_workflow_config(db)
    rule_map = get_or_create_workflow_step_rules(db, config)
    query = (
        db.query(WorkflowInstance, SystemInfo.system_name)
        .join(SystemInfo, SystemInfo.id == WorkflowInstance.system_id)
        .filter(WorkflowInstance.status == "in_progress", WorkflowInstance.due_at.is_not(None))
    )
    rows = query.order_by(WorkflowInstance.due_at.asc()).all()
    owner_email_map: dict[str, str] = {}
    if send and channel in {"email", "both"}:
        owner_candidates: set[str] = set()
        for instance, _ in rows:
            owner_name = get_workflow_step_owner(db, config, instance.current_step_index, rule_map=rule_map)
            if owner_name and "@" not in owner_name:
                owner_candidates.add(owner_name)
        if owner_candidates:
            users = db.query(UserAccount).filter(UserAccount.username.in_(owner_candidates)).all()
            owner_email_map = {
                user.username: user.username
                for user in users
                if user.username and "@" in user.username
            }
    items: list[dict[str, Any]] = []
    now = datetime.now()
    for instance, system_name in rows:
        remaining_seconds = int((instance.due_at - now).total_seconds())
        is_overdue = remaining_seconds < 0
        if mode == "overdue" and not is_overdue:
            continue
        if mode == "due" and (is_overdue or remaining_seconds > within_hours * 3600):
            continue
        step_name = ""
        if 0 <= instance.current_step_index < len(config.steps_json):
            step_name = config.steps_json[instance.current_step_index]
        owner = get_workflow_step_owner(db, config, instance.current_step_index, rule_map=rule_map)
        item = {
            "instance_id": instance.id,
            "system_id": instance.system_id,
            "system_name": system_name,
            "step_name": step_name,
            "owner": owner,
            "due_at": instance.due_at,
            "remaining_seconds": remaining_seconds,
            "status": "overdue" if is_overdue else "due_soon",
        }
        items.append(item)
        if send:
            msg = f"系统[{system_name}]节点[{step_name}]将于{instance.due_at}到期，请及时处理。"
            if is_overdue:
                msg = f"系统[{system_name}]节点[{step_name}]已超时，请立即处理。"
            if channel in {"in_app", "both"}:
                db.add(
                    WorkflowReminder(
                        instance_id=instance.id,
                        step_name=step_name or "未知节点",
                        receiver=owner or "system",
                        reminder_type=item["status"],
                        channel="in_app",
                        content=msg,
                    )
                )
            if channel in {"email", "both"}:
                to_email = ""
                if owner and "@" in owner:
                    to_email = owner
                else:
                    to_email = owner_email_map.get(owner or "", "")
                if not to_email:
                    to_email = WORKFLOW_EMAIL_DEFAULT_TO
                ok = False
                err = "无可用邮箱地址"
                if to_email:
                    ok, err = send_workflow_email(
                        to_email=to_email,
                        subject=f"[流程提醒]{system_name}-{step_name or '未知节点'}",
                        body=msg,
                    )
                log_msg = msg if ok else f"{msg} (邮件发送失败: {err})"
                db.add(
                    WorkflowReminder(
                        instance_id=instance.id,
                        step_name=step_name or "未知节点",
                        receiver=to_email or (owner or "system"),
                        reminder_type=item["status"],
                        channel="email",
                        content=log_msg,
                    )
                )
    if send and items:
        db.commit()
    return {"total": len(items), "items": items}


@app.get("/api/workflow/reminder-logs")
def workflow_reminder_logs(db: Session = Depends(get_db)) -> dict[str, Any]:
    rows = db.query(WorkflowReminder).order_by(WorkflowReminder.created_at.desc()).limit(200).all()
    return {
        "total": len(rows),
        "items": [
            {
                "instance_id": r.instance_id,
                "step_name": r.step_name,
                "receiver": r.receiver,
                "reminder_type": r.reminder_type,
                "channel": r.channel,
                "content": r.content,
                "created_at": r.created_at,
            }
            for r in rows
        ],
    }


@app.get("/api/workflow/instances/{system_id}")
def workflow_instance(system_id: int, db: Session = Depends(get_db)) -> dict[str, Any]:
    get_system_or_404(db, system_id)
    config = get_workflow_config(db)
    get_or_create_workflow_step_rules(db, config)
    instance = db.query(WorkflowInstance).filter(WorkflowInstance.system_id == system_id).first()
    if not instance:
        instance = WorkflowInstance(system_id=system_id, current_step_index=0, status="in_progress")
        db.add(instance)
        recalc_workflow_due_at(db, instance, config)
        db.commit()
        db.refresh(instance)
    elif instance.status == "in_progress" and not instance.due_at:
        recalc_workflow_due_at(db, instance, config)
        db.commit()
        db.refresh(instance)
    step_name = ""
    if 0 <= instance.current_step_index < len(config.steps_json):
        step_name = config.steps_json[instance.current_step_index]
    owner = get_workflow_step_owner(db, config, instance.current_step_index)
    remaining_seconds = None
    overdue = False
    if instance.due_at:
        delta = int((instance.due_at - datetime.now()).total_seconds())
        remaining_seconds = delta
        overdue = delta < 0
    logs = (
        db.query(WorkflowAction)
        .filter(WorkflowAction.instance_id == instance.id)
        .order_by(WorkflowAction.created_at.desc())
        .all()
    )
    return {
        "system_id": system_id,
        "status": instance.status,
        "current_step_index": instance.current_step_index,
        "current_step_name": step_name,
        "current_owner": owner,
        "due_at": instance.due_at,
        "remaining_seconds": remaining_seconds,
        "overdue": overdue,
        "steps": config.steps_json,
        "logs": [
            {
                "step_name": l.step_name,
                "actor": l.actor,
                "action": l.action,
                "comment": l.comment,
                "created_at": l.created_at,
            }
            for l in logs
        ],
    }


@app.post("/api/workflow/instances/{system_id}/advance")
def workflow_advance(
    request: Request,
    system_id: int,
    action: str = Query("complete"),
    comment: str = Query(""),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor_name, _ = require_roles(request, db, {"admin", "reviewer", "evaluator"}, legacy_admin=True)
    assert_safe_text(comment, "comment")
    get_system_or_404(db, system_id)
    config = get_workflow_config(db)
    instance = db.query(WorkflowInstance).filter(WorkflowInstance.system_id == system_id).first()
    if not instance:
        instance = WorkflowInstance(system_id=system_id, current_step_index=0, status="in_progress")
        db.add(instance)
        db.flush()
    steps = config.steps_json
    if not steps:
        raise HTTPException(status_code=400, detail="流程步骤为空。")
    current_idx = instance.current_step_index
    current_name = steps[current_idx] if current_idx < len(steps) else "已完成"
    action = action.lower()
    if action in {"complete", "approve"}:
        if current_idx + 1 >= len(steps):
            instance.status = "completed"
        else:
            instance.current_step_index = current_idx + 1
            instance.status = "in_progress"
    elif action in {"reject", "abnormal"}:
        instance.status = "abnormal"
    elif action == "reset":
        instance.current_step_index = 0
        instance.status = "in_progress"
    else:
        raise HTTPException(status_code=400, detail="action 仅支持 complete/approve/reject/abnormal/reset。")
    db.add(
        WorkflowAction(
            instance_id=instance.id,
            step_name=current_name,
            actor=actor_name,
            action=action,
            comment=comment,
        )
    )
    recalc_workflow_due_at(db, instance, config)
    db.commit()
    db.refresh(instance)
    return {
        "message": "流程已更新",
        "status": instance.status,
        "current_step_index": instance.current_step_index,
        "current_step_name": steps[instance.current_step_index]
        if instance.current_step_index < len(steps)
        else "已完成",
    }


@app.get("/api/dashboard/summary")
def dashboard_summary(
    request: Request,
    start_date: date | None = None,
    end_date: date | None = None,
    industry: str | None = None,
    city: str | None = None,
    level: int | None = None,
    evaluator: str | None = None,
    project_status: str | None = None,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    current_user = get_current_user_optional(request, db)
    role = current_user.role if current_user else ""
    username = current_user.username if current_user else ""
    org_q = db.query(Organization).filter(Organization.deleted_at.is_(None))
    if start_date:
        org_q = org_q.filter(Organization.created_at >= datetime.combine(start_date, datetime.min.time()))
    if end_date:
        org_q = org_q.filter(Organization.created_at <= datetime.combine(end_date, datetime.max.time()))
    if industry:
        org_q = org_q.filter(Organization.industry == industry)
    if city:
        org_q = org_q.filter(Organization.filing_region.like(f"%{escape_like(city)}%"))
    if evaluator:
        org_q = org_q.filter(Organization.created_by.like(f"%{escape_like(evaluator)}%"))
    if role == "evaluator":
        org_q = org_q.filter(Organization.created_by == username)

    sys_q = db.query(SystemInfo).join(Organization, Organization.id == SystemInfo.organization_id).filter(
        SystemInfo.deleted_at.is_(None), Organization.deleted_at.is_(None)
    )
    if industry:
        sys_q = sys_q.filter(Organization.industry == industry)
    if city:
        sys_q = sys_q.filter(Organization.filing_region.like(f"%{escape_like(city)}%"))
    if start_date:
        sys_q = sys_q.filter(SystemInfo.created_at >= datetime.combine(start_date, datetime.min.time()))
    if end_date:
        sys_q = sys_q.filter(SystemInfo.created_at <= datetime.combine(end_date, datetime.max.time()))
    if level:
        sys_q = sys_q.filter(SystemInfo.proposed_level == level)
    if evaluator:
        sys_q = sys_q.filter(SystemInfo.created_by.like(f"%{escape_like(evaluator)}%"))
    if role == "evaluator":
        sys_q = sys_q.filter(SystemInfo.created_by == username)

    if project_status:
        sys_q = sys_q.join(WorkflowInstance, WorkflowInstance.system_id == SystemInfo.id).filter(
            WorkflowInstance.status == project_status
        )

    region_stats = (
        org_q.with_entities(Organization.filing_region, func.count(Organization.id)).group_by(Organization.filing_region).all()
    )
    industry_stats = (
        org_q.with_entities(Organization.industry, func.count(Organization.id)).group_by(Organization.industry).all()
    )
    level_stats = sys_q.with_entities(SystemInfo.proposed_level, func.count(SystemInfo.id)).group_by(SystemInfo.proposed_level).all()
    system_ids_query = sys_q.with_entities(SystemInfo.id).distinct()
    pending_reports = (
        db.query(Report)
        .filter(Report.status == "submitted", Report.system_id.in_(system_ids_query))
        .count()
    )
    in_progress_projects = (
        db.query(WorkflowInstance)
        .filter(WorkflowInstance.status == "in_progress", WorkflowInstance.system_id.in_(system_ids_query))
        .count()
    )
    archived_system_count = sys_q.filter(SystemInfo.archived.is_(True)).count()
    return {
        "totals": {
            "organization_count": org_q.count(),
            "system_count": sys_q.count(),
            "archived_organization_count": org_q.filter(Organization.archived.is_(True)).count(),
            "archived_system_count": archived_system_count,
            "pending_review_reports": pending_reports,
            "in_progress_projects": in_progress_projects,
        },
        "region_distribution": [{"name": k or "未填写", "value": v} for k, v in region_stats],
        "industry_distribution": [{"name": k or "未填写", "value": v} for k, v in industry_stats],
        "level_distribution": [{"name": f"{k}级", "value": v} for k, v in level_stats],
    }


@app.get("/api/alerts/summary")
def alerts_summary(request: Request, db: Session = Depends(get_db)) -> dict[str, Any]:
    current_user = get_current_user_optional(request, db)
    role = current_user.role if current_user else ""
    username = current_user.username if current_user else ""

    delete_query = db.query(DeleteRequest).filter(DeleteRequest.status == "pending")
    if role == "evaluator":
        delete_query = delete_query.filter(DeleteRequest.requested_by == username)

    pending_org_delete_count = delete_query.filter(DeleteRequest.entity_type == "organization").count()
    pending_sys_delete_count = delete_query.filter(DeleteRequest.entity_type == "system").count()

    org_recycle_query = db.query(Organization).filter(Organization.deleted_at.is_not(None))
    sys_recycle_query = db.query(SystemInfo).filter(SystemInfo.deleted_at.is_not(None))
    if role == "evaluator":
        org_recycle_query = org_recycle_query.filter(Organization.created_by == username)
        sys_recycle_query = sys_recycle_query.filter(SystemInfo.created_by == username)

    org_recycle_count = org_recycle_query.count()
    sys_recycle_count = sys_recycle_query.count()
    total_attention_count = pending_org_delete_count + pending_sys_delete_count + org_recycle_count + sys_recycle_count

    entries: list[dict[str, Any]] = []
    if pending_org_delete_count:
        entries.append({
            "key": "org-requests",
            "label": "单位删除申请待处理",
            "count": pending_org_delete_count,
            "href": "/organizations?attention=org-requests#workspaceMaintenanceSection",
        })
    if pending_sys_delete_count:
        entries.append({
            "key": "sys-requests",
            "label": "系统删除申请待处理",
            "count": pending_sys_delete_count,
            "href": "/organizations?attention=sys-requests#workspaceMaintenanceSection",
        })
    if org_recycle_count:
        entries.append({
            "key": "org-recycle",
            "label": "单位回收站",
            "count": org_recycle_count,
            "href": "/organizations?attention=org-recycle#workspaceMaintenanceSection",
        })
    if sys_recycle_count:
        entries.append({
            "key": "sys-recycle",
            "label": "系统回收站",
            "count": sys_recycle_count,
            "href": "/organizations?attention=sys-recycle#workspaceMaintenanceSection",
        })

    items: list[str] = []
    for entry in entries:
        items.append(f"{entry['label']} {entry['count']} 条")
    if not items:
        items.append("当前没有待处理删除申请和回收站对象")

    return {
        "total_attention_count": total_attention_count,
        "pending_org_delete_count": pending_org_delete_count,
        "pending_sys_delete_count": pending_sys_delete_count,
        "org_recycle_count": org_recycle_count,
        "sys_recycle_count": sys_recycle_count,
        "entries": entries,
        "items": items,
    }


@app.get("/api/dashboard/drilldown")
def dashboard_drilldown(
    request: Request,
    dimension: str = Query(...),
    value: str = Query(""),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    current_user = get_current_user_optional(request, db)
    role = current_user.role if current_user else ""
    username = current_user.username if current_user else ""
    dimension = dimension.lower()
    if dimension not in {"region", "industry", "level"}:
        raise HTTPException(status_code=400, detail="dimension 仅支持 region/industry/level。")
    org_query = db.query(Organization).filter(Organization.deleted_at.is_(None))
    sys_query = db.query(SystemInfo).join(Organization, Organization.id == SystemInfo.organization_id).filter(
        SystemInfo.deleted_at.is_(None), Organization.deleted_at.is_(None)
    )
    if dimension == "region":
        org_query = org_query.filter(Organization.filing_region.like(f"%{escape_like(value)}%"))
        sys_query = sys_query.filter(Organization.filing_region.like(f"%{escape_like(value)}%"))
    elif dimension == "industry":
        org_query = org_query.filter(Organization.industry.like(f"%{escape_like(value)}%"))
        sys_query = sys_query.filter(Organization.industry.like(f"%{escape_like(value)}%"))
    else:
        try:
            level = int(value)
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"level 取值非法: {exc}") from exc
        sys_query = sys_query.filter(SystemInfo.proposed_level == level)
    if role == "evaluator":
        org_query = org_query.filter(Organization.created_by == username)
        sys_query = sys_query.filter(SystemInfo.created_by == username)

    org_items = org_query.order_by(Organization.created_at.desc()).limit(500).all()
    sys_items = sys_query.order_by(SystemInfo.created_at.desc()).limit(500).all()
    return {
        "dimension": dimension,
        "value": value,
        "organizations": [
            {
                "id": o.id,
                "name": o.name,
                "industry": o.industry,
                "filing_region": o.filing_region,
                "created_by": o.created_by,
            }
            for o in org_items
        ],
        "systems": [
            {
                "id": s.id,
                "system_name": s.system_name,
                "system_code": s.system_code,
                "proposed_level": s.proposed_level,
                "organization_id": s.organization_id,
            }
            for s in sys_items
        ],
    }


@app.get("/api/dashboard/trend")
def dashboard_trend(
    request: Request,
    months: int = Query(12, ge=1, le=60),
    industry: str | None = None,
    city: str | None = None,
    level: int | None = None,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    """按月统计备案数量变化趋势，返回最近N个月的组织和系统新增数量。"""
    current_user = get_current_user_optional(request, db)
    role = current_user.role if current_user else ""
    username = current_user.username if current_user else ""

    # 计算起始时间
    now = datetime.now()
    start_year = now.year
    start_month = now.month - (months - 1)
    while start_month <= 0:
        start_month += 12
        start_year -= 1
    start_dt = datetime(start_year, start_month, 1)

    org_year = extract("year", Organization.created_at)
    org_month = extract("month", Organization.created_at)
    org_q = db.query(
        org_year.label("year"),
        org_month.label("month"),
        func.count(Organization.id).label("count"),
    ).filter(
        Organization.deleted_at.is_(None),
        Organization.created_at >= start_dt,
    )
    if industry:
        org_q = org_q.filter(Organization.industry == industry)
    if city:
        org_q = org_q.filter(Organization.filing_region.like(f"%{escape_like(city)}%"))
    if role == "evaluator":
        org_q = org_q.filter(Organization.created_by == username)
    org_q = org_q.group_by(org_year, org_month).order_by(org_year, org_month)

    sys_year = extract("year", SystemInfo.created_at)
    sys_month = extract("month", SystemInfo.created_at)
    sys_q = db.query(
        sys_year.label("year"),
        sys_month.label("month"),
        func.count(SystemInfo.id).label("count"),
    ).join(Organization, Organization.id == SystemInfo.organization_id).filter(
        SystemInfo.deleted_at.is_(None),
        Organization.deleted_at.is_(None),
        SystemInfo.created_at >= start_dt,
    )
    if industry:
        sys_q = sys_q.filter(Organization.industry == industry)
    if city:
        sys_q = sys_q.filter(Organization.filing_region.like(f"%{escape_like(city)}%"))
    if level:
        sys_q = sys_q.filter(SystemInfo.proposed_level == level)
    if role == "evaluator":
        sys_q = sys_q.filter(SystemInfo.created_by == username)
    sys_q = sys_q.group_by(sys_year, sys_month).order_by(sys_year, sys_month)

    org_by_month: dict[str, int] = {}
    for row in org_q.all():
        label = f"{int(row.year):04d}-{int(row.month):02d}"
        org_by_month[label] = int(row.count)

    sys_by_month: dict[str, int] = {}
    for row in sys_q.all():
        label = f"{int(row.year):04d}-{int(row.month):02d}"
        sys_by_month[label] = int(row.count)

    # 生成连续月份列表
    result = []
    cur = start_dt
    for _ in range(months):
        label = cur.strftime("%Y-%m")
        result.append({
            "month": label,
            "organization_count": org_by_month.get(label, 0),
            "system_count": sys_by_month.get(label, 0),
        })
        # 移到下个月
        if cur.month == 12:
            cur = datetime(cur.year + 1, 1, 1)
        else:
            cur = datetime(cur.year, cur.month + 1, 1)

    return {"months": months, "trend": result}


@app.get("/api/dashboard/export/excel")
def export_dashboard_excel(
    request: Request,
    start_date: date | None = None,
    end_date: date | None = None,
    industry: str | None = None,
    city: str | None = None,
    level: int | None = None,
    evaluator: str | None = None,
    project_status: str | None = None,
    # auth injected below
    db: Session = Depends(get_db),
) -> FileResponse:
    require_roles(request, db, {"admin", "evaluator"}, legacy_admin=True)
    data = dashboard_summary(
        request=request,
        start_date=start_date,
        end_date=end_date,
        industry=industry,
        city=city,
        level=level,
        evaluator=evaluator,
        project_status=project_status,
        db=db,
    )
    wb = Workbook()
    ws = wb.active
    ws.title = "看板汇总"
    ws.append(["指标", "数值"])
    for k, v in data["totals"].items():
        ws.append([k, v])
    ws2 = wb.create_sheet("地区分布")
    ws2.append(["地区", "数量"])
    for row in data["region_distribution"]:
        ws2.append([row["name"], row["value"]])
    ws3 = wb.create_sheet("行业分布")
    ws3.append(["行业", "数量"])
    for row in data["industry_distribution"]:
        ws3.append([row["name"], row["value"]])
    ws4 = wb.create_sheet("级别分布")
    ws4.append(["级别", "数量"])
    for row in data["level_distribution"]:
        ws4.append([row["name"], row["value"]])
    path = EXPORT_DIR / f"dashboard_{datetime.now():%Y%m%d%H%M%S}.xlsx"
    wb.save(path)
    return FileResponse(path=str(path), filename=path.name, background=BackgroundTask(_cleanup_export_file, path))


@app.get("/api/dashboard/export/pdf")
def export_dashboard_pdf(
    request: Request,
    start_date: date | None = None,
    end_date: date | None = None,
    industry: str | None = None,
    city: str | None = None,
    level: int | None = None,
    evaluator: str | None = None,
    project_status: str | None = None,
    db: Session = Depends(get_db),
) -> FileResponse:
    require_roles(request, db, {"admin", "evaluator"}, legacy_admin=True)
    data = dashboard_summary(
        request=request,
        start_date=start_date,
        end_date=end_date,
        industry=industry,
        city=city,
        level=level,
        evaluator=evaluator,
        project_status=project_status,
        db=db,
    )
    if not HAS_REPORTLAB:
        raise HTTPException(status_code=503, detail="当前环境未安装 reportlab，暂不支持导出 PDF。")
    path = EXPORT_DIR / f"dashboard_{datetime.now():%Y%m%d%H%M%S}.pdf"
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    c = canvas.Canvas(str(path), pagesize=A4)
    c.setFont("STSong-Light", 15)
    c.drawString(36, 800, "数据看板导出")
    c.setFont("STSong-Light", 11)
    y = 770
    for k, v in data["totals"].items():
        c.drawString(36, y, f"{k}: {v}")
        y -= 18
    y -= 8
    c.drawString(36, y, "地区分布")
    y -= 18
    for row in data["region_distribution"][:15]:
        c.drawString(50, y, f"{row['name']}: {row['value']}")
        y -= 16
    if y < 120:
        c.showPage()
        c.setFont("STSong-Light", 11)
        y = 800
    y -= 8
    c.drawString(36, y, "行业分布")
    y -= 18
    for row in data["industry_distribution"][:15]:
        c.drawString(50, y, f"{row['name']}: {row['value']}")
        y -= 16
    if y < 120:
        c.showPage()
        c.setFont("STSong-Light", 11)
        y = 800
    y -= 8
    c.drawString(36, y, "级别分布")
    y -= 18
    for row in data["level_distribution"][:15]:
        c.drawString(50, y, f"{row['name']}: {row['value']}")
        y -= 16
    c.save()
    return FileResponse(path=str(path), filename=path.name, background=BackgroundTask(_cleanup_export_file, path))


@app.post("/api/knowledge/upload")
async def upload_knowledge(
    request: Request,
    title: str = Form(...),
    doc_type: str = Form(...),
    city: str = Form(""),
    district: str = Form(""),
    keywords: str = Form(""),
    protection_level: str = Form(""),
    version_no: int = Form(1),
    source_doc_id: int | None = Form(None),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor_name, _ = require_roles(request, db, {"admin"}, legacy_admin=True)
    assert_safe_text(title, "title")
    assert_safe_text(doc_type, "doc_type")
    assert_safe_text(city, "city")
    assert_safe_text(district, "district")
    assert_safe_text(keywords, "keywords")
    assert_safe_text(protection_level, "protection_level")
    content = await file.read()
    if len(content) > MAX_KNOWLEDGE_FILE:
        raise HTTPException(status_code=400, detail="文件过大。")
    safe_name = Path(file.filename).name
    save_name = f"{uuid.uuid4().hex}_{safe_name}"
    path = UPLOAD_DIR / "knowledge" / save_name
    path.parent.mkdir(parents=True, exist_ok=True)
    final_version = version_no
    if source_doc_id:
        source = db.query(KnowledgeDocument).filter(KnowledgeDocument.id == source_doc_id).first()
        if not source:
            raise HTTPException(status_code=404, detail="source_doc_id 对应文档不存在。")
        title = title or source.title
        doc_type = doc_type or source.doc_type
        city = city or source.city or ""
        district = district or source.district or ""
        keywords = keywords or source.keywords or ""
        protection_level = protection_level or source.protection_level or ""
        final_version = max(source.version_no + 1, version_no)
    path.write_bytes(content)
    row = KnowledgeDocument(
        title=title,
        keywords=keywords,
        city=city,
        district=district,
        doc_type=doc_type,
        protection_level=protection_level,
        version_no=final_version,
        status="enabled",
        file_name=safe_name,
        file_path=str(path),
        file_size=len(content),
        uploaded_by=actor_name,
    )
    db.add(row)
    db.flush()
    log_knowledge_version(db, row, "upload", actor_name)
    db.commit()
    db.refresh(row)
    return {"message": "上传成功", "data": {"id": row.id, "title": row.title}}


@app.get("/api/knowledge")
def list_knowledge(
    keyword: str | None = None,
    match_mode: str = Query("fuzzy"),
    city: str | None = None,
    district: str | None = None,
    doc_type: str | None = None,
    protection_level: str | None = None,
    start_date: date | None = None,
    end_date: date | None = None,
    enabled_only: bool = True,
    page: int | None = Query(None, ge=1),
    page_size: int | None = Query(None, ge=1, le=200),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    q = db.query(KnowledgeDocument)
    match_mode = match_mode.strip().lower()
    if match_mode not in {"fuzzy", "exact"}:
        raise HTTPException(status_code=400, detail="match_mode 仅支持 fuzzy/exact。")
    if enabled_only:
        q = q.filter(KnowledgeDocument.status == "enabled")
    if keyword:
        if match_mode == "exact":
            q = q.filter(or_(KnowledgeDocument.title == keyword, KnowledgeDocument.keywords == keyword))
        else:
            q = q.filter(
                or_(
                    KnowledgeDocument.title.like(f"%{escape_like(keyword)}%"),
                    KnowledgeDocument.keywords.like(f"%{escape_like(keyword)}%"),
                )
            )
    if city:
        q = q.filter(KnowledgeDocument.city.like(f"%{escape_like(city)}%"))
    if district:
        q = q.filter(KnowledgeDocument.district.like(f"%{escape_like(district)}%"))
    if doc_type:
        q = q.filter(KnowledgeDocument.doc_type.like(f"%{escape_like(doc_type)}%"))
    if protection_level:
        q = q.filter(KnowledgeDocument.protection_level.like(f"%{escape_like(protection_level)}%"))
    if start_date:
        q = q.filter(KnowledgeDocument.uploaded_at >= datetime.combine(start_date, datetime.min.time()))
    if end_date:
        q = q.filter(KnowledgeDocument.uploaded_at <= datetime.combine(end_date, datetime.max.time()))
    total = q.count()
    rows_query = (
        q.outerjoin(KnowledgePin, KnowledgePin.document_id == KnowledgeDocument.id)
        .add_columns(KnowledgePin.pinned_at.label("pinned_at"))
        .order_by(
            KnowledgePin.pinned_at.is_(None),
            KnowledgePin.pinned_at.desc(),
            KnowledgeDocument.uploaded_at.desc(),
        )
    )
    if page is not None or page_size is not None:
        effective_page = page or 1
        effective_page_size = page_size or 50
        rows_query = rows_query.offset((effective_page - 1) * effective_page_size).limit(effective_page_size)
    rows = rows_query.all()
    return {
        "total": total,
        "items": [
            {
                "id": doc.id,
                "title": doc.title,
                "doc_type": doc.doc_type,
                "city": doc.city,
                "district": doc.district,
                "protection_level": doc.protection_level,
                "version_no": doc.version_no,
                "status": doc.status,
                "file_name": doc.file_name,
                "uploaded_by": doc.uploaded_by,
                "uploaded_at": doc.uploaded_at,
                "pinned": pinned_at is not None,
                "pinned_at": pinned_at,
            }
            for doc, pinned_at in rows
        ],
    }


@app.get("/api/knowledge/download-logs")
def knowledge_download_logs(request: Request, db: Session = Depends(get_db)) -> dict[str, Any]:
    # 仅管理员可查看下载日志
    require_roles(request, db, {"admin"})
    rows = db.query(KnowledgeDownloadLog).order_by(KnowledgeDownloadLog.downloaded_at.desc()).limit(500).all()
    return {
        "total": len(rows),
        "items": [
            {
                "document_id": r.document_id,
                "download_by": r.download_by,
                "downloaded_at": r.downloaded_at,
            }
            for r in rows
        ],
    }


@app.post("/api/knowledge/batch-download")
def batch_download_knowledge(
    request: Request,
    doc_ids: list[int],
    db: Session = Depends(get_db),
) -> FileResponse:
    # 批量下载知识库文档（打包为zip）
    actor_name, _ = require_roles(request, db, {"admin", "reviewer", "evaluator"})
    if not doc_ids:
        raise HTTPException(status_code=400, detail="doc_ids 不能为空。")
    rows = db.query(KnowledgeDocument).filter(KnowledgeDocument.id.in_(doc_ids)).all()
    if not rows:
        raise HTTPException(status_code=404, detail="未找到文档。")
    disabled = [row.id for row in rows if row.status != "enabled"]
    if disabled:
        raise HTTPException(status_code=403, detail=f"包含已下架文档，禁止下载: {disabled}")
    zip_path = EXPORT_DIR / f"knowledge_batch_{datetime.now():%Y%m%d%H%M%S}.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for doc in rows:
            src = Path(doc.file_path)
            if src.exists():
                zf.write(src, arcname=doc.file_name)
                db.add(KnowledgeDownloadLog(document_id=doc.id, download_by=actor_name))
    db.commit()
    return FileResponse(path=str(zip_path), filename=zip_path.name, background=BackgroundTask(_cleanup_export_file, zip_path))


@app.put("/api/knowledge/{doc_id}")
def update_knowledge(
    request: Request,
    doc_id: int,
    payload: dict[str, Any],
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor_name, _ = require_roles(request, db, {"admin"}, legacy_admin=True)
    assert_safe_payload(payload, "payload")
    doc = db.query(KnowledgeDocument).filter(KnowledgeDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在。")
    log_knowledge_version(db, doc, "update_before", actor_name)
    fields = {
        "title",
        "keywords",
        "city",
        "district",
        "doc_type",
        "protection_level",
        "status",
        "version_no",
    }
    for k, v in payload.items():
        if k in fields:
            if k in {"title", "doc_type", "status", "version_no"} and v is None:
                raise HTTPException(status_code=400, detail=f"{k} 不能为空。")
            setattr(doc, k, v)
    db.commit()
    return {"message": "更新成功"}


@app.post("/api/knowledge/{doc_id}/new-version")
async def new_knowledge_version(
    request: Request,
    doc_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor_name, _ = require_roles(request, db, {"admin"}, legacy_admin=True)
    doc = db.query(KnowledgeDocument).filter(KnowledgeDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在。")
    content = await file.read()
    if len(content) > MAX_KNOWLEDGE_FILE:
        raise HTTPException(status_code=400, detail="文件过大。")
    log_knowledge_version(db, doc, "new_version_before", actor_name)
    safe_name = Path(file.filename).name
    save_name = f"{uuid.uuid4().hex}_{safe_name}"
    path = UPLOAD_DIR / "knowledge" / save_name
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(content)
    doc.file_name = safe_name
    doc.file_path = str(path)
    doc.file_size = len(content)
    doc.version_no = int(doc.version_no or 0) + 1
    doc.uploaded_by = actor_name
    db.commit()
    return {"message": "新版本上传成功", "version_no": doc.version_no}


@app.get("/api/knowledge/{doc_id}/versions")
def knowledge_versions(doc_id: int, db: Session = Depends(get_db)) -> dict[str, Any]:
    _ = db.query(KnowledgeDocument).filter(KnowledgeDocument.id == doc_id).first()
    if not _:
        raise HTTPException(status_code=404, detail="文档不存在。")
    rows = (
        db.query(KnowledgeDocumentVersion)
        .filter(KnowledgeDocumentVersion.document_id == doc_id)
        .order_by(KnowledgeDocumentVersion.changed_at.desc())
        .all()
    )
    return {
        "total": len(rows),
        "items": [
            {
                "id": r.id,
                "action": r.action,
                "changed_by": r.changed_by,
                "changed_at": r.changed_at,
                "snapshot": r.snapshot,
            }
            for r in rows
        ],
    }


@app.post("/api/knowledge/{doc_id}/rollback/{version_id}")
def rollback_knowledge(
    request: Request,
    doc_id: int,
    version_id: int,
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor_name, _ = require_roles(request, db, {"admin"}, legacy_admin=True)
    doc = db.query(KnowledgeDocument).filter(KnowledgeDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在。")
    version = (
        db.query(KnowledgeDocumentVersion)
        .filter(KnowledgeDocumentVersion.id == version_id, KnowledgeDocumentVersion.document_id == doc_id)
        .first()
    )
    if not version:
        raise HTTPException(status_code=404, detail="版本记录不存在。")
    snap = dict(version.snapshot or {})
    if not snap:
        raise HTTPException(status_code=400, detail="版本快照为空。")
    target_file = ensure_file_exists(snap.get("file_path"), "版本文件不存在，无法回滚。")
    log_knowledge_version(db, doc, "rollback_before", actor_name)
    for key in ["title", "keywords", "city", "district", "doc_type", "protection_level", "status", "file_name", "file_path", "file_size"]:
        if key in snap:
            setattr(doc, key, snap[key])
    doc.version_no = int(doc.version_no or 0) + 1
    doc.uploaded_by = actor_name
    db.commit()
    return {"message": "回滚成功", "version_no": doc.version_no}


@app.post("/api/knowledge/{doc_id}/pin")
def pin_knowledge(
    request: Request,
    doc_id: int,
    enabled: bool = Query(True),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor_name, _ = require_roles(request, db, {"admin"}, legacy_admin=True)
    doc = db.query(KnowledgeDocument).filter(KnowledgeDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在。")
    row = db.query(KnowledgePin).filter(KnowledgePin.document_id == doc_id).first()
    if enabled and not row:
        row = KnowledgePin(document_id=doc_id, pinned_by=actor_name)
        db.add(row)
    if not enabled and row:
        db.delete(row)
    db.commit()
    return {"message": "置顶状态已更新", "pinned": enabled}


@app.post("/api/knowledge/{doc_id}/toggle")
def toggle_knowledge(
    request: Request,
    doc_id: int,
    enabled: bool = Query(True),
    db: Session = Depends(get_db),
) -> dict[str, Any]:
    actor_name, _ = require_roles(request, db, {"admin"}, legacy_admin=True)
    doc = db.query(KnowledgeDocument).filter(KnowledgeDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在。")
    log_knowledge_version(db, doc, "toggle_before", actor_name)
    doc.status = "enabled" if enabled else "disabled"
    db.commit()
    return {"message": "状态更新成功", "status": doc.status}


@app.get("/api/knowledge/{doc_id}/download")
def download_knowledge(request: Request, doc_id: int, db: Session = Depends(get_db)) -> FileResponse:
    actor_name, _ = require_roles(request, db, {"admin", "reviewer", "evaluator"})
    doc = db.query(KnowledgeDocument).filter(KnowledgeDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在。")
    if doc.status != "enabled":
        raise HTTPException(status_code=403, detail="文档已下架。")
    path = ensure_file_exists(doc.file_path, "文档文件不存在或已被移除。")
    db.add(KnowledgeDownloadLog(document_id=doc.id, download_by=actor_name))
    db.commit()
    return FileResponse(path=str(path), filename=doc.file_name)


@app.delete("/api/knowledge/{doc_id}")
def delete_knowledge(request: Request, doc_id: int, db: Session = Depends(get_db)) -> dict[str, Any]:
    actor_name, _ = require_roles(request, db, {"admin"}, legacy_admin=True)
    doc = db.query(KnowledgeDocument).filter(KnowledgeDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在。")
    log_knowledge_version(db, doc, "delete_before", actor_name)
    safe_delete_uploaded_file(doc.file_path)
    pin = db.query(KnowledgePin).filter(KnowledgePin.document_id == doc_id).first()
    if pin:
        db.delete(pin)
    db.delete(doc)
    db.commit()
    return {"message": "文档已删除。"}
