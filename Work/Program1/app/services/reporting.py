import io
import json
import logging
import re
import zipfile
from datetime import datetime, date
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape as xml_escape

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

try:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfgen import canvas
    from reportlab.lib.pdfencrypt import StandardEncryption

    HAS_REPORTLAB = True
except Exception:
    HAS_REPORTLAB = False

logger = logging.getLogger(__name__)


# ──────────────────────────────────────────────────────────────────────────────
# 数据完整性检查
# ──────────────────────────────────────────────────────────────────────────────

class ReportDataError(ValueError):
    """报告数据不完整或不合法时抛出的异常。"""


def _require_attr(obj: Any, attr: str, label: str) -> str:
    """从对象中取属性值；若为空则抛出 ReportDataError。"""
    value = getattr(obj, attr, None)
    if value is None or (isinstance(value, str) and not value.strip()):
        raise ReportDataError(f"报告生成失败：{label}（{attr}）为空，请先补全数据。")
    return str(value).strip()


def check_report_data_integrity(org: Any, system: Any) -> None:
    """
    检查生成报告所需的关键字段完整性。

    必填字段：
    - 单位：name、credit_code、filing_region、industry、legal_representative
    - 系统：system_name、system_code、proposed_level
    """
    _require_attr(org, "name", "单位名称")
    _require_attr(org, "credit_code", "统一社会信用代码")
    _require_attr(org, "filing_region", "备案地区")
    _require_attr(org, "industry", "所属行业")
    _require_attr(org, "legal_representative", "法定代表人")
    _require_attr(system, "system_name", "系统名称")
    _require_attr(system, "system_code", "系统编号")

    level = getattr(system, "proposed_level", None)
    if level is None:
        raise ReportDataError("报告生成失败：拟定等级（proposed_level）为空，请先补全数据。")
    try:
        level_int = int(level)
    except (TypeError, ValueError):
        raise ReportDataError(f"报告生成失败：拟定等级值 '{level}' 不合法，应为 1-5 的整数。")
    if not (1 <= level_int <= 5):
        raise ReportDataError(f"报告生成失败：拟定等级值 {level_int} 超出范围（1-5）。")


# ──────────────────────────────────────────────────────────────────────────────
# 报告 payload 生成
# ──────────────────────────────────────────────────────────────────────────────

def generate_report_payload(report_type: str, org: Any, system: Any) -> dict[str, Any]:
    """
    根据报告类型生成结构化 payload 字典。

    在生成前会执行数据完整性检查，字段缺失时抛出 ReportDataError。
    """
    check_report_data_integrity(org, system)

    base = {
        "单位名称": org.name,
        "统一社会信用代码": org.credit_code,
        "系统名称": system.system_name,
        "系统编号": system.system_code,
        "拟定等级": f"{system.proposed_level}级",
        "备案地区": org.filing_region,
        "所属行业": org.industry,
        "部署方式": system.deployment_mode or "",
        "定级依据": system.level_basis or "依据《网络安全等级保护定级指南》及业务影响范围综合判定。",
        "定级理由": system.level_reason or "根据业务重要性、服务对象规模、数据敏感程度综合确定。",
    }
    if report_type == "filing_form":
        return {
            "标题": "网络安全等级保护定级备案表",
            "章节": [
                {"名称": "基础信息", "内容": base},
                {
                    "名称": "备案意见",
                    "内容": {"测评师意见": "", "主管审核意见": ""},
                },
            ],
        }
    if report_type == "expert_review_form":
        return {
            "标题": "专家评审意见表",
            "章节": [
                {"名称": "系统信息", "内容": base},
                {
                    "名称": "专家意见",
                    "内容": {"专家结论": "", "签字": "", "日期": ""},
                },
            ],
        }
    return {
        "标题": "网络安全等级保护定级报告",
        "章节": [
            {"名称": "责任主体", "内容": {"单位": org.name, "负责人": org.legal_representative}},
            {"名称": "定级对象构成", "内容": {"系统": system.system_name, "边界": system.boundary or ""}},
            {
                "名称": "承载业务与数据",
                "内容": {
                    "业务描述": system.business_description or "",
                    "数据类别": system.data_category or "",
                    "数据级别": system.data_level or "",
                },
            },
            {"名称": "安全责任", "内容": {"网络安全责任部门": org.cybersecurity_dept or ""}},
            {
                "名称": "等级确定过程",
                "内容": {
                    "影响范围": system.impact_scope or "",
                    "定级依据": base["定级依据"],
                    "定级理由": base["定级理由"],
                    "结论": f"建议定级为第{system.proposed_level}级",
                },
            },
        ],
    }


# ──────────────────────────────────────────────────────────────────────────────
# 内部工具函数
# ──────────────────────────────────────────────────────────────────────────────

def _iter_report_lines(content: dict[str, Any]) -> list[str]:
    lines: list[str] = [str(content.get("标题", ""))]
    for section in content.get("章节", []):
        title = section.get("名称", "")
        if title:
            lines.append(f"[{title}]")
        section_content = section.get("内容", {})
        if isinstance(section_content, dict):
            for k, v in section_content.items():
                lines.append(f"{k}: {v}")
        else:
            lines.append(str(section_content))
        lines.append("")
    return lines


def _replace_paragraph_text(paragraph: Any, value: str) -> None:
    """将整段文字替换为 value，同时保留第一个 run 的格式。"""
    if not paragraph.runs:
        paragraph.text = value
        return
    paragraph.runs[0].text = value
    for run in paragraph.runs[1:]:
        run.text = ""


# ──────────────────────────────────────────────────────────────────────────────
# Word 导出（无模板）
# ──────────────────────────────────────────────────────────────────────────────

def export_report_docx(title: str, content: dict[str, Any], output_path: Path) -> None:
    """
    生成不依赖官方模板的 Word 报告文件。

    参数：
        title: 报告标题
        content: 报告内容 dict，包含 '章节' 列表
        output_path: 输出文件路径（Path 对象）
    """
    if not title or not title.strip():
        raise ValueError("export_report_docx: title 不能为空。")
    if not isinstance(content, dict):
        raise TypeError("export_report_docx: content 必须为 dict 类型。")

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # 设置正文字体为宋体 12pt
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "宋体"
    try:
        normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    except AttributeError:
        # 部分环境下 rPr 可能为 None，容错处理
        logger.warning("无法设置正文东亚字体，跳过。")
    normal_style.font.size = Pt(12)

    # 添加页脚页码："第 {PAGE} 页"
    footer = doc.sections[0].footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = 2  # 右对齐
    p.add_run("第 ")
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    p.add_run(" 页")

    doc.add_heading(title, level=1)
    for section in content.get("章节", []):
        section_title = section.get("名称", "")
        if section_title:
            doc.add_heading(section_title, level=2)
        section_content = section.get("内容", {})
        if isinstance(section_content, dict):
            for key, value in section_content.items():
                para = doc.add_paragraph(f"{key}: {value}")
                para.paragraph_format.line_spacing = 1.5
        else:
            para = doc.add_paragraph(str(section_content))
            para.paragraph_format.line_spacing = 1.5

    doc.save(output_path)
    logger.info("Word 报告已生成：%s", output_path)


# ──────────────────────────────────────────────────────────────────────────────
# 官方模板脱敏处理
# ──────────────────────────────────────────────────────────────────────────────

def _load_template_replacements() -> dict[str, str]:
    """从外部配置文件加载模板脱敏替换规则。"""
    config_path = Path(__file__).resolve().parents[2] / "config" / "template_replacements.json"
    if config_path.exists():
        try:
            with open(config_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, dict):
                return {str(k): str(v) for k, v in data.items()}
        except (json.JSONDecodeError, OSError):
            logger.warning("模板替换配置文件解析失败：%s", config_path)
    else:
        logger.warning("模板替换配置文件不存在：%s，脱敏替换将不生效", config_path)
    return {}


OFFICIAL_TEMPLATE_TEXT_REPLACEMENTS = _load_template_replacements()


def _normalize_text_for_replace(text: str) -> str:
    result = text
    for old, new in OFFICIAL_TEMPLATE_TEXT_REPLACEMENTS.items():
        result = result.replace(old, new)
    result = re.sub(r"\b[0-9A-Z]{18}\b", "【统一社会信用代码】", result)
    result = re.sub(r"\b1[3-9]\d{9}\b", "【联系电话】", result)
    result = re.sub(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", "【邮箱】", result)
    result = re.sub(r"\b\d{17}[0-9Xx]\b", "【证件号】", result)
    result = re.sub(r"\b\d{4}年\d{1,2}月\d{1,2}日\b", "____年__月__日", result)
    result = re.sub(r"https?://[^\s，。；;]+", "【访问地址】", result)
    return result


def sanitize_template_docx_content(content: bytes) -> bytes:
    """对 Word 模板文件进行脱敏处理，替换敏感文本为占位符。"""
    if not content:
        raise ValueError("sanitize_template_docx_content: content 不能为空。")

    doc = Document(io.BytesIO(content))
    for p in doc.paragraphs:
        raw = p.text or ""
        cleaned = _normalize_text_for_replace(raw)
        if cleaned != raw:
            _replace_paragraph_text(p, cleaned)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    raw = p.text or ""
                    cleaned = _normalize_text_for_replace(raw)
                    if cleaned != raw:
                        _replace_paragraph_text(p, cleaned)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


# ──────────────────────────────────────────────────────────────────────────────
# Word 导出（基于官方模板填充）
# ──────────────────────────────────────────────────────────────────────────────

def _replace_tokens(text: str, field_map: dict[str, str]) -> str:
    """将文本中的 {{key}} 和 【key】 占位符替换为对应值。"""
    result = text
    for k, v in field_map.items():
        if not k:
            continue
        result = result.replace(f"{{{{{k}}}}}", v)
        result = result.replace(f"【{k}】", v)
    return result


def _fill_row_by_label(row: Any, field_map: dict[str, str]) -> None:
    """
    按表格行中的标签列填充相邻值列。

    规则：找到包含 field_map key 的单元格，将其右侧相邻单元格设置为对应值。
    """
    cells = list(row.cells)
    labels = [c.text.strip().replace("\n", " ") for c in cells]
    for idx, label in enumerate(labels):
        if not label:
            continue
        for key, value in field_map.items():
            if not value:
                continue
            if key in label and idx + 1 < len(cells):
                cells[idx + 1].text = value


def export_report_docx_with_template(
    template_path: Path,
    field_map: dict[str, Any],
    output_path: Path,
) -> None:
    """
    基于 Word 模板填充字段并导出报告。

    参数：
        template_path: 模板文件路径（.docx）
        field_map: 字段映射 {占位符名称: 填充值}
        output_path: 输出文件路径

    说明：
        支持 {{key}} 和 【key】 两种占位符格式，同时支持按标签列填充表格。
        None 值的字段不参与替换，避免写入字面量 "None"。
    """
    template_path = Path(template_path)
    if not template_path.exists():
        raise FileNotFoundError(f"模板文件不存在：{template_path}")
    if not template_path.suffix.lower() == ".docx":
        raise ValueError(f"模板文件必须为 .docx 格式，当前：{template_path.suffix}")

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(template_path))

    # 过滤掉 None 值，防止写入字符串 "None"
    normalized_map: dict[str, str] = {
        str(k).strip(): str(v)
        for k, v in (field_map or {}).items()
        if str(k).strip() and v is not None
    }

    # 替换正文段落
    for p in doc.paragraphs:
        raw = p.text or ""
        replaced = _replace_tokens(raw, normalized_map)
        if replaced != raw:
            _replace_paragraph_text(p, replaced)

    # 替换表格内容（先按标签填充，再做 token 替换）
    for table in doc.tables:
        for row in table.rows:
            _fill_row_by_label(row, normalized_map)
            for cell in row.cells:
                for p in cell.paragraphs:
                    raw = p.text or ""
                    replaced = _replace_tokens(raw, normalized_map)
                    if replaced != raw:
                        _replace_paragraph_text(p, replaced)

    doc.save(output_path)
    logger.info("模板 Word 报告已生成：%s", output_path)


# ──────────────────────────────────────────────────────────────────────────────
# PDF 导出
# ──────────────────────────────────────────────────────────────────────────────

def export_report_pdf(title: str, content: dict[str, Any], output_path: Path, password: str | None = None) -> None:
    """
    生成 PDF 报告文件。

    优先使用 reportlab（支持中文和加密）；若 reportlab 不可用，
    则降级输出仅含 ASCII 内容的最小合规 PDF。

    参数：
        title: 报告标题
        content: 报告内容 dict
        output_path: 输出文件路径
        password: 可选 PDF 加密密码（仅 reportlab 可用时生效）
    """
    if not title or not title.strip():
        raise ValueError("export_report_pdf: title 不能为空。")
    if not isinstance(content, dict):
        raise TypeError("export_report_pdf: content 必须为 dict 类型。")

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if HAS_REPORTLAB:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        encrypt = None
        if password:
            encrypt = StandardEncryption(
                userPassword=password,
                ownerPassword=password,
                canPrint=1,
                canModify=0,
                canCopy=0,
            )
        c = canvas.Canvas(str(output_path), encrypt=encrypt)
        page_no = 1

        def draw_page_footer() -> None:
            c.setFont("STSong-Light", 9)
            c.drawString(520, 20, f"第{page_no}页")

        c.setFont("STSong-Light", 14)
        y = 800
        c.drawString(60, y, title)
        y -= 30
        c.setFont("STSong-Light", 11)
        for line in _iter_report_lines(content):
            if y < 50:
                draw_page_footer()
                c.showPage()
                page_no += 1
                c.setFont("STSong-Light", 11)
                y = 800
            # 截断过长行，防止单行溢出页面
            c.drawString(60, y, line[:120])
            y -= 18
        draw_page_footer()
        c.save()
        logger.info("PDF 报告已生成（reportlab）：%s", output_path)
        return

    # reportlab 不可用时的降级输出（仅 ASCII 内容）
    logger.warning("reportlab 不可用，降级生成最小 PDF，中文内容将被忽略。")
    lines = [title] + _iter_report_lines(content)
    ascii_lines = [line.encode("ascii", errors="ignore").decode("ascii") or " " for line in lines[:35]]
    stream_parts = ["BT", "/F1 12 Tf", "50 800 Td"]
    for idx, line in enumerate(ascii_lines):
        escaped = line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        stream_parts.append(f"({escaped}) Tj" if idx == 0 else f"T* ({escaped}) Tj")
    stream_parts.append("ET")
    stream = "\n".join(stream_parts).encode("latin-1", errors="ignore")

    objects = [
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n",
        b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n",
        b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >> endobj\n",
        b"4 0 obj << /Length " + str(len(stream)).encode("ascii") + b" >> stream\n" + stream + b"\nendstream endobj\n",
        b"5 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n",
    ]
    header = b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n"
    offsets: list[int] = []
    body = b""
    current = len(header)
    for obj in objects:
        offsets.append(current)
        body += obj
        current += len(obj)
    xref_pos = len(header) + len(body)
    xref = [b"xref\n0 6\n0000000000 65535 f \n"]
    for off in offsets:
        xref.append(f"{off:010d} 00000 n \n".encode("ascii"))
    trailer = b"trailer << /Root 1 0 R /Size 6 >>\nstartxref\n" + str(xref_pos).encode("ascii") + b"\n%%EOF\n"
    output_path.write_bytes(header + body + b"".join(xref) + trailer)
    logger.info("PDF 报告已生成（降级模式）：%s", output_path)


# ──────────────────────────────────────────────────────────────────────────────
# JSON 工具
# ──────────────────────────────────────────────────────────────────────────────

def dump_json(data: dict[str, Any]) -> str:
    return json.dumps(data, ensure_ascii=False, indent=2)


# ──────────────────────────────────────────────────────────────────────────────
# XML 级 Word 模板填充（避免 python-docx 丢失方框等符号）
# ──────────────────────────────────────────────────────────────────────────────

CN_LEVEL_MAP = {1: "一", 2: "二", 3: "三", 4: "四", 5: "五"}

SAMPLE_ORG_NAME = "山西晋深交易有限公司"
SAMPLE_SYSTEM_NAME = "山西省省属企业采购与供应链信息管理系统"
SAMPLE_OWNER_NAME_GRADING = "郭丽娟"
SAMPLE_OWNER_ID = "140521198912207028"
SAMPLE_OWNER_NAME_REVIEW = "张宇阳"
SAMPLE_OWNER_PHONE = "15034099805"
SAMPLE_EXPERTS = "孔庆花、马丽娜、郭煜"
SAMPLE_REVIEW_DATE = "2026年1月21日"
SAMPLE_GRADING_DATE_LINE = "2026年 1 月 20 日"
SAMPLE_BUILDER = "金蝶软件（中国）有限公司"


def _cn_level(level: Any) -> str:
    try:
        n = int(level)
    except (TypeError, ValueError):
        return "三"
    return CN_LEVEL_MAP.get(n, "三")


def _today_cn() -> str:
    today = datetime.now()
    return f"{today.year}年 {today.month} 月 {today.day} 日"


def _today_short() -> str:
    today = datetime.now()
    return f"{today.year}年{today.month}月{today.day}日"


def _read_docx_xml(template_path: Path) -> tuple[str, list[tuple[str, bytes]]]:
    """读取 .docx，返回 (document.xml 文本, 其余文件列表)。"""
    with zipfile.ZipFile(template_path, "r") as zin:
        xml_content = zin.read("word/document.xml").decode("utf-8")
        others = [(n, zin.read(n)) for n in zin.namelist() if n != "word/document.xml"]
    return xml_content, others


def _write_docx_xml(output_path: Path, xml_content: str, others: list[tuple[str, bytes]]) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zout:
        zout.writestr("word/document.xml", xml_content.encode("utf-8"))
        for name, data in others:
            zout.writestr(name, data)


def _xml_replace(xml: str, replacements: dict[str, str]) -> str:
    """XML 文本替换，自动 xml-escape 键值。None 值会被当作空字符串。"""
    for old, new in replacements.items():
        if not old:
            continue
        xml = xml.replace(xml_escape(old), xml_escape("" if new is None else str(new)))
    return xml


def _replace_last_n(xml: str, needle: str, replacement: str, n: int) -> str:
    """将 xml 中最后 n 个 needle（xml 转义后）替换为 replacement。"""
    escaped_needle = xml_escape(needle)
    escaped_repl = xml_escape(replacement)
    positions = [m.start() for m in re.finditer(re.escape(escaped_needle), xml)]
    if not positions:
        return xml
    for pos in reversed(positions[-n:]):
        xml = xml[:pos] + escaped_repl + xml[pos + len(escaped_needle):]
    return xml


def export_grading_report_docx(
    template_path: Path,
    org: Any,
    system: Any,
    output_path: Path,
    content: dict[str, Any] | None = None,
    topology_path: Path | None = None,
    business_matrix: list[list[int]] | None = None,
    service_matrix: list[list[int]] | None = None,
) -> None:
    """基于 02-定级报告.docx 模板生成定级报告。

    - 优先使用用户填写的 content（各节段落文本）覆盖样例；未填字段保留模板样例。
    - topology_path 指向拓扑图，若提供则插入到"定级对象构成"段落末尾。
    - business_matrix / service_matrix 为 3x3 命中矩阵，对对应单元格底色涂灰。
    """
    template_path = Path(template_path)
    if not template_path.exists():
        raise FileNotFoundError(f"定级报告模板不存在：{template_path}")

    cn = _cn_level(getattr(system, "proposed_level", 3))
    org_name = getattr(org, "name", "") or SAMPLE_ORG_NAME
    sys_name = getattr(system, "system_name", "") or SAMPLE_SYSTEM_NAME
    content = content or {}

    xml, others = _read_docx_xml(template_path)

    # 基础替换：单位/系统名
    xml = _xml_replace(xml, {SAMPLE_ORG_NAME: org_name, SAMPLE_SYSTEM_NAME: sys_name})

    # 用户填写的段落内容替换整段（定位到原样例段落；跨 run 合并后匹配）
    paragraph_map = {
        "responsible_subject": "定级对象于2025年12月",
        "object_composition": "部署于山西省政务云",
        "carried_business": "该系统包含信息采集",
        "carried_data": "该定级对象承载的数据主要包括",
        "security_responsibility": "按照网络安全法的相关要求",
        "business_info_description": "该定级对象承载的数据主要为采购",
        "business_info_damage_object": "该业务信息遭到破坏后",
        "business_info_damage_degree": "侵害的客观方面表现为",
        "service_description": "该系统主要负责省属企业",
        "service_damage_object": "该系统服务遭到破坏后",
        "service_damage_degree": "对客体的侵害结果为",
    }
    for key, match_prefix in paragraph_map.items():
        new_text = str(content.get(key, "") or "").strip()
        if new_text:
            xml = _replace_paragraph_starts_with(xml, match_prefix, new_text)

    # 填表日期（签章行） - 跨 run 合并匹配
    fill_date = str(content.get("fill_date") or "").strip() or _today_cn()
    xml = _replace_paragraph_starts_with(xml, "2026年 1 月 20 日", fill_date)

    # 结论句（这些在单段内连续，用 _xml_replace 即可）
    cn_label = f"第{cn}级"
    xml = _xml_replace(
        xml,
        {
            "业务信息安全保护等级为第三级": f"业务信息安全保护等级为{cn_label}",
            "系统服务安全保护等级为第三级": f"系统服务安全保护等级为{cn_label}",
            "网络安全保护等级为第三级": f"网络安全保护等级为{cn_label}",
        },
    )

    # 汇总表最后 3 个"第三级"单元格
    if cn != "三":
        xml = _replace_last_n(xml, "第三级", cn_label, 3)

    # 矩阵涂灰（业务信息矩阵 = 第一张表；系统服务矩阵 = 第二张表）
    if business_matrix or service_matrix:
        xml = _shade_matrix_cells(xml, business_matrix, service_matrix)

    # 拓扑图插入：在"定级对象构成"段落之后
    if topology_path and Path(topology_path).exists():
        xml, others = _insert_topology_image(xml, others, Path(topology_path))

    _write_docx_xml(output_path, xml, others)
    logger.info("定级报告已生成：%s", output_path)


def _paragraph_plain_text(p_seg: str) -> str:
    """拼接 <w:p> 内所有 <w:t> 的文本（不解转义），用于跨 run 匹配。"""
    parts = re.findall(r"<w:t\b[^>]*>(.*?)</w:t>", p_seg, re.DOTALL)
    return "".join(parts)


def _replace_paragraph_starts_with(xml: str, match_substring: str, new_text: str) -> str:
    """定位文本（跨 run 合并后）包含 match_substring 的 <w:p>，保留首 run 属性并替换所有 w:t 内容为 new_text。"""
    escaped_sub = xml_escape(match_substring)
    p_pattern = re.compile(r"<w:p\b[^>]*>.*?</w:p>", re.DOTALL)
    done = [False]

    def repl(m: re.Match) -> str:
        if done[0]:
            return m.group(0)
        seg = m.group(0)
        plain = _paragraph_plain_text(seg)
        if escaped_sub not in plain:
            return seg
        done[0] = True
        first_run_match = re.search(r"<w:r\b[^>]*>(.*?)</w:r>", seg, re.DOTALL)
        rpr = ""
        if first_run_match:
            rpr_match = re.search(r"<w:rPr\b[^>]*>.*?</w:rPr>", first_run_match.group(1), re.DOTALL)
            if rpr_match:
                rpr = rpr_match.group(0)
        ppr_match = re.search(r"<w:pPr\b[^>]*>.*?</w:pPr>", seg, re.DOTALL)
        ppr = ppr_match.group(0) if ppr_match else ""
        new_run = f"<w:r>{rpr}<w:t xml:space=\"preserve\">{xml_escape(new_text)}</w:t></w:r>"
        open_match = re.match(r"<w:p\b[^>]*>", seg)
        open_tag = open_match.group(0) if open_match else "<w:p>"
        return f"{open_tag}{ppr}{new_run}</w:p>"

    return p_pattern.sub(repl, xml)


def _replace_paragraph_contains(xml: str, match_substring: str, new_text: str) -> str:
    """同 _replace_paragraph_starts_with 语义（保留以兼容之前命名）。"""
    return _replace_paragraph_starts_with(xml, match_substring, new_text)


def _remove_sample_numbered_paragraphs(xml: str, max_remove: int = 5) -> str:
    """移除紧随 '形成如下评审意见' 段落之后、到 '评审专家组组长' 段落之前的所有内容段。"""
    pattern = re.compile(r"<w:p\b[^>]*>.*?</w:p>", re.DOTALL)
    matches = list(pattern.finditer(xml))
    opinion_end = -1
    stop_start = -1
    for i, m in enumerate(matches):
        plain = _paragraph_plain_text(m.group(0))
        if opinion_end < 0 and "形成如下评审意见" in plain:
            opinion_end = m.end()
            continue
        if opinion_end > 0 and ("评审专家组组长" in plain or "评审专家组成员" in plain):
            stop_start = m.start()
            break
    if opinion_end < 0 or stop_start < 0 or stop_start <= opinion_end:
        return xml
    return xml[:opinion_end] + xml[stop_start:]


def _shade_matrix_cells(
    xml: str,
    business_matrix: list[list[int]] | None,
    service_matrix: list[list[int]] | None,
) -> str:
    """在 02 定级报告的两张 3x3 矩阵表中，根据命中矩阵给对应单元格加灰底。

    矩阵表结构（前两张表）均为 5 行 4 列：
      row 0~1: 表头
      row 2~4: 数据行（对应行索引 0/1/2）
      col 1~3: 数据列（对应列索引 0/1/2）
    """
    shade_xml = '<w:shd w:val="clear" w:color="auto" w:fill="BFBFBF"/>'

    tbls = _find_top_level_tables(xml)
    if len(tbls) < 2:
        return xml

    def apply_shade(table_seg: str, matrix: list[list[int]] | None) -> str:
        if not matrix:
            return table_seg
        rows = _split_rows(table_seg)
        if len(rows) < 5:
            return table_seg
        for ri in range(3):
            for ci in range(3):
                if not matrix[ri][ci]:
                    continue
                row_idx = 2 + ri
                col_idx = 1 + ci
                rows[row_idx] = _shade_row_cell(rows[row_idx], col_idx, shade_xml)
        return _assemble_table(table_seg, rows)

    # 第一张矩阵 = 业务信息
    start, end = tbls[0]
    new_seg = apply_shade(xml[start:end], business_matrix)
    xml = xml[:start] + new_seg + xml[end:]
    # 第二张矩阵 = 系统服务（重新定位）
    tbls2 = _find_top_level_tables(xml)
    if len(tbls2) >= 2:
        start, end = tbls2[1]
        new_seg = apply_shade(xml[start:end], service_matrix)
        xml = xml[:start] + new_seg + xml[end:]
    return xml


def _find_top_level_tables(xml: str) -> list[tuple[int, int]]:
    """返回所有顶层 <w:tbl>...</w:tbl> 的 (start, end) 偏移列表。"""
    result = []
    idx = 0
    while True:
        start = xml.find("<w:tbl>", idx)
        if start == -1:
            start2 = re.search(r"<w:tbl[\s>]", xml[idx:])
            if not start2:
                break
            start = idx + start2.start()
        # 找到匹配的 </w:tbl>
        depth = 1
        scan = start + len("<w:tbl")
        end = -1
        while depth > 0:
            next_open = xml.find("<w:tbl", scan)
            next_close = xml.find("</w:tbl>", scan)
            if next_close == -1:
                return result
            # 只计算 <w:tbl 后紧跟空白或 > 的（避免匹配到 <w:tblPr 等）
            if next_open != -1 and next_open < next_close:
                ch = xml[next_open + len("<w:tbl"):next_open + len("<w:tbl") + 1]
                if ch in (" ", ">", "\n", "\t", "\r"):
                    depth += 1
                scan = next_open + len("<w:tbl")
                continue
            depth -= 1
            end = next_close + len("</w:tbl>")
            scan = end
        result.append((start, end))
        idx = end
    return result


def _split_rows(table_seg: str) -> list[str]:
    rows: list[str] = []
    idx = 0
    while True:
        start = table_seg.find("<w:tr", idx)
        if start == -1:
            break
        close = table_seg.find("</w:tr>", start)
        if close == -1:
            break
        end = close + len("</w:tr>")
        rows.append(table_seg[start:end])
        idx = end
    return rows


def _assemble_table(orig_table_seg: str, rows: list[str]) -> str:
    """用新 rows 替换表格中所有 <w:tr>..</w:tr> 片段，保留 <w:tblPr>/<w:tblGrid>。"""
    # 定位第一个 <w:tr 和最后一个 </w:tr>
    first_tr = orig_table_seg.find("<w:tr")
    last_close = orig_table_seg.rfind("</w:tr>") + len("</w:tr>")
    if first_tr == -1 or last_close == -1:
        return orig_table_seg
    return orig_table_seg[:first_tr] + "".join(rows) + orig_table_seg[last_close:]


def _shade_row_cell(row_xml: str, col_idx: int, shade_xml: str) -> str:
    """给 row_xml 中第 col_idx 个 <w:tc> 加 shade。"""
    cells = []
    idx = 0
    while True:
        start = row_xml.find("<w:tc", idx)
        if start == -1:
            break
        close = row_xml.find("</w:tc>", start)
        if close == -1:
            break
        end = close + len("</w:tc>")
        cells.append((start, end))
        idx = end
    if col_idx >= len(cells):
        return row_xml
    start, end = cells[col_idx]
    cell_seg = row_xml[start:end]
    # 在 <w:tcPr>...</w:tcPr> 里插入 shade；若无 tcPr，则新建
    if re.search(r"<w:tcPr\b[^>]*>.*?</w:tcPr>", cell_seg, re.DOTALL):
        cell_seg = re.sub(
            r"(<w:tcPr\b[^>]*>)(.*?)(</w:tcPr>)",
            lambda m: m.group(1) + re.sub(r"<w:shd\b[^/]*/>", "", m.group(2)) + shade_xml + m.group(3),
            cell_seg,
            count=1,
            flags=re.DOTALL,
        )
    else:
        open_tc = re.match(r"<w:tc\b[^>]*>", cell_seg)
        if open_tc:
            cell_seg = cell_seg[:open_tc.end()] + f"<w:tcPr>{shade_xml}</w:tcPr>" + cell_seg[open_tc.end():]
    return row_xml[:start] + cell_seg + row_xml[end:]


def _insert_topology_image(
    xml: str,
    others: list[tuple[str, bytes]],
    image_path: Path,
) -> tuple[str, list[tuple[str, bytes]]]:
    """将图片作为媒体资源加入 .docx 并在"定级对象构成"段落后插入一个图片段落。"""
    ext = image_path.suffix.lower().lstrip(".") or "png"
    if ext == "jpg":
        ext = "jpeg"
    media_name = f"word/media/topology.{ext}"
    img_bytes = image_path.read_bytes()

    # 更新 others：添加或替换媒体文件
    others = [(n, d) for n, d in others if n != media_name]
    others.append((media_name, img_bytes))

    # 更新 [Content_Types].xml 加入扩展
    ctypes_name = "[Content_Types].xml"
    for i, (n, d) in enumerate(others):
        if n == ctypes_name:
            ct_xml = d.decode("utf-8")
            if f'Extension="{ext}"' not in ct_xml:
                insert = f'<Default Extension="{ext}" ContentType="image/{ext}"/>'
                ct_xml = ct_xml.replace("<Types ", "<Types ").replace("</Types>", insert + "</Types>")
            others[i] = (n, ct_xml.encode("utf-8"))
            break

    # 更新 word/_rels/document.xml.rels 加入图片关系
    rels_name = "word/_rels/document.xml.rels"
    rel_id = "rIdTopology"
    for i, (n, d) in enumerate(others):
        if n == rels_name:
            rels_xml = d.decode("utf-8")
            if rel_id not in rels_xml:
                new_rel = f'<Relationship Id="{rel_id}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/topology.{ext}"/>'
                rels_xml = rels_xml.replace("</Relationships>", new_rel + "</Relationships>")
            others[i] = (n, rels_xml.encode("utf-8"))
            break

    # 构造 drawing 段落，插到"定级对象构成"段落之后
    drawing_xml = _make_drawing_paragraph(rel_id, image_path.name)
    anchor_prefix = xml_escape("山西省省属企业采购与供应链信息管理系统部署于")  # 原样例
    # 无论样例是否还在，都尝试找到"定级对象构成"标题段落后插入
    anchor_title = xml_escape("（二）定级对象构成")
    anchor_pos = xml.find(anchor_title)
    if anchor_pos != -1:
        # 插到该段落后
        p_close = xml.find("</w:p>", anchor_pos)
        if p_close != -1:
            insert_at = p_close + len("</w:p>")
            xml = xml[:insert_at] + drawing_xml + xml[insert_at:]
    return xml, others


def _make_drawing_paragraph(rel_id: str, name: str) -> str:
    """构造一个居中对齐、嵌入式图片的段落 XML。图片显示宽度约 14cm x 高度自适应（这里固定 10cm）。"""
    # EMU: 1 cm = 360000 EMU; 14 cm = 5040000, 10 cm = 3600000
    cx, cy = 5040000, 3600000
    return (
        '<w:p>'
        '<w:pPr><w:jc w:val="center"/></w:pPr>'
        '<w:r><w:drawing>'
        f'<wp:inline xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" distT="0" distB="0" distL="0" distR="0">'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:docPr id="1" name="{xml_escape(name)}"/>'
        '<wp:cNvGraphicFramePr><a:graphicFrameLocks xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" noChangeAspect="1"/></wp:cNvGraphicFramePr>'
        '<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
        '<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        '<pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        '<pic:nvPicPr><pic:cNvPr id="1" name="topology"/><pic:cNvPicPr/></pic:nvPicPr>'
        f'<pic:blipFill><a:blip xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" r:embed="{rel_id}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill>'
        f'<pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm><a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr>'
        '</pic:pic></a:graphicData></a:graphic></wp:inline></w:drawing></w:r></w:p>'
    )


# ──────────────────────────────────────────────────────────────────────────────
# Word 导入解析（XML 级，保留方框/符号等 python-docx 会略过的内容）
# ──────────────────────────────────────────────────────────────────────────────

_WORD_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _parse_docx_structure(content: bytes) -> dict[str, Any]:
    """返回 {'paragraphs': [text,...], 'tables': [[[cell_text]]...]}。"""
    import xml.etree.ElementTree as ET

    with zipfile.ZipFile(io.BytesIO(content)) as z:
        if "word/document.xml" not in z.namelist():
            raise ValueError("上传的文件不是有效的 Word (.docx)。")
        raw = z.read("word/document.xml")

    root = ET.fromstring(raw)

    def element_text(el: Any) -> str:
        parts: list[str] = []
        for node in el.iter():
            tag = node.tag
            if tag == _WORD_NS + "t":
                parts.append(node.text or "")
            elif tag == _WORD_NS + "tab":
                parts.append("\t")
            elif tag == _WORD_NS + "br":
                parts.append("\n")
        return "".join(parts)

    paragraphs: list[str] = []
    tables: list[list[list[str]]] = []

    def walk(el: Any, in_table: bool = False) -> None:
        for child in list(el):
            tag = child.tag
            if tag == _WORD_NS + "p":
                if not in_table:
                    paragraphs.append(element_text(child).strip())
            elif tag == _WORD_NS + "tbl":
                rows: list[list[str]] = []
                for tr in child.findall(_WORD_NS + "tr"):
                    cells: list[str] = []
                    for tc in tr.findall(_WORD_NS + "tc"):
                        cells.append(element_text(tc).strip())
                    rows.append(cells)
                tables.append(rows)
                walk(child, in_table=True)
            else:
                walk(child, in_table)

    body = root.find(_WORD_NS + "body")
    if body is not None:
        walk(body)

    return {"paragraphs": paragraphs, "tables": tables}


# 定级报告段落标题 → 字段 key 的映射
_GRADING_HEADINGS = [
    ("（一）责任主体", "responsible_subject"),
    ("（二）定级对象构成", "object_composition"),
    ("（三）承载业务", "carried_business"),
    ("（四）承载数据", "carried_data"),
    ("（五）安全责任", "security_responsibility"),
    ("1、业务信息描述", "business_info_description"),
    ("2、业务信息受到破坏时所侵害客体的确定", "business_info_damage_object"),
    ("3、业务信息受到破坏时对侵害客体的侵害程度的确定", "business_info_damage_degree"),
    ("1、系统服务描述", "service_description"),
    ("2、系统服务受到破坏时所侵害客体的确定", "service_damage_object"),
    ("3、系统服务受到破坏时对侵害客体的侵害程度的确定", "service_damage_degree"),
]


def parse_grading_report_docx(content: bytes) -> dict[str, Any]:
    """从已填写的定级报告 Word 提取段落内容 → {字段 key: 文本}。"""
    struct = _parse_docx_structure(content)
    paragraphs = [p for p in struct["paragraphs"] if p]

    result: dict[str, str] = {}
    # 识别段落，标题段落的下一个（非空、非表格数据）段落即为内容
    # 支持不严格匹配：忽略数字序号后的格式差异
    def norm(s: str) -> str:
        return re.sub(r"\s+", "", s)

    heading_norm = [(norm(h), k) for h, k in _GRADING_HEADINGS]

    i = 0
    n = len(paragraphs)
    while i < n:
        para = paragraphs[i]
        para_n = norm(para)
        matched_key: str | None = None
        for h_n, key in heading_norm:
            if para_n == h_n or para_n.startswith(h_n):
                matched_key = key
                break
        if matched_key:
            # 向下取第一个非空、非下一个标题段
            collected: list[str] = []
            j = i + 1
            while j < n:
                nxt = paragraphs[j]
                nxt_n = norm(nxt)
                if not nxt:
                    j += 1
                    continue
                # 若遇到下一个已知标题 / 二级章节，则停止
                if any(nxt_n == h_n or nxt_n.startswith(h_n) for h_n, _ in heading_norm):
                    break
                if re.match(r"^[一二三四五六七八九十]+[、.]", nxt) or re.match(r"^[（(][一二三四五六七八九十][）)]", nxt):
                    break
                # 过滤"表一"、"表二"之类的图表注释行
                if re.match(r"^表[一二三四五]", nxt):
                    break
                collected.append(nxt)
                j += 1
                # 仅取第一段（通常用户填写单段）；如需多段可扩展
                break
            if collected:
                result[matched_key] = collected[0]
            i = j if collected else i + 1
            continue
        # 尝试提取填表日期：以"年"和"日"结尾且含数字的段落
        if re.match(r"^\d{4}年\s*\d{1,2}\s*月\s*\d{1,2}\s*日", para):
            result.setdefault("fill_date", para)
        i += 1

    return result


def parse_expert_review_city_docx(content: bytes) -> dict[str, Any]:
    """从市级专家评审意见表 Word 提取字段。"""
    struct = _parse_docx_structure(content)
    result: dict[str, Any] = {}
    for table in struct["tables"]:
        for row in table:
            if not row:
                continue
            first = row[0]
            first_n = re.sub(r"\s+", "", first)
            value = row[1] if len(row) > 1 else ""
            if "单位名称" == first_n and value and value != first:
                result["unit_name"] = value
            elif "项目负责人" == first_n:
                result["project_owner"] = value
                # 若本行还有"联系电话"
                for idx, cell in enumerate(row):
                    if re.sub(r"\s+", "", cell) == "联系电话" and idx + 1 < len(row):
                        result["contact_phone"] = row[idx + 1]
                        break
            elif "信息系统名称" == first_n and value and value != first:
                result["system_name"] = value
            elif "系统自定安全级别" == first_n and value and value != first:
                result["self_level"] = value
            elif "专家组成员" == first_n and value and value != first:
                result["experts"] = value
            elif "评审专家组组长（签字）" in first_n or first_n.startswith("评审专家组组长"):
                if value and value != first:
                    result["leader_signature"] = value
            elif "评审专家组成员（签字）" in first_n or first_n.startswith("评审专家组成员"):
                if value and value != first:
                    result["member_signatures"] = value
            # 审核意见：整行内容跨 4 列相同的长段
            elif "形成如下评审意见" in first or "评审意见：" in first:
                # 段落文本紧随"形成如下评审意见："后面
                match = re.search(r"形成如下评审意见[：:](.*)", first, re.DOTALL)
                if match and match.group(1).strip():
                    result["review_opinion"] = match.group(1).strip()
            # 尝试提取评审日期："YYYY年MM月DD日"
            for cell in row:
                m = re.search(r"(\d{4}年\s*\d{1,2}\s*月\s*\d{1,2}\s*日)", cell)
                if m and "review_date" not in result:
                    result["review_date"] = m.group(1)
    return result


def parse_expert_review_department_docx(content: bytes) -> dict[str, Any]:
    """从厅级专家评审意见表 Word 提取字段。"""
    struct = _parse_docx_structure(content)
    result: dict[str, Any] = {}
    labels = [
        ("信息系统运营、使用单位名称", "unit_name"),
        ("信息系统运营、使用单位地址", "unit_address"),
        ("项目负责人", "project_owner"),
        ("联系电话", "contact_phone"),
        ("邮件地址", "email"),
        ("信息系统名称", "system_name"),
        ("系统自定安全级别", "self_level"),
    ]
    for table in struct["tables"]:
        for row in table:
            for cell in row:
                for label, key in labels:
                    # 支持格式 "信息系统名称：XXX" 或单独出现
                    if cell.startswith(label):
                        rest = cell[len(label):].lstrip("：: ").strip()
                        if rest:
                            result[key] = rest
            # 评审专家组意见
            for cell in row:
                if cell.startswith("评审专家组意见"):
                    rest = cell[len("评审专家组意见"):].lstrip("：: ").strip()
                    if rest:
                        result["review_opinion"] = rest
    # 填表时间
    for p in struct["paragraphs"]:
        m = re.search(r"填表时间[：:]\s*(.+)", p)
        if m:
            value = m.group(1).strip()
            if value and not value.startswith("年"):
                result["review_date"] = value
                break
    return result


def parse_filing_form_docx(content: bytes) -> dict[str, Any]:
    """从备案表 Word 提取核心文本字段（仅文本，选项/附件不处理）。"""
    struct = _parse_docx_structure(content)
    result: dict[str, Any] = {}

    exact_labels = {
        "单位名称": "org_name",
        "单位地址": "org_address",
        "业务描述": "business_description",
        "定级时间": "grading_date",
        "邮政编码": "postal_code",
        "行政区划代码": "district_code",
        "网络安全责任部门": "cybersecurity_dept",
        "数据安全管理部门": "data_security_dept",
    }
    contains_labels = {
        "统一社会信用代码": "credit_code",
        "何时投入运行使用": "go_live_date",
    }

    for table in struct["tables"]:
        for row in table:
            for idx, cell in enumerate(row):
                cell_n = re.sub(r"\s+", "", cell)
                next_val = row[idx + 1].strip() if idx + 1 < len(row) else ""
                for label, key in exact_labels.items():
                    if cell_n == label and next_val and next_val != label:
                        result.setdefault(key, next_val)
                for label, key in contains_labels.items():
                    if label in cell_n and next_val and next_val != cell_n:
                        result.setdefault(key, next_val)
                if cell_n == "定级对象" and next_val and next_val != "定级对象":
                    result.setdefault("system_name", next_val)
                if cell_n == "单位负责人":
                    for si, sc in enumerate(row[idx + 1:], idx + 1):
                        if re.sub(r"\s+", "", sc) == "姓名" and si + 1 < len(row) and row[si + 1].strip():
                            result.setdefault("legal_representative", row[si + 1].strip())
                            break
                if cell_n.startswith("网络安全责任部门联系人"):
                    for si, sc in enumerate(row[idx + 1:], idx + 1):
                        if re.sub(r"\s+", "", sc) == "姓名" and si + 1 < len(row) and row[si + 1].strip():
                            result.setdefault("cybersecurity_owner_name", row[si + 1].strip())
                            break
                if cell_n.startswith("填表人") and "：" in cell_n:
                    val = cell_n.split("：", 1)[1].strip()
                    if val:
                        result.setdefault("filler_name", val)
                if cell_n.startswith("填表日期") and "：" in cell_n:
                    val = re.sub(r"\s+", "", cell).split("：", 1)[1].strip()
                    if val:
                        result.setdefault("filled_date", cell.split("：", 1)[1].strip())

    return result


def export_expert_review_docx(
    template_path: Path,
    org: Any,
    system: Any,
    output_path: Path,
    variant: str,
    content: dict[str, Any] | None = None,
) -> None:
    """基于 03-专家评审意见表 模板生成专家评审意见表（市级/厅级）。

    content 为用户填写的表单数据；缺省时使用当前单位/系统 prefill。
    """
    template_path = Path(template_path)
    if not template_path.exists():
        raise FileNotFoundError(f"专家评审意见表模板不存在：{template_path}")

    cn = _cn_level(getattr(system, "proposed_level", 3))
    org_name = getattr(org, "name", "") or SAMPLE_ORG_NAME
    sys_name = getattr(system, "system_name", "") or SAMPLE_SYSTEM_NAME
    owner_name = getattr(org, "cybersecurity_owner_name", "") or "________"
    owner_phone = getattr(org, "cybersecurity_owner_phone", "") or "____________"
    owner_email = getattr(org, "cybersecurity_owner_email", "") or "____________"
    address = getattr(org, "address", "") or "____________"

    c = content or {}
    # 表单内容优先，其次系统/单位兜底
    f_unit = (c.get("unit_name") or "").strip() or org_name
    f_sys = (c.get("system_name") or "").strip() or sys_name
    f_owner = (c.get("project_owner") or "").strip() or owner_name
    f_phone = (c.get("contact_phone") or "").strip() or owner_phone
    f_level = (c.get("self_level") or "").strip() or f"第{cn}级"
    f_experts = (c.get("experts") or "").strip() or "________"
    f_date = (c.get("review_date") or "").strip() or _today_short()

    xml, others = _read_docx_xml(template_path)

    if variant == "city":
        # 基础单元格（单行文本）
        repl = {
            SAMPLE_ORG_NAME: f_unit,
            SAMPLE_SYSTEM_NAME: f_sys,
            SAMPLE_OWNER_NAME_REVIEW: f_owner,
            SAMPLE_OWNER_PHONE: f_phone,
            SAMPLE_EXPERTS: f_experts,
            "第三级": f_level,
        }
        xml = _xml_replace(xml, repl)

        # 评审意见长段：整段替换（日期 + 单位 + 系统 + 意见）
        opinion = (c.get("review_opinion") or "").strip() or "________"
        full_paragraph = (
            f"{f_date}，{f_unit}组织网络安全等级保护专家对本单位拟备案的{f_sys}"
            f"进行了安全保护等级定级评审。与会专家听取了该单位对系统情况的介绍，"
            f"审阅了有关材料，经过讨论质询，形成如下评审意见：{opinion}"
        )
        xml = _replace_paragraph_contains(xml, "形成如下评审意见", full_paragraph)
        # 移除紧随其后的模板样例段落（数字开头："1.定级文档..."）
        xml = _remove_sample_numbered_paragraphs(xml, max_remove=5)

        # 签字行
        leader = (c.get("leader_signature") or "").strip()
        if leader:
            xml = _replace_paragraph_contains(xml, "评审专家组组长（签字）", f"评审专家组组长（签字）：{leader}")
        members = (c.get("member_signatures") or "").strip()
        if members:
            xml = _replace_paragraph_contains(xml, "评审专家组成员（签字）", f"评审专家组成员（签字）：{members}")
    else:
        f_address = (c.get("unit_address") or "").strip() or address
        f_email = (c.get("email") or "").strip() or owner_email
        repl = {
            "第*级": f_level,
            "****": f_sys,
            "***": f_level,
            "填表时间：    年    月    日": f"填表时间：{f_date}",
        }
        xml = _xml_replace(xml, repl)
        xml = _append_label_value(xml, "信息系统运营、使用单位名称：", f_unit)
        xml = _append_label_value(xml, "信息系统运营、使用单位地址：", f_address)
        xml = _append_label_value(xml, "项目负责人：", f_owner)
        xml = _append_label_value(xml, "联系电话：", f_phone)
        xml = _append_label_value(xml, "邮件地址：", f_email)
        xml = _append_label_value(xml, "信息系统名称：", f_sys)
        xml = _append_label_value(xml, "系统自定安全级别：", f_level)
        opinion = (c.get("review_opinion") or "").strip()
        if opinion:
            xml = _append_label_value(xml, "评审专家组意见：", opinion)

    _write_docx_xml(output_path, xml, others)
    logger.info("专家评审意见表已生成（%s）：%s", variant, output_path)


def _append_label_value(xml: str, label: str, value: str) -> str:
    """在包含 label 的 <w:t> 后追加一个带 value 的 <w:t>（同段落）。"""
    if not value:
        return xml
    escaped_label = xml_escape(label)
    escaped_value = xml_escape(value)
    # 将 <w:t ...>label</w:t> 变为 <w:t ...>label value</w:t>
    pattern = re.compile(r"(<w:t\b[^>]*>)(" + re.escape(escaped_label) + r")(</w:t>)")
    return pattern.sub(lambda m: f"{m.group(1)}{m.group(2)}{escaped_value}{m.group(3)}", xml, count=1)


def export_expert_review_merged_docx(
    template_path: Path,
    org: Any,
    systems: list[Any],
    output_path: Path,
    variant: str,
) -> None:
    """合并多个系统到同一份专家评审意见表：以首个系统为主，将其它系统名追加展示。"""
    if not systems:
        raise ValueError("systems 不能为空")
    primary = systems[0]
    # 以第一个系统为主生成基础文档，然后在系统名称处合并追加
    export_expert_review_docx(template_path, org, primary, output_path, variant)

    joined_names = "、".join(getattr(s, "system_name", "") or "" for s in systems)
    primary_name = getattr(primary, "system_name", "") or ""
    levels = "、".join(f"第{_cn_level(getattr(s, 'proposed_level', 3))}级" for s in systems)

    xml, others = _read_docx_xml(output_path)
    if primary_name and joined_names != primary_name:
        xml = xml.replace(xml_escape(primary_name), xml_escape(joined_names))
    cn_primary = _cn_level(getattr(primary, "proposed_level", 3))
    if len({_cn_level(getattr(s, "proposed_level", 3)) for s in systems}) > 1:
        xml = xml.replace(xml_escape(f"第{cn_primary}级"), xml_escape(levels))
    _write_docx_xml(output_path, xml, others)
    logger.info("合并专家评审意见表已生成：%s", output_path)
