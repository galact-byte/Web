import asyncio
import importlib
import hashlib
import io
import json
import os
import unittest
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document
import httpx
from fastapi.testclient import TestClient as FastAPITestClient
from openpyxl import Workbook


class CompatibleTestClient:
    def __init__(self, app, base_url: str = 'http://testserver', follow_redirects: bool = True):
        self.base_url = base_url
        self.follow_redirects = follow_redirects
        self.transport = httpx.ASGITransport(app=app)
        self.cookies = httpx.Cookies()

    async def _request(self, method: str, url: str, **kwargs):
        follow_redirects = kwargs.pop('follow_redirects', self.follow_redirects)
        async with httpx.AsyncClient(
            transport=self.transport,
            base_url=self.base_url,
            cookies=self.cookies,
            follow_redirects=follow_redirects,
        ) as client:
            response = await client.request(method, url, **kwargs)
            self.cookies.update(client.cookies)
            self.cookies.update(response.cookies)
            return response

    def request(self, method: str, url: str, **kwargs):
        return asyncio.run(self._request(method, url, **kwargs))

    def get(self, url: str, **kwargs):
        return self.request('GET', url, **kwargs)

    def post(self, url: str, **kwargs):
        return self.request('POST', url, **kwargs)

    def put(self, url: str, **kwargs):
        return self.request('PUT', url, **kwargs)

    def delete(self, url: str, **kwargs):
        return self.request('DELETE', url, **kwargs)

    def close(self):
        asyncio.run(self.transport.aclose())


def build_test_client(app):
    try:
        return FastAPITestClient(app)
    except TypeError as exc:
        if "unexpected keyword argument 'app'" not in str(exc):
            raise
        return CompatibleTestClient(app)


class ApiFlowTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        os.environ['DATABASE_URL'] = 'sqlite:///:memory:'
        os.environ['API_AUTH_REQUIRED'] = '0'

        import app.db as db_module
        import app.main as main_module
        import app.models as models_module

        importlib.reload(db_module)
        importlib.reload(models_module)
        importlib.reload(main_module)
        db_module.init_db()

        db = db_module.SessionLocal()
        try:
            admin = db.query(models_module.UserAccount).filter(models_module.UserAccount.username == 'admin').first()
            if not admin:
                db.add(
                    models_module.UserAccount(
                        username='admin',
                        password_hash=hashlib.sha256('admin123'.encode('utf-8')).hexdigest(),
                        role='admin',
                        enabled=True,
                    )
                )
                db.commit()
        finally:
            db.close()

        cls.client = build_test_client(main_module.app)
        cls.main_module = main_module
        cls.db_module = db_module
        login_resp = cls.client.post('/api/auth/login', json={'username': 'admin', 'password': 'admin123'})
        assert login_resp.status_code == 200, login_resp.text
        cls.admin_token = login_resp.json()['token']
        cls.admin_headers = {'X-Auth-Token': cls.admin_token}

    @classmethod
    def tearDownClass(cls):
        cls.client.close()

    def build_customer_filing_excel(
        self,
        *,
        org_name: str,
        credit_code: str,
        system_name: str,
        level: int,
        business_description: str = '客户填报业务描述',
        cybersecurity_dept: str = '信息安全部',
        go_live_date: str = '2024-03-01',
        leader_name: str = '负责人甲',
        leader_title: str = '主任',
        leader_phone: str = '010-88886666',
        leader_email: str = 'leader@example.com',
        cyber_owner_name: str = '网安联系人',
        cyber_owner_title: str = '安全主管',
        cyber_owner_phone: str = '010-66668888',
        cyber_owner_mobile: str = '13900139099',
        cyber_owner_email: str = 'cyber@example.com',
        data_dept: str = '数据治理部',
        data_owner_name: str = '数安联系人',
        data_owner_title: str = '数据主管',
        data_owner_phone: str = '010-55557777',
        data_owner_mobile: str = '13900139100',
        data_owner_email: str = 'data@example.com',
        unit_internet_addresses: str = 'https://unit.example.com',
        public_network: str = '是',
        public_ip: str = '203.0.113.10',
        protocol_ports: str = 'HTTPS/443',
        access_address: str = 'https://sys.example.com',
        room_location: str = 'A 栋 3 层机房',
        room_owner: str = '机房管理员',
        room_owner_mobile: str = '13900139101',
    ) -> bytes:
        wb = Workbook()
        ws_org = wb.active
        ws_org.title = '单位信息'
        ws_sys = wb.create_sheet('系统信息')
        wb.create_sheet('Sheet3')

        ws_org['B2'] = org_name
        ws_org['B3'] = credit_code
        ws_org['B4'] = leader_name
        ws_org['E4'] = leader_title
        ws_org['B5'] = leader_phone
        ws_org['E5'] = leader_email
        ws_org['B6'] = cybersecurity_dept
        ws_org['B7'] = cyber_owner_name
        ws_org['E7'] = cyber_owner_title
        ws_org['B8'] = cyber_owner_phone
        ws_org['B9'] = cyber_owner_mobile
        ws_org['E8'] = cyber_owner_email
        ws_org['B10'] = data_dept
        ws_org['B11'] = data_owner_name
        ws_org['E11'] = data_owner_title
        ws_org['B12'] = data_owner_phone
        ws_org['B13'] = data_owner_mobile
        ws_org['E12'] = data_owner_email
        ws_org['B14'] = unit_internet_addresses

        ws_sys['B2'] = system_name
        ws_sys['D2'] = f'第{level}级'
        ws_sys['B3'] = business_description
        ws_sys['B8'] = go_live_date
        ws_sys['B9'] = public_network
        ws_sys['D9'] = public_ip
        ws_sys['B10'] = protocol_ports
        ws_sys['B11'] = access_address
        ws_sys['B12'] = room_location
        ws_sys['B13'] = room_owner
        ws_sys['D13'] = room_owner_mobile

        bio = io.BytesIO()
        wb.save(bio)
        return bio.getvalue()

    def test_01_org_system_report_flow(self):
        org_payload = {
            'name': '测试单位A',
            'credit_code': '91350100M000100Y43',
            'legal_representative': '张三',
            'address': '福建省福州市',
            'office_phone': '0591-1234567',
            'mobile_phone': '13800138000',
            'email': 'a@example.com',
            'industry': '金融',
            'organization_type': '企业',
            'filing_region': '福州',
            'created_by': 'tester',
        }
        org_resp = self.client.post('/api/organizations', json=org_payload)
        self.assertEqual(org_resp.status_code, 200)
        org_id = org_resp.json()['data']['id']

        system_payload = {
            'organization_id': org_id,
            'system_name': '核心业务系统',
            'proposed_level': 3,
            'deployment_mode': '混合部署',
            'created_by': 'tester',
        }
        sys_resp = self.client.post('/api/systems', json=system_payload)
        self.assertEqual(sys_resp.status_code, 200)
        system_id = sys_resp.json()['data']['id']

        report_resp = self.client.post(f'/api/reports/generate?system_id={system_id}&report_type=grading_report&actor=tester')
        self.assertEqual(report_resp.status_code, 200)
        report_id = report_resp.json()['data']['id']

        submit_resp = self.client.post(
            f'/api/reports/{report_id}/submit?actor=tester&reviewer=leader',
            headers=self.admin_headers,
        )
        self.assertEqual(submit_resp.status_code, 200)
        self.assertEqual(submit_resp.json()['status'], 'submitted')

        review_resp = self.client.post(
            f'/api/reports/{report_id}/review?actor=leader&action=approve&comment=通过',
            headers=self.admin_headers,
        )
        self.assertEqual(review_resp.status_code, 200)
        self.assertEqual(review_resp.json()['status'], 'approved')

    def test_02_validation_and_dashboard(self):
        bad_payload = {
            'name': '测试单位B',
            'credit_code': 'BAD',
            'legal_representative': '李四',
            'address': '北京市',
            'mobile_phone': '13800138000',
            'email': 'b@example.com',
            'industry': '政府',
            'organization_type': '机关单位',
            'filing_region': '北京',
            'created_by': 'tester',
        }
        bad_resp = self.client.post('/api/organizations', json=bad_payload)
        self.assertEqual(bad_resp.status_code, 400)

        good_payload = {
            'name': '测试单位C',
            'credit_code': '91350100M000100Y44',
            'legal_representative': '王五',
            'address': '北京市',
            'mobile_phone': '13900139000',
            'email': 'c@example.com',
            'industry': '政府',
            'organization_type': '机关单位',
            'filing_region': '北京',
            'created_by': 'tester',
        }
        ok_resp = self.client.post('/api/organizations', json=good_payload)
        self.assertEqual(ok_resp.status_code, 200)

        bad_name_payload = {
            'name': '/',
            'credit_code': '91350100M000100Y53',
            'legal_representative': '赵六',
            'address': '天津市',
            'mobile_phone': '13900139001',
            'email': 'd@example.com',
            'industry': '政府',
            'organization_type': '机关单位',
            'filing_region': '天津',
            'created_by': 'tester',
        }
        bad_name_resp = self.client.post('/api/organizations', json=bad_name_payload)
        self.assertEqual(bad_name_resp.status_code, 400)

        bad_format_payload = {
            'name': '测试单位D',
            'credit_code': '/',
            'legal_representative': '/',
            'address': '/',
            'mobile_phone': '/',
            'email': '/',
            'industry': '/',
            'organization_type': '/',
            'filing_region': '/',
            'created_by': 'tester',
        }
        bad_format_resp = self.client.post('/api/organizations', json=bad_format_payload)
        self.assertEqual(bad_format_resp.status_code, 400)

        placeholder_payload = {
            'name': '测试单位E',
            'credit_code': '91350100M000100Y53',
            'legal_representative': '/',
            'address': '/',
            'mobile_phone': '13900139001',
            'email': 'placeholder@example.com',
            'industry': '/',
            'organization_type': '/',
            'filing_region': '/',
            'created_by': 'tester',
        }
        placeholder_resp = self.client.post('/api/organizations', json=placeholder_payload)
        self.assertEqual(placeholder_resp.status_code, 200)

        dash_resp = self.client.get('/api/dashboard/summary?city=北京')
        self.assertEqual(dash_resp.status_code, 200)
        self.assertIn('totals', dash_resp.json())

    def test_02a_damage_level_row_selected_accepts_detailed_items(self):
        helper = self.main_module.damage_level_row_selected
        self.assertTrue(
            helper(
                '对公民、法人和其他组织的合法权益造成严重损害 / 对公民、法人和其他组织的合法权益造成特别严重损害 / 对社会秩序和公共利益造成一般损害',
                {'对社会秩序和公共利益造成一般损害'},
            )
        )
        self.assertTrue(
            helper(
                '对国家安全或地区安全、国计民生造成严重损害 / 对国家安全或地区安全、国计民生造成特别严重损害',
                {'对国家安全或地区安全、国计民生造成特别严重损害'},
            )
        )
        self.assertFalse(
            helper(
                '对社会秩序和公共利益造成严重损害',
                {'仅对公民、法人和其他组织的合法权益造成一般损害'},
            )
        )

    def test_02ad_table4_record_export_name_respects_status_flag(self):
        helper = self.main_module.table4_record_export_name
        prefix = '测试单位-测试系统'
        self.assertEqual(
            helper({'record_status': 'has'}, prefix, '云平台备案证明'),
            f'{prefix}-云平台备案证明',
        )
        self.assertEqual(
            helper({'record_status': 'none', 'record_file_name': 'legacy.pdf', 'attachment_ids': [1]}, prefix, '云平台备案证明'),
            '',
        )
        self.assertEqual(
            helper({'record_file_name': 'legacy.pdf', 'attachment_ids': [1]}, prefix, '云平台备案证明'),
            f'{prefix}-云平台备案证明',
        )

    def test_02ae_table6_data_total_check_indices_allow_both_rows(self):
        helper = self.main_module.table6_data_total_check_indices
        self.assertEqual(helper({'data_total_gb': '3.35'}), {0})
        self.assertEqual(helper({'data_total_tb': '1.2'}), {0})
        self.assertEqual(helper({'data_total_records': '0.5'}), {1})
        self.assertEqual(helper({'data_total_gb': '3.35', 'data_total_records': '0.5'}), {0, 1})

    def test_02a_filing_workspace_flow_and_systems_redirect(self):
        org_resp = self.client.post('/api/organizations', json={
            'name': '备案工作台单位',
            'credit_code': '91350100M000100Y92',
            'legal_representative': '赵七',
            'address': '太原市',
            'mobile_phone': '13900139002',
            'email': 'workspace@example.com',
            'industry': '能源',
            'organization_type': '企业',
            'filing_region': '太原',
            'created_by': 'tester',
        })
        self.assertEqual(org_resp.status_code, 200, org_resp.text)
        org_id = org_resp.json()['data']['id']

        system_resp = self.client.post('/api/systems', json={
            'organization_id': org_id,
            'system_name': '备案工作台系统',
            'proposed_level': 2,
            'go_live_date': '2024-01-15',
            'created_by': 'tester',
        })
        self.assertEqual(system_resp.status_code, 200, system_resp.text)
        system_id = system_resp.json()['data']['id']

        redirect_resp = self.client.get('/systems', follow_redirects=False)
        self.assertEqual(redirect_resp.status_code, 307, redirect_resp.text)
        self.assertEqual(redirect_resp.headers.get('location'), '/organizations?focus=systems')

        overview_resp = self.client.get('/api/filing-workspace/overview', headers=self.admin_headers)
        self.assertEqual(overview_resp.status_code, 200, overview_resp.text)
        self.assertGreaterEqual(overview_resp.json().get('total', 0), 1)

        detail_resp = self.client.get(f'/api/filing-workspace/systems/{system_id}', headers=self.admin_headers)
        self.assertEqual(detail_resp.status_code, 200, detail_resp.text)
        detail = detail_resp.json()
        self.assertEqual(detail['organization']['id'], org_id)
        self.assertEqual(detail['system']['id'], system_id)
        self.assertEqual(detail['organization']['filing_detail']['table1']['unit_internet_addresses'], '无')
        self.assertEqual(detail['system']['filing_detail']['table2']['running_status'], '已运行')

        save_payload = {
            'organization': {
                'name': '备案工作台单位',
                'credit_code': '91350100M000100Y92',
                'filing_region': '太原',
                'filing_detail': {
                    'table1': {
                        'unit_internet_addresses': 'https://example.com',
                        'address_parts': {'province': '山西', 'city': '太原', 'district': '小店区', 'detail': '测试地址'},
                    }
                },
            },
            'system': {
                'system_name': '备案工作台系统',
                'filing_detail': {
                    'table2': {
                        'running_status': '在建设',
                        'object_types': ['信息系统', '数据资源'],
                        'technology_types': ['云计算技术'],
                        'business_types': ['公众服务'],
                        'business_description': '统一工作台保存测试',
                        'network_service': {
                            'scope_code': '10',
                            'scope_label': '全国',
                            'service_targets': ['两者均包括'],
                        },
                        'network_platform': {
                            'coverage': {'code': '3', 'label': '广域网'},
                            'interconnection': {'items': ['与本单位其他系统连接']},
                        },
                    },
                    'table3': {
                        'business_security_damage_items': ['对社会秩序和公共利益造成严重损害'],
                        'business_security_level': 3,
                        'service_security_damage_items': ['仅对公民、法人和其他组织的合法权益造成一般损害'],
                        'service_security_level': 1,
                    },
                    'table6': {
                        'items': [
                            {
                                'data_name': '交易数据',
                                'data_level_code': '2',
                                'data_category': '业务数据',
                            }
                        ]
                    },
                },
            },
        }
        save_resp = self.client.put(f'/api/filing-workspace/systems/{system_id}', json=save_payload, headers=self.admin_headers)
        self.assertEqual(save_resp.status_code, 200, save_resp.text)
        saved = save_resp.json()['data']
        self.assertEqual(saved['organization']['filing_detail']['table1']['unit_internet_addresses'], 'https://example.com')
        self.assertEqual(saved['system']['filing_detail']['table2']['running_status'], '在建设')
        self.assertEqual(saved['system']['filing_detail']['table2']['business_description'], '统一工作台保存测试')
        self.assertEqual(saved['system']['filing_detail']['table3']['final_level'], 3)
        self.assertEqual(saved['system']['service_scope'], '全国')
        self.assertEqual(saved['system']['service_object'], '两者均包括')
        self.assertEqual(saved['system']['deployment_mode'], '广域网')
        self.assertEqual(saved['system']['boundary'], '与本单位其他系统连接')
        self.assertEqual(saved['system']['system_type'], '信息系统（采用：云计算技术） / 数据资源')
        self.assertEqual(saved['system']['data_name'], '交易数据')
        self.assertEqual(saved['system']['data_level'], '重要及以上数据')
        self.assertEqual(saved['system']['data_category'], '业务数据')

    def test_02a1_workspace_overview_supports_inline_edit_fields(self):
        org_resp = self.client.post('/api/organizations', json={
            'name': '工作台可编辑单位',
            'credit_code': '91350100M000100Y72',
            'legal_representative': '张三',
            'address': '太原市迎泽区原地址',
            'mobile_phone': '13900139072',
            'email': 'inline-edit@example.com',
            'industry': '能源',
            'organization_type': '企业',
            'filing_region': '太原',
            'created_by': 'tester',
        })
        self.assertEqual(org_resp.status_code, 200, org_resp.text)
        org_id = org_resp.json()['data']['id']

        system_resp = self.client.post('/api/systems', json={
            'organization_id': org_id,
            'system_name': '工作台可编辑系统',
            'proposed_level': 2,
            'business_description': '原业务描述',
            'system_type': '信息系统',
            'deployment_mode': '本地部署',
            'created_by': 'tester',
        })
        self.assertEqual(system_resp.status_code, 200, system_resp.text)
        system_id = system_resp.json()['data']['id']

        update_org_resp = self.client.put(f'/api/organizations/{org_id}', json={
            'name': '工作台已修改单位',
            'credit_code': '91350100M000100Y73',
            'filing_region': '晋中',
        }, headers=self.admin_headers)
        self.assertEqual(update_org_resp.status_code, 200, update_org_resp.text)
        self.assertEqual(update_org_resp.json()['data']['credit_code'], '91350100M000100Y73')

        update_system_resp = self.client.put(f'/api/systems/{system_id}', json={
            'system_name': '工作台已修改系统',
            'proposed_level': 3,
            'business_description': '已修改业务描述',
            'deployment_mode': '云部署',
        }, headers=self.admin_headers)
        self.assertEqual(update_system_resp.status_code, 200, update_system_resp.text)

        overview_resp = self.client.get('/api/filing-workspace/overview?keyword=工作台已修改单位', headers=self.admin_headers)
        self.assertEqual(overview_resp.status_code, 200, overview_resp.text)
        item = overview_resp.json()['items'][0]
        self.assertEqual(item['organization']['name'], '工作台已修改单位')
        self.assertEqual(item['organization']['credit_code'], '91350100M000100Y73')
        self.assertEqual(item['organization']['legal_representative'], '张三')
        self.assertEqual(item['organization']['address'], '太原市迎泽区原地址')
        self.assertEqual(item['organization']['organization_type'], '企业')
        self.assertEqual(item['systems'][0]['system_name'], '工作台已修改系统')
        self.assertEqual(item['systems'][0]['business_description'], '已修改业务描述')
        self.assertEqual(item['systems'][0]['deployment_mode'], '云部署')

    def test_02aa_filing_workspace_excel_preview_returns_candidate(self):
        org_resp = self.client.post('/api/organizations', json={
            'name': 'Excel预检单位',
            'credit_code': '91350100M000100Y96',
            'legal_representative': '李四',
            'address': '太原市小店区',
            'mobile_phone': '13900139006',
            'email': 'excel-preview@example.com',
            'industry': '能源',
            'organization_type': '企业',
            'filing_region': '太原',
            'created_by': 'tester',
        })
        self.assertEqual(org_resp.status_code, 200, org_resp.text)
        org_id = org_resp.json()['data']['id']

        system_resp = self.client.post('/api/systems', json={
            'organization_id': org_id,
            'system_name': 'Excel预检系统',
            'proposed_level': 3,
            'created_by': 'tester',
        })
        self.assertEqual(system_resp.status_code, 200, system_resp.text)
        system_id = system_resp.json()['data']['id']

        excel_bytes = self.build_customer_filing_excel(
            org_name='Excel预检单位',
            credit_code='91350100M000100Y96',
            system_name='Excel预检系统',
            level=3,
            business_description='来自客户 Excel 的业务描述',
            cybersecurity_dept='网络安全处',
            go_live_date='2024-03-01',
            leader_name='李四',
            leader_email='excel-preview@example.com',
            cyber_owner_mobile='13900139006',
            public_ip='203.0.113.88',
            access_address='https://preview.example.com',
        )
        resp = self.client.post(
            f'/api/filing-workspace/systems/{system_id}/import-excel/preview',
            files={'file': ('customer.xlsx', excel_bytes, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')},
            headers=self.admin_headers,
        )
        self.assertEqual(resp.status_code, 200, resp.text)
        data = resp.json()['data']
        self.assertEqual(data['candidate']['organization']['cybersecurity_dept'], '网络安全处')
        self.assertEqual(data['candidate']['system']['filing_detail']['table2']['go_live_date'], '2024-03-01')
        self.assertEqual(data['candidate']['system']['filing_detail']['table2']['network_platform']['network_nature']['source_ip_range'], '203.0.113.88')
        self.assertEqual(data['candidate']['grading_report_content']['carried_business'], '来自客户 Excel 的业务描述')
        self.assertEqual(data['candidate']['grading_report_content']['object_composition'], '系统连接公网，公网IP地址：203.0.113.88；主要协议/端口：HTTPS/443；系统访问地址：https://preview.example.com；机房位置：A 栋 3 层机房；机房负责人：机房管理员（13900139101）')
        self.assertEqual(data['conflicts'], [])
        self.assertIn('organization.cybersecurity_dept', data['direct_update_keys'])

    def test_02ab_filing_workspace_excel_preview_blocks_identity_mismatch(self):
        org_resp = self.client.post('/api/organizations', json={
            'name': 'Excel拦截单位',
            'credit_code': '91350100M000100Y97',
            'legal_representative': '王五',
            'address': '太原市迎泽区',
            'mobile_phone': '13900139007',
            'email': 'excel-block@example.com',
            'industry': '能源',
            'organization_type': '企业',
            'filing_region': '太原',
            'created_by': 'tester',
        })
        self.assertEqual(org_resp.status_code, 200, org_resp.text)
        org_id = org_resp.json()['data']['id']

        system_resp = self.client.post('/api/systems', json={
            'organization_id': org_id,
            'system_name': 'Excel拦截系统',
            'proposed_level': 2,
            'created_by': 'tester',
        })
        self.assertEqual(system_resp.status_code, 200, system_resp.text)
        system_id = system_resp.json()['data']['id']

        excel_bytes = self.build_customer_filing_excel(
            org_name='Excel拦截单位',
            credit_code='91350100M000100Y97',
            system_name='Excel拦截系统',
            level=3,
        )
        resp = self.client.post(
            f'/api/filing-workspace/systems/{system_id}/import-excel/preview',
            files={'file': ('customer.xlsx', excel_bytes, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')},
            headers=self.admin_headers,
        )
        self.assertEqual(resp.status_code, 400, resp.text)
        self.assertIn('备案等级', resp.json()['detail'])

    def test_02ac_filing_workspace_excel_apply_only_overrides_selected_conflicts(self):
        org_resp = self.client.post('/api/organizations', json={
            'name': 'Excel覆盖单位',
            'credit_code': '91350100M000100Y98',
            'legal_representative': '赵六',
            'address': '太原市杏花岭区',
            'mobile_phone': '13900139008',
            'email': 'excel-apply@example.com',
            'industry': '能源',
            'organization_type': '企业',
            'filing_region': '太原',
            'created_by': 'tester',
        })
        self.assertEqual(org_resp.status_code, 200, org_resp.text)
        org_id = org_resp.json()['data']['id']

        system_resp = self.client.post('/api/systems', json={
            'organization_id': org_id,
            'system_name': 'Excel覆盖系统',
            'proposed_level': 2,
            'business_description': '旧业务描述',
            'created_by': 'tester',
        })
        self.assertEqual(system_resp.status_code, 200, system_resp.text)
        system_id = system_resp.json()['data']['id']

        save_resp = self.client.put(
            f'/api/filing-workspace/systems/{system_id}',
            json={
                'organization': {
                    'name': 'Excel覆盖单位',
                    'credit_code': '91350100M000100Y98',
                    'filing_region': '太原',
                    'cybersecurity_dept': '旧网安部门',
                },
                'system': {
                    'system_name': 'Excel覆盖系统',
                    'business_description': '旧业务描述',
                    'filing_detail': {
                        'table2': {
                            'business_description': '旧业务描述',
                        }
                    },
                },
            },
            headers=self.admin_headers,
        )
        self.assertEqual(save_resp.status_code, 200, save_resp.text)

        report_save_resp = self.client.put(
            f'/api/systems/{system_id}/grading-report',
            json={'content': {'carried_business': '旧业务描述', 'security_responsibility': '旧责任描述'}},
            headers=self.admin_headers,
        )
        self.assertEqual(report_save_resp.status_code, 200, report_save_resp.text)

        excel_bytes = self.build_customer_filing_excel(
            org_name='Excel覆盖单位',
            credit_code='91350100M000100Y98',
            system_name='Excel覆盖系统',
            level=2,
            business_description='新业务描述',
            cybersecurity_dept='新网安部门',
        )
        preview_resp = self.client.post(
            f'/api/filing-workspace/systems/{system_id}/import-excel/preview',
            files={'file': ('customer.xlsx', excel_bytes, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')},
            headers=self.admin_headers,
        )
        self.assertEqual(preview_resp.status_code, 200, preview_resp.text)
        preview_data = preview_resp.json()['data']
        conflict_keys = {item['key'] for item in preview_data['conflicts']}
        self.assertIn('organization.cybersecurity_dept', conflict_keys)
        self.assertIn('system.filing_detail.table2.business_description', conflict_keys)
        self.assertIn('grading_report_content.carried_business', conflict_keys)

        apply_resp = self.client.post(
            f'/api/filing-workspace/systems/{system_id}/import-excel/apply',
            json={
                'candidate': preview_data['candidate'],
                'override_keys': ['organization.cybersecurity_dept'],
            },
            headers=self.admin_headers,
        )
        self.assertEqual(apply_resp.status_code, 200, apply_resp.text)

        detail_resp = self.client.get(f'/api/filing-workspace/systems/{system_id}', headers=self.admin_headers)
        self.assertEqual(detail_resp.status_code, 200, detail_resp.text)
        detail = detail_resp.json()
        self.assertEqual(detail['organization']['cybersecurity_dept'], '新网安部门')
        self.assertEqual(detail['system']['business_description'], '旧业务描述')
        self.assertEqual(detail['system']['filing_detail']['table2']['business_description'], '旧业务描述')

        report_resp = self.client.get(f'/api/systems/{system_id}/grading-report', headers=self.admin_headers)
        self.assertEqual(report_resp.status_code, 200, report_resp.text)
        self.assertEqual(report_resp.json()['data']['content']['carried_business'], '旧业务描述')

    def test_02b_alerts_summary_counts_pending_and_recycle_items(self):
        org_pending_resp = self.client.post('/api/organizations', json={
            'name': '提醒待处理单位',
            'credit_code': '91350100M000100Y93',
            'legal_representative': '提醒人',
            'address': '太原市',
            'mobile_phone': '13900139003',
            'email': 'alert@example.com',
            'industry': '能源',
            'organization_type': '企业',
            'filing_region': '太原',
            'created_by': 'tester',
        })
        self.assertEqual(org_pending_resp.status_code, 200, org_pending_resp.text)
        pending_org_id = org_pending_resp.json()['data']['id']
        org_del = self.client.post(f'/api/organizations/{pending_org_id}/delete-request?actor=tester&reason=提醒测试', headers=self.admin_headers)
        self.assertEqual(org_del.status_code, 200, org_del.text)

        org_with_system_resp = self.client.post('/api/organizations', json={
            'name': '提醒待处理系统所属单位',
            'credit_code': '91350100M000100Y95',
            'legal_representative': '提醒人',
            'address': '太原市',
            'mobile_phone': '13900139005',
            'email': 'alert_system@example.com',
            'industry': '能源',
            'organization_type': '企业',
            'filing_region': '太原',
            'created_by': 'tester',
        })
        self.assertEqual(org_with_system_resp.status_code, 200, org_with_system_resp.text)
        org_with_system_id = org_with_system_resp.json()['data']['id']

        sys_pending_resp = self.client.post('/api/systems', json={
            'organization_id': org_with_system_id,
            'system_name': '提醒待处理系统',
            'proposed_level': 2,
            'created_by': 'tester',
        })
        self.assertEqual(sys_pending_resp.status_code, 200, sys_pending_resp.text)
        pending_system_id = sys_pending_resp.json()['data']['id']

        sys_del = self.client.post(f'/api/systems/{pending_system_id}/delete-request?actor=tester&reason=提醒测试', headers=self.admin_headers)
        self.assertEqual(sys_del.status_code, 200, sys_del.text)

        org_recycle_resp = self.client.post('/api/organizations', json={
            'name': '提醒回收站单位',
            'credit_code': '91350100M000100Y94',
            'legal_representative': '提醒人',
            'address': '太原市',
            'mobile_phone': '13900139004',
            'email': 'alert_recycle@example.com',
            'industry': '能源',
            'organization_type': '企业',
            'filing_region': '太原',
            'created_by': 'tester',
        })
        self.assertEqual(org_recycle_resp.status_code, 200, org_recycle_resp.text)
        recycle_org_id = org_recycle_resp.json()['data']['id']

        sys_recycle_resp = self.client.post('/api/systems', json={
            'organization_id': recycle_org_id,
            'system_name': '提醒回收站系统',
            'proposed_level': 2,
            'created_by': 'tester',
        })
        self.assertEqual(sys_recycle_resp.status_code, 200, sys_recycle_resp.text)
        recycle_system_id = sys_recycle_resp.json()['data']['id']

        recycle_sys_delete = self.client.delete(f'/api/systems/{recycle_system_id}', headers=self.admin_headers)
        self.assertEqual(recycle_sys_delete.status_code, 200, recycle_sys_delete.text)
        recycle_org_delete = self.client.delete(f'/api/organizations/{recycle_org_id}', headers=self.admin_headers)
        self.assertEqual(recycle_org_delete.status_code, 200, recycle_org_delete.text)

        summary_resp = self.client.get('/api/alerts/summary', headers=self.admin_headers)
        self.assertEqual(summary_resp.status_code, 200, summary_resp.text)
        summary = summary_resp.json()
        self.assertGreaterEqual(summary['pending_org_delete_count'], 1)
        self.assertGreaterEqual(summary['pending_sys_delete_count'], 1)
        self.assertGreaterEqual(summary['org_recycle_count'], 1)
        self.assertGreaterEqual(summary['sys_recycle_count'], 1)
        self.assertTrue(summary['items'])
        self.assertEqual(len(summary.get('entries') or []), 4)
        self.assertTrue(all((row.get('href') or '').startswith('/organizations?attention=') for row in summary['entries']))

    def test_02c_grading_and_expert_review_save_persist_content(self):
        org_resp = self.client.post('/api/organizations', json={
            'name': '表单保存测试单位',
            'credit_code': '91350100M000100Y81',
            'legal_representative': '保存人',
            'address': '太原市小店区',
            'mobile_phone': '13900139081',
            'email': 'form-save@example.com',
            'industry': '制造业',
            'organization_type': '企业',
            'filing_region': '太原',
            'created_by': 'tester',
        })
        self.assertEqual(org_resp.status_code, 200, org_resp.text)
        org_id = org_resp.json()['data']['id']

        sys_resp = self.client.post('/api/systems', json={
            'organization_id': org_id,
            'system_name': '表单保存测试系统',
            'business_category': '生产控制',
            'service_type': '内部服务',
            'deployment_type': '本地部署',
            'proposed_level': 3,
            'created_by': 'tester',
        }, headers=self.admin_headers)
        self.assertEqual(sys_resp.status_code, 200, sys_resp.text)
        system_id = sys_resp.json()['data']['id']

        grading_save_resp = self.client.put(
            f'/api/systems/{system_id}/grading-report',
            json={'content': {'responsible_subject': '责任主体A', 'carried_business': '承载业务B'}},
            headers=self.admin_headers,
        )
        self.assertEqual(grading_save_resp.status_code, 200, grading_save_resp.text)
        self.assertEqual(grading_save_resp.json().get('message'), '已保存')

        grading_get_resp = self.client.get(f'/api/systems/{system_id}/grading-report', headers=self.admin_headers)
        self.assertEqual(grading_get_resp.status_code, 200, grading_get_resp.text)
        grading_content = grading_get_resp.json()['data']['content']
        self.assertEqual(grading_content['responsible_subject'], '责任主体A')
        self.assertEqual(grading_content['carried_business'], '承载业务B')

        expert_save_resp = self.client.put(
            f'/api/systems/{system_id}/expert-review?variant=city',
            json={'content': {'unit_name': '单位甲', 'review_opinion': '评审意见乙'}},
            headers=self.admin_headers,
        )
        self.assertEqual(expert_save_resp.status_code, 200, expert_save_resp.text)
        self.assertEqual(expert_save_resp.json().get('message'), '已保存')

        expert_get_resp = self.client.get(
            f'/api/systems/{system_id}/expert-review?variant=city',
            headers=self.admin_headers,
        )
        self.assertEqual(expert_get_resp.status_code, 200, expert_get_resp.text)
        expert_content = expert_get_resp.json()['data']['content']
        self.assertEqual(expert_content['unit_name'], '单位甲')
        self.assertEqual(expert_content['review_opinion'], '评审意见乙')

    def test_03_word_import_and_collection_review(self):
        doc = Document()
        doc.add_paragraph('单位名称: 客户填报单位')
        doc.add_paragraph('统一社会信用代码: 91350100M000100Y45')
        doc.add_paragraph('单位负责人: 赵六')
        doc.add_paragraph('单位地址: 上海市浦东新区')
        doc.add_paragraph('移动电话: 13700137000')
        doc.add_paragraph('邮箱: d@example.com')
        doc.add_paragraph('所属行业: 能源')
        doc.add_paragraph('单位类型: 企业')
        doc.add_paragraph('备案地区: 上海')
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        files = {'file': ('org.docx', bio.getvalue(), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        import_resp = self.client.post('/api/organizations/import/word?actor=tester', files=files)
        self.assertEqual(import_resp.status_code, 200)
        imported_org_id = import_resp.json()['data']['id']

        link_resp = self.client.post(
            f'/api/organizations/collection-links?organization_id={imported_org_id}&expires_days=7&actor=tester',
            headers=self.admin_headers,
        )
        self.assertEqual(link_resp.status_code, 200)
        token = link_resp.json()['data']['token']

        submit_payload = {
            'address': '上海市黄浦区',
            'mobile_phone': '13600136000',
            'email': 'new@example.com',
            'submitter': '客户A',
        }
        submit_resp = self.client.post(f'/api/public/organizations/collect/{token}', json=submit_payload)
        self.assertEqual(submit_resp.status_code, 200)
        submission_id = submit_resp.json()['data']['submission_id']

        review_resp = self.client.post(
            f'/api/organizations/submissions/{submission_id}/review?action=approve&actor=tester',
            headers=self.admin_headers,
        )
        self.assertEqual(review_resp.status_code, 200)
        self.assertEqual(review_resp.json()['status'], 'approved')

    def test_04_system_word_import_and_pdf_encrypt_export(self):
        org_payload = {
            'name': '系统导入单位',
            'credit_code': '91350100M000100Y46',
            'legal_representative': '周七',
            'address': '广州市',
            'mobile_phone': '13500135000',
            'email': 'sys@example.com',
            'industry': '交通',
            'organization_type': '企业',
            'filing_region': '广州',
            'created_by': 'tester',
        }
        org_resp = self.client.post('/api/organizations', json=org_payload)
        self.assertEqual(org_resp.status_code, 200)
        org_id = org_resp.json()['data']['id']

        doc = Document()
        doc.add_paragraph('系统名称: 轨道调度系统')
        doc.add_paragraph(f'单位ID: {org_id}')
        doc.add_paragraph('拟定等级: 3级')
        doc.add_paragraph('部署方式: 混合部署')
        doc.add_paragraph('系统类型: 生产控制')
        doc.add_paragraph('上线时间: 2025-01-01')
        doc.add_paragraph('定级依据: 按照指南执行')
        doc.add_paragraph('定级理由: 影响范围较大')
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        files = {'file': ('sys.docx', bio.getvalue(), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        import_resp = self.client.post('/api/systems/import/word?actor=tester', files=files, headers=self.admin_headers)
        self.assertEqual(import_resp.status_code, 200)
        system_id = import_resp.json()['data']['id']

        report_resp = self.client.post(f'/api/reports/generate?system_id={system_id}&report_type=grading_report&actor=tester')
        self.assertEqual(report_resp.status_code, 200)
        report_id = report_resp.json()['data']['id']

        pdf_resp = self.client.get(f'/api/reports/{report_id}/export/pdf?password=123456')
        self.assertEqual(pdf_resp.status_code, 200)
        self.assertIn('application/pdf', pdf_resp.headers.get('content-type', ''))

    def test_05_report_sections_and_dashboard_export(self):
        org_payload = {
            'name': '\u7ae0\u8282\u6d4b\u8bd5\u5355\u4f4d',
            'credit_code': '91350100M000100Y47',
            'legal_representative': '\u94b1\u516b',
            'address': '\u6210\u90fd\u5e02',
            'mobile_phone': '13400134000',
            'email': 'sec@example.com',
            'industry': '\u533b\u7597',
            'organization_type': '\u4e8b\u4e1a\u5355\u4f4d',
            'filing_region': '\u6210\u90fd',
            'created_by': 'tester',
        }
        org_resp = self.client.post('/api/organizations', json=org_payload)
        self.assertEqual(org_resp.status_code, 200)
        org_id = org_resp.json()['data']['id']

        sys_payload = {
            'organization_id': org_id,
            'system_name': '\u7ae0\u8282\u7ba1\u7406\u7cfb\u7edf',
            'proposed_level': 2,
            'created_by': 'tester',
        }
        sys_resp = self.client.post('/api/systems', json=sys_payload)
        self.assertEqual(sys_resp.status_code, 200)
        system_id = sys_resp.json()['data']['id']

        report_resp = self.client.post(f'/api/reports/generate?system_id={system_id}&report_type=grading_report&actor=tester')
        self.assertEqual(report_resp.status_code, 200)
        report_id = report_resp.json()['data']['id']

        add_section_resp = self.client.post(
            f'/api/reports/{report_id}/sections?actor=tester',
            json={'name': '\u8865\u5145\u7ae0\u8282', 'content': {'\u8bf4\u660e': '\u81ea\u52a8\u5316\u6d4b\u8bd5'}},
        )
        self.assertEqual(add_section_resp.status_code, 200)

        reorder_resp = self.client.post(
            f'/api/reports/{report_id}/sections/reorder?from_index=0&to_index=1&actor=tester',
            headers=self.admin_headers,
        )
        self.assertEqual(reorder_resp.status_code, 200)

        drill_resp = self.client.get('/api/dashboard/drilldown?dimension=region&value=\u6210\u90fd')
        self.assertEqual(drill_resp.status_code, 200)
        self.assertIn('organizations', drill_resp.json())

        excel_resp = self.client.get('/api/dashboard/export/excel?city=\u6210\u90fd')
        self.assertEqual(excel_resp.status_code, 200)
        self.assertIn('application/vnd.openxmlformats', excel_resp.headers.get('content-type', ''))

        pdf_resp = self.client.get('/api/dashboard/export/pdf?city=\u6210\u90fd')
        if getattr(self.__class__.main_module, 'HAS_REPORTLAB', False):
            self.assertEqual(pdf_resp.status_code, 200, pdf_resp.text)
            self.assertIn('application/pdf', pdf_resp.headers.get('content-type', ''))
        else:
            self.assertEqual(pdf_resp.status_code, 503, pdf_resp.text)
            self.assertIn('reportlab', str(pdf_resp.json().get('detail', '')))

    def test_06_workflow_rules_and_knowledge_versioning(self):
        org_payload = {
            'name': '流程测试单位',
            'credit_code': '91350100M000100Y48',
            'legal_representative': '孙九',
            'address': '杭州市',
            'mobile_phone': '13300133000',
            'email': 'wf@example.com',
            'industry': '教育',
            'organization_type': '事业单位',
            'filing_region': '杭州',
            'created_by': 'tester',
        }
        org_resp = self.client.post('/api/organizations', json=org_payload)
        self.assertEqual(org_resp.status_code, 200)
        org_id = org_resp.json()['data']['id']

        sys_payload = {
            'organization_id': org_id,
            'system_name': '流程时限系统',
            'proposed_level': 2,
            'created_by': 'tester',
        }
        sys_resp = self.client.post('/api/systems', json=sys_payload)
        self.assertEqual(sys_resp.status_code, 200)
        system_id = sys_resp.json()['data']['id']

        rule_payload = {
            'updated_by': 'admin',
            'rules': [
                {'step_name': '信息收集', 'owner': 'collector', 'time_limit_hours': 1, 'enabled': True},
                {'step_name': '信息审核', 'owner': 'reviewer', 'time_limit_hours': 2, 'enabled': True},
            ],
        }
        rule_resp = self.client.put('/api/workflow/rules', json=rule_payload, headers=self.admin_headers)
        self.assertEqual(rule_resp.status_code, 200)

        instance_resp = self.client.get(f'/api/workflow/instances/{system_id}')
        self.assertEqual(instance_resp.status_code, 200)
        self.assertIn('due_at', instance_resp.json())

        remind_resp = self.client.get('/api/workflow/reminders?mode=due&within_hours=24&send=true')
        self.assertEqual(remind_resp.status_code, 200)
        self.assertIn('items', remind_resp.json())

        files = {'file': ('a.docx', b'v1', 'application/octet-stream')}
        upload_resp = self.client.post(
            '/api/knowledge/upload',
            data={'title': '模板A', 'doc_type': '备案表模板', 'actor': 'admin'},
            files=files,
            headers=self.admin_headers,
        )
        self.assertEqual(upload_resp.status_code, 200)
        doc_id = upload_resp.json()['data']['id']

        pin_resp = self.client.post(f'/api/knowledge/{doc_id}/pin?enabled=true&actor=admin', headers=self.admin_headers)
        self.assertEqual(pin_resp.status_code, 200)

        update_resp = self.client.put(
            f'/api/knowledge/{doc_id}?actor=admin',
            json={'city': '杭州', 'district': '西湖区'},
            headers=self.admin_headers,
        )
        self.assertEqual(update_resp.status_code, 200)

        files_v2 = {'file': ('a_v2.docx', b'v2', 'application/octet-stream')}
        new_version_resp = self.client.post(
            f'/api/knowledge/{doc_id}/new-version',
            data={'actor': 'admin'},
            files=files_v2,
            headers=self.admin_headers,
        )
        self.assertEqual(new_version_resp.status_code, 200)

        versions_resp = self.client.get(f'/api/knowledge/{doc_id}/versions')
        self.assertEqual(versions_resp.status_code, 200)
        self.assertGreaterEqual(versions_resp.json()['total'], 2)

        version_id = versions_resp.json()['items'][-1]['id']
        rollback_resp = self.client.post(
            f'/api/knowledge/{doc_id}/rollback/{version_id}?actor=admin',
            headers=self.admin_headers,
        )
        self.assertEqual(rollback_resp.status_code, 200)

    def test_07_auth_and_template_flow(self):
        create_user_resp = self.client.post(
            '/api/auth/users',
            json={'username': 'template_admin', 'password': 'secret123', 'role': 'admin'},
            headers=self.admin_headers,
        )
        self.assertIn(create_user_resp.status_code, [200, 409])

        local_tpl_docs = [
            ('01-自动测试备案表.docx', '网络安全等级保护备案表'),
            ('02-自动测试定级报告.docx', '网络安全等级保护定级报告'),
            ('03-自动测试专家评审意见表.docx', '专家评审意见表'),
        ]

        def _safe_remove(path: Path) -> None:
            try:
                path.unlink(missing_ok=True)
            except PermissionError:
                pass

        for filename, title in local_tpl_docs:
            p = Path(filename)
            d = Document()
            d.add_paragraph(title)
            t = d.add_table(rows=2, cols=2)
            t.cell(0, 0).text = '单位名称'
            t.cell(0, 1).text = '山西晋深交易有限公司'
            t.cell(1, 0).text = '系统名称'
            t.cell(1, 1).text = '山西省省属企业采购与供应链信息管理系统'
            d.save(p)
            self.addCleanup(_safe_remove, p)

        import_local_resp = self.client.post('/api/templates/import-local-official?actor=admin', headers=self.admin_headers)
        self.assertEqual(import_local_resp.status_code, 200)
        self.assertGreaterEqual(len(import_local_resp.json().get('imported', [])), 3)

        tpl_config = {
            'title': '福州市专家评审意见表模板',
            'sections': [
                {'名称': '基础信息', '内容': {'地区': '福州', '说明': '地市模板测试'}},
                {'名称': '专家结论', '内容': {'意见': ''}},
            ],
            'signatures': {'公司签章': '', '测评师签章': ''},
        }
        tpl_file = {'file': ('expert_fz.docx', b'expert-template', 'application/octet-stream')}
        upload_tpl_resp = self.client.post(
            '/api/templates/upload',
            data={
                'template_name': '福州专家模板',
                'report_type': 'expert_review_form',
                'city': '福州',
                'protection_level': '3',
                'is_default': 'true',
                'config_json': json.dumps(tpl_config, ensure_ascii=False),
            },
            files=tpl_file,
            headers=self.admin_headers,
        )
        self.assertEqual(upload_tpl_resp.status_code, 200)
        template_id = upload_tpl_resp.json()['data']['id']

        list_tpl_resp = self.client.get('/api/templates?report_type=expert_review_form')
        self.assertEqual(list_tpl_resp.status_code, 200)
        self.assertGreaterEqual(list_tpl_resp.json()['total'], 1)

        test_fill_resp = self.client.post(f'/api/templates/{template_id}/test-fill')
        self.assertEqual(test_fill_resp.status_code, 200)
        self.assertIn('preview', test_fill_resp.json())

        org_payload = {
            'name': '模板匹配单位',
            'credit_code': '91350100M000100Y49',
            'legal_representative': '吴十',
            'address': '福州市',
            'mobile_phone': '13200132000',
            'email': 'tplmatch@example.com',
            'industry': '政府',
            'organization_type': '机关单位',
            'filing_region': '福州',
            'created_by': 'tester',
        }
        org_resp = self.client.post('/api/organizations', json=org_payload)
        self.assertEqual(org_resp.status_code, 200)
        org_id = org_resp.json()['data']['id']

        sys_payload = {
            'organization_id': org_id,
            'system_name': '专家模板匹配系统',
            'proposed_level': 3,
            'created_by': 'tester',
        }
        sys_resp = self.client.post('/api/systems', json=sys_payload)
        self.assertEqual(sys_resp.status_code, 200)
        system_id = sys_resp.json()['data']['id']

        report_resp = self.client.post(
            f'/api/reports/generate?system_id={system_id}&report_type=expert_review_form&actor=tester'
        )
        self.assertEqual(report_resp.status_code, 200)
        content = report_resp.json()['data']['content']
        self.assertIn('模板信息', content)
        self.assertEqual(content['模板信息']['template_id'], template_id)

    def test_08_recycle_bin_and_report_compare(self):
        org_payload = {
            'name': '回收站测试单位',
            'credit_code': '91350100M000100Y50',
            'legal_representative': '郑十一',
            'address': '南京市',
            'mobile_phone': '13100131000',
            'email': 'recycle@example.com',
            'industry': '企业',
            'organization_type': '企业',
            'filing_region': '南京',
            'created_by': 'tester',
        }
        org_resp = self.client.post('/api/organizations', json=org_payload)
        self.assertEqual(org_resp.status_code, 200)
        org_id = org_resp.json()['data']['id']

        sys_payload = {
            'organization_id': org_id,
            'system_name': '回收站系统',
            'proposed_level': 2,
            'created_by': 'tester',
        }
        sys_resp = self.client.post('/api/systems', json=sys_payload)
        self.assertEqual(sys_resp.status_code, 200)
        system_id = sys_resp.json()['data']['id']

        delete_org_fail = self.client.delete(
            f'/api/organizations/{org_id}?actor=admin&is_admin=true',
            headers=self.admin_headers,
        )
        self.assertEqual(delete_org_fail.status_code, 400)
        self.assertIn('关联系统', str(delete_org_fail.json().get('detail', '')))

        delete_sys_ok = self.client.delete(
            f'/api/systems/{system_id}?actor=admin&is_admin=true',
            headers=self.admin_headers,
        )
        self.assertEqual(delete_sys_ok.status_code, 200)

        sys_recycle = self.client.get('/api/systems/recycle-bin/list')
        self.assertEqual(sys_recycle.status_code, 200)
        self.assertTrue(any(i['id'] == system_id for i in sys_recycle.json()['items']))

        restore_sys_ok = self.client.post(f'/api/systems/{system_id}/restore?actor=tester', headers=self.admin_headers)
        self.assertEqual(restore_sys_ok.status_code, 200)

        report1_resp = self.client.post(
            f'/api/reports/generate?system_id={system_id}&report_type=grading_report&actor=tester'
        )
        self.assertEqual(report1_resp.status_code, 200)
        report1_id = report1_resp.json()['data']['id']

        content_resp = self.client.get(f'/api/reports/{report1_id}')
        self.assertEqual(content_resp.status_code, 200)
        edited_content = content_resp.json()['content']
        edited_content['备注'] = '版本1修改内容'
        edit_resp = self.client.put(
            f'/api/reports/{report1_id}?actor=tester',
            json={'content': edited_content},
        )
        self.assertEqual(edit_resp.status_code, 200)

        report2_resp = self.client.post(
            f'/api/reports/generate?system_id={system_id}&report_type=grading_report&actor=tester'
        )
        self.assertEqual(report2_resp.status_code, 200)
        report2_id = report2_resp.json()['data']['id']

        compare_resp = self.client.get(f'/api/reports/{report1_id}/compare/{report2_id}')
        self.assertEqual(compare_resp.status_code, 200)
        self.assertIn('summary', compare_resp.json())
        self.assertGreaterEqual(compare_resp.json()['summary']['field_change_count'], 1)

        delete_sys_blocked = self.client.post(f'/api/systems/{system_id}/delete-request?actor=tester')
        self.assertEqual(delete_sys_blocked.status_code, 400)
        self.assertIn('关联报告', str(delete_sys_blocked.json().get('detail', '')))

    def test_09_delete_request_attachment_batch_and_knowledge_exact(self):
        org_payload = {
            'name': '删除申请单位',
            'credit_code': '91350100M000100Y51',
            'legal_representative': '王十二',
            'address': '厦门市',
            'mobile_phone': '13000130000',
            'email': 'delreq@example.com',
            'industry': '企业',
            'organization_type': '企业',
            'filing_region': '厦门',
            'created_by': 'tester',
        }
        org_resp = self.client.post('/api/organizations', json=org_payload)
        self.assertEqual(org_resp.status_code, 200)
        org_id = org_resp.json()['data']['id']

        req_resp = self.client.post(
            f'/api/organizations/{org_id}/delete-request?actor=tester&reason=重复录入',
        )
        self.assertEqual(req_resp.status_code, 200)
        request_id = req_resp.json()['data']['request_id']

        review_resp = self.client.post(
            f'/api/delete-requests/{request_id}/review?action=approve&actor=admin',
            headers=self.admin_headers,
        )
        self.assertEqual(review_resp.status_code, 200)
        self.assertEqual(review_resp.json()['status'], 'approved')

        recycle_resp = self.client.get('/api/organizations/recycle-bin/list')
        self.assertEqual(recycle_resp.status_code, 200)
        self.assertTrue(any(i['id'] == org_id for i in recycle_resp.json()['items']))

        org2_payload = {
            'name': '附件批量单位',
            'credit_code': '91350100M000100Y52',
            'legal_representative': '赵十三',
            'address': '泉州市',
            'mobile_phone': '13900129000',
            'email': 'att@example.com',
            'industry': '企业',
            'organization_type': '企业',
            'filing_region': '泉州',
            'created_by': 'tester',
        }
        org2_resp = self.client.post('/api/organizations', json=org2_payload)
        self.assertEqual(org2_resp.status_code, 200)
        org2_id = org2_resp.json()['data']['id']

        files = [
            ('files', ('a.pdf', b'%PDF-1.4 test', 'application/pdf')),
            ('files', ('b.png', b'\x89PNG\r\n\x1a\n', 'image/png')),
        ]
        batch_resp = self.client.post(
            f'/api/attachments/organization/{org2_id}/batch?actor=tester',
            files=files,
        )
        self.assertEqual(batch_resp.status_code, 200)
        self.assertEqual(batch_resp.json()['uploaded'], 2)

        list_att = self.client.get(f'/api/attachments/organization/{org2_id}')
        self.assertEqual(list_att.status_code, 200)
        self.assertGreaterEqual(list_att.json()['total'], 2)
        attachment_id = list_att.json()['items'][0]['id']

        preview_resp = self.client.get(f'/api/attachment-files/{attachment_id}/preview')
        self.assertEqual(preview_resp.status_code, 200)
        self.assertIn('content-type', preview_resp.headers)

        upload_resp = self.client.post(
            '/api/knowledge/upload',
            data={'title': '精确检索文档', 'doc_type': '政策文件', 'actor': 'admin', 'keywords': '精确关键字'},
            files={'file': ('exact.docx', b'content', 'application/octet-stream')},
            headers=self.admin_headers,
        )
        self.assertEqual(upload_resp.status_code, 200)
        list_exact = self.client.get('/api/knowledge?keyword=精确关键字&match_mode=exact')
        self.assertEqual(list_exact.status_code, 200)
        self.assertTrue(any(i['title'] == '精确检索文档' for i in list_exact.json()['items']))

    def test_09a_admin_can_purge_recycle_org_and_release_credit_code(self):
        org_resp = self.client.post(
            '/api/organizations',
            json={
                'name': '永久删除级联单位',
                'credit_code': '91350100M000100Y96',
                'legal_representative': '测试人',
                'address': '太原市迎泽区',
                'mobile_phone': '13100131996',
                'email': 'purge-org@example.com',
                'industry': '企业',
                'organization_type': '企业',
                'filing_region': '太原',
                'created_by': 'tester',
            },
            headers=self.admin_headers,
        )
        self.assertEqual(org_resp.status_code, 200, org_resp.text)
        org_id = org_resp.json()['data']['id']

        sys_resp = self.client.post(
            '/api/systems',
            json={
                'organization_id': org_id,
                'system_name': '永久删除级联系统',
                'proposed_level': 2,
                'created_by': 'tester',
            },
            headers=self.admin_headers,
        )
        self.assertEqual(sys_resp.status_code, 200, sys_resp.text)
        system_id = sys_resp.json()['data']['id']

        report_resp = self.client.post(
            f'/api/reports/generate?system_id={system_id}&report_type=grading_report&actor=tester',
            headers=self.admin_headers,
        )
        self.assertEqual(report_resp.status_code, 200, report_resp.text)

        delete_sys_resp = self.client.delete(f'/api/systems/{system_id}', headers=self.admin_headers)
        self.assertEqual(delete_sys_resp.status_code, 200, delete_sys_resp.text)

        delete_org_resp = self.client.delete(f'/api/organizations/{org_id}', headers=self.admin_headers)
        self.assertEqual(delete_org_resp.status_code, 200, delete_org_resp.text)

        purge_resp = self.client.post(f'/api/organizations/{org_id}/purge', headers=self.admin_headers)
        self.assertEqual(purge_resp.status_code, 200, purge_resp.text)
        self.assertIn('级联清理', purge_resp.json().get('message', ''))

        db = self.__class__.db_module.SessionLocal()
        try:
            org = db.query(self.__class__.main_module.Organization).filter_by(id=org_id).first()
            system = db.query(self.__class__.main_module.SystemInfo).filter_by(id=system_id).first()
            reports = db.query(self.__class__.main_module.Report).filter_by(system_id=system_id).count()
            workflow_instances = db.query(self.__class__.main_module.WorkflowInstance).filter_by(system_id=system_id).count()
            self.assertIsNone(org)
            self.assertIsNone(system)
            self.assertEqual(reports, 0)
            self.assertEqual(workflow_instances, 0)
        finally:
            db.close()

        recreate_resp = self.client.post(
            '/api/organizations',
            json={
                'name': '永久删除后重建单位',
                'credit_code': '91350100M000100Y96',
                'legal_representative': '测试人',
                'address': '太原市迎泽区',
                'mobile_phone': '13100131997',
                'email': 'purge-org-recreate@example.com',
                'industry': '企业',
                'organization_type': '企业',
                'filing_region': '太原',
                'created_by': 'tester',
            },
            headers=self.admin_headers,
        )
        self.assertEqual(recreate_resp.status_code, 200, recreate_resp.text)

    def test_09b_non_admin_cannot_purge_recycle_org(self):
        create_user_resp = self.client.post(
            '/api/auth/users',
            json={
                'username': 'purge_eval_user',
                'password': 'purgeEval123',
                'role': 'evaluator',
                'require_password_change': False,
            },
            headers=self.admin_headers,
        )
        self.assertIn(create_user_resp.status_code, [200, 409], create_user_resp.text)
        login_resp = self.client.post('/api/auth/login', json={'username': 'purge_eval_user', 'password': 'purgeEval123'})
        self.assertEqual(login_resp.status_code, 200, login_resp.text)
        eval_headers = {'X-Auth-Token': login_resp.json()['token']}

        org_resp = self.client.post(
            '/api/organizations',
            json={
                'name': '非管理员永久删除限制单位',
                'credit_code': '91350100M000100Y97',
                'legal_representative': '测试人',
                'address': '限制城',
                'mobile_phone': '13100131998',
                'email': 'purge-limit@example.com',
                'industry': '企业',
                'organization_type': '企业',
                'filing_region': '限制城',
                'created_by': 'tester',
            },
            headers=self.admin_headers,
        )
        self.assertEqual(org_resp.status_code, 200, org_resp.text)
        org_id = org_resp.json()['data']['id']

        delete_org_resp = self.client.delete(f'/api/organizations/{org_id}', headers=self.admin_headers)
        self.assertEqual(delete_org_resp.status_code, 200, delete_org_resp.text)

        purge_resp = self.client.post(f'/api/organizations/{org_id}/purge', headers=eval_headers)
        self.assertEqual(purge_resp.status_code, 403, purge_resp.text)

    def test_10_permission_boundaries(self):
        saved_cookies = dict(self.client.cookies)
        self.client.cookies.clear()
        unauth_collection = self.client.get('/api/organizations/collection-links')
        self.assertEqual(unauth_collection.status_code, 401)
        self.client.cookies.update(saved_cookies)

        create_user_resp = self.client.post(
            '/api/auth/users',
            json={'username': 'ev_user', 'password': 'evpass123', 'role': 'evaluator', 'require_password_change': False},
            headers=self.admin_headers,
        )
        self.assertIn(create_user_resp.status_code, [200, 409])

        evaluator_login = self.client.post('/api/auth/login', json={'username': 'ev_user', 'password': 'evpass123'})
        self.assertEqual(evaluator_login.status_code, 200)
        evaluator_headers = {'X-Auth-Token': evaluator_login.json()['token']}

        org_payload = {
            'name': '权限边界单位',
            'credit_code': '91350100M000100Y54',
            'legal_representative': '周十',
            'address': '福州市',
            'mobile_phone': '13100131000',
            'email': 'perm@example.com',
            'industry': '政府',
            'organization_type': '机关单位',
            'filing_region': '福州',
            'created_by': 'tester',
        }
        org_resp = self.client.post('/api/organizations', json=org_payload)
        self.assertEqual(org_resp.status_code, 200)
        org_id = org_resp.json()['data']['id']

        system_payload = {
            'organization_id': org_id,
            'system_name': '权限边界系统',
            'proposed_level': 3,
            'created_by': 'tester',
        }
        sys_resp = self.client.post('/api/systems', json=system_payload)
        self.assertEqual(sys_resp.status_code, 200)
        system_id = sys_resp.json()['data']['id']

        report_resp = self.client.post(f'/api/reports/generate?system_id={system_id}&report_type=grading_report&actor=tester')
        self.assertEqual(report_resp.status_code, 200)
        report_id = report_resp.json()['data']['id']

        submit_resp = self.client.post(
            f'/api/reports/{report_id}/submit?actor=tester&reviewer=leader',
            headers=self.admin_headers,
        )
        self.assertEqual(submit_resp.status_code, 200)

        eval_review_resp = self.client.post(
            f'/api/reports/{report_id}/review?actor=ev_user&action=approve&comment=越权审核测试',
            headers=evaluator_headers,
        )
        self.assertEqual(eval_review_resp.status_code, 403)

        evaluator_collection = self.client.get('/api/organizations/collection-links', headers=evaluator_headers)
        self.assertEqual(evaluator_collection.status_code, 200)

    def test_11_change_password_revokes_old_sessions(self):
        create_user_resp = self.client.post(
            '/api/auth/users',
            json={'username': 'pw_rotate_user', 'password': 'oldpass123', 'role': 'evaluator'},
            headers=self.admin_headers,
        )
        self.assertIn(create_user_resp.status_code, [200, 409])

        login_old = self.client.post('/api/auth/login', json={'username': 'pw_rotate_user', 'password': 'oldpass123'})
        self.assertEqual(login_old.status_code, 200)
        old_token = login_old.json()['token']
        old_headers = {'X-Auth-Token': old_token}

        me_before = self.client.get('/api/auth/me', headers=old_headers)
        self.assertEqual(me_before.status_code, 200)

        change_resp = self.client.post(
            '/api/auth/change-password',
            json={'current_password': 'oldpass123', 'new_password': 'newpass123'},
            headers=old_headers,
        )
        self.assertEqual(change_resp.status_code, 200)

        me_after = self.client.get('/api/auth/me', headers=old_headers)
        self.assertEqual(me_after.status_code, 401)

        login_old_again = self.client.post('/api/auth/login', json={'username': 'pw_rotate_user', 'password': 'oldpass123'})
        self.assertEqual(login_old_again.status_code, 401)

        login_new = self.client.post('/api/auth/login', json={'username': 'pw_rotate_user', 'password': 'newpass123'})
        self.assertEqual(login_new.status_code, 200)

    def test_12_must_change_password_is_enforced_server_side(self):
        create_user_resp = self.client.post(
            '/api/auth/users',
            json={'username': 'must_change_user', 'password': 'temp12345', 'role': 'evaluator'},
            headers=self.admin_headers,
        )
        self.assertIn(create_user_resp.status_code, [200, 409])

        login_resp = self.client.post('/api/auth/login', json={'username': 'must_change_user', 'password': 'temp12345'})
        self.assertEqual(login_resp.status_code, 200)
        self.assertTrue(login_resp.json().get('must_change_password'))
        old_token = login_resp.json()['token']
        old_headers = {'X-Auth-Token': old_token}

        me_resp = self.client.get('/api/auth/me', headers=old_headers)
        self.assertEqual(me_resp.status_code, 200)
        self.assertTrue(me_resp.json().get('must_change_password'))

        blocked_resp = self.client.get('/api/organizations/collection-links', headers=old_headers)
        self.assertEqual(blocked_resp.status_code, 403)

        change_resp = self.client.post(
            '/api/auth/change-password',
            json={'current_password': 'temp12345', 'new_password': 'temp12345_new'},
            headers=old_headers,
        )
        self.assertEqual(change_resp.status_code, 200)

        old_token_resp = self.client.get('/api/auth/me', headers=old_headers)
        self.assertEqual(old_token_resp.status_code, 401)

        login_new = self.client.post('/api/auth/login', json={'username': 'must_change_user', 'password': 'temp12345_new'})
        self.assertEqual(login_new.status_code, 200)
        self.assertFalse(login_new.json().get('must_change_password'))
        new_headers = {'X-Auth-Token': login_new.json()['token']}

        allowed_resp = self.client.get('/api/organizations/collection-links', headers=new_headers)
        self.assertEqual(allowed_resp.status_code, 200)

    def test_13_require_password_change_explicit_bool_parse(self):
        create_false_resp = self.client.post(
            '/api/auth/users',
            json={
                'username': 'bool_parse_user',
                'password': 'boolpass123',
                'role': 'evaluator',
                'require_password_change': 'false',
            },
            headers=self.admin_headers,
        )
        self.assertEqual(create_false_resp.status_code, 200)
        self.assertFalse(create_false_resp.json()['data']['must_change_password'])

        invalid_resp = self.client.post(
            '/api/auth/users',
            json={
                'username': 'bool_parse_user_bad',
                'password': 'boolpass123',
                'role': 'evaluator',
                'require_password_change': 'not-bool',
            },
            headers=self.admin_headers,
        )
        self.assertEqual(invalid_resp.status_code, 400)

    def test_14_is_admin_param_cannot_bypass_lock_and_report_freeze(self):
        create_eval_resp = self.client.post(
            '/api/auth/users',
            json={'username': 'non_admin_editor', 'password': 'editor123', 'role': 'evaluator', 'require_password_change': False},
            headers=self.admin_headers,
        )
        self.assertIn(create_eval_resp.status_code, [200, 409])
        eval_login = self.client.post('/api/auth/login', json={'username': 'non_admin_editor', 'password': 'editor123'})
        self.assertEqual(eval_login.status_code, 200)
        eval_headers = {'X-Auth-Token': eval_login.json()['token']}

        org_resp = self.client.post(
            '/api/organizations',
            json={
                'name': '权限绕过修复单位',
                'credit_code': '91350100M000100Y55',
                'legal_representative': '赵一',
                'address': '修复城A',
                'mobile_phone': '13100131001',
                'email': 'permfix@example.com',
                'industry': '政府',
                'organization_type': '机关单位',
                'filing_region': '修复城A',
                'created_by': 'tester',
            },
        )
        self.assertEqual(org_resp.status_code, 200)
        org_id = org_resp.json()['data']['id']

        sys_resp = self.client.post(
            '/api/systems',
            json={
                'organization_id': org_id,
                'system_name': '权限绕过修复系统',
                'proposed_level': 3,
                'created_by': 'tester',
            },
        )
        self.assertEqual(sys_resp.status_code, 200)
        system_id = sys_resp.json()['data']['id']

        report_resp = self.client.post(f'/api/reports/generate?system_id={system_id}&report_type=grading_report&actor=tester')
        self.assertEqual(report_resp.status_code, 200)
        report_id = report_resp.json()['data']['id']

        submit_resp = self.client.post(
            f'/api/reports/{report_id}/submit?actor=tester&reviewer=leader',
            headers=self.admin_headers,
        )
        self.assertEqual(submit_resp.status_code, 200)

        bypass_edit_resp = self.client.put(
            f'/api/reports/{report_id}?actor=non_admin_editor&is_admin=true',
            json={'content': {'标题': '尝试绕过'}},
            headers=eval_headers,
        )
        self.assertEqual(bypass_edit_resp.status_code, 403)

        admin_edit_resp = self.client.put(
            f'/api/reports/{report_id}?actor=admin&is_admin=true',
            json={'content': {'标题': '管理员可编辑'}},
            headers=self.admin_headers,
        )
        self.assertEqual(admin_edit_resp.status_code, 200)

        archive_org_resp = self.client.post(f'/api/organizations/{org_id}/archive', headers=self.admin_headers)
        self.assertEqual(archive_org_resp.status_code, 200)
        bypass_org_update_resp = self.client.put(
            f'/api/organizations/{org_id}',
            json={'address': '被绕过地址'},
            headers=eval_headers,
        )
        self.assertEqual(bypass_org_update_resp.status_code, 403)
        admin_org_update_resp = self.client.put(
            f'/api/organizations/{org_id}',
            json={'address': '管理员地址'},
            headers=self.admin_headers,
        )
        self.assertEqual(admin_org_update_resp.status_code, 200)

        archive_sys_resp = self.client.post(f'/api/systems/{system_id}/archive', headers=self.admin_headers)
        self.assertEqual(archive_sys_resp.status_code, 200)
        bypass_sys_update_resp = self.client.put(
            f'/api/systems/{system_id}',
            json={'system_name': '绕过系统名'},
            headers=eval_headers,
        )
        self.assertEqual(bypass_sys_update_resp.status_code, 403)
        admin_sys_update_resp = self.client.put(
            f'/api/systems/{system_id}',
            json={'system_name': '管理员系统名'},
            headers=self.admin_headers,
        )
        self.assertEqual(admin_sys_update_resp.status_code, 200)

    def test_15_batch_download_and_dashboard_filter_consistency(self):
        upload_resp = self.client.post(
            '/api/knowledge/upload',
            data={'title': '批量下载下架校验', 'doc_type': '政策文件', 'actor': 'admin'},
            files={'file': ('disabled.docx', b'content', 'application/octet-stream')},
            headers=self.admin_headers,
        )
        self.assertEqual(upload_resp.status_code, 200)
        disabled_doc_id = upload_resp.json()['data']['id']

        disable_resp = self.client.post(
            f'/api/knowledge/{disabled_doc_id}/toggle?enabled=false&actor=admin',
            headers=self.admin_headers,
        )
        self.assertEqual(disable_resp.status_code, 200)

        batch_resp = self.client.post(
            '/api/knowledge/batch-download?actor=admin',
            json=[disabled_doc_id],
            headers=self.admin_headers,
        )
        self.assertEqual(batch_resp.status_code, 403)

        city_a = '筛选一致城A'
        city_b = '筛选一致城B'
        org_a_resp = self.client.post(
            '/api/organizations',
            json={
                'name': '看板筛选A单位',
                'credit_code': '91350100M000100Y56',
                'legal_representative': '甲',
                'address': city_a,
                'mobile_phone': '13100131002',
                'email': 'citya@example.com',
                'industry': '教育',
                'organization_type': '事业单位',
                'filing_region': city_a,
                'created_by': 'tester',
            },
        )
        self.assertEqual(org_a_resp.status_code, 200)
        org_a_id = org_a_resp.json()['data']['id']

        org_b_resp = self.client.post(
            '/api/organizations',
            json={
                'name': '看板筛选B单位',
                'credit_code': '91350100M000100Y57',
                'legal_representative': '乙',
                'address': city_b,
                'mobile_phone': '13100131003',
                'email': 'cityb@example.com',
                'industry': '教育',
                'organization_type': '事业单位',
                'filing_region': city_b,
                'created_by': 'tester',
            },
        )
        self.assertEqual(org_b_resp.status_code, 200)
        org_b_id = org_b_resp.json()['data']['id']

        sys_a_resp = self.client.post(
            '/api/systems',
            json={'organization_id': org_a_id, 'system_name': '看板系统A', 'proposed_level': 2, 'created_by': 'tester'},
        )
        self.assertEqual(sys_a_resp.status_code, 200)
        sys_a_id = sys_a_resp.json()['data']['id']

        sys_b_resp = self.client.post(
            '/api/systems',
            json={'organization_id': org_b_id, 'system_name': '看板系统B', 'proposed_level': 2, 'created_by': 'tester'},
        )
        self.assertEqual(sys_b_resp.status_code, 200)
        sys_b_id = sys_b_resp.json()['data']['id']

        report_a = self.client.post(f'/api/reports/generate?system_id={sys_a_id}&report_type=grading_report&actor=tester')
        self.assertEqual(report_a.status_code, 200)
        report_b = self.client.post(f'/api/reports/generate?system_id={sys_b_id}&report_type=grading_report&actor=tester')
        self.assertEqual(report_b.status_code, 200)
        report_b_id = report_b.json()['data']['id']

        submit_b = self.client.post(
            f'/api/reports/{report_b_id}/submit?actor=tester&reviewer=leader',
            headers=self.admin_headers,
        )
        self.assertEqual(submit_b.status_code, 200)

        archive_b = self.client.post(f'/api/systems/{sys_b_id}/archive', headers=self.admin_headers)
        self.assertEqual(archive_b.status_code, 200)

        summary_resp = self.client.get(f'/api/dashboard/summary?city={city_a}', headers=self.admin_headers)
        self.assertEqual(summary_resp.status_code, 200)
        totals = summary_resp.json()['totals']
        self.assertEqual(totals['archived_system_count'], 0)
        self.assertEqual(totals['pending_review_reports'], 0)
        self.assertEqual(totals['in_progress_projects'], 1)

    def test_16_lite_mode_can_edit_frozen_report_with_is_admin_flag(self):
        org_resp = self.client.post(
            '/api/organizations',
            json={
                'name': '无鉴权演示单位',
                'credit_code': '91350100M000100Y58',
                'legal_representative': '丙',
                'address': '演示城',
                'mobile_phone': '13100131004',
                'email': 'lite@example.com',
                'industry': '政务',
                'organization_type': '机关单位',
                'filing_region': '演示城',
                'created_by': 'tester',
            },
        )
        self.assertEqual(org_resp.status_code, 200)
        org_id = org_resp.json()['data']['id']

        sys_resp = self.client.post(
            '/api/systems',
            json={'organization_id': org_id, 'system_name': '无鉴权演示系统', 'proposed_level': 2, 'created_by': 'tester'},
        )
        self.assertEqual(sys_resp.status_code, 200)
        system_id = sys_resp.json()['data']['id']

        report_resp = self.client.post(f'/api/reports/generate?system_id={system_id}&report_type=grading_report&actor=tester')
        self.assertEqual(report_resp.status_code, 200)
        report_id = report_resp.json()['data']['id']

        submit_resp = self.client.post(
            f'/api/reports/{report_id}/submit?actor=tester&reviewer=leader',
            headers=self.admin_headers,
        )
        self.assertEqual(submit_resp.status_code, 200)

        lite_edit_resp = self.client.put(
            f'/api/reports/{report_id}?actor=demo_admin',
            json={'content': {'标题': '无鉴权维护'}},
            headers=self.admin_headers,
        )
        self.assertEqual(lite_edit_resp.status_code, 200)

    def test_17_import_excel_flush_error_does_not_break_whole_transaction(self):
        wb = Workbook()
        ws = wb.active
        ws.append(['name', 'credit_code', 'legal_representative', 'address', 'office_phone', 'mobile_phone', 'email', 'industry', 'organization_type', 'filing_region'])
        ws.append(['导入A', '91350100M000100Y59', '张一', 'A市', '', '13100131005', 'a59@example.com', '教育', '事业单位', 'A市'])
        ws.append(['导入B', '91350100M000100Y59', '张二', 'B市', '', '13100131006', 'b59@example.com', '教育', '事业单位', 'B市'])
        bio = io.BytesIO()
        wb.save(bio)
        bio.seek(0)

        original_checker = self.__class__.main_module.assert_credit_code_available
        self.__class__.main_module.assert_credit_code_available = lambda *args, **kwargs: None
        try:
            resp = self.client.post(
                '/api/organizations/import/excel',
                files={'file': ('orgs.xlsx', bio.getvalue(), 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')},
                headers=self.admin_headers,
            )
        finally:
            self.__class__.main_module.assert_credit_code_available = original_checker

        self.assertEqual(resp.status_code, 200)
        data = resp.json()
        self.assertEqual(data['imported'], 1)
        self.assertGreaterEqual(len(data['skipped']), 1)

    def test_18_update_template_empty_required_fields_returns_400(self):
        upload_resp = self.client.post(
            '/api/templates/upload',
            data={
                'template_name': '必填校验模板',
                'report_type': 'expert_review_form',
                'category': '测试',
                'city': '测试城',
                'protection_level': '2',
                'is_default': 'false',
                'config_json': '{}',
            },
            files={'file': ('req_template.docx', b'template', 'application/octet-stream')},
            headers=self.admin_headers,
        )
        self.assertEqual(upload_resp.status_code, 200)
        tpl_id = upload_resp.json()['data']['id']

        bad_status = self.client.put(
            f'/api/templates/{tpl_id}',
            json={'status': ''},
            headers=self.admin_headers,
        )
        self.assertEqual(bad_status.status_code, 400)

        bad_name = self.client.put(
            f'/api/templates/{tpl_id}',
            json={'template_name': ''},
            headers=self.admin_headers,
        )
        self.assertEqual(bad_name.status_code, 400)

    def test_19_generate_system_code_not_reused_without_insert(self):
        db = self.__class__.db_module.SessionLocal()
        try:
            code1 = self.__class__.main_module.generate_system_code(db)
            code2 = self.__class__.main_module.generate_system_code(db)
        finally:
            db.close()
        self.assertNotEqual(code1, code2)

    def test_20_restore_deleted_system_requires_alive_org(self):
        org_resp = self.client.post(
            '/api/organizations',
            json={
                'name': '恢复校验单位',
                'credit_code': '91350100M000100Y60',
                'legal_representative': '丁',
                'address': '恢复城',
                'mobile_phone': '13100131006',
                'email': 'restore60@example.com',
                'industry': '政务',
                'organization_type': '机关单位',
                'filing_region': '恢复城',
                'created_by': 'tester',
            },
        )
        self.assertEqual(org_resp.status_code, 200)
        org_id = org_resp.json()['data']['id']

        sys_resp = self.client.post(
            '/api/systems',
            json={'organization_id': org_id, 'system_name': '恢复校验系统', 'proposed_level': 2, 'created_by': 'tester'},
        )
        self.assertEqual(sys_resp.status_code, 200)
        system_id = sys_resp.json()['data']['id']

        delete_sys_resp = self.client.delete(
            f'/api/systems/{system_id}?actor=admin&is_admin=true',
            headers=self.admin_headers,
        )
        self.assertEqual(delete_sys_resp.status_code, 200)

        delete_org_resp = self.client.delete(
            f'/api/organizations/{org_id}?actor=admin&is_admin=true',
            headers=self.admin_headers,
        )
        self.assertEqual(delete_org_resp.status_code, 200)

        restore_sys_resp = self.client.post(f'/api/systems/{system_id}/restore?actor=tester', headers=self.admin_headers)
        self.assertEqual(restore_sys_resp.status_code, 400)
        self.assertIn('先恢复单位', str(restore_sys_resp.json().get('detail', '')))

    def test_21_system_excel_import_accepts_datetime_cell(self):
        org_resp = self.client.post(
            '/api/organizations',
            json={
                'name': 'Excel日期单位',
                'credit_code': '91350100M000100Y61',
                'legal_representative': '戊',
                'address': 'Excel城',
                'mobile_phone': '13100131007',
                'email': 'excel61@example.com',
                'industry': '教育',
                'organization_type': '事业单位',
                'filing_region': 'Excel城',
                'created_by': 'tester',
            },
        )
        self.assertEqual(org_resp.status_code, 200)
        org_id = org_resp.json()['data']['id']

        wb = Workbook()
        ws = wb.active
        ws.append(['系统名称', '系统编号', '单位ID', '拟定等级', '部署方式', '系统类型', '上线时间', '录入人'])
        ws.append(['Excel日期系统', '', org_id, 3, '混合部署', '管理系统', datetime(2025, 1, 1, 0, 0, 0), 'tester'])
        bio = io.BytesIO()
        wb.save(bio)
        bio.seek(0)

        resp = self.client.post(
            '/api/systems/import/excel',
            files={'file': ('systems_datetime.xlsx', bio.getvalue(), 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')},
            headers=self.admin_headers,
        )
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.json()['imported'], 1)
        self.assertEqual(len(resp.json()['skipped']), 0)

        list_resp = self.client.get('/api/systems?system_name=Excel日期系统')
        self.assertEqual(list_resp.status_code, 200)
        self.assertTrue(
            any((item.get('go_live_date') or '').startswith('2025-01-01') for item in list_resp.json().get('items', []))
        )

    def test_22_system_import_validates_proposed_level_range(self):
        org_resp = self.client.post(
            '/api/organizations',
            json={
                'name': '等级校验单位',
                'credit_code': '91350100M000100Y62',
                'legal_representative': '己',
                'address': '等级城',
                'mobile_phone': '13100131008',
                'email': 'level62@example.com',
                'industry': '医疗',
                'organization_type': '企业',
                'filing_region': '等级城',
                'created_by': 'tester',
            },
        )
        self.assertEqual(org_resp.status_code, 200)
        org_id = org_resp.json()['data']['id']

        wb = Workbook()
        ws = wb.active
        ws.append(['系统名称', '系统编号', '单位ID', '拟定等级', '部署方式', '系统类型', '上线时间', '录入人'])
        ws.append(['非法等级系统', '', org_id, 10, '混合部署', '业务系统', '', 'tester'])
        bio = io.BytesIO()
        wb.save(bio)
        bio.seek(0)

        excel_resp = self.client.post(
            '/api/systems/import/excel',
            files={'file': ('systems_invalid_level.xlsx', bio.getvalue(), 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')},
            headers=self.admin_headers,
        )
        self.assertEqual(excel_resp.status_code, 200)
        self.assertEqual(excel_resp.json()['imported'], 0)
        self.assertTrue(any('1-5级范围' in msg for msg in excel_resp.json()['skipped']))

        doc = Document()
        doc.add_paragraph('系统名称: Word非法等级系统')
        doc.add_paragraph(f'单位ID: {org_id}')
        doc.add_paragraph('拟定等级: 10级')
        bio_doc = io.BytesIO()
        doc.save(bio_doc)
        bio_doc.seek(0)
        word_resp = self.client.post(
            '/api/systems/import/word',
            files={'file': ('sys_invalid_level.docx', bio_doc.getvalue(), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')},
            headers=self.admin_headers,
        )
        self.assertEqual(word_resp.status_code, 400)
        self.assertIn('1-5级范围', str(word_resp.json().get('detail', '')))

    def test_23_knowledge_download_returns_404_when_file_missing(self):
        upload_resp = self.client.post(
            '/api/knowledge/upload',
            data={'title': '缺失文件下载校验', 'doc_type': '政策文件', 'actor': 'admin'},
            files={'file': ('missing.docx', b'missing-content', 'application/octet-stream')},
            headers=self.admin_headers,
        )
        self.assertEqual(upload_resp.status_code, 200)
        doc_id = upload_resp.json()['data']['id']

        db = self.__class__.db_module.SessionLocal()
        try:
            row = db.query(self.__class__.main_module.KnowledgeDocument).filter_by(id=doc_id).first()
            self.assertIsNotNone(row)
            row.file_path = f"{row.file_path}.missing"
            db.commit()
        finally:
            db.close()

        download_resp = self.client.get(f'/api/knowledge/{doc_id}/download', headers=self.admin_headers)
        self.assertEqual(download_resp.status_code, 404)
        self.assertIn('文件不存在', str(download_resp.json().get('detail', '')))

    def test_24_workflow_reminders_reuse_rule_map(self):
        org_resp = self.client.post(
            '/api/organizations',
            json={
                'name': '提醒性能单位',
                'credit_code': '91350100M000100Y63',
                'legal_representative': '庚',
                'address': '提醒城',
                'mobile_phone': '13100131009',
                'email': 'wf63@example.com',
                'industry': '能源',
                'organization_type': '企业',
                'filing_region': '提醒城',
                'created_by': 'tester',
            },
        )
        self.assertEqual(org_resp.status_code, 200)
        org_id = org_resp.json()['data']['id']

        for idx in range(2):
            sys_resp = self.client.post(
                '/api/systems',
                json={
                    'organization_id': org_id,
                    'system_name': f'提醒性能系统{idx}',
                    'proposed_level': 2,
                    'created_by': 'tester',
                },
            )
            self.assertEqual(sys_resp.status_code, 200)

        original_func = self.__class__.main_module.get_or_create_workflow_step_rules
        call_count = {'value': 0}

        def wrapped(db, config):
            call_count['value'] += 1
            return original_func(db, config)

        self.__class__.main_module.get_or_create_workflow_step_rules = wrapped
        try:
            resp = self.client.get('/api/workflow/reminders?mode=all')
        finally:
            self.__class__.main_module.get_or_create_workflow_step_rules = original_func

        self.assertEqual(resp.status_code, 200)
        self.assertEqual(call_count['value'], 1)

    def test_25_template_download_returns_404_when_file_missing(self):
        upload_resp = self.client.post(
            '/api/templates/upload',
            data={
                'template_name': '缺失模板下载校验',
                'report_type': 'grading_report',
                'category': '测试',
                'city': '测试城',
                'protection_level': '2',
                'is_default': 'false',
                'config_json': '{}',
            },
            files={'file': ('missing_template.docx', b'template-content', 'application/octet-stream')},
            headers=self.admin_headers,
        )
        self.assertEqual(upload_resp.status_code, 200)
        tpl_id = upload_resp.json()['data']['id']

        db = self.__class__.db_module.SessionLocal()
        try:
            row = db.query(self.__class__.main_module.ReportTemplate).filter_by(id=tpl_id).first()
            self.assertIsNotNone(row)
            row.file_path = f"{row.file_path}.missing"
            db.commit()
        finally:
            db.close()

        download_resp = self.client.get(f'/api/templates/{tpl_id}/download')
        self.assertEqual(download_resp.status_code, 404)
        self.assertIn('模板文件不存在', str(download_resp.json().get('detail', '')))

    def test_26_attachment_file_endpoints_return_404_when_file_missing(self):
        org_resp = self.client.post(
            '/api/organizations',
            json={
                'name': '附件缺失校验单位',
                'credit_code': '91350100M000100Y64',
                'legal_representative': '辛',
                'address': '附件城',
                'mobile_phone': '13100131010',
                'email': 'att64@example.com',
                'industry': '教育',
                'organization_type': '事业单位',
                'filing_region': '附件城',
                'created_by': 'tester',
            },
        )
        self.assertEqual(org_resp.status_code, 200)
        org_id = org_resp.json()['data']['id']

        upload_resp = self.client.post(
            f'/api/attachments/organization/{org_id}?actor=tester',
            files={'file': ('attach.pdf', b'%PDF-1.4 attachment-content', 'application/pdf')},
        )
        self.assertEqual(upload_resp.status_code, 200)
        attachment_id = upload_resp.json()['data']['id']

        db = self.__class__.db_module.SessionLocal()
        try:
            row = db.query(self.__class__.main_module.Attachment).filter_by(id=attachment_id).first()
            self.assertIsNotNone(row)
            row.file_path = f"{row.file_path}.missing"
            db.commit()
        finally:
            db.close()

        preview_resp = self.client.get(f'/api/attachment-files/{attachment_id}/preview')
        self.assertEqual(preview_resp.status_code, 404)
        self.assertIn('附件文件不存在', str(preview_resp.json().get('detail', '')))

        download_resp = self.client.get(f'/api/attachment-files/{attachment_id}/download')
        self.assertEqual(download_resp.status_code, 404)
        self.assertIn('附件文件不存在', str(download_resp.json().get('detail', '')))

    def test_27_update_workflow_rules_invalid_time_limit_returns_400(self):
        bad_rule_payload = {
            'updated_by': 'admin',
            'rules': [
                {'step_name': '信息收集', 'owner': 'collector', 'time_limit_hours': 'abc', 'enabled': True},
            ],
        }
        resp = self.client.put('/api/workflow/rules', json=bad_rule_payload, headers=self.admin_headers)
        self.assertEqual(resp.status_code, 400)
        self.assertIn('time_limit_hours', str(resp.json().get('detail', '')))

    def test_28_knowledge_list_pagination_and_pinned_order(self):
        keyword = '分页置顶专用关键字_001'
        doc_ids: list[int] = []
        for idx in range(3):
            upload_resp = self.client.post(
                '/api/knowledge/upload',
                data={
                    'title': f'分页置顶文档{idx}',
                    'doc_type': '政策文件',
                    'actor': 'admin',
                    'keywords': keyword,
                },
                files={'file': (f'page_pin_{idx}.docx', f'content-{idx}'.encode('utf-8'), 'application/octet-stream')},
                headers=self.admin_headers,
            )
            self.assertEqual(upload_resp.status_code, 200)
            doc_ids.append(upload_resp.json()['data']['id'])

        pin_resp = self.client.post(f'/api/knowledge/{doc_ids[-1]}/pin?enabled=true&actor=admin', headers=self.admin_headers)
        self.assertEqual(pin_resp.status_code, 200)

        list_resp = self.client.get('/api/knowledge', params={'keyword': keyword, 'match_mode': 'exact', 'page': 1, 'page_size': 2})
        self.assertEqual(list_resp.status_code, 200)
        data = list_resp.json()
        self.assertEqual(data['total'], 3)
        self.assertEqual(len(data['items']), 2)
        self.assertTrue(data['items'][0]['pinned'])
        self.assertEqual(data['items'][0]['id'], doc_ids[-1])

    def test_29_knowledge_list_without_pagination_returns_all_for_compatibility(self):
        keyword = f"兼容全量关键字_{datetime.now().strftime('%Y%m%d%H%M%S%f')}"
        db = self.__class__.db_module.SessionLocal()
        try:
            for idx in range(55):
                db.add(
                    self.__class__.main_module.KnowledgeDocument(
                        title=f'兼容全量文档{idx}',
                        keywords=keyword,
                        city='兼容城',
                        district='兼容区',
                        doc_type='政策文件',
                        protection_level='2',
                        version_no=1,
                        status='enabled',
                        file_name=f'compat_{idx}.docx',
                        file_path=f'/tmp/compat_{idx}.docx',
                        file_size=10,
                        uploaded_by='tester',
                    )
                )
            db.commit()
        finally:
            db.close()

        list_resp = self.client.get('/api/knowledge', params={'keyword': keyword, 'match_mode': 'exact'})
        self.assertEqual(list_resp.status_code, 200)
        data = list_resp.json()
        self.assertEqual(data['total'], 55)
        self.assertEqual(len(data['items']), 55)

    def test_30_login_page_is_accessible(self):
        resp = self.client.get('/login')
        self.assertEqual(resp.status_code, 200)
        self.assertIn('\u8d26\u53f7\u767b\u5f55', resp.text)
        self.assertIn('\u8bf7\u4f7f\u7528\u7ba1\u7406\u5458\u5206\u53d1\u7684\u8d26\u53f7\u8fdb\u5165\u5de5\u4f5c\u53f0', resp.text)
        self.assertNotIn('admin123', resp.text)

    def test_31_register_then_login_success(self):
        register_resp = self.client.post(
            '/api/auth/register',
            json={
                'username': 'self_register_user',
                'password': 'register123',
                'confirm_password': 'register123',
            },
        )
        self.assertEqual(register_resp.status_code, 403)
        self.assertIn('未开放自助注册', str(register_resp.json().get('detail', '')))

    def test_32_admin_can_open_users_page(self):
        resp = self.client.get('/users', headers=self.admin_headers)
        self.assertEqual(resp.status_code, 200)
        self.assertIn('用户管理', resp.text)

    def test_33_admin_can_open_backup_page(self):
        resp = self.client.get('/backup', headers=self.admin_headers)
        self.assertEqual(resp.status_code, 200)
        self.assertIn('数据备份与恢复', resp.text)

    def test_34_backup_create_returns_400_for_in_memory_database(self):
        resp = self.client.post('/api/backup/create', headers=self.admin_headers)
        self.assertEqual(resp.status_code, 400)
        self.assertIn('内存数据库', str(resp.json().get('detail', '')))

    def test_35_organization_word_import_invalid_docx_returns_400(self):
        bad_docx = b'not-a-valid-docx'
        resp = self.client.post(
            '/api/organizations/import/word',
            files={'file': ('bad.docx', bad_docx, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')},
            headers=self.admin_headers,
        )
        self.assertEqual(resp.status_code, 400)
        self.assertIn('无法解析', str(resp.json().get('detail', '')))

    def test_36_parse_docx_fallback_supports_missing_package_relationship(self):
        template_root = self.__class__.main_module.LOCAL_OFFICIAL_TEMPLATE_DIR
        candidates = sorted(template_root.glob('01-*.docx'))
        if not candidates:
            self.skipTest('missing fixture: 01-*.docx')
        src = candidates[0].read_bytes()
        in_buf = io.BytesIO(src)
        out_buf = io.BytesIO()
        with zipfile.ZipFile(in_buf, 'r') as zin, zipfile.ZipFile(out_buf, 'w', zipfile.ZIP_DEFLATED) as zout:
            for name in zin.namelist():
                if name == '_rels/.rels':
                    continue
                zout.writestr(name, zin.read(name))
        kv = self.__class__.main_module.parse_docx_key_values(out_buf.getvalue())
        self.assertGreater(len(kv), 0)

    def test_37_organization_word_import_supports_official_new_form_docx(self):
        template_root = self.__class__.main_module.LOCAL_OFFICIAL_TEMPLATE_DIR
        candidates = sorted(template_root.glob('01-*.docx'))
        if not candidates:
            self.skipTest('missing fixture: 01-*.docx')
        src = candidates[0]
        resp = self.client.post(
            '/api/organizations/import/word',
            files={'file': (src.name, src.read_bytes(), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')},
            headers=self.admin_headers,
        )
        self.assertEqual(resp.status_code, 200, resp.text)
        body = resp.json()
        self.assertEqual(body.get('message'), 'Word\u5bfc\u5165\u6210\u529f')
        data = body.get('data') or {}
        self.assertTrue(str(data.get('name') or '').strip())
        self.assertTrue(str(data.get('credit_code') or '').strip())
        self.assertTrue(str(data.get('mobile_phone') or '').strip())
        self.assertTrue(str(data.get('email') or '').strip())

    def test_38_update_template_null_required_fields_returns_400(self):
        upload_resp = self.client.post(
            '/api/templates/upload',
            data={
                'template_name': 'Null必填模板',
                'report_type': 'grading_report',
                'category': '测试',
                'city': '测试城',
                'protection_level': '2',
                'is_default': 'false',
                'config_json': '{}',
            },
            files={'file': ('null_required_template.docx', b'template', 'application/octet-stream')},
            headers=self.admin_headers,
        )
        self.assertEqual(upload_resp.status_code, 200, upload_resp.text)
        tpl_id = upload_resp.json()['data']['id']

        bad_status = self.client.put(
            f'/api/templates/{tpl_id}',
            json={'status': None},
            headers=self.admin_headers,
        )
        self.assertEqual(bad_status.status_code, 400, bad_status.text)

        bad_name = self.client.put(
            f'/api/templates/{tpl_id}',
            json={'template_name': None},
            headers=self.admin_headers,
        )
        self.assertEqual(bad_name.status_code, 400, bad_name.text)

    def test_39_update_knowledge_null_status_returns_400(self):
        upload_resp = self.client.post(
            '/api/knowledge/upload',
            data={'title': 'Null状态知识库', 'doc_type': '制度文档', 'actor': 'admin'},
            files={'file': ('null_status_knowledge.docx', b'knowledge', 'application/octet-stream')},
            headers=self.admin_headers,
        )
        self.assertEqual(upload_resp.status_code, 200, upload_resp.text)
        doc_id = upload_resp.json()['data']['id']

        update_resp = self.client.put(
            f'/api/knowledge/{doc_id}?actor=admin',
            json={'status': None},
            headers=self.admin_headers,
        )
        self.assertEqual(update_resp.status_code, 400, update_resp.text)

    def test_40_dashboard_trend_uses_calendar_month_window(self):
        original_datetime = self.__class__.main_module.datetime

        class FixedDateTime(datetime):
            @classmethod
            def now(cls, tz=None):
                base = cls(2026, 3, 15, 10, 0, 0)
                if tz is not None:
                    return base.replace(tzinfo=tz)
                return base

        self.__class__.main_module.datetime = FixedDateTime
        try:
            resp = self.client.get('/api/dashboard/trend?months=2')
        finally:
            self.__class__.main_module.datetime = original_datetime

        self.assertEqual(resp.status_code, 200, resp.text)
        months = [item['month'] for item in resp.json().get('trend', [])]
        self.assertEqual(months, ['2026-02', '2026-03'])

    def test_41_workflow_reminders_email_user_lookup_not_n_plus_one(self):
        db = self.__class__.db_module.SessionLocal()
        try:
            org = self.__class__.main_module.Organization(
                name='提醒N1优化单位',
                credit_code=f'N1TEST{datetime.now():%Y%m%d%H%M%S}',
                legal_representative='辛',
                address='提醒优化城',
                mobile_phone='13100131011',
                email='wf65@example.com',
                industry='互联网',
                organization_type='企业',
                filing_region='提醒优化城',
                created_by='tester',
            )
            db.add(org)
            db.commit()
            db.refresh(org)
            org_id = org.id
        finally:
            db.close()

        for idx in range(2):
            sys_resp = self.client.post(
                '/api/systems',
                json={
                    'organization_id': org_id,
                    'system_name': f'提醒N1系统{idx}',
                    'proposed_level': 2,
                    'created_by': 'tester',
                },
            )
            self.assertEqual(sys_resp.status_code, 200, sys_resp.text)

        rule_resp = self.client.put(
            '/api/workflow/rules',
            json={
                'updated_by': 'admin',
                'rules': [
                    {'step_name': '信息收集', 'owner': 'owner_n1', 'time_limit_hours': 24, 'enabled': True},
                    {'step_name': '信息审核', 'owner': 'reviewer_n1', 'time_limit_hours': 24, 'enabled': True},
                ],
            },
            headers=self.admin_headers,
        )
        self.assertEqual(rule_resp.status_code, 200, rule_resp.text)

        session_cls = self.__class__.main_module.Session
        original_query = session_cls.query
        user_query_count = {'value': 0}

        def wrapped_query(session_self, *entities, **kwargs):
            if len(entities) == 1 and entities[0] is self.__class__.main_module.UserAccount:
                user_query_count['value'] += 1
            return original_query(session_self, *entities, **kwargs)

        session_cls.query = wrapped_query
        try:
            remind_resp = self.client.get('/api/workflow/reminders?mode=all&send=true&channel=email')
        finally:
            session_cls.query = original_query

        self.assertEqual(remind_resp.status_code, 200, remind_resp.text)
        self.assertLessEqual(user_query_count['value'], 1, f"UserAccount查询次数过多: {user_query_count['value']}")

    def test_42_legacy_admin_blocked_when_strict_auth_enabled(self):
        db = self.__class__.db_module.SessionLocal()
        try:
            org = self.__class__.main_module.Organization(
                name='严格鉴权绕过校验单位',
                credit_code=f'STRICT{datetime.now():%y%m%d%H%M%S%f}'[:18],
                legal_representative='壬',
                address='严格鉴权城',
                mobile_phone='13100131012',
                email='strict-auth@example.com',
                industry='政务',
                organization_type='机关单位',
                filing_region='严格鉴权城',
                created_by='tester',
            )
            db.add(org)
            db.commit()
            db.refresh(org)
            org_id = org.id
        finally:
            db.close()

        original_strict = self.__class__.main_module.STRICT_AUTH
        self.__class__.main_module.STRICT_AUTH = True
        self.client.cookies.clear()
        try:
            resp = self.client.post(f'/api/organizations/{org_id}/unlock?actor=admin')
        finally:
            self.__class__.main_module.STRICT_AUTH = original_strict
            self.client.post('/api/auth/login', json={'username': 'admin', 'password': 'admin123'})

        self.assertEqual(resp.status_code, 401, resp.text)

    def test_43_system_excel_import_row_failure_should_not_break_whole_commit(self):
        db = self.__class__.db_module.SessionLocal()
        try:
            org = self.__class__.main_module.Organization(
                name='系统导入回滚校验单位',
                credit_code=f'RLBK{datetime.now():%y%m%d%H%M%S%f}'[:18],
                legal_representative='癸',
                address='回滚城',
                mobile_phone='13100131013',
                email='rollback-import@example.com',
                industry='制造',
                organization_type='企业',
                filing_region='回滚城',
                created_by='tester',
            )
            db.add(org)
            db.flush()
            db.add(
                self.__class__.main_module.SystemInfo(
                    organization_id=org.id,
                    system_name='已有重复编号系统',
                    system_code='DUP-SYSTEM-CODE-001',
                    proposed_level=2,
                    created_by='tester',
                )
            )
            db.commit()
            db.refresh(org)
            org_id = org.id
        finally:
            db.close()

        wb = Workbook()
        ws = wb.active
        ws.append(['系统名称', '系统编号', '单位ID', '拟定等级', '部署方式', '系统类型', '上线时间', '录入人'])
        ws.append(['重复编号导入系统', '', org_id, 2, '本地部署', '业务系统', '', 'tester'])
        bio = io.BytesIO()
        wb.save(bio)
        bio.seek(0)

        original_generate_code = self.__class__.main_module.generate_system_code
        self.__class__.main_module.generate_system_code = lambda _db: 'DUP-SYSTEM-CODE-001'
        try:
            resp = self.client.post(
                '/api/systems/import/excel',
                files={'file': ('systems_duplicate_code.xlsx', bio.getvalue(), 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')},
                headers=self.admin_headers,
            )
        finally:
            self.__class__.main_module.generate_system_code = original_generate_code

        self.assertEqual(resp.status_code, 200, resp.text)
        self.assertEqual(resp.json().get('imported'), 0)
        self.assertGreaterEqual(len(resp.json().get('skipped') or []), 1)

    def test_44_knowledge_upload_invalid_source_should_not_leave_orphan_file(self):
        knowledge_dir = self.__class__.main_module.UPLOAD_DIR / 'knowledge'
        before = {p.name for p in knowledge_dir.glob('*') if p.is_file()}

        resp = self.client.post(
            '/api/knowledge/upload',
            data={
                'title': '无效来源文档上传',
                'doc_type': '制度文档',
                'actor': 'admin',
                'source_doc_id': 99999999,
            },
            files={'file': ('orphan_source_check.docx', b'orphan-check', 'application/octet-stream')},
            headers=self.admin_headers,
        )
        self.assertEqual(resp.status_code, 404, resp.text)

        after = {p.name for p in knowledge_dir.glob('*') if p.is_file()}
        self.assertEqual(after, before, '无效 source_doc_id 不应在 uploads/knowledge 留下孤儿文件')

    def test_45_template_upload_invalid_config_should_not_leave_orphan_file(self):
        template_dir = self.__class__.main_module.UPLOAD_DIR / 'templates'
        before = {p.name for p in template_dir.glob('*') if p.is_file()}

        resp = self.client.post(
            '/api/templates/upload',
            data={
                'template_name': '非法配置模板',
                'report_type': 'grading_report',
                'category': '测试',
                'city': '测试城',
                'protection_level': '2',
                'is_default': 'false',
                'config_json': '{"bad_json":',
            },
            files={'file': ('orphan_config_check.docx', b'orphan-config-check', 'application/octet-stream')},
            headers=self.admin_headers,
        )
        self.assertEqual(resp.status_code, 400, resp.text)
        self.assertIn('config_json', str(resp.json().get('detail', '')))

        after = {p.name for p in template_dir.glob('*') if p.is_file()}
        self.assertEqual(after, before, '非法 config_json 不应在 uploads/templates 留下孤儿文件')

    def test_46_update_workflow_rules_string_false_should_disable_step(self):
        get_resp = self.client.get('/api/workflow/rules', headers=self.admin_headers)
        self.assertEqual(get_resp.status_code, 200, get_resp.text)
        rules = get_resp.json().get('rules') or []
        self.assertTrue(rules, 'workflow rules 不能为空')

        first = dict(rules[0])
        first['enabled'] = 'false'
        put_resp = self.client.put(
            '/api/workflow/rules',
            json={'updated_by': 'admin', 'rules': [first]},
            headers=self.admin_headers,
        )
        self.assertEqual(put_resp.status_code, 200, put_resp.text)

        after_resp = self.client.get('/api/workflow/rules', headers=self.admin_headers)
        self.assertEqual(after_resp.status_code, 200, after_resp.text)
        after_rules = after_resp.json().get('rules') or []
        target = next((r for r in after_rules if r.get('step_name') == first.get('step_name')), None)
        self.assertIsNotNone(target, '未找到目标流程节点')
        self.assertFalse(target.get('enabled'))

    def test_47_update_template_is_default_string_false_should_not_set_true(self):
        upload_1 = self.client.post(
            '/api/templates/upload',
            data={
                'template_name': '布尔解析模板1',
                'report_type': 'filing_form',
                'is_default': 'true',
            },
            files={'file': ('bool_tpl_1.docx', b'a', 'application/octet-stream')},
            headers=self.admin_headers,
        )
        self.assertEqual(upload_1.status_code, 200, upload_1.text)
        tpl_id_1 = upload_1.json()['data']['id']

        upload_2 = self.client.post(
            '/api/templates/upload',
            data={
                'template_name': '布尔解析模板2',
                'report_type': 'filing_form',
                'is_default': 'false',
            },
            files={'file': ('bool_tpl_2.docx', b'b', 'application/octet-stream')},
            headers=self.admin_headers,
        )
        self.assertEqual(upload_2.status_code, 200, upload_2.text)
        tpl_id_2 = upload_2.json()['data']['id']

        update_resp = self.client.put(
            f'/api/templates/{tpl_id_2}',
            json={'is_default': 'false'},
            headers=self.admin_headers,
        )
        self.assertEqual(update_resp.status_code, 200, update_resp.text)

        db = self.__class__.db_module.SessionLocal()
        try:
            row_1 = db.query(self.__class__.main_module.ReportTemplate).filter_by(id=tpl_id_1).first()
            row_2 = db.query(self.__class__.main_module.ReportTemplate).filter_by(id=tpl_id_2).first()
            self.assertIsNotNone(row_1)
            self.assertIsNotNone(row_2)
            self.assertTrue(bool(row_1.is_default))
            self.assertFalse(bool(row_2.is_default))
        finally:
            db.close()

    def test_48_direct_delete_organization_requires_admin_review_path(self):
        create_user_resp = self.client.post(
            '/api/auth/users',
            json={
                'username': 'delete_req_eval_user',
                'password': 'deleteReq123',
                'role': 'evaluator',
                'require_password_change': False,
            },
            headers=self.admin_headers,
        )
        self.assertIn(create_user_resp.status_code, [200, 409], create_user_resp.text)
        login_resp = self.client.post('/api/auth/login', json={'username': 'delete_req_eval_user', 'password': 'deleteReq123'})
        self.assertEqual(login_resp.status_code, 200, login_resp.text)
        eval_headers = {'X-Auth-Token': login_resp.json()['token']}

        create_resp = self.client.post(
            '/api/organizations',
            json={
                'name': '直删限制单位',
                'credit_code': '91350100M000100Y70',
                'legal_representative': '测试人',
                'address': '限制城',
                'mobile_phone': '13100131014',
                'email': 'org70@example.com',
                'industry': '教育',
                'organization_type': '企业',
                'filing_region': '限制城',
                'created_by': 'tester',
            },
        )
        self.assertEqual(create_resp.status_code, 200, create_resp.text)
        org_id = create_resp.json()['data']['id']

        delete_resp = self.client.delete(f'/api/organizations/{org_id}?actor=tester&is_admin=false', headers=eval_headers)
        self.assertEqual(delete_resp.status_code, 403, delete_resp.text)
        self.assertIn('删除申请', str(delete_resp.json().get('detail', '')))

    def test_49_direct_delete_system_requires_admin_review_path(self):
        create_user_resp = self.client.post(
            '/api/auth/users',
            json={
                'username': 'delete_req_eval_user_sys',
                'password': 'deleteReqSys123',
                'role': 'evaluator',
                'require_password_change': False,
            },
            headers=self.admin_headers,
        )
        self.assertIn(create_user_resp.status_code, [200, 409], create_user_resp.text)
        login_resp = self.client.post('/api/auth/login', json={'username': 'delete_req_eval_user_sys', 'password': 'deleteReqSys123'})
        self.assertEqual(login_resp.status_code, 200, login_resp.text)
        eval_headers = {'X-Auth-Token': login_resp.json()['token']}

        org_resp = self.client.post(
            '/api/organizations',
            json={
                'name': '直删限制系统单位',
                'credit_code': '91350100M000100Y71',
                'legal_representative': '测试人',
                'address': '限制城',
                'mobile_phone': '13100131015',
                'email': 'sys71@example.com',
                'industry': '教育',
                'organization_type': '企业',
                'filing_region': '限制城',
                'created_by': 'tester',
            },
        )
        self.assertEqual(org_resp.status_code, 200, org_resp.text)
        org_id = org_resp.json()['data']['id']

        sys_resp = self.client.post(
            '/api/systems',
            json={
                'organization_id': org_id,
                'system_name': '直删限制系统',
                'proposed_level': 2,
                'created_by': 'tester',
            },
        )
        self.assertEqual(sys_resp.status_code, 200, sys_resp.text)
        system_id = sys_resp.json()['data']['id']

        delete_resp = self.client.delete(f'/api/systems/{system_id}?actor=tester&is_admin=false', headers=eval_headers)
        self.assertEqual(delete_resp.status_code, 403, delete_resp.text)
        self.assertIn('删除申请', str(delete_resp.json().get('detail', '')))

    def test_50_knowledge_upload_should_create_missing_knowledge_dir(self):
        original_upload_dir = self.__class__.main_module.UPLOAD_DIR
        temp_upload_dir = Path(__file__).resolve().parent.parent / '.tmp_upload_dir_for_test'
        if temp_upload_dir.exists():
            import shutil
            shutil.rmtree(temp_upload_dir, ignore_errors=True)

        self.__class__.main_module.UPLOAD_DIR = temp_upload_dir
        self.addCleanup(setattr, self.__class__.main_module, 'UPLOAD_DIR', original_upload_dir)
        self.addCleanup(lambda: __import__('shutil').rmtree(temp_upload_dir, ignore_errors=True))

        upload_resp = self.client.post(
            '/api/knowledge/upload',
            data={'title': '目录自愈测试', 'doc_type': '制度文档', 'actor': 'admin'},
            files={'file': ('mkdir_knowledge.docx', b'content', 'application/octet-stream')},
            headers=self.admin_headers,
        )
        self.assertEqual(upload_resp.status_code, 200, upload_resp.text)
        doc_id = upload_resp.json()['data']['id']

        db = self.__class__.db_module.SessionLocal()
        try:
            row = db.query(self.__class__.main_module.KnowledgeDocument).filter_by(id=doc_id).first()
            self.assertIsNotNone(row)
            self.assertTrue(Path(row.file_path).exists(), '上传成功后文件应实际落盘')
        finally:
            db.close()

    def test_51_workflow_rules_should_cleanup_duplicate_rows_and_keep_latest(self):
        db = self.__class__.db_module.SessionLocal()
        try:
            db.add(
                self.__class__.main_module.WorkflowStepRule(
                    config_name='default',
                    step_name='信息收集',
                    owner='dup_old_owner',
                    time_limit_hours=24,
                    enabled=True,
                    updated_by='tester',
                )
            )
            db.flush()
            db.add(
                self.__class__.main_module.WorkflowStepRule(
                    config_name='default',
                    step_name='信息收集',
                    owner='dup_new_owner',
                    time_limit_hours=12,
                    enabled=False,
                    updated_by='tester',
                )
            )
            db.commit()
        finally:
            db.close()

        list_resp = self.client.get('/api/workflow/rules', headers=self.admin_headers)
        self.assertEqual(list_resp.status_code, 200, list_resp.text)
        rules = list_resp.json().get('rules') or []
        target = next((r for r in rules if r.get('step_name') == '信息收集'), None)
        self.assertIsNotNone(target, '未返回信息收集规则')
        self.assertEqual(target.get('owner'), 'dup_new_owner')
        self.assertFalse(target.get('enabled'))

        db = self.__class__.db_module.SessionLocal()
        try:
            count = (
                db.query(self.__class__.main_module.WorkflowStepRule)
                .filter_by(config_name='default', step_name='信息收集')
                .count()
            )
            self.assertEqual(count, 1)
        finally:
            db.close()

    def test_52_attachment_download_should_block_upload_dir_prefix_bypass(self):
        base_dir = Path(__file__).resolve().parent.parent
        bypass_dir = base_dir / 'uploads_shadow'
        bypass_dir.mkdir(parents=True, exist_ok=True)
        self.addCleanup(lambda: __import__('shutil').rmtree(bypass_dir, ignore_errors=True))

        bypass_file = bypass_dir / 'prefix_bypass.txt'
        bypass_file.write_text('prefix-bypass-check', encoding='utf-8')

        db = self.__class__.db_module.SessionLocal()
        try:
            row = self.__class__.main_module.Attachment(
                entity_type='organization',
                entity_id=999999,
                file_name='prefix_bypass.txt',
                file_path=str(bypass_file),
                file_ext='txt',
                file_size=bypass_file.stat().st_size,
                uploaded_by='tester',
            )
            db.add(row)
            db.commit()
            db.refresh(row)
            attachment_id = row.id
        finally:
            db.close()

        download_resp = self.client.get(f'/api/attachment-files/{attachment_id}/download')
        self.assertEqual(download_resp.status_code, 403, download_resp.text)

    def test_53_restore_template_default_should_keep_single_default(self):
        upload_1 = self.client.post(
            '/api/templates/upload',
            data={
                'template_name': '恢复默认冲突模板A',
                'report_type': 'grading_report',
                'is_default': 'true',
            },
            files={'file': ('restore_default_a.docx', b'a', 'application/octet-stream')},
            headers=self.admin_headers,
        )
        self.assertEqual(upload_1.status_code, 200, upload_1.text)
        tpl_id_1 = upload_1.json()['data']['id']

        upload_2 = self.client.post(
            '/api/templates/upload',
            data={
                'template_name': '恢复默认冲突模板B',
                'report_type': 'grading_report',
                'is_default': 'false',
            },
            files={'file': ('restore_default_b.docx', b'b', 'application/octet-stream')},
            headers=self.admin_headers,
        )
        self.assertEqual(upload_2.status_code, 200, upload_2.text)
        tpl_id_2 = upload_2.json()['data']['id']

        set_tpl2_default_resp = self.client.put(
            f'/api/templates/{tpl_id_2}',
            json={'is_default': 'true'},
            headers=self.admin_headers,
        )
        self.assertEqual(set_tpl2_default_resp.status_code, 200, set_tpl2_default_resp.text)

        new_version_resp = self.client.post(
            f'/api/templates/{tpl_id_2}/new-version',
            files={'file': ('restore_default_b_v2.docx', b'bb', 'application/octet-stream')},
            headers=self.admin_headers,
        )
        self.assertEqual(new_version_resp.status_code, 200, new_version_resp.text)

        versions_resp = self.client.get(f'/api/templates/{tpl_id_2}/versions')
        self.assertEqual(versions_resp.status_code, 200, versions_resp.text)
        versions = versions_resp.json().get('items') or []
        default_snapshot = next((item for item in versions if bool((item.get('snapshot') or {}).get('is_default'))), None)
        self.assertIsNotNone(default_snapshot, '未找到 is_default=True 的历史快照')
        restore_version_id = default_snapshot['id']

        set_tpl1_default_resp = self.client.post(
            f'/api/templates/{tpl_id_1}/set-default',
            headers=self.admin_headers,
        )
        self.assertEqual(set_tpl1_default_resp.status_code, 200, set_tpl1_default_resp.text)

        restore_resp = self.client.post(
            f'/api/templates/{tpl_id_2}/restore/{restore_version_id}',
            headers=self.admin_headers,
        )
        self.assertEqual(restore_resp.status_code, 200, restore_resp.text)

        db = self.__class__.db_module.SessionLocal()
        try:
            row_1 = db.query(self.__class__.main_module.ReportTemplate).filter_by(id=tpl_id_1).first()
            row_2 = db.query(self.__class__.main_module.ReportTemplate).filter_by(id=tpl_id_2).first()
            self.assertIsNotNone(row_1)
            self.assertIsNotNone(row_2)
            self.assertFalse(bool(row_1.is_default))
            self.assertTrue(bool(row_2.is_default))
        finally:
            db.close()

    def test_54_safe_extract_zip_should_reject_too_many_entries(self):
        main_module = self.__class__.main_module
        original_limit = getattr(main_module, 'MAX_BACKUP_ZIP_ENTRIES', None)
        main_module.MAX_BACKUP_ZIP_ENTRIES = 2
        temp_root = Path(__file__).resolve().parent.parent / '.tmp_safe_extract_zip_limit'
        if temp_root.exists():
            import shutil
            shutil.rmtree(temp_root, ignore_errors=True)
        temp_root.mkdir(parents=True, exist_ok=True)

        self.addCleanup(lambda: __import__('shutil').rmtree(temp_root, ignore_errors=True))
        if original_limit is None:
            self.addCleanup(lambda: delattr(main_module, 'MAX_BACKUP_ZIP_ENTRIES'))
        else:
            self.addCleanup(setattr, main_module, 'MAX_BACKUP_ZIP_ENTRIES', original_limit)

        bio = io.BytesIO()
        with zipfile.ZipFile(bio, mode='w', compression=zipfile.ZIP_DEFLATED) as zf:
            zf.writestr('a.txt', 'a')
            zf.writestr('b.txt', 'b')
            zf.writestr('c.txt', 'c')
        bio.seek(0)

        with zipfile.ZipFile(bio, mode='r') as zf:
            with self.assertRaises(main_module.HTTPException) as ctx:
                main_module.safe_extract_zip(zf, temp_root)

        self.assertEqual(ctx.exception.status_code, 400)
        self.assertIn('文件数量', str(ctx.exception.detail))

    def test_55_safe_extract_zip_should_reject_oversized_uncompressed_content(self):
        main_module = self.__class__.main_module
        original_limit = getattr(main_module, 'MAX_BACKUP_UNCOMPRESSED', None)
        main_module.MAX_BACKUP_UNCOMPRESSED = 2
        temp_root = Path(__file__).resolve().parent.parent / '.tmp_safe_extract_zip_size'
        if temp_root.exists():
            import shutil
            shutil.rmtree(temp_root, ignore_errors=True)
        temp_root.mkdir(parents=True, exist_ok=True)

        self.addCleanup(lambda: __import__('shutil').rmtree(temp_root, ignore_errors=True))
        if original_limit is None:
            self.addCleanup(lambda: delattr(main_module, 'MAX_BACKUP_UNCOMPRESSED'))
        else:
            self.addCleanup(setattr, main_module, 'MAX_BACKUP_UNCOMPRESSED', original_limit)

        bio = io.BytesIO()
        with zipfile.ZipFile(bio, mode='w', compression=zipfile.ZIP_DEFLATED) as zf:
            zf.writestr('big.txt', 'abcd')
        bio.seek(0)

        with zipfile.ZipFile(bio, mode='r') as zf:
            with self.assertRaises(main_module.HTTPException) as ctx:
                main_module.safe_extract_zip(zf, temp_root)

        self.assertEqual(ctx.exception.status_code, 400)
        self.assertIn('总大小', str(ctx.exception.detail))

    def test_56_direct_delete_should_auto_close_pending_delete_request(self):
        org_resp = self.client.post(
            '/api/organizations',
            json={
                'name': '自动结案单位',
                'credit_code': '91350100M000100Y90',
                'legal_representative': '结案人',
                'address': '厦门市思明区',
                'mobile_phone': '13100131990',
                'email': 'autoclose@example.com',
                'industry': '企业',
                'organization_type': '企业',
                'filing_region': '厦门',
                'created_by': 'tester',
            },
            headers=self.admin_headers,
        )
        self.assertEqual(org_resp.status_code, 200, org_resp.text)
        org_id = org_resp.json()['data']['id']

        db = self.__class__.db_module.SessionLocal()
        try:
            pending = self.__class__.main_module.DeleteRequest(
                entity_type='organization',
                entity_id=org_id,
                reason='旧申请仍待审',
                status='pending',
                requested_by='tester',
            )
            db.add(pending)
            db.commit()
            db.refresh(pending)
            request_id = pending.id
        finally:
            db.close()

        delete_resp = self.client.delete(f'/api/organizations/{org_id}', headers=self.admin_headers)
        self.assertEqual(delete_resp.status_code, 200, delete_resp.text)

        list_resp = self.client.get('/api/delete-requests?entity_type=organization', headers=self.admin_headers)
        self.assertEqual(list_resp.status_code, 200, list_resp.text)
        item = next((row for row in list_resp.json()['items'] if row['id'] == request_id), None)
        self.assertIsNotNone(item)
        self.assertEqual(item['status'], 'approved')
        self.assertEqual(item['reviewed_by'], 'admin')
        self.assertIn('自动结案', item.get('review_comment') or '')

    def test_57_list_delete_requests_should_backfill_auto_closed_comment_for_old_rows(self):
        org_resp = self.client.post(
            '/api/organizations',
            json={
                'name': '旧结案单位',
                'credit_code': '91350100M000100Y91',
                'legal_representative': '回填人',
                'address': '太原市',
                'mobile_phone': '13100131991',
                'email': 'backfill@example.com',
                'industry': '企业',
                'organization_type': '企业',
                'filing_region': '太原',
                'created_by': 'tester',
            },
            headers=self.admin_headers,
        )
        self.assertEqual(org_resp.status_code, 200, org_resp.text)
        org_id = org_resp.json()['data']['id']

        db = self.__class__.db_module.SessionLocal()
        try:
            org = db.query(self.__class__.main_module.Organization).filter_by(id=org_id).first()
            org.deleted_by = 'admin'
            org.deleted_at = datetime.now()
            row = self.__class__.main_module.DeleteRequest(
                entity_type='organization',
                entity_id=org_id,
                reason='历史记录',
                status='approved',
                requested_by='tester',
                reviewed_by='admin',
                reviewed_at=org.deleted_at,
                review_comment=None,
            )
            db.add(row)
            db.commit()
            db.refresh(row)
            request_id = row.id
        finally:
            db.close()

        list_resp = self.client.get('/api/delete-requests?entity_type=organization', headers=self.admin_headers)
        self.assertEqual(list_resp.status_code, 200, list_resp.text)
        item = next((row for row in list_resp.json()['items'] if row['id'] == request_id), None)
        self.assertIsNotNone(item)
        self.assertEqual(item['reviewed_by'], 'admin')
        self.assertIn('自动结案', item.get('review_comment') or '')


if __name__ == '__main__':
    unittest.main()
