"""CodeAudit 报告导出模块 —— Markdown / HTML / JSON / Word / Excel"""

import html
import io
import json
import re
from datetime import datetime

import markdown

SEVERITY_MAP = {
    "critical": "严重", "high": "高风险", "medium": "中风险",
    "low": "低风险", "info": "安全",
}


# ---------------------------------------------------------------------------
#  漏洞解析器 —— 从 LLM 的 Markdown 输出中提取结构化数据
# ---------------------------------------------------------------------------

def parse_vulnerabilities(result_text: str) -> list:
    """从审计结果 Markdown 中提取结构化漏洞列表。"""
    if not result_text:
        return []

    vulns = []
    # 定位"发现的安全问题"段落
    sections = re.split(r"^##\s+", result_text, flags=re.MULTILINE)
    vuln_section = ""
    for s in sections:
        first_line = s.split("\n")[0].strip()
        if "安全问题" in first_line or "发现" in first_line:
            vuln_section = s
            break

    if not vuln_section:
        return vulns

    # 按 ### 标题拆分单个漏洞
    items = re.split(r"^###\s+", vuln_section, flags=re.MULTILINE)

    for item in items[1:]:
        lines = item.strip().split("\n")
        if not lines:
            continue
        title = lines[0].strip().lstrip("0123456789.、 ")
        body = "\n".join(lines[1:])

        vuln = {"title": title}
        field_patterns = {
            "description":    r"\*\*问题描述\*\*[：:]\s*(.*?)(?=\n\s*[-*]\s*\*\*|\Z)",
            "severity":       r"\*\*风险等级\*\*[：:]\s*(.*?)(?=\n|$)",
            "location":       r"\*\*问题位置\*\*[：:]\s*(.*?)(?=\n\s*[-*]\s*\*\*|\Z)",
            "category":       r"\*\*(?:漏洞分类|标准条款)\*\*[：:]\s*(.*?)(?=\n\s*[-*]\s*\*\*|\Z)",
            "analysis":       r"\*\*详细分析\*\*[：:]\s*(.*?)(?=\n\s*[-*]\s*\*\*|\Z)",
            "recommendation": r"\*\*修复建议\*\*[：:]\s*(.*?)(?=\n\s*[-*]\s*\*\*|\Z)",
        }

        for key, pattern in field_patterns.items():
            m = re.search(pattern, body, re.DOTALL)
            vuln[key] = m.group(1).strip() if m else ""

        if vuln.get("title"):
            vulns.append(vuln)

    return vulns


def _extract_section(result_text: str, *keywords) -> str:
    """从审计结果中提取某个 ## 段落的内容。"""
    if not result_text:
        return ""
    sections = re.split(r"^##\s+", result_text, flags=re.MULTILINE)
    for s in sections:
        first_line = s.split("\n")[0].strip()
        if any(kw in first_line for kw in keywords):
            lines = s.strip().split("\n")
            return "\n".join(lines[1:]).strip()
    return ""


# ---------------------------------------------------------------------------
#  Markdown 导出
# ---------------------------------------------------------------------------

def export_markdown(audit) -> str:
    ts = audit.created_at.strftime("%Y-%m-%d %H:%M:%S")
    severity_label = SEVERITY_MAP.get(audit.severity, "未知")

    return (
        f"# 代码审计报告\n\n"
        f"## 基本信息\n\n"
        f"| 项目 | 内容 |\n"
        f"| :--- | :--- |\n"
        f"| **标题** | {audit.title} |\n"
        f"| **审计时间** | {ts} |\n"
        f"| **使用模型** | {audit.model_used or '-'} |\n"
        f"| **代码语言** | {audit.language} |\n"
        f"| **风险等级** | {severity_label} |\n\n"
        f"---\n\n"
        f"## 审计代码\n\n```{audit.language if audit.language != 'auto' else ''}\n"
        f"{audit.code_content}\n```\n\n"
        f"---\n\n"
        f"{audit.result or '无结果'}\n"
    )


# ---------------------------------------------------------------------------
#  JSON 导出
# ---------------------------------------------------------------------------

def export_json(audit) -> str:
    vulns = parse_vulnerabilities(audit.result or "")
    return json.dumps({
        "id": audit.id,
        "title": audit.title,
        "language": audit.language,
        "model": audit.model_used,
        "severity": audit.severity,
        "status": audit.status,
        "created_at": audit.created_at.strftime("%Y-%m-%d %H:%M:%S"),
        "code": audit.code_content,
        "result": audit.result,
        "vulnerabilities": vulns,
    }, ensure_ascii=False, indent=2)


# ---------------------------------------------------------------------------
#  HTML 导出
# ---------------------------------------------------------------------------

def export_html(audit) -> str:
    ts = audit.created_at.strftime("%Y-%m-%d %H:%M:%S")
    severity_map = {
        "critical": ("严重", "#ff4d6a", "rgba(255,77,106,0.12)"),
        "high": ("高风险", "#ff8c42", "rgba(255,140,66,0.12)"),
        "medium": ("中风险", "#f0c040", "rgba(240,192,64,0.12)"),
        "low": ("低风险", "#00d4aa", "rgba(0,212,170,0.12)"),
        "info": ("安全", "#8b949e", "rgba(139,148,158,0.12)"),
    }
    sev_label, sev_color, sev_bg = severity_map.get(
        audit.severity, ("未知", "#8b949e", "rgba(139,148,158,0.12)")
    )

    result_body = markdown.markdown(
        audit.result or "", extensions=["fenced_code", "tables"]
    )

    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>审计报告 - {html.escape(audit.title)}</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link href="https://fonts.googleapis.com/css2?family=Space+Grotesk:wght@500;600;700&family=DM+Sans:wght@400;500&family=JetBrains+Mono:wght@400&display=swap" rel="stylesheet">
<link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/highlight.js/11.9.0/styles/github-dark.min.css">
<style>
:root {{
  --bg: #0a0a0f;
  --surface: #141419;
  --border: #2a2a35;
  --text: #e4e4e7;
  --text-muted: #8b8b9e;
  --primary: #00d4aa;
  --accent: #ff6b9d;
}}
* {{ box-sizing: border-box; margin: 0; padding: 0; }}
body {{
  font-family: 'DM Sans', 'Microsoft YaHei', sans-serif;
  background: var(--bg); color: var(--text);
  line-height: 1.7; max-width: 900px; margin: 0 auto; padding: 48px 32px;
}}
h1,h2,h3 {{ font-family: 'Space Grotesk', 'Microsoft YaHei', sans-serif; color: #f0f0f5; font-weight: 600; }}
h1 {{ font-size: 1.6rem; margin-bottom: 8px; }}
h2 {{ font-size: 1.15rem; margin-top: 32px; margin-bottom: 16px; padding-bottom: 8px; border-bottom: 1px solid var(--border); }}
h3 {{ font-size: 1rem; margin-top: 20px; margin-bottom: 8px; }}
a {{ color: var(--primary); }}
hr {{ border: none; border-top: 1px solid var(--border); margin: 24px 0; }}
code,pre {{ font-family: 'JetBrains Mono', Consolas, monospace; font-size: 0.9rem; }}
pre {{ background: var(--surface); border: 1px solid var(--border); border-radius: 6px; padding: 16px; overflow-x: auto; margin-bottom: 16px; }}
code {{ background: var(--surface); padding: 2px 6px; border-radius: 3px; }}
pre code {{ background: none; padding: 0; }}
p {{ margin-bottom: 12px; }}
ul,ol {{ margin-bottom: 12px; padding-left: 24px; }}
li {{ margin-bottom: 6px; }}
strong {{ color: #f0f0f5; }}
table {{ width: 100%; border-collapse: collapse; margin-bottom: 16px; }}
th,td {{ padding: 8px 14px; border: 1px solid var(--border); text-align: left; font-size: 0.93rem; }}
th {{ background: var(--surface); font-weight: 600; color: #f0f0f5; }}
.report-header {{ border-bottom: 1px solid var(--border); padding-bottom: 24px; margin-bottom: 32px; }}
.report-brand {{ display: flex; align-items: center; gap: 8px; font-family: 'Space Grotesk', sans-serif; font-weight: 700; font-size: 0.88rem; color: var(--primary); margin-bottom: 16px; text-transform: uppercase; letter-spacing: 0.06em; }}
.meta-grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(180px, 1fr)); gap: 12px; margin-top: 16px; }}
.meta-item {{ font-size: 0.88rem; }}
.meta-item .label {{ color: var(--text-muted); font-size: 0.78rem; text-transform: uppercase; letter-spacing: 0.04em; margin-bottom: 2px; }}
.severity-badge {{ display: inline-block; padding: 3px 12px; border-radius: 6px; font-size: 0.85rem; font-weight: 600; color: {sev_color}; background: {sev_bg}; }}
.report-footer {{ margin-top: 48px; padding-top: 16px; border-top: 1px solid var(--border); font-size: 0.78rem; color: var(--text-muted); text-align: center; }}
@media print {{
  body {{ background: #fff; color: #1a1a1a; padding: 20px; }}
  pre {{ background: #f5f5f5; border-color: #ddd; }}
  code {{ background: #f5f5f5; }}
  th {{ background: #f0f0f0; }}
  .severity-badge {{ border: 1px solid {sev_color}; background: transparent; }}
}}
</style>
</head>
<body>
<div class="report-header">
  <div class="report-brand">
    <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="m18 16 4-4-4-4"/><path d="m6 8-4 4 4 4"/><path d="m14.5 4-5 16"/></svg>
    CodeAudit Report
  </div>
  <h1>{html.escape(audit.title)}</h1>
  <div class="meta-grid">
    <div class="meta-item"><div class="label">审计时间</div>{ts}</div>
    <div class="meta-item"><div class="label">使用模型</div>{html.escape(audit.model_used or '-')}</div>
    <div class="meta-item"><div class="label">代码语言</div>{html.escape(audit.language)}</div>
    <div class="meta-item"><div class="label">风险等级</div><span class="severity-badge">{sev_label}</span></div>
  </div>
</div>

<h2>审计代码</h2>
<pre><code>{html.escape(audit.code_content)}</code></pre>
<hr>
{result_body}
<div class="report-footer">Generated by CodeAudit &middot; {ts}</div>
<script src="https://cdnjs.cloudflare.com/ajax/libs/highlight.js/11.9.0/highlight.min.js"></script>
<script>hljs.highlightAll();</script>
</body>
</html>"""


# ---------------------------------------------------------------------------
#  Word (.docx) 导出
# ---------------------------------------------------------------------------

def export_docx(audit) -> bytes:
    """生成 Word (.docx) 格式的专业审计报告。"""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # -- 默认样式 --
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.5

    def _add_run(paragraph, text, bold=False, size=10.5, color=None):
        run = paragraph.add_run(text)
        run.font.name = "微软雅黑"
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color
        return run

    def _set_cell(cell, text, bold=False, size=10.5):
        cell.text = ""
        p = cell.paragraphs[0]
        _add_run(p, str(text), bold=bold, size=size)

    # -- 封面 --
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _add_run(p, "源代码安全审计报告", bold=True, size=28)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _add_run(p, audit.title, size=16, color=RGBColor(80, 80, 80))

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _add_run(p, f"报告时间：{audit.created_at.strftime('%Y年%m月%d日')}",
             size=12, color=RGBColor(100, 100, 100))

    doc.add_page_break()

    # -- 1. 基本信息 --
    doc.add_heading("基本信息", level=1)

    info_data = [
        ("审计标题", audit.title),
        ("审计时间", audit.created_at.strftime("%Y-%m-%d %H:%M:%S")),
        ("使用模型", audit.model_used or "-"),
        ("代码语言", audit.language),
        ("风险等级", SEVERITY_MAP.get(audit.severity, "未知")),
    ]
    table = doc.add_table(rows=len(info_data), cols=2, style="Table Grid")
    for i, (label, value) in enumerate(info_data):
        _set_cell(table.cell(i, 0), label, bold=True)
        _set_cell(table.cell(i, 1), value)
        table.cell(i, 0).width = Cm(4)
        table.cell(i, 1).width = Cm(12)

    doc.add_paragraph()

    # -- 2. 审计依据 --
    doc.add_heading("审计依据", level=1)

    standards = [
        "GB/T 39412-2020 信息安全技术 代码安全审计规范",
        "OWASP Top 10 2021",
        "CWE（Common Weakness Enumeration）",
    ]
    lang = (audit.language or "").lower()
    if "java" in lang:
        standards.insert(1, "GB/T 34944-2017 Java语言源代码漏洞测试规范")
    elif "c++" in lang or lang in ("c", "c/c++"):
        standards.insert(1, "GB/T 34943-2017 C/C++语言源代码漏洞测试规范")
    elif "c#" in lang:
        standards.insert(1, "GB/T 34946-2017 C#语言源代码漏洞测试规范")
    for s in standards:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_paragraph()

    # -- 3. 总体评估 --
    doc.add_heading("总体评估", level=1)
    assessment = _extract_section(audit.result or "", "总体评估")
    if assessment:
        for line in assessment.split("\n"):
            line = line.strip()
            if not line:
                continue
            line = re.sub(r"\*\*(.*?)\*\*", r"\1", line)
            doc.add_paragraph(line)
    else:
        doc.add_paragraph("请查看详细审计结果。")

    doc.add_paragraph()

    # -- 4. 漏洞详情 --
    doc.add_heading("漏洞详情", level=1)
    vulns = parse_vulnerabilities(audit.result or "")

    if vulns:
        # 按严重程度分组
        severity_groups = {"高危": [], "中危": [], "低危": [], "信息": []}
        for v in vulns:
            sev = v.get("severity", "")
            if "严重" in sev or "高" in sev:
                severity_groups["高危"].append(v)
            elif "中" in sev:
                severity_groups["中危"].append(v)
            elif "低" in sev:
                severity_groups["低危"].append(v)
            else:
                severity_groups["信息"].append(v)

        sec = 1
        for group_name, group_vulns in severity_groups.items():
            if not group_vulns:
                continue
            doc.add_heading(f"4.{sec} {group_name}漏洞", level=2)

            for vi, v in enumerate(group_vulns, 1):
                vuln_title = v.get("title", f"漏洞 {vi}")
                doc.add_heading(f"4.{sec}.{vi} {vuln_title}", level=3)

                # 漏洞信息表
                rows = []
                if v.get("category"):
                    rows.append(("漏洞分类", v["category"]))
                rows.append(("风险等级", v.get("severity", "-")))
                rows.append(("问题位置", v.get("location", "-")))

                t = doc.add_table(rows=len(rows), cols=2, style="Table Grid")
                for ri, (lbl, val) in enumerate(rows):
                    _set_cell(t.cell(ri, 0), lbl, bold=True)
                    _set_cell(t.cell(ri, 1), val)
                    t.cell(ri, 0).width = Cm(4)
                    t.cell(ri, 1).width = Cm(12)
                doc.add_paragraph()

                # 问题描述
                if v.get("description"):
                    p = doc.add_paragraph()
                    _add_run(p, "问题描述：", bold=True)
                    _add_run(p, v["description"])

                # 详细分析
                if v.get("analysis"):
                    p = doc.add_paragraph()
                    _add_run(p, "详细分析：", bold=True)
                    _add_run(p, v["analysis"])

                # 修复建议
                if v.get("recommendation"):
                    p = doc.add_paragraph()
                    _add_run(p, "修复建议：", bold=True)
                    _add_run(p, v["recommendation"])

                doc.add_paragraph()
            sec += 1
    else:
        # 无法解析结构化数据时，将原始结果写入
        doc.add_paragraph("以下为审计原始结果：")
        doc.add_paragraph()
        for line in (audit.result or "无结果").split("\n"):
            line = line.strip()
            if not line:
                continue
            clean = re.sub(r"\*\*(.*?)\*\*", r"\1", line)
            if line.startswith("## "):
                doc.add_heading(clean[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(clean[4:], level=3)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(clean[2:], style="List Bullet")
            else:
                doc.add_paragraph(clean)

    # -- 5. 安全建议 --
    doc.add_heading("安全建议", level=1)
    suggestions = _extract_section(audit.result or "", "安全建议")
    if suggestions:
        for line in suggestions.split("\n"):
            line = line.strip()
            if not line:
                continue
            line = re.sub(r"\*\*(.*?)\*\*", r"\1", line)
            if line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif re.match(r"^\d+\.\s", line):
                doc.add_paragraph(re.sub(r"^\d+\.\s*", "", line), style="List Number")
            else:
                doc.add_paragraph(line)
    else:
        doc.add_paragraph("请参考审计结果中的建议。")

    # -- 保存到内存 --
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
#  Excel (.xlsx) 导出
# ---------------------------------------------------------------------------

def export_xlsx(audit) -> bytes:
    """生成 Excel (.xlsx) 漏洞清单，参照 DeepAudit 导出格式。"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

    wb = Workbook()
    ws = wb.active
    ws.title = "漏洞清单"

    # -- 样式 --
    header_font = Font(name="微软雅黑", bold=True, size=11, color="FFFFFF")
    header_fill = PatternFill(start_color="2F5496", end_color="2F5496", fill_type="solid")
    header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    cell_font = Font(name="微软雅黑", size=10)
    cell_align = Alignment(vertical="top", wrap_text=True)
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )
    severity_fills = {
        "严重": ("C00000", "FFFFFF"), "极高": ("C00000", "FFFFFF"),
        "高":   ("FF4D6A", "FFFFFF"), "高风险": ("FF4D6A", "FFFFFF"),
        "中":   ("F0C040", "000000"), "中风险": ("F0C040", "000000"),
        "低":   ("92D050", "000000"), "低风险": ("92D050", "000000"),
        "信息": ("BDD7EE", "000000"), "安全":   ("BDD7EE", "000000"),
    }

    # -- 标题行 --
    title_cell = ws.cell(row=1, column=1, value="代码安全审计 - 漏洞清单")
    title_cell.font = Font(name="微软雅黑", bold=True, size=14)
    title_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.merge_cells("A1:H1")

    # -- 项目信息 --
    ts = audit.created_at.strftime("%Y-%m-%d %H:%M:%S")
    info_rows = [
        f"审计标题: {audit.title}",
        f"审计时间: {ts}    使用模型: {audit.model_used or '-'}    代码语言: {audit.language}",
    ]
    for i, text in enumerate(info_rows, 2):
        ws.cell(row=i, column=1, value=text).font = Font(name="微软雅黑", size=10)
        ws.merge_cells(start_row=i, start_column=1, end_row=i, end_column=8)

    # -- 表头 --
    header_row = len(info_rows) + 3
    headers = ["序号", "漏洞名称", "漏洞分类", "风险等级", "漏洞位置", "问题描述", "危害分析", "修复建议"]
    col_widths = [6, 25, 20, 10, 35, 40, 40, 45]

    for ci, (h, w) in enumerate(zip(headers, col_widths), 1):
        cell = ws.cell(row=header_row, column=ci, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = thin_border
        ws.column_dimensions[cell.column_letter].width = w

    # -- 漏洞数据 --
    vulns = parse_vulnerabilities(audit.result or "")

    if not vulns:
        cell = ws.cell(row=header_row + 1, column=1, value="未发现安全问题或无法解析审计结果")
        ws.merge_cells(start_row=header_row + 1, start_column=1,
                       end_row=header_row + 1, end_column=len(headers))
        cell.font = cell_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
    else:
        for ri, v in enumerate(vulns, 1):
            row = header_row + ri
            row_data = [
                ri,
                v.get("title", ""),
                v.get("category", ""),
                v.get("severity", ""),
                v.get("location", ""),
                v.get("description", ""),
                v.get("analysis", ""),
                v.get("recommendation", ""),
            ]
            for ci, val in enumerate(row_data, 1):
                cell = ws.cell(row=row, column=ci, value=val)
                cell.font = cell_font
                cell.alignment = cell_align
                cell.border = thin_border

                # 风险等级着色
                if ci == 4:
                    sev_text = str(val)
                    for key, (bg, fg) in severity_fills.items():
                        if key in sev_text:
                            cell.fill = PatternFill(start_color=bg, end_color=bg, fill_type="solid")
                            cell.font = Font(name="微软雅黑", size=10, bold=True, color=fg)
                            break

    # -- 审计概要 sheet --
    ws2 = wb.create_sheet("审计概要")
    summary = [
        ("审计标题", audit.title),
        ("审计时间", ts),
        ("使用模型", audit.model_used or "-"),
        ("代码语言", audit.language),
        ("风险等级", SEVERITY_MAP.get(audit.severity, "未知")),
        ("漏洞总数", str(len(vulns))),
    ]
    ws2.column_dimensions["A"].width = 15
    ws2.column_dimensions["B"].width = 40
    for ri, (label, value) in enumerate(summary, 1):
        c1 = ws2.cell(row=ri, column=1, value=label)
        c1.font = Font(name="微软雅黑", bold=True, size=10)
        c1.border = thin_border
        c2 = ws2.cell(row=ri, column=2, value=value)
        c2.font = Font(name="微软雅黑", size=10)
        c2.border = thin_border

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
